import os
import shutil
import sqlite3
import sys
import threading
from datetime import datetime

import pandas as pd
import win32com.client as win32
from PyQt5 import QtWidgets, QtCore, QtGui
from PyQt5.QtCore import QCoreApplication, QPropertyAnimation, QRect, QSize, QUrl, QSequentialAnimationGroup
from PyQt5.QtCore import QDate, QDateTime, QThread, QSortFilterProxyModel
from PyQt5.QtCore import Qt, QTimer
from PyQt5.QtGui import QPixmap, QColor, QDesktopServices, QMovie, QIcon
from PyQt5.QtWidgets import QApplication, QWidget, QFileDialog
from PyQt5.QtWidgets import QDialog, QLabel
from PyQt5.QtWidgets import QFileSystemModel, QGraphicsDropShadowEffect
from PyQt5.uic import loadUi
from docx import Document
from pandas import *
from docxcompose.composer import *
from docxtpl import DocxTemplate
# from goto import goto, label  # optional, for linter purpose
#from goto import with_goto
#from goto import goto, label
from workalendar.europe import Russia
from openpyxl.workbook import *
from time import sleep
from openpyxl import load_workbook
from openpyxl import *
from openpyxl.utils import get_column_letter

# from win32api import GetSystemMetrics

global user
global loginUse
global FullFioLogin
global FindNumGU
global TipDop
global Deadline
global DeadlineDay
global ZY
global version
global UPR
global MsgError
global pkud
global sqllong2
global UseTableGU
global ALLUser
version = "v.0.0.56"
FindNumGU = "*"
# BaseFrom1C = 'YSourceGitHub/DB/BaseFrom1C.db'
# BaseFrom1C = '/Portal/UseFile/BaseFrom1C.db'
GUFolders = "SourceGitHub/UI/login.ui"
if os.path.exists(GUFolders):
    ZY = "Z"
else:
    ZY = "Y"

class LoadScreen(QWidget, threading.Thread):
    def __init__(self):
        global ZY

        super(LoadScreen, self).__init__()
        loadUi("SourceGitHub/UI/Loading.ui", self)

        self.setFixedSize(5760, 1080)
        # self.setAttribute(Qt.WA_TranslucentBackground, True)
        self.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.CustomizeWindowHint)
        self.label_animation = QLabel(self)
        # self.setGeometry(-1920, 0, 5760, 1080)
        self.label_animation.move(2750, 400)
        self.label_animation.setScaledContents(True)
        self.label_animation.setFixedSize(100, 100)
        self.label_animation.setStyleSheet("background-color: transparent")

        try:
            self.movie = QMovie("C:/Portal/UseFile/Pictures/XOsX.gif")
        except:
            self.movie = QMovie("SourceGitHub/Pictures/XOsX.gif")

        self.label_animation.setMovie(self.movie)
        # self.movie.start()

        self.timer = QTimer(self)
        self.startAnimation()
        self.timer.singleShot(2000, self.stopAnimation)

        self.show()

    def startAnimation(self):
        self.movie.start()

    def stopAnimation(self):
        create = LoginScreen()
        widget.addWidget(create)
        widget.setCurrentIndex(widget.currentIndex() + 1)
        widget.setGeometry(-1920, 0, 5760, 1080)
        widget.show()
        Load = 2
        # self.movie.stop()


class LoginScreen(QDialog):
    def __init__(self):
        global version
        global UPR
        global ALLUser
        super(LoginScreen, self).__init__()
        # self.loading_screen = LoadScreen()
        loadUi("SourceGitHub/UI/login.ui", self)
        self.resize(500, 400)
        self.setWindowFlags(Qt.FramelessWindowHint)
        self.passwordfield.setEchoMode(QtWidgets.QLineEdit.Password)
        self.login.clicked.connect(self.loginfunction)
        self.create.clicked.connect(self.gotocreate)
        self.UpdaterB.clicked.connect(self.DownloadUpdate)
        self.bgwidget.clicked = False
        self.version.setText(version)
        self.bgwidget.move(2570, 200)
        self.UpdaterB.hide()
        self.show()

        icon = QtGui.QIcon()
        icon.addPixmap(QtGui.QPixmap("SourceGitHub/Pictures/IconP.ico"), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        widget.setWindowIcon(icon)

        conn = sqlite3.connect("SourceGitHub/DB/shop_data.db")
        cur = conn.cursor()
        query = 'SELECT Version FROM Settings'
        cur.execute(query)
        NewVersion = cur.fetchone()
        query = 'SELECT Download FROM Settings'
        cur.execute(query)

        if NewVersion[0] != version:
            self.UpdaterB.show()
            self.login.setEnabled(False)

        #Download = cur.fetchone()
        #if NewVersion[0] != version:
            #self.Link.setText('<a href=' + Download + '>Скачать новую версию</a>')
            #self.Link.setOpenExternalLinks(True)
            #self.login.setEnabled(False)

        try:
            connl = sqlite3.connect("/Portal/UseFile/Settings.db")
            curl = connl.cursor()
            query = 'SELECT LastLogin FROM SettingsUser'
            curl.execute(query)
            LastLogin = curl.fetchone()
            self.emailfield.setText(LastLogin[0])
        except:
            print("Нужно скачать Sittings.db")

        self.login.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.create.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.emailfield.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=1, yOffset=2))
        self.passwordfield.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=1, yOffset=2))

        self.Closeer.clicked.connect(QCoreApplication.instance().quit)
        self.Mini.clicked.connect(self.Minamal)



    def DownloadUpdate(self):  # Проверить и создать папку
        GUFolder = "SourceGitHub/Актуальная версия Программы"
        if os.path.exists(GUFolder):
            os.startfile(GUFolder)
        else:
            os.startfile(GUFolder)

    def Minamal(self):
        widget.showMinimized()

    def mousePressEvent(self, event):
        self.old_pos = event.screenPos()

    def mouseMoveEvent(self, event):
        if self.bgwidget.clicked:
            dx = self.old_pos.x() - event.screenPos().x()
            dy = self.old_pos.y() - event.screenPos().y()
            self.move(self.pos().x() - dx, self.pos().y() - dy)
        self.old_pos = event.screenPos()
        self.bgwidget.clicked = True

        # return QWidgets.mouseMoveEvent(self.bgwidget, event)

    def gotocreate(self):
        create = CreateAccScreen()
        widget.addWidget(create)
        widget.setCurrentIndex(widget.currentIndex() + 1)

    def loginfunction(self):
        global loginUser
        global TipDop
        global UPR
        global user
        global ALLUser

        user = self.emailfield.text()
        loginUser = user
        password = self.passwordfield.text()

        if len(user) == 0 or len(password) == 0:
            self.error.setText("Пустые поля")

        else:
            conn = sqlite3.connect("SourceGitHub/DB/shop_data.db")
            cur = conn.cursor()
            query = 'SELECT password FROM Login WHERE username =\'' + user + "\'"
            cur.execute(query)
            result_pass = cur.fetchone()[0]
            Tip = 'SELECT tipdop FROM Login WHERE username =\'' + user + "\'"
            cur.execute(Tip)
            result_tip = cur.fetchone()[0]
            TipDop = result_tip
            TakeUPR = 'SELECT otdel FROM Login WHERE username =\'' + user + "\'"
            cur.execute(TakeUPR)
            UPR = cur.fetchone()[0]
            PatchNoteRes = 'SELECT PatchNote FROM Login WHERE username =\'' + user + "\'"
            cur.execute(PatchNoteRes)
            result_PatchNoteRes = cur.fetchone()[0]
            login = self.emailfield.text()
            DownloadDB = 'SELECT DownloadDB FROM Login WHERE username =\'' + user + "\'"
            cur.execute(DownloadDB)
            DownloadDB = cur.fetchone()
            cur.execute('SELECT firstname, myname, lastname FROM Login WHERE otdel = ?', (UPR,))
            ALLUser = cur.fetchall()



            try:
                connl = sqlite3.connect("/Portal/UseFile/Settings.db")
                curl = connl.cursor()
                curl.execute('UPDATE SettingsUser SET LastLogin = ?', (login,))
                connl.commit()
            except:
                print("Нужно скачать Settings.db")


            if result_PatchNoteRes is None or result_PatchNoteRes == "No" or result_PatchNoteRes == "":
                cur.execute("UPDATE Login SET PatchNote = ? WHERE username = ?", ("Yes", user,))
                conn.commit()
                if result_pass == password:
                    GUFolder = "C:/Portal"
                    if os.path.exists(GUFolder):
                        What = "Папка уже создана"
                    else:
                        os.mkdir(GUFolder)
                    GUFolder = "C:/Portal/UseFile"
                    if os.path.exists(GUFolder):
                        What = "Папка уже создана"
                    else:
                        os.mkdir(GUFolder)
                    GUFolder = "C:/Portal/UseFile/Pictures"
                    if os.path.exists(GUFolder):
                        What = "Папка уже создана"
                    else:
                        os.mkdir(GUFolder)
                    #cur.execute("UPDATE Login SET DownloadDB = ? WHERE username = ?", ("Обновил", user,))
                    conn.commit()
                    if DownloadDB[0] != "Обновил":
                        shutil.copy(r'SourceGitHub/DB/BaseFrom1C.db',
                                    r'C:/Portal/UseFile/BaseFrom1C.db')
                    FindSettings = shutil.which(r'C:/Portal/UseFile/Settings.db')
                    if FindSettings == None:
                        shutil.copy(r'SourceGitHub/DB/Settings.db',
                                    r'C:/Portal/UseFile/Settings.db')

                    FindArrow = shutil.which(r'C:/Portal/UseFile/Pictures/25623.png')
                    if FindArrow == None:
                        shutil.copy(r'SourceGitHub/Pictures/25623.png',
                                    r'C:/Portal/UseFile/Pictures/25623.png')

                    FindGus = shutil.which(r'C:/Portal/UseFile/Pictures/XOsX.gif')
                    if FindGus == None:
                        shutil.copy(r'SourceGitHub/Pictures/XOsX.gif',
                                    r'C:/Portal/UseFile/Pictures/XOsX.gif')


                    create = PatchNote()
                    widget.addWidget(create)
                    widget.setCurrentIndex(widget.currentIndex() + 1)
                    widget.setGeometry(-1920, 0, 5760, 1080)
                    widget.show()
                    self.error.setText("")
                else:
                    self.error.setText("Неправильный логин или пароль")
            else:
                if result_pass == password:
                    GUFolder = "C:/Portal"
                    if os.path.exists(GUFolder):
                        What = "Папка уже создана"
                    else:
                        os.mkdir(GUFolder)
                    GUFolder = "C:/Portal/UseFile"
                    if os.path.exists(GUFolder):
                        What = "Папка уже создана"
                    else:
                        os.mkdir(GUFolder)
                    GUFolder = "C:/Portal/UseFile/Pictures"
                    if os.path.exists(GUFolder):
                        What = "Папка уже создана"
                    else:
                        os.mkdir(GUFolder)
                    #cur.execute("UPDATE Login SET DownloadDB = ? WHERE username = ?", ("Обновил", user,))
                    conn.commit()
                    if DownloadDB[0] != "Обновил":
                        shutil.copy(r'SourceGitHub/DB/BaseFrom1C.db',
                                    r'C:/Portal/UseFile/BaseFrom1C.db')
                    FindSettings = shutil.which(r'C:/Portal/UseFile/Settings.db')
                    if FindSettings == None:
                        shutil.copy(r'SourceGitHub/DB/Settings.db',
                                    r'C:/Portal/UseFile/Settings.db')
                    FindArrow = shutil.which(r'C:/Portal/UseFile/Pictures/25623.png')
                    if FindArrow == None:
                        shutil.copy(r'SourceGitHub/Pictures/25623.png',
                                    r'C:/Portal/UseFile/Pictures/25623.png')

                    FindGus = shutil.which(r'C:/Portal/UseFile/Pictures/XOsX.gif')
                    if FindGus == None:
                        shutil.copy(r'SourceGitHub/Pictures/XOsX.gif',
                                    r'C:/Portal/UseFile/Pictures/XOsX.gif')

                    create = MainWindow()
                    widget.addWidget(create)
                    widget.setCurrentIndex(widget.currentIndex() + 1)
                    widget.setGeometry(-1920, 0, 5760, 1080)
                    widget.show()
                    self.error.setText("")

                else:
                    self.error.setText("Неправильный логин или пароль")
            conn.commit()

    def prescript(self):
        create = Loading()
        widget.addWidget(create)
        widget.setCurrentIndex(widget.currentIndex() + 1)
        widget.setGeometry(-1920, 0, 5760, 1080)
        widget.show()


class Loading(QWidget, threading.Thread):
    def __init__(self):
        super(Loading, self).__init__()
        # self.setAttribute(Qt.WA_TranslucentBackground, True)
        self.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.CustomizeWindowHint)
        loadUi("SourceGitHub/UI/Loading2.ui", self)
        # self.label_animation.move(2420, 300)
        # self.label_animation.setStyleSheet("background-color: transparent")
        self.movie = QMovie("/Portal/UseFile/Gear.gif")
        self.label_animation.setMovie(self.movie)
        self.movie.start()
        self.show()


class CreateAccScreen(QDialog):
    def __init__(self):
        super(CreateAccScreen, self).__init__()
        loadUi("SourceGitHub/UI/createacc.ui", self)
        widget.setFixedHeight(700)
        widget.setFixedWidth(600)
        self.passwordfield.setEchoMode(QtWidgets.QLineEdit.Password)
        self.confirmpasswordfield.setEchoMode(QtWidgets.QLineEdit.Password)
        self.signup.clicked.connect(self.signupfunction)

    def signupfunction(self):
        global user

        user = self.emailfield.text()
        password = self.passwordfield.text()
        confirmpassword = self.confirmpasswordfield.text()

        if len(user) == 0 or len(password) == 0 or len(confirmpassword) == 0:
            self.error.setText("Please fill in all inputs.")

        elif password != confirmpassword:
            self.error.setText("Passwords do not match.")
        else:
            conn = sqlite3.connect("SourceGitHub/DB/shop_data.db")
            cur = conn.cursor()

            user_info = [user, password]
            cur.execute('INSERT INTO Login (username, password) VALUES (?,?)', user_info)

            conn.commit()
            conn.close()

            fillprofile = FillProfileScreen()
            widget.addWidget(fillprofile)
            widget.setCurrentIndex(widget.currentIndex() + 1)


class FillProfileScreen(QDialog):
    def __init__(self):
        super(FillProfileScreen, self).__init__()
        loadUi("SourceGitHub/UI/fillprofile.ui", self)
        widget.setFixedHeight(600)
        widget.setFixedWidth(1000)
        # widget.setGeometry(100, 100)

        self.image.setPixmap(QPixmap('placeholder.png'))
        self.registration.clicked.connect(self.signupfunction)
        self.registration.clicked.connect(self.gotologin)

    def gotologin(self):
        create = LoginScreen()
        widget.addWidget(create)
        widget.setFixedHeight(700)
        widget.setFixedWidth(600)
        widget.setCurrentIndex(widget.currentIndex() + 1)

    def signupfunction(self):
        global user
        myname = self.username.text()
        firstname = self.firstname.text()
        lastname = self.lastname.text()
        birthday = self.birthday.text()
        pol = self.pol.currentText()
        email = self.email.text()
        upravlenie = self.upravlenie.currentText()
        otdel = self.otdel.currentText()
        if self.Rb1.isChecked() == True:
            tipDop = "1"
        elif self.Rb2.isChecked() == True:
            tipDop = "2"
        elif self.Rb3.isChecked() == True:
            tipDop = "3"
        elif self.Rb4.isChecked() == True:
            tipDop = "4"
        elif self.Rb5.isChecked() == True:
            tipDop = "5"
        elif self.Rb6.isChecked() == True:
            tipDop = "6"
        elif self.Rb7.isChecked() == True:
            tipDop = "7"
        elif self.Rb8.isChecked() == True:
            tipDop = "8"
        ips = myname

        conn = sqlite3.connect("SourceGitHub/DB/shop_data.db")
        cur = conn.cursor()
        user_update = [myname]
        cur.execute(
            "UPDATE Login SET myname = ?, firstname = ?, lastname = ?, birthday = ?, pol = ?, email = ?, upr = ?, otdel = ?, tipdop = ?  WHERE username = ?",
            (myname, firstname, lastname, birthday, pol, email, upravlenie, otdel, tipDop, user))
        # cur.execute("UPDATE login_info SET firstname = ? WHERE username = ?", (firstname, user))

        conn.commit()
        conn.close()

        fillprofile = FillProfileScreen()
        widget.addWidget(fillprofile)
        widget.setCurrentIndex(widget.currentIndex() + 1)


class EditButtonsWidget(QtWidgets.QWidget):  # Создание своего аиджета с двумя кнопками

    def __init__(self, parent=None):
        super(EditButtonsWidget, self).__init__(parent)

        # add your buttons
        layout = QtWidgets.QVBoxLayout()

        # adjust spacings to your needs
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # add your buttons
        layout.addWidget(QtWidgets.QPushButton('Сохранить'))
        layout.addWidget(QtWidgets.QPushButton('Отменить'))
        self.setLayout(layout)


def gotoGroup1GU(self):
    create = CreateGroup1GU()
    create.exec_()


class CreateAlarm(QDialog):
    def __init__(self):
        global MsgError
        super(CreateAlarm, self).__init__()
        self.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.CustomizeWindowHint)
        #self.setAttribute(Qt.WA_TranslucentBackground, True)
        loadUi("SourceGitHub/UI/Alarm.ui", self)
        self.MsgErrorText.setText(MsgError)
        self.movie = QMovie("SourceGitHub/Pictures/Darg1.gif")
        self.ErrorDrag1.setMovie(self.movie)
        self.movie.start()
        self.movie = QMovie("SourceGitHub/Pictures/Darg2.gif")
        self.ErrorDrag2.setMovie(self.movie)
        self.movie.start()
        self.ButtonErrorNext.clicked.connect(self.CloseError)

    def CloseError(self):
        self.close()

class CreateWhat(QDialog):
    def __init__(self):
        super(CreateWhat, self).__init__()
        self.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.CustomizeWindowHint)

        loadUi("SourceGitHub/UI/ExcelDownload.ui", self)
        self.NewPortalExlB.clicked.connect(self.NewPortalExl)
        self.NewPortalExlB2.clicked.connect(self.NewPortalExl2)
        self.NewPortalExlB3.clicked.connect(self.NewPortalExl3)
        self.NewPortalExlB4.clicked.connect(self.NewPortalExl4)
        self.NewPortalExlB5.clicked.connect(self.NewPortalExl5)
        self.NewPortalExlB6.clicked.connect(self.NewPortalExl6)
        self.NewPortalExlB7.clicked.connect(self.NewPortalExl7)

        self.OldPortalExlB.clicked.connect(self.OldPortalExl)
        self.Closeer.clicked.connect(self.DoCloseer)

    def NewPortalExl(self):
        try:
            datenow = QDate.currentDate().toPyDate()
            datenow = str(datenow)
            i = 1
            VigFile = 'SourceGitHub/Выгрузки/' + datenow
            if os.path.exists(VigFile):
                What = "Папка уже создана"
            else:
                os.mkdir(VigFile)
            connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
            # reo = connection.cursor()
            query = "SELECT * FROM GU1Group"
            df = pd.read_sql(query, connection)
            VigFolder = VigFile
            VigFile = VigFile + "/" + "(Новый портал) Все ГУ " + "(v."+ str(i) + ")" + ".xlsx"
            writer = pd.ExcelWriter(VigFile)
            df.to_excel(writer, sheet_name='All_GU', index=False)
            for column in df:
                column_width = max(df[column].astype(str).map(len).max(), len(column))
                col_idx = df.columns.get_loc(column)
                writer.sheets['All_GU'].set_column(col_idx, col_idx, column_width)

            writer.save()
            os.startfile(VigFolder)
            connection.commit()
            connection.close()
            self.close()
        except:
            i = i + 1
            self.NewPortalExl()
    def NewPortalExl2(self):
        try:
            datenow = QDate.currentDate().toPyDate()
            datenow = str(datenow)
            i = 1
            VigFile = 'SourceGitHub/Выгрузки/' + datenow
            if os.path.exists(VigFile):
                What = "Папка уже создана"
            else:
                os.mkdir(VigFile)
            connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
            # reo = connection.cursor()
            query = 'SELECT ТипГУ, Регистрационныйномер, Датарегистрации, РегистрационныйномерАЛвРАЛ, Наименованиезаявителя, ' \
                    'Ответственныйисполнитель, СтатусГУ, Датавыбораэксперта, ФИОЭкспертапоаккредитации, НаименованиеЭО, ' \
                    'ПредложениепосоставуЭГ, ДатаприказаосоставеЭГ, Договор1Г, ДатаЭЗ, ВыводЭЗ, ДатаПриказа2гр, ДатаАкта, ' \
                    'ВыводАкта, ДатаПриказа3гр, СопроводИтог FROM GU1Group WHERE ' \
                    '(СтатусГУ != "Возврат без рассмотрения" AND СтатусГУ != "Отзыв ГУ" AND ' \
                    'СтатусГУ != "Отказ ГУ (договор)" AND СтатусГУ != "Приказ об отказе") AND ' \
                    '(ОтделУП = "Отдел аккредитации испытательных лабораторий" OR ОтделУП = "Отдел аккредитации в отдельных сферах")'
            df = pd.read_sql(query, connection)
            VigFolder = VigFile
            VigFile = VigFile + "/" + "(Сокр. Новый портал) Все ГУ " + "(v."+ str(i) + ")" + ".xlsx"
            writer = pd.ExcelWriter(VigFile)
            df.to_excel(writer, sheet_name='All_GU', startrow=1, header=False, index=False)
            workbook = writer.book
            worksheet = writer.sheets['All_GU']
            column_settings = [{'header': column} for column in df.columns]
            (max_row, max_col) = df.shape
            worksheet.set_column('A:C', 15)
            worksheet.set_column('D:F', 25)
            worksheet.set_column('G:H', 15)
            worksheet.set_column('I:L', 25)
            worksheet.set_column('M:T', 15)

            CellsDate = [2, 7, 12, 13, 15, 16, 18]
            for ii in range(len(CellsDate)):
                for i in range(len(df)):
                    try:
                        DateText = df.iat[i, CellsDate[ii]]
                        print(DateText)
                        date_time = datetime.strptime(DateText, '%d.%m.%Y')
                        date_format  = workbook.add_format({'num_format': 'dd.mm.yyyy'})
                        worksheet.write_datetime(i + 1, CellsDate[ii], date_time, date_format)
                    except:
                        pass

            worksheet.add_table(0, 0, max_row, max_col - 1, {'columns': column_settings})
            #worksheet.autofilter(0, 0, max_row, max_col - 1)
            writer.save()
            os.startfile(VigFolder)
            connection.commit()
            connection.close()
            self.close()
        except:
            i = i + 1
            self.NewPortalExl2()

    def NewPortalExl3(self):
        try:
            datenow = QDate.currentDate().toPyDate()
            datenow = str(datenow)
            i = 1
            VigFile = 'SourceGitHub/Выгрузки/' + datenow
            if os.path.exists(VigFile):
                What = "Папка уже создана"
            else:
                os.mkdir(VigFile)
            connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
            # reo = connection.cursor()
            query = 'SELECT ТипГУ, Регистрационныйномер, Датарегистрации, РегистрационныйномерАЛвРАЛ, Наименованиезаявителя, ' \
                    'Ответственныйисполнитель, СтатусГУ, Датавыбораэксперта, ФИОЭкспертапоаккредитации, НаименованиеЭО, ' \
                    'ПредложениепосоставуЭГ, ДатаприказаосоставеЭГ, Договор1Г, ДатаЭЗ, ВыводЭЗ, ДатаПриказа2гр, ДатаАкта, ' \
                    'ВыводАкта, ДатаПриказа3гр, СопроводИтог FROM GU1Group WHERE ' \
                    '(СтатусГУ != "Возврат без рассмотрения" AND СтатусГУ != "Отзыв ГУ" AND ' \
                    'СтатусГУ != "Отказ ГУ (договор)" AND СтатусГУ != "Приказ об отказе") AND ' \
                    '(ОтделУП = "Отдел аккредитации испытательных лабораторий" OR ОтделУП = "Отдел аккредитации в отдельных сферах")' \
                    'AND (ДатаприказаосоставеЭГ IS NULL OR ДатаприказаосоставеЭГ = "Не зарегистрирован")'
            df = pd.read_sql(query, connection)
            VigFolder = VigFile
            VigFile = VigFile + "/" + "(Новый портал) 1 Группа " + "(v."+ str(i) + ")" + ".xlsx"
            writer = pd.ExcelWriter(VigFile)
            df.to_excel(writer, sheet_name='All_GU', startrow=1, header=False, index=False)
            workbook = writer.book
            worksheet = writer.sheets['All_GU']
            column_settings = [{'header': column} for column in df.columns]
            (max_row, max_col) = df.shape
            worksheet.set_column('A:C', 15)
            worksheet.set_column('D:F', 25)
            worksheet.set_column('G:H', 15)
            worksheet.set_column('I:L', 25)
            worksheet.set_column('M:T', 15)

            CellsDate = [2, 7, 12, 13, 15, 16, 18]
            for ii in range(len(CellsDate)):
                for i in range(len(df)):
                    try:
                        DateText = df.iat[i, CellsDate[ii]]
                        print(DateText)
                        date_time = datetime.strptime(DateText, '%d.%m.%Y')
                        date_format  = workbook.add_format({'num_format': 'dd.mm.yyyy'})
                        worksheet.write_datetime(i + 1, CellsDate[ii], date_time, date_format)
                    except:
                        pass

            worksheet.add_table(0, 0, max_row, max_col - 1, {'columns': column_settings})
            #worksheet.autofilter(0, 0, max_row, max_col - 1)
            writer.save()
            os.startfile(VigFolder)
            connection.commit()
            connection.close()
            self.close()
        except:
            i = i + 1
            self.NewPortalExl3()
    def NewPortalExl4(self):
        try:
            datenow = QDate.currentDate().toPyDate()
            datenow = str(datenow)
            i = 1
            VigFile = 'SourceGitHub/Выгрузки/' + datenow
            if os.path.exists(VigFile):
                What = "Папка уже создана"
            else:
                os.mkdir(VigFile)
            connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
            # reo = connection.cursor()
            query = 'SELECT ТипГУ, Регистрационныйномер, Датарегистрации, РегистрационныйномерАЛвРАЛ, Наименованиезаявителя, ' \
                    'Ответственныйисполнитель, СтатусГУ, Датавыбораэксперта, ФИОЭкспертапоаккредитации, НаименованиеЭО, ' \
                    'ПредложениепосоставуЭГ, ДатаприказаосоставеЭГ, Договор1Г, ДатаЭЗ, ВыводЭЗ, ДатаПриказа2гр, ДатаАкта, ' \
                    'ВыводАкта, ДатаПриказа3гр, СопроводИтог FROM GU1Group WHERE ' \
                    '(СтатусГУ != "Возврат без рассмотрения" AND СтатусГУ != "Отзыв ГУ" AND ' \
                    'СтатусГУ != "Отказ ГУ (договор)" AND СтатусГУ != "Приказ об отказе") AND ' \
                    '(ОтделУП = "Отдел аккредитации испытательных лабораторий" OR ОтделУП = "Отдел аккредитации в отдельных сферах")' \
                    'AND ДатаприказаосоставеЭГ IS NOT NULL AND (Договор1Г IS NULL OR Договор1Г = "")'
            df = pd.read_sql(query, connection)
            VigFolder = VigFile
            VigFile = VigFile + "/" + "(Новый портал) Ждем договор " + "(v."+ str(i) + ")" + ".xlsx"
            writer = pd.ExcelWriter(VigFile)
            df.to_excel(writer, sheet_name='All_GU', startrow=1, header=False, index=False)
            workbook = writer.book
            worksheet = writer.sheets['All_GU']
            column_settings = [{'header': column} for column in df.columns]
            (max_row, max_col) = df.shape
            worksheet.set_column('A:C', 15)
            worksheet.set_column('D:F', 25)
            worksheet.set_column('G:H', 15)
            worksheet.set_column('I:L', 25)
            worksheet.set_column('M:T', 15)

            CellsDate = [2, 7, 12, 13, 15, 16, 18]
            for ii in range(len(CellsDate)):
                for i in range(len(df)):
                    try:
                        DateText = df.iat[i, CellsDate[ii]]
                        print(DateText)
                        date_time = datetime.strptime(DateText, '%d.%m.%Y')
                        date_format  = workbook.add_format({'num_format': 'dd.mm.yyyy'})
                        worksheet.write_datetime(i + 1, CellsDate[ii], date_time, date_format)
                    except:
                        pass

            worksheet.add_table(0, 0, max_row, max_col - 1, {'columns': column_settings})
            #worksheet.autofilter(0, 0, max_row, max_col - 1)
            writer.save()
            os.startfile(VigFolder)
            connection.commit()
            connection.close()
            self.close()
        except:
            i = i + 1
            self.NewPortalExl4()
    def NewPortalExl5(self):
        try:
            datenow = QDate.currentDate().toPyDate()
            datenow = str(datenow)
            i = 1
            VigFile = 'SourceGitHub/Выгрузки/' + datenow
            if os.path.exists(VigFile):
                What = "Папка уже создана"
            else:
                os.mkdir(VigFile)
            connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
            # reo = connection.cursor()
            query = 'SELECT ТипГУ, Регистрационныйномер, Датарегистрации, РегистрационныйномерАЛвРАЛ, Наименованиезаявителя, ' \
                    'Ответственныйисполнитель, СтатусГУ, Датавыбораэксперта, ФИОЭкспертапоаккредитации, НаименованиеЭО, ' \
                    'ПредложениепосоставуЭГ, ДатаприказаосоставеЭГ, Договор1Г, ДатаЭЗ, ВыводЭЗ, ДатаПриказа2гр, ДатаАкта, ' \
                    'ВыводАкта, ДатаПриказа3гр, СопроводИтог FROM GU1Group WHERE ' \
                    '(СтатусГУ != "Возврат без рассмотрения" AND СтатусГУ != "Отзыв ГУ" AND ' \
                    'СтатусГУ != "Отказ ГУ (договор)" AND СтатусГУ != "Приказ об отказе") AND ' \
                    '(ОтделУП = "Отдел аккредитации испытательных лабораторий" OR ОтделУП = "Отдел аккредитации в отдельных сферах")' \
                    'AND (ТипГУ = "АК" OR ТипГУ = "РОА") AND (Договор1Г IS NOT NULL AND Договор1Г != "") ' \
                    'AND (ДатаПриказа2гр = "" OR ДатаПриказа2гр IS NULL)'
            df = pd.read_sql(query, connection)
            VigFolder = VigFile
            VigFile = VigFile + "/" + "(Новый портал) 2 Группа " + "(v."+ str(i) + ")" + ".xlsx"
            writer = pd.ExcelWriter(VigFile)
            df.to_excel(writer, sheet_name='All_GU', startrow=1, header=False, index=False)
            workbook = writer.book
            worksheet = writer.sheets['All_GU']
            column_settings = [{'header': column} for column in df.columns]
            (max_row, max_col) = df.shape
            worksheet.set_column('A:C', 15)
            worksheet.set_column('D:F', 25)
            worksheet.set_column('G:H', 15)
            worksheet.set_column('I:L', 25)
            worksheet.set_column('M:T', 15)

            CellsDate = [2, 7, 12, 13, 15, 16, 18]
            for ii in range(len(CellsDate)):
                for i in range(len(df)):
                    try:
                        DateText = df.iat[i, CellsDate[ii]]
                        print(DateText)
                        date_time = datetime.strptime(DateText, '%d.%m.%Y')
                        date_format  = workbook.add_format({'num_format': 'dd.mm.yyyy'})
                        worksheet.write_datetime(i + 1, CellsDate[ii], date_time, date_format)
                    except:
                        pass

            worksheet.add_table(0, 0, max_row, max_col - 1, {'columns': column_settings})
            #worksheet.autofilter(0, 0, max_row, max_col - 1)
            writer.save()
            os.startfile(VigFolder)
            connection.commit()
            connection.close()
            self.close()
        except:
            i = i + 1
            self.NewPortalExl5()
    def NewPortalExl6(self):
        try:
            datenow = QDate.currentDate().toPyDate()
            datenow = str(datenow)
            i = 1
            VigFile = 'SourceGitHub/Выгрузки/' + datenow
            if os.path.exists(VigFile):
                What = "Папка уже создана"
            else:
                os.mkdir(VigFile)
            connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
            # reo = connection.cursor()
            query = 'SELECT ТипГУ, Регистрационныйномер, Датарегистрации, РегистрационныйномерАЛвРАЛ, Наименованиезаявителя, ' \
                    'Исполнитель3Гр, СтатусГУ, Датавыбораэксперта, ФИОЭкспертапоаккредитации, НаименованиеЭО, ' \
                    'ПредложениепосоставуЭГ, ДатаприказаосоставеЭГ, Договор1Г, ДатаЭЗ, ВыводЭЗ, ДатаПриказа2гр, ДатаАкта, ' \
                    'ВыводАкта, ДатаПриказа3гр, СопроводИтог FROM GU1Group WHERE ((Исполнитель3Гр != "" AND ' \
                    'Исполнитель3Гр is not NULL)) AND ((ДатаПриказа3гр = "" OR ДатаПриказа3гр is NULL) OR ' \
                    '(ДатаПриказа3гр != "" AND ДатаПриказа3гр is not NULL AND ТУ = "Центральный федеральный округ" AND ' \
                    'СопроводИтог = "") OR (ПриостановкаАКилиРОА = "Да"))'
            df = pd.read_sql(query, connection)
            VigFolder = VigFile
            VigFile = VigFile + "/" + "(Новый портал) 3 Группа " + "(v."+ str(i) + ")" + ".xlsx"
            writer = pd.ExcelWriter(VigFile)
            df.to_excel(writer, sheet_name='All_GU', startrow=1, header=False, index=False)
            workbook = writer.book
            worksheet = writer.sheets['All_GU']
            column_settings = [{'header': column} for column in df.columns]
            (max_row, max_col) = df.shape
            worksheet.set_column('A:C', 15)
            worksheet.set_column('D:F', 25)
            worksheet.set_column('G:H', 15)
            worksheet.set_column('I:L', 25)
            worksheet.set_column('M:T', 15)

            CellsDate = [2, 7, 12, 13, 15, 16, 18]
            for ii in range(len(CellsDate)):
                for i in range(len(df)):
                    try:
                        DateText = df.iat[i, CellsDate[ii]]
                        print(DateText)
                        date_time = datetime.strptime(DateText, '%d.%m.%Y')
                        date_format  = workbook.add_format({'num_format': 'dd.mm.yyyy'})
                        worksheet.write_datetime(i + 1, CellsDate[ii], date_time, date_format)
                    except:
                        pass

            worksheet.add_table(0, 0, max_row, max_col - 1, {'columns': column_settings})
            #worksheet.autofilter(0, 0, max_row, max_col - 1)
            writer.save()
            os.startfile(VigFolder)
            connection.commit()
            connection.close()
            self.close()
        except:
            i = i + 1
            self.NewPortalExl6()
    def NewPortalExl7(self):
        try:
            datenow = QDate.currentDate().toPyDate()
            datenow = str(datenow)
            i = 1
            VigFile = 'SourceGitHub/Выгрузки/' + datenow
            if os.path.exists(VigFile):
                What = "Папка уже создана"
            else:
                os.mkdir(VigFile)
            connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
            # reo = connection.cursor()
            query = 'SELECT ТипГУ, Регистрационныйномер, Датарегистрации, РегистрационныйномерАЛвРАЛ, Наименованиезаявителя, ' \
                    'Исполнитель3Гр, СтатусГУ, Датавыбораэксперта, ФИОЭкспертапоаккредитации, НаименованиеЭО, ' \
                    'ПредложениепосоставуЭГ, ДатаприказаосоставеЭГ, Договор1Г, ДатаЭЗ, ВыводЭЗ, ДатаПриказа2гр, ДатаАкта, ' \
                    'ВыводАкта, ДатаПриказа3гр, СопроводИтог FROM GU1Group WHERE (ДатаПриказа3гр != "" AND ДатаПриказа3гр is not NULL)'
            df = pd.read_sql(query, connection)
            VigFolder = VigFile
            VigFile = VigFile + "/" + "(Новый портал) Завершенные " + "(v."+ str(i) + ")" + ".xlsx"
            writer = pd.ExcelWriter(VigFile)
            df.to_excel(writer, sheet_name='All_GU', startrow=1, header=False, index=False)
            workbook = writer.book
            worksheet = writer.sheets['All_GU']
            column_settings = [{'header': column} for column in df.columns]
            (max_row, max_col) = df.shape
            worksheet.set_column('A:C', 15)
            worksheet.set_column('D:F', 25)
            worksheet.set_column('G:H', 15)
            worksheet.set_column('I:L', 25)
            worksheet.set_column('M:T', 15)

            CellsDate = [2, 7, 12, 13, 15, 16, 18]
            for ii in range(len(CellsDate)):
                for i in range(len(df)):
                    try:
                        DateText = df.iat[i, CellsDate[ii]]
                        print(DateText)
                        date_time = datetime.strptime(DateText, '%d.%m.%Y')
                        date_format  = workbook.add_format({'num_format': 'dd.mm.yyyy'})
                        worksheet.write_datetime(i + 1, CellsDate[ii], date_time, date_format)
                    except:
                        pass

            worksheet.add_table(0, 0, max_row, max_col - 1, {'columns': column_settings})
            #worksheet.autofilter(0, 0, max_row, max_col - 1)
            writer.save()
            os.startfile(VigFolder)
            connection.commit()
            connection.close()
            self.close()
        except:
            i = i + 1
            self.NewPortalExl7()

    def OldPortalExl(self):
        try:
            datenow = QDate.currentDate().toPyDate()
            datenow = str(datenow)
            i = 1
            VigFile = 'SourceGitHub/Выгрузки/' + datenow
            if os.path.exists(VigFile):
                What = "Папка уже создана"
            else:
                os.mkdir(VigFile)
            connection = sqlite3.connect('SourceGitHub/DB/PortalKu.db')
            # reo = connection.cursor()
            query = "SELECT * FROM ALLIN"
            df = pd.read_sql(query, connection)
            VigFolder = VigFile
            VigFile = VigFile + "/" + "(Cтарый портал) Все ГУ " + "(v."+ str(i) + ")" + ".xlsx"
            writer = pd.ExcelWriter(VigFile)
            df.to_excel(writer, sheet_name='All_GU', startrow=1, header=False, index=False)
            workbook = writer.book
            worksheet = writer.sheets['All_GU']
            column_settings = [{'header': column} for column in df.columns]
            (max_row, max_col) = df.shape
            #for column in df:
                #column_width = max(df[column].astype(str).map(len).max(), len(column))
                #col_idx = df.columns.get_loc(column)
                #writer.sheets['All_GU'].set_column(col_idx, col_idx, column_width)
            worksheet.set_column('A:E', 12)
            worksheet.set_column('F:F', 25)
            worksheet.set_column('G:H', 20)
            worksheet.set_column('I:P', 12)
            worksheet.set_column('Q:Q', 20)
            worksheet.set_column('R:R', 40)

            CellsDate = [4, 8, 9, 10, 11, 13, 14]
            for ii in range(len(CellsDate)):
                for i in range(len(df)):
                    try:
                        DateText = df.iat[i, CellsDate[ii]]
                        print(DateText)
                        date_time = datetime.strptime(DateText, '%d.%m.%Y')
                        date_format = workbook.add_format({'num_format': 'dd.mm.yyyy'})
                        worksheet.write_datetime(i + 1, CellsDate[ii], date_time, date_format)
                    except:
                        pass



            worksheet.add_table(0, 0, max_row, max_col - 1, {'columns': column_settings})
            writer.save()
            os.startfile(VigFolder)
            connection.commit()
            connection.close()
            self.close()
        except:
            i = i + 1
            self.OldPortalExl()

    def DoCloseer(self):
        self.close()








class CreateGroup1GU(QDialog):

    #@with_goto
    def __init__(self):
        connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
        cur = connection.cursor()
        global z, EmailExpert
        global globalNomerGU
        global FindNumGU
        global Deadline
        global DeadlineDay
        global TipDop
        global FullNaimKogo
        global UPR
        global ALLUser
        global UseTableGU

        FindGU = globalNomerGU
        cur.execute("SELECT * FROM GU1Group WHERE Регистрационныйномер = ?", (FindGU,))
        result = cur.fetchall()
        connection.commit()
        connection.close()

        super(CreateGroup1GU, self).__init__()
        loadUi("SourceGitHub/UI/Group1GU.ui", self)


        self.OpenGU.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.SaveALL.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.openRAL.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Ispolnitel2.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.StatusGU.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.PriznakSK.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.NomerGU2.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.TipGU_3.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Comment.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.DataRospIsp.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.DeadLine.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.FioExpert2.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.DataRospExp.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Perevibor.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.DatePredlog.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.PereviborDate.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.DataPricaza.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.DogovorCheck.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Url1CGU.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.URLCloud.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Open1C.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.OpenCloud.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.treeView.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.TipGU_12.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.NomerGU_8.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.NomerGU_9.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.EmailRAL.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.ShortName_2.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.YrAdres.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.FullNaim.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.NomerGU_19.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.NomerGU_18.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.NomerGU_20.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.dateEdit_23.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.NomerGU.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.NomerGU.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.ObshObl.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.pushButton_3.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.FioExpert.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.EmailExpert.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.pushButton_4.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        #self.DataYvedT.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.REO.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.MailREO.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        #self.DataYvedEGT.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.NumberTeh.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.DataPricaz1GR.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.NomerPricaz1Gr.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.SrokD.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.YvedEG.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.FioExpert_5.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.CreatePricaz1Group.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.FioExpert_3.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Teh_1.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.EoTeh_1.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Teh_2.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.EoTeh_2.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Teh_3.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.EoTeh_3.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Teh_4.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.EoTeh_4.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Teh_5.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.EoTeh_5.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Teh_6.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.EoTeh_6.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Teh_7.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.EoTeh_7.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Teh_8.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.EoTeh_8.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Teh_9.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.EoTeh_9.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Teh_10.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.EoTeh_10.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.TipGU.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.VozvratCheck.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.DogovorCheck.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Vibor_TU_2.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.KommentPricaz.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Vibor_TU.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Rucovod_TU.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.OpenWiki.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Oblast_Button.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Oblast_1.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Oblast_2.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Oblast_3.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Oblast_4.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Oblast_5.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Oblast_6.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Oblast_7.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Oblast_8.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Oblast_9.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Oblast_10.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Oblast_11.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Oblast_12.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Oblast_13.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Oblast_14.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Oblast_Close.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Oblast_Close.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Priastanovka.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.OtzivCheck.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.PcazObOtcaze.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.OtcazCheck.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Comment.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Comment_2.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Comment_3.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))


        self.Oblasti.hide()
        self.Oblast_Button.clicked.connect(self.TakeOblast)
        self.VozvratCheck.clicked.connect(self.DoVozvrat)

        NomerGU = result[0][1]
        DataGU = result[0][2]
        FullGU = NomerGU + " от " + DataGU
        self.NomerGU.setPlainText(FullGU)

        TipGU = result[0][0]
        if TipGU == "Подтверждение компетентности":
            self.TipGU.setCurrentText("ПК")
            self.TipGU_3.setCurrentText("ПК")
        elif TipGU == "Расширение области аккредитации":
            self.TipGU.setCurrentText("РОА")
            self.TipGU_3.setCurrentText("РОА")
        elif TipGU == "Аккредитация":
            self.TipGU.setCurrentText("АК")
            self.TipGU_3.setCurrentText("АК")
        else:
            self.TipGU.setCurrentText(TipGU)
            self.TipGU_3.setCurrentText(TipGU)

        ShortName = result[0][4]
        self.ShortName_2.setPlainText(ShortName)
        self.ShortName.setPlainText(ShortName)
        # self.FullNaim.setPlainText(result[0][5])

        NomerRAL = result[0][3]
        self.NomerRAL.setPlainText(NomerRAL)

        #Deadline = result[0][81]
        #Deadline = datetime.strptime(Deadline, '%d.%m.%Y')
        #self.DeadLine.setDate(Deadline)

        # Дата росписи ГУ на Испонителя
        if result[0][14] != None:
            DataRospIsp = result[0][14]
            DataRospIsp = DataRospIsp[:10]
            Data_ISP_date = datetime.strptime(DataRospIsp, '%d.%m.%Y')
            self.DataRospIsp.setDate(Data_ISP_date)
        # Дата выбора эксперта
        if result[0][18] == None:
            self.DataRospExp.setStyleSheet(
                'background-color: rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; color: rgb(255, 255, 255);')
        else:
            DataRospExp = result[0][18]
            DataRospExp = DataRospExp[:10]
            DataRospExp = datetime.strptime(DataRospExp, '%d.%m.%Y')
            self.DataRospExp.setDate(DataRospExp)
        self.DataRospExp.dateChanged.connect(self.DataRospExpChanged)
        # Дата предложения
        if result[0][24] == None:
            self.DatePredlog.setStyleSheet(
                'background-color: rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; color: rgb(255, 255, 255);')
        else:
            DatePredl = result[0][24]
            try:
                DatePredlS = DatePredl.split()
                DatePredlS = DatePredlS[2]
                DatePredlS = DatePredlS[:10]
                DatePredlS = datetime.strptime(DatePredlS, '%d.%m.%Y')
                self.DatePredlog.setDate(DatePredlS)
            except:
                if DatePredl != "Не зарегистрировано":
                    DatePredl = datetime.strptime(DatePredl, '%d.%m.%Y')
                    self.DatePredlog.setDate(DatePredl)
                elif DatePredl == "Не зарегистрировано":
                    datenow = QDate.currentDate().toPyDate()
                    self.DatePredlog.setDate(datenow)

        self.DatePredlog.dateChanged.connect(self.DatePredlogChanged)
        # Дата итгового приказа 1 Группа
        if result[0][25] == None or result[0][25] == "Не зарегистрирован":
            self.DataPricaza.setStyleSheet(
                'background-color: rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; color: rgb(255, 255, 255);')

        else:
            try:
                DataPricaza = result[0][25]
                DataPricaza = DataPricaza.split()
                DataPricaza = DataPricaza[2]
                DataPricaza = DataPricaza[:10]
                DataPricaza = datetime.strptime(DataPricaza, '%d.%m.%Y')
                self.DataPricaza.setDate(DataPricaza)
            except:
                DataPricaza = result[0][25]
                DataPricaza = datetime.strptime(DataPricaza, '%d.%m.%Y')
                self.DataPricaza.setDate(DataPricaza)
        self.DataPricaza.dateChanged.connect(self.DataPricazaChanged)

        # данные из РАЛ
        connection = sqlite3.connect('/Portal/UseFile/BaseFrom1C.db')
        cur = connection.cursor()
        cur.execute("SELECT * FROM RalReestr WHERE Номераттестатааккредитации = ?", (NomerRAL,))
        resultRAL = cur.fetchall()
        connection.commit()
        connection.close()
        # ЭЛ почта из РАЛ
        if resultRAL != []:
            self.EmailRAL.setPlainText(resultRAL[0][19])
        # Юридический адрес
        if resultRAL != []:
            self.YrAdres.setPlainText(resultRAL[0][27])

        # Тех. Эксперты
        connection = sqlite3.connect('/Portal/UseFile/BaseFrom1C.db')
        teh = connection.cursor()
        teh.execute('SELECT COUNT(0) FROM ReestTeh')  # Запрос в sql на кол-ва ячеек в базе данных
        tehlong = teh.fetchone()  # Находим длинну базы данных в кортеже
        tehlong2 = tehlong[0]  # Вытаскиваем длину баззы данных из первого картежа
        teh.execute("SELECT * FROM ReestTeh ")
        FullresultTeh = teh.fetchall()
        for i in range(tehlong2):
            self.Teh_1.addItems([FullresultTeh[i][2]])
            self.Teh_2.addItems([FullresultTeh[i][2]])
            self.Teh_3.addItems([FullresultTeh[i][2]])
            self.Teh_4.addItems([FullresultTeh[i][2]])
            self.Teh_5.addItems([FullresultTeh[i][2]])
            self.Teh_6.addItems([FullresultTeh[i][2]])
            self.Teh_7.addItems([FullresultTeh[i][2]])
            self.Teh_8.addItems([FullresultTeh[i][2]])
            self.Teh_9.addItems([FullresultTeh[i][2]])
            self.Teh_10.addItems([FullresultTeh[i][2]])
        self.Teh_1.currentIndexChanged[str].connect(self.ChangedTeh)
        self.Teh_2.currentIndexChanged[str].connect(self.ChangedTeh)
        self.Teh_3.currentIndexChanged[str].connect(self.ChangedTeh)
        self.Teh_4.currentIndexChanged[str].connect(self.ChangedTeh)
        self.Teh_5.currentIndexChanged[str].connect(self.ChangedTeh)
        self.Teh_6.currentIndexChanged[str].connect(self.ChangedTeh)
        self.Teh_7.currentIndexChanged[str].connect(self.ChangedTeh)
        self.Teh_8.currentIndexChanged[str].connect(self.ChangedTeh)
        self.Teh_9.currentIndexChanged[str].connect(self.ChangedTeh)
        self.Teh_10.currentIndexChanged[str].connect(self.ChangedTeh)

        self.TehSlider.valueChanged[int].connect(self.TehSliderChanged)
        self.NumberTeh.currentIndexChanged.connect(self.TehNumberChanged)

        Number = 1
        if Number == 1:
            self.Teh_1.show()
            self.Teh_2.hide()
            self.Teh_3.hide()
            self.Teh_4.hide()
            self.Teh_5.hide()
            self.Teh_6.hide()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.hide()
            self.Teh_10.hide()
            self.EoTeh_1.show()
            self.EoTeh_2.hide()
            self.EoTeh_3.hide()
            self.EoTeh_4.hide()
            self.EoTeh_5.hide()
            self.EoTeh_6.hide()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.hide()
            self.EoTeh_10.hide()
        elif Number == 2:
            self.Teh_1.hide()
            self.Teh_2.show()
            self.Teh_3.hide()
            self.Teh_4.hide()
            self.Teh_5.hide()
            self.Teh_6.hide()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.hide()
            self.Teh_10.hide()
            self.EoTeh_1.hide()
            self.EoTeh_2.show()
            self.EoTeh_3.hide()
            self.EoTeh_4.hide()
            self.EoTeh_5.hide()
            self.EoTeh_6.hide()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.hide()
            self.EoTeh_10.hide()
        elif Number == 3:
            self.Teh_1.hide()
            self.Teh_2.hide()
            self.Teh_3.show()
            self.Teh_4.hide()
            self.Teh_5.hide()
            self.Teh_6.hide()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.hide()
            self.Teh_10.hide()
            self.EoTeh_1.hide()
            self.EoTeh_2.hide()
            self.EoTeh_3.show()
            self.EoTeh_4.hide()
            self.EoTeh_5.hide()
            self.EoTeh_6.hide()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.hide()
            self.EoTeh_10.hide()
        elif Number == 4:
            self.Teh_1.hide()
            self.Teh_2.hide()
            self.Teh_3.hide()
            self.Teh_4.show()
            self.Teh_5.hide()
            self.Teh_6.hide()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.hide()
            self.Teh_10.hide()
            self.EoTeh_1.hide()
            self.EoTeh_2.hide()
            self.EoTeh_3.hide()
            self.EoTeh_4.show()
            self.EoTeh_5.hide()
            self.EoTeh_6.hide()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.hide()
            self.EoTeh_10.hide()
        elif Number == 5:
            self.Teh_1.hide()
            self.Teh_2.hide()
            self.Teh_3.hide()
            self.Teh_4.hide()
            self.Teh_5.show()
            self.Teh_6.hide()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.hide()
            self.Teh_10.hide()
            self.EoTeh_1.hide()
            self.EoTeh_2.hide()
            self.EoTeh_3.hide()
            self.EoTeh_4.hide()
            self.EoTeh_5.show()
            self.EoTeh_6.hide()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.hide()
            self.EoTeh_10.hide()
        elif Number == 6:
            self.Teh_1.hide()
            self.Teh_2.hide()
            self.Teh_3.hide()
            self.Teh_4.hide()
            self.Teh_5.hide()
            self.Teh_6.show()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.hide()
            self.Teh_10.hide()
            self.EoTeh_1.hide()
            self.EoTeh_2.hide()
            self.EoTeh_3.hide()
            self.EoTeh_4.hide()
            self.EoTeh_5.hide()
            self.EoTeh_6.show()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.hide()
            self.EoTeh_10.hide()
        elif Number == 7:
            self.Teh_1.hide()
            self.Teh_2.hide()
            self.Teh_3.hide()
            self.Teh_4.hide()
            self.Teh_5.hide()
            self.Teh_6.show()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.hide()
            self.Teh_10.hide()
            self.EoTeh_1.hide()
            self.EoTeh_2.hide()
            self.EoTeh_3.hide()
            self.EoTeh_4.hide()
            self.EoTeh_5.hide()
            self.EoTeh_6.show()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.hide()
            self.EoTeh_10.hide()
        elif Number == 8:
            self.Teh_1.hide()
            self.Teh_2.hide()
            self.Teh_3.hide()
            self.Teh_4.hide()
            self.Teh_5.hide()
            self.Teh_6.hide()
            self.Teh_7.hide()
            self.Teh_8.show()
            self.Teh_9.hide()
            self.Teh_10.hide()
            self.EoTeh_1.hide()
            self.EoTeh_2.hide()
            self.EoTeh_3.hide()
            self.EoTeh_4.hide()
            self.EoTeh_5.hide()
            self.EoTeh_6.hide()
            self.EoTeh_7.hide()
            self.EoTeh_8.show()
            self.EoTeh_9.hide()
            self.EoTeh_10.hide()
        elif Number == 9:
            self.Teh_1.hide()
            self.Teh_2.hide()
            self.Teh_3.hide()
            self.Teh_4.hide()
            self.Teh_5.hide()
            self.Teh_6.hide()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.show()
            self.Teh_10.hide()
            self.EoTeh_1.hide()
            self.EoTeh_2.hide()
            self.EoTeh_3.hide()
            self.EoTeh_4.hide()
            self.EoTeh_5.hide()
            self.EoTeh_6.hide()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.show()
            self.EoTeh_10.hide()
        elif Number == 10:
            self.Teh_1.hide()
            self.Teh_2.hide()
            self.Teh_3.hide()
            self.Teh_4.hide()
            self.Teh_5.hide()
            self.Teh_6.hide()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.hide()
            self.Teh_10.show()
            self.EoTeh_1.hide()
            self.EoTeh_2.hide()
            self.EoTeh_3.hide()
            self.EoTeh_4.hide()
            self.EoTeh_5.hide()
            self.EoTeh_6.hide()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.hide()
            self.EoTeh_10.show()

        # _____________________

        NomerGU2 = FullGU
        self.NomerGU2.setPlainText(NomerGU2)

        #self.Perevibor.currentIndexChanged[str].connect(self.HidePerevibor)
        if self.Perevibor.currentText() == "Нет":
            self.PereviborDate.hide()
            self.label_PereviborDate.hide()

        '''else:
            self.PereviborDate.show()
            self.label_PereviborDate.show()'''

        Ispolnitel2 = result[0][7]
        connection = sqlite3.connect('/Portal/UseFile/BaseFrom1C.db')
        cur2 = connection.cursor()
        cur2.execute('SELECT DISTINCT Ответственныйисполнитель FROM TableFrom1C')
        resultIlsList = cur2.fetchall()
        cur2.execute('SELECT COUNT(DISTINCT Ответственныйисполнитель) FROM TableFrom1C')
        Colvo = cur2.fetchone()
        for i in range(Colvo[0] - 1):
            if resultIlsList[i][0] is not None:
                self.Ispolnitel2.addItems([resultIlsList[i][0]])
        self.Ispolnitel2.setCurrentText(Ispolnitel2)

        connection = sqlite3.connect('/Portal/UseFile/BaseFrom1C.db')
        exp = connection.cursor()
        exp.execute('SELECT COUNT(0) FROM Experts')  # Запрос в sql на кол-ва ячеек в базе данных
        sqllong = exp.fetchone()  # Находим длинну базы данных в кортеже
        sqllong2 = sqllong[0]  # Вытаскиваем длину баззы данных из первого картежа
        exp.execute("SELECT ФИОэкспертапоаккредитации FROM Experts")
        Expert = exp.fetchall()
        for i in range(sqllong2):
            self.FioExpert.addItems([Expert[i][0]])
            self.FioExpert2.addItems([Expert[i][0]])
        FioExpert = result[0][20]
        self.FioExpert.setCurrentText(FioExpert)
        self.FioExpert2.setCurrentText(FioExpert)
        exp.execute("SELECT * FROM Experts WHERE ФИОэкспертапоаккредитации = ?", (FioExpert,))
        OneExpert = exp.fetchall()
        if OneExpert == []:
            self.EmailExpert.setPlainText("-------")
            self.FioExpert.setCurrentText("Эксперт не выбран")
            self.FioExpert2.setCurrentText("Эксперт не выбран")
        else:
            EmailExpert = OneExpert[0][8]
            self.EmailExpert.setPlainText(EmailExpert)
        connection.commit()
        connection.close()

        if self.FioExpert.currentText() == "Эксперт не выбран":
            self.VozvratCheck.show()
        else:
            self.VozvratCheck.hide()
            #self.Priastanovka.move(210, 390)
        #if self.DataPricaza.text() != "01.01.2000":
            #self.Priastanovka.hide()



        connection = sqlite3.connect('/Portal/UseFile/BaseFrom1C.db')
        reo = connection.cursor()
        reo.execute('SELECT COUNT(0) FROM ReestrREO')  # Запрос в sql на кол-ва ячеек в базе данных
        reolong = reo.fetchone()  # Находим длинну базы данных в кортеже
        reolong2 = reolong[0]  # Вытаскиваем длину баззы данных из первого картежа
        reo.execute("SELECT * FROM ReestrREO ")
        FullresultREO = reo.fetchall()
        if OneExpert == []:
            self.FioExpert.setCurrentText("Эксперт не найден")
        else:
            reo.execute("SELECT * FROM ReestrREO WHERE Адресэлектроннойпочты = ?", (OneExpert[0][7],))
            resultREO = reo.fetchall()

            for i in range(reolong2):
                self.REO.addItems([FullresultREO[i][9]])
            if resultREO == []:
                REOreserv = OneExpert[0][5]
                MailREOreserv = OneExpert[0][8]
                self.REO.setCurrentText(REOreserv)
                self.MailREO.setPlainText(MailREOreserv)
            else:
                self.REO.setCurrentText(resultREO[0][9])
                self.MailREO.setPlainText(resultREO[0][12])
        connection.commit()
        connection.close()
        self.FioExpert.currentIndexChanged[str].connect(self.ChangedExp)
        self.REO.currentIndexChanged[str].connect(self.ChangedREO)

        # работа Word приказ 1 группа
        self.CreatePricaz1Group.clicked.connect(self.Create_Pricaz_1Group)
        self.CreatePricaz1Group_2.clicked.connect(self.CreatePolojGroup2)
        self.CreatePricaz1Group_3.clicked.connect(self.CreatePolojGroup2)
        self.CreatePricaz1Group_7.clicked.connect(self.CreateSzOtricGroup2)
        self.CreatePricaz3Group.clicked.connect(self.CreatePricaz3GR)
        self.CreateSZGroup3.clicked.connect(self.CreateSzPolojGroup3)
        self.CreatePricaz1Group_9.clicked.connect(self.CreateSoprPolojGroup3)
        #  Проверить и создать папку
        self.OpenGU.clicked.connect(self.OpenGUFolder)

        # ________________________________________
        # GUFolder = "C:/Users/Love/PycharmProjects/PortalFSA/ГУ/1 Группа/" + FullGU
        GUFolder = "SourceGitHub/ГУ/1 Группа/" + FullGU
        if os.path.exists(GUFolder):
            What = "Папка уже создана"
        else:
            os.mkdir(GUFolder)
        dirPath = r":\Управление Аккредитации\Программа\PortalFSA\ГУ\1 Группа" + r"\""
        # dirPath = r"C:\Users\Love\PycharmProjects\PortalFSA\ГУ\1 Группа" + r"\""
        dirPath = dirPath[:59] + FullGU
        # dirPath.replace(r'\\', r"\" )
        self.model = QFileSystemModel()
        self.model.setRootPath(dirPath)
        self.treeView.setModel(self.model)
        self.treeView.setRootIndex(self.model.index(dirPath))
        self.treeView.setColumnWidth(0, 250)
        self.treeView.customContextMenuRequested.connect(self.open_file)
        menu = QtWidgets.QMenu()
        open = menu.addAction("Open in new maya")
        open.triggered.connect(self.open_file)
        open_file = menu.addAction("Open file")
        self.treeView.clicked.connect(self.onClicked)
        self.SaveALL.clicked.connect(self.SaveALLS)  # кнопка сохранить
        self.YvedEG.clicked.connect(self.CreateYvedEG)  # кнопка Дата уведомленния ЭГ

        # Заполнение
        Find_z = result[0][5]
        if Find_z is not None:
            self.FullNaim.setPlainText(Find_z)
            self.FixFullNaim()
        Find_z = result[0][45]
        if Find_z is not None:
            self.EmailRAL.setPlainText(Find_z)
        Find_z = result[0][46]
        if Find_z is not None:
            self.YrAdres.setPlainText(Find_z)
        Find_z = result[0][20]
        if Find_z is not None:
            self.FioExpert.setCurrentText(Find_z)
        Find_z = result[0][47]
        if Find_z is not None:
            self.EmailExpert.setPlainText(Find_z)
        Find_z = result[0][48]
        if Find_z is not None:
            self.REO.setCurrentText(Find_z)
        Find_z = result[0][49]
        if Find_z is not None:
            self.MailREO.setPlainText(Find_z)
        Find_z = result[0][50]
        if Find_z is not None:
            self.NumberTeh.setCurrentText(Find_z)
        Find_z = result[0][51]
        if Find_z is not None:
            self.Teh_1.setCurrentText(Find_z)
        Find_z = result[0][52]
        if Find_z is not None:
            self.EoTeh_1.setPlainText(Find_z)
        Find_z = result[0][53]
        if Find_z is not None:
            self.Teh_2.setCurrentText(Find_z)
        Find_z = result[0][54]
        if Find_z is not None:
            self.EoTeh_2.setPlainText(Find_z)
        Find_z = result[0][55]
        if Find_z is not None:
            self.Teh_3.setCurrentText(Find_z)
        Find_z = result[0][56]
        if Find_z is not None:
            self.EoTeh_3.setPlainText(Find_z)
        Find_z = result[0][57]
        if Find_z is not None:
            self.Teh_4.setCurrentText(Find_z)
        Find_z = result[0][58]
        if Find_z is not None:
            self.EoTeh_4.setPlainText(Find_z)
        Find_z = result[0][59]
        if Find_z is not None:
            self.Teh_5.setCurrentText(Find_z)
        Find_z = result[0][60]
        if Find_z is not None:
            self.EoTeh_5.setPlainText(Find_z)
        Find_z = result[0][61]
        if Find_z is not None:
            self.Teh_6.setCurrentText(Find_z)
        Find_z = result[0][62]
        if Find_z is not None:
            self.EoTeh_6.setPlainText(Find_z)
        Find_z = result[0][63]
        if Find_z is not None:
            self.Teh_7.setCurrentText(Find_z)
        Find_z = result[0][64]
        if Find_z is not None:
            self.EoTeh_7.setPlainText(Find_z)
        Find_z = result[0][65]
        if Find_z is not None:
            self.Teh_8.setCurrentText(Find_z)
        Find_z = result[0][66]
        if Find_z is not None:
            self.EoTeh_8.setPlainText(Find_z)
        Find_z = result[0][67]
        if Find_z is not None:
            self.Teh_9.setCurrentText(Find_z)
        Find_z = result[0][68]
        if Find_z is not None:
            self.EoTeh_9.setPlainText(Find_z)
        Find_z = result[0][69]
        if Find_z is not None:
            self.Teh_10.setCurrentText(Find_z)
        Find_z = result[0][70]
        if Find_z is not None:
            self.EoTeh_10.setPlainText(Find_z)
        Find_z = result[0][71]
        if Find_z is not None:
            self.ObshObl.setCurrentText(Find_z)
        Find_z = result[0][72]
        if Find_z is not None or Find_z != "":
            self.DataYvedT.setText(Find_z)
        Find_z = result[0][73]
        if Find_z is not None:
            self.SrokD.setPlainText(Find_z)
        Find_z = result[0][74]
        if Find_z is not None:
            Data2 = datetime.strptime(Find_z, '%d.%m.%Y')
            self.DataPricaz1GR.setDate(Data2)
        Find_z = result[0][75]
        if Find_z is not None:
            self.NomerPricaz1Gr.setPlainText(Find_z)
        Find_z = result[0][76]
        if Find_z is not None:
            self.PriznakSK.setCurrentText(Find_z)
        Find_z = result[0][77]
        if Find_z is not None:
            self.Comment.setPlainText(Find_z)
        Find_z = result[0][78]
        if Find_z is not None:
            self.Perevibor.setCurrentText(Find_z)
        Find_z = result[0][79]
        try:
            if Find_z != "" or Find_z is not None:
                Data3 = datetime.strptime(Find_z, '%d.%m.%Y')
                self.PereviborDate.setDate(Data3)
        except:
            print("проблема с PereviborDate")
        Find_z = result[0][10]
        if Find_z is not None:
            self.StatusGU.setCurrentText(Find_z)
        Find_z = result[0][91]
        if Find_z is not None:
            self.Vibor_TU.setCurrentText(Find_z)
        Find_z = result[0][92]
        if Find_z is not None:
            self.Rucovod_TU.setPlainText(Find_z)
        Find_z = result[0][80]
        if Find_z is not None or Find_z != "":
            self.DataYvedEGT.setText(Find_z)
        Find_z = result[0][5]
        if Find_z == "Не отправлен" or Find_z is None or Find_z == "":
            self.DogovorCheck.setText("Сопровод не отправлен")
            self.DogovorCheck.setStyleSheet("background-color: rgb(255, 90, 78)")
        elif Find_z == "Отправлен":
            self.DogovorCheck.setText("Сопровод отправлен")
            self.DogovorCheck.setStyleSheet("background-color: rgb(204, 255, 185)")
        Find_z = result[0][88]
        if Find_z is not None:
            self.Oblast_Button.setText(Find_z)
        Find_z = result[0][89]
        if Find_z is not None:
            self.Url1CGU.setPlainText(Find_z)
        Find_z = result[0][90]
        if Find_z is not None:
            self.URLCloud.setPlainText(Find_z)
        Find_z = result[0][93]
        if Find_z is None or Find_z == "Да":
            self.VKS.setChecked(True)
        elif Find_z == "Нет":
            self.VKS.setChecked(False)
        Find_z = result[0][94]
        if Find_z is None or Find_z == "Нет":
            self.VMeropr.setChecked(False)
        elif Find_z == "Да":
            self.VMeropr.setChecked(True)
        Find_z = result[0][107]
        if Find_z is not None:
            self.Comment_3.setPlainText(Find_z)
        Find_z = result[0][113]
        if Find_z is not None:
            self.Comment_2.setPlainText(Find_z)

        self.Vibor_TU_Changed()
        if UPR == "Отдел аккредитации испытательных лабораторий":
            self.ObshObl.setCurrentText("ИЛ")
            self.Oblast_Button.setText("")
        elif UPR == "Отдел аккредитации в отдельных сферах":
            self.ObshObl.setCurrentText("Метрология")

        if self.PereviborDate.text() == "01.01.2000":
            self.PereviborDate.setStyleSheet(
                'background-color: rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; color: rgb(255, 255, 255);')
        if self.DataPricaz1GR.text() == "01.01.2000":
            self.DataPricaz1GR.setStyleSheet(
                'background-color: rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; color: rgb(255, 255, 255);')

        self.DataYved.dateChanged.connect(self.DataYvedTChanged)
        self.DataYvedEG.dateChanged.connect(self.DataYvedEGTChanged)
        self.DataDogD.dateChanged.connect(self.DataDogChanged)
        self.PereviborDate.dateChanged.connect(self.PereviborDateChanged)
        self.DataPricaz1GR.dateChanged.connect(self.DataPricaz1GRChanged)
        self.DogovorCheck.clicked.connect(self.DogovorCheckB)
        if self.DataPricaza.text() == "01.01.2000":
            self.DogovorCheck.hide()
        else:
            self.DogovorCheck.show()

        self.openRAL.clicked.connect(self.OpenRAL)
        self.VozvratCheck_2.hide()
        self.VozvratCheck_3.hide()

        self.VozvratCheck_2.clicked.connect(self.VozvratAccept)
        self.VozvratCheck_3.clicked.connect(self.VozvratDenide)
        if result[0][85] == "":
            Viborg = "NO"
        else:
            Viborg = "YES"
        if result[0][85] is None:
            Viborg = "NO"
        else:
            Viborg = "YES"

        if Viborg == "YES":
            self.PereviborDate.show()

            self.VozvratCheck.hide()
            self.VozvratCheck_2.show()
            self.VozvratCheck_3.show()

            self.label_PereviborDate.show()
            self.label_PereviborDate.setText("Дата Возврата")
            # datenow2 = datetime.strptime(datenow, '%d.%m.%Y')
            dataPereviborDate = datetime.strptime(result[0][85], '%Y-%m-%d')
            self.PereviborDate.setDate(dataPereviborDate)
            self.label_32.hide()
            self.FioExpert2.hide()
            self.label_36.hide()
            self.DataRospExp.hide()
            self.label_35.hide()
            self.Perevibor.hide()
            self.label_33.hide()
            self.DatePredlog.hide()
            self.label_34.hide()
            self.DataPricaza.hide()
            self.label_31.hide()
            self.DeadLine.hide()
            self.Priastanovka.hide()

        if self.StatusGU.currentText() == "Приостановка":
            self.PereviborDate.hide()
            self.VozvratCheck_2.setText("Возобновить ГУ")
            self.VozvratCheck.hide()
            self.VozvratCheck_2.show()
            self.VozvratCheck_3.hide()
            self.label_PereviborDate.hide()
            self.label_32.hide()
            self.FioExpert2.hide()
            self.label_36.hide()
            self.DataRospExp.hide()
            self.label_35.hide()
            self.Perevibor.hide()
            self.label_33.hide()
            self.DatePredlog.hide()
            self.label_34.hide()
            self.DataPricaza.hide()
            self.label_31.hide()
            self.DeadLine.hide()
            self.Priastanovka.hide()
            self.DogovorCheck.hide()

        if self.StatusGU.currentText() == "Отзыв ГУ":
            self.PereviborDate.hide()
            self.VozvratCheck_2.setText("Подтвердить Отзыв")
            self.VozvratCheck_3.setText("Вернуть в работу")
            self.VozvratCheck.hide()
            self.VozvratCheck_2.show()
            self.VozvratCheck_3.show()
            self.label_PereviborDate.hide()
            self.label_32.hide()
            self.FioExpert2.hide()
            self.label_36.hide()
            self.DataRospExp.hide()
            self.label_35.hide()
            self.Perevibor.hide()
            self.label_33.hide()
            self.DatePredlog.hide()
            self.label_34.hide()
            self.DataPricaza.hide()
            self.label_31.hide()
            self.DeadLine.hide()
            self.Priastanovka.hide()
            self.DogovorCheck.hide()
            self.OtzivCheck.hide()
            self.PcazObOtcaze.hide()
            self.OtcazCheck.hide()
        if self.StatusGU.currentText() == "Отказ ГУ (договор)":
            self.PereviborDate.hide()
            self.VozvratCheck_2.setText("Подтвердить Отказ")
            self.VozvratCheck_3.setText("Вернуть в работу")
            self.VozvratCheck.hide()
            self.VozvratCheck_2.show()
            self.VozvratCheck_3.show()
            self.label_PereviborDate.hide()
            self.label_32.hide()
            self.FioExpert2.hide()
            self.label_36.hide()
            self.DataRospExp.hide()
            self.label_35.hide()
            self.Perevibor.hide()
            self.label_33.hide()
            self.DatePredlog.hide()
            self.label_34.hide()
            self.DataPricaza.hide()
            self.label_31.hide()
            self.DeadLine.hide()
            self.Priastanovka.hide()
            self.DogovorCheck.hide()
            self.OtzivCheck.hide()
            self.PcazObOtcaze.hide()
            self.OtcazCheck.hide()


        self.TipGU_3.currentIndexChanged[str].connect(self.TipChanged1)
        self.TipGU.currentIndexChanged[str].connect(self.TipChanged2)


        self.OpenWiki.clicked.connect(self.OpenWikiURL)
        self.Open1C.clicked.connect(self.Open1CURL)
        self.OpenCloud.clicked.connect(self.OpenCloudURL)


        self.Vibor_TU.currentIndexChanged[str].connect(self.Vibor_TUChanged)

        self.pushButton_3.clicked.connect(self.SendMailYved)
        self.pushButton_4.clicked.connect(self.SendMailYvedNapom)

        self.Vibor_TU.currentIndexChanged[str].connect(self.Vibor_TU_Changed)
        self.Vibor_TU_2.currentIndexChanged[str].connect(self.Vibor_TU_Changed2)

        self.ObshObl.currentIndexChanged[str].connect(self.ObshObl_Changed)

        self.Priastanovka.clicked.connect(self.Priastanovka_Check)

        if TipDop == 10:
            self.StatusGU.setEnabled(True)

        if TipGU == "ПК1":
            self.DO.setEnabled(True)
            self.VO.setEnabled(True)
            self.DO.setChecked(True)
        else:
            self.DO.setEnabled(False)
            self.VO.setEnabled(False)
        if TipGU == "ПК2":
            self.VO.setChecked(True)
        if TipGU == "ПК1":
            self.DO_VO.setChecked(False)
            self.DO_VO.setEnabled(False)
        else:
            self.DO_VO.setChecked(True)
            self.DO_VO.setEnabled(True)

        self.DO.toggled.connect(self.DO_Check)
        if self.DO.isChecked() == True:
            self.VKS.setChecked(False)
        else:
            self.VKS.setChecked(True)

        self.VO.toggled.connect(self.VO_Check)
        self.DO_VO.toggled.connect(self.DO_VO_Check)

        TipGU = result[0][0]
        if TipGU == "ПК1" or TipGU == "ПК1+ИМОД":
            self.SrokD.setPlainText("20")
        elif TipGU == "ПК2" or TipGU == "ПК2+ИМОД":
            self.SrokD.setPlainText("20")
        elif TipGU == "АК" or TipGU == "РОА":
            self.SrokD.setPlainText("33")
        else:
            self.SrokD.setPlainText("28")

        self.FixFullNaimKogo()
        Srokid = self.SrokD.toPlainText()
        if Srokid == "":
            Srokid = "________"
        self.KommentPricaz.setPlainText(
            self.EmailExpert.toPlainText() + " " + self.MailREO.toPlainText() + " " + self.EmailRAL.toPlainText() + " " + "\nОбращаем "
                                                                                                                          "внимание эксперта по аккредитации о необходимости направления в Росаккредитацию акта экспертизы и "
                                                                                                                          "документа, содержащего описание области аккредитации аккредитованного лица, прилагаемого к этому акту, "
                                                                                                                          "в электронном виде не позднее " + Srokid + " рабочих дней со дня направления эксперту по аккредитации "
                                                                                                                                                                      "прилагаемого приказа Росаккредитации о проведении процедуры подтверждения компетентности, расширения "
                                                                                                                                                                      "области аккредитации " + FullNaimKogo)
        self.label_save.hide()

        Dogovor = result[0][84]
        if Dogovor == "" or Dogovor is None:
            self.DogovorCheck.setText("Указать дату договора")
            self.DataDogT.hide()
            self.DataDogD.hide()
        elif Dogovor != "":
            self.DogovorCheck.setText("Договор подписан")
            self.DataDogT.show()
            self.DataDogD.show()

        self.DataDogT.setText(Dogovor)

        self.OtzivCheck.clicked.connect(self.OtzivCheck_Check)
        self.PcazObOtcaze.clicked.connect(self.PcazObOtcaze_Check)
        self.OtcazCheck.clicked.connect(self.OtcazCheck_Check)

        TakeVivodEZNew = result[0][97]
        TakeVivodEZ1C = result[0][26]
        if TakeVivodEZNew == "" or TakeVivodEZNew == None:
            if TakeVivodEZ1C != "" and TakeVivodEZ1C != None:
                TakeVivodEZ1C = TakeVivodEZ1C.split()
                try:
                    TakeVivodEZ1C = TakeVivodEZ1C[2].replace("г.", "")
                    self.DataEZT.setText(TakeVivodEZ1C)
                except:
                    self.DataEZT.setText("")
        else:
            self.DataEZT.setText(TakeVivodEZNew)

        self.VivodEZ.setCurrentText(result[0][98])
        self.Itog2Gr.setCurrentText(result[0][99])

        self.DataEZ.dateChanged.connect(self.DataEZChanged)
        self.DataPricaz2GR.dateChanged.connect(self.DataPricaz2GRChanged)

        DataPricaz2GRNew = result[0][100]
        DataPricaz2GR1Cbad = result[0][31]
        DataPricaz2GR1Cbad2 = result[0][44]
        DataPricaz2GR1Cgood = result[0][30]
        if DataPricaz2GRNew == "" or DataPricaz2GRNew == None:
            if DataPricaz2GR1Cgood != "" and DataPricaz2GR1Cgood != None:
                DataPricaz2GR1Cgood = DataPricaz2GR1Cgood.split()
                try:
                    DataPricaz2GR1Cgood = DataPricaz2GR1Cgood[2].replace("г.", "")
                    self.DataPricaz2GRT.setText(DataPricaz2GR1Cgood)
                except:
                    self.DataPricaz2GRT.setText("")
            elif DataPricaz2GR1Cbad != "" and DataPricaz2GR1Cbad != None:
                self.Itog2Gr.setCurrentText("Отриц. отказ")
                self.VivodEZ.setCurrentText("Отрицательный")
                DataPricaz2GR1Cbad = DataPricaz2GR1Cbad.split()
                try:
                    DataPricaz2GR1Cbad = DataPricaz2GR1Cbad[2].replace("г.", "")
                    self.DataPricaz2GRT.setText(DataPricaz2GR1Cbad)
                except:
                    self.DataPricaz2GRT.setText("")
            elif DataPricaz2GR1Cbad2 != "" and DataPricaz2GR1Cbad2 != None:
                self.Itog2Gr.setCurrentText("Отриц. отказ")
                self.VivodEZ.setCurrentText("Отрицательный")
                DataPricaz2GR1Cbad2 = DataPricaz2GR1Cbad2.split()
                try:
                    DataPricaz2GR1Cbad2 = DataPricaz2GR1Cbad2[2].replace("г.", "")
                    self.DataPricaz2GRT.setText(DataPricaz2GR1Cbad2)
                except:
                    self.DataPricaz2GRT.setText("")
        else:
            self.DataPricaz2GRT.setText(TakeVivodEZNew)

        self.LitsDoIt2GR.clicked.connect(self.Magic2GR)

        self.TextIn39.setPlainText(result[0][102])
        self.TextNoIn39.setPlainText(result[0][103])
        self.TextGost.setPlainText(result[0][104])
        self.ItogText2GR.setPlainText(result[0][105])
        self.ItogVivod.setPlainText(result[0][106])

        if self.VivodEZ.currentText() == "Положительный":
            self.groupBox_7.setEnabled(False)
            self.groupBox_8.setEnabled(False)
            self.groupBox_9.setEnabled(False)
            self.groupBox_20.setEnabled(False)
            self.groupBox_19.setEnabled(False)
            self.LitsDoIt2GR.setEnabled(False)
        elif self.VivodEZ.currentText() == "Отрицательный":
            self.groupBox_7.setEnabled(True)
            self.groupBox_8.setEnabled(True)
            self.groupBox_9.setEnabled(True)
            self.groupBox_20.setEnabled(True)
            self.groupBox_19.setEnabled(True)
            self.LitsDoIt2GR.setEnabled(True)

        if self.Itog2Gr.currentText() == "Не выбран":
            self.groupBox_4.setEnabled(False)
            self.groupBox_5.setEnabled(False)
            self.groupBox_6.setEnabled(False)
        elif self.Itog2Gr.currentText() == "Полож.":
            self.groupBox_4.setEnabled(True)
            self.groupBox_5.setEnabled(False)
            self.groupBox_6.setEnabled(False)
        elif self.Itog2Gr.currentText() == "Отриц. не отказ":
            self.groupBox_4.setEnabled(False)
            self.groupBox_5.setEnabled(True)
            self.groupBox_6.setEnabled(False)
        elif self.Itog2Gr.currentText() == "Отриц. отказ":
            self.groupBox_4.setEnabled(False)
            self.groupBox_5.setEnabled(False)
            self.groupBox_6.setEnabled(True)

        self.Itog2Gr.currentIndexChanged[str].connect(self.Itog2GrChange)

        if TipDop == 10:
            self.Ispolnitel3Gr.setEnabled(True)
            self.Ispolnitel2Gr.setEnabled(True)
            self.DataActaT.setEnabled(True)
            self.DataActa.setEnabled(True)
            self.VivodActa.setEnabled(True)
            self.Risk.setEnabled(True)

        ALLUserLong = len(ALLUser)
        for i in range(ALLUserLong):
            BigALLUser = ALLUser[i][0] + " " + ALLUser[i][1] + " " + ALLUser[i][2]
            self.Ispolnitel2Gr.addItems([BigALLUser])
            self.Ispolnitel3Gr.addItems([BigALLUser])



        self.Ispolnitel3Gr.setCurrentText(result[0][108])
        self.DataActaT.setText(result[0][111])
        self.VivodActa.setCurrentText(result[0][109])
        self.Risk.setCurrentText(result[0][110])
        self.DataActa.dateChanged.connect(self.DataActChanged)

        #self.Ispolnitel3Gr.currentIndexChanged[str].connect(self.Ispol2GrChange)

        if result[0][112] == "" or result[0][112] is None:
            self.Ispolnitel2Gr.setCurrentText(result[0][7])
        else:
            self.Ispolnitel2Gr.setCurrentText(result[0][112])

        if UseTableGU == "1 Группа":
            self.Comment.show()
            self.Comment_2.hide()
            self.Comment_3.hide()
        elif UseTableGU == "2 Группа":
            self.Comment.hide()
            self.Comment_2.show()
            self.Comment_3.hide()
        elif UseTableGU == "3 Группа":
            self.Comment.hide()
            self.Comment_2.hide()
            self.Comment_3.show()

        if result[0][0] != "РОА" and result[0][0] != "АК":
            self.tabWidget.setTabVisible(3, False)
            self.DataPricaz2GRT.hide()
            self.DataPricaz2GR.hide()
            self.label_37.hide()
        self.SaveGif.hide()

        self.Etap3GR.setCurrentText(result[0][114])
        self.Etap3GR.currentIndexChanged[str].connect(self.Etap3GRChange)
        self.ShortName_3.setPlainText(result[0][4])
        self.ShortName_3.textChanged.connect(self.ShortNameChanged)
        self.FullNaim_2.setPlainText(result[0][5])
        self.FullNaim_2.textChanged.connect(self.FullNaimChanged)
        self.EmailRAL_2.setPlainText(result[0][45])
        self.EmailRAL_2.textChanged.connect(self.EmailRALChanged)
        self.YrAdres_2.setPlainText(result[0][46])
        self.YrAdres_2.textChanged.connect(self.YrAdresChanged)
        self.NomerRAL_6.setPlainText(result[0][3])
        self.NomerRAL_6.textChanged.connect(self.NomerRALChanged)
        self.TipGU_text.setPlainText(result[0][0])
        self.ObshObl_text.setPlainText(result[0][71])
        self.ObshObl.setCurrentText(result[0][71])
        self.Obl_text.setPlainText(result[0][88])
        if result[0][91] == "Выбрать территориальное Управление":
            TU = "Не  выбрано"
            self.Vibor_TU_text.setPlainText(TU)
        else:
            TU = result[0][91]

        Pricaz1GR = result[0][25]
        try:
            if Pricaz1GR != "" and Pricaz1GR is not None and Pricaz1GR != "Не зарегистрирован":
                Pricaz1GR = Pricaz1GR.split()
                DatrPricaz1GR = Pricaz1GR[2].replace("г.", "")
                NomerPricaz1GR = Pricaz1GR[0].replace("№", "")
                self.NomerEG.setPlainText(NomerPricaz1GR)
                self.DataEG.setText(DatrPricaz1GR)
        except:
            pass
        DataActa = result[0][111]
        NomerActa2 = result[0][116]
        NomerActa = result[0][32]
        self.NomerAct.setPlainText(NomerActa2)
        self.DataAct.setText(DataActa)
        if NomerActa != "" and NomerActa is not None:
            NomerActa = NomerActa.split()
            DatrNomerActa = NomerActa[2].replace("г.", "")
            NomerNomerActa = NomerActa[0].replace("№", "")
            self.NomerAct.setPlainText(NomerNomerActa)
            self.DataAct.setText(DatrNomerActa)

        self.DataPricaz3GRT.setText(result[0][101])
        self.DataPricaz3GR.dateChanged.connect(self.DataPricaz3GRChanged)

        self.VivodEZ.currentIndexChanged[str].connect(self.VivodEZChange)

        NomerEZzz = result[0][26]
        try:
            NomerEZzzSplit = NomerEZzz.split()
            NomerEZzz = NomerEZzzSplit[0].replace("№", "")
        except:
            pass
        self.NomerEZ.setPlainText(NomerEZzz)

        Find_z = result[0][121]
        if Find_z == "Да":
            self.DopZapr.setChecked(True)
        Find_z = result[0][118]
        if Find_z == "Да":
            self.Voz.setChecked(True)
        Find_z = result[0][119]
        if Find_z == "Да":
            self.Sogl.setChecked(True)
        Find_z = result[0][120]
        if Find_z == "Да":
            self.Sokr.setChecked(True)
        self.DataAct_2.setText(result[0][117])
        self.DataEZ_2.setText(result[0][122])

        self.SoprovodItog.setCurrentText(result[0][123])

        if UPR == "Отдел аккредитации в отдельных сферах":
            self.FioExpert_5.addItems(["Шкабура В.В."])
            self.FioExpert_5.addItems(["Белогуров С.И."])
            self.FioExpert_5.setCurrentText("Шкабура В.В.")
            self.FioExpert_3.addItems(["Золотаревский С.Ю."])
            self.FioExpert_3.addItems(["Гоголев Д.В."])
            self.FioExpert_3.setCurrentText("Золотаревский С.Ю.")

            self.FioExpert_12.addItems(["Шкабура В.В."])
            self.FioExpert_12.addItems(["Белогуров С.И."])
            self.FioExpert_12.setCurrentText("Шкабура В.В.")
            self.FioExpert_11.addItems(["Золотаревский С.Ю."])
            self.FioExpert_11.addItems(["Гоголев Д.В."])
            self.FioExpert_11.setCurrentText("Золотаревский С.Ю.")

            self.FioExpert_14.addItems(["Шкабура В.В."])
            self.FioExpert_14.addItems(["Белогуров С.И."])
            self.FioExpert_14.setCurrentText("Шкабура В.В.")
            self.FioExpert_19.addItems(["Золотаревский С.Ю."])
            self.FioExpert_19.addItems(["Гоголев Д.В."])
            self.FioExpert_19.setCurrentText("Золотаревский С.Ю.")
        elif UPR == "Отдел аккредитации испытательных лабораторий":
            self.FioExpert_5.addItems(["Бухарова А.В."])
            self.FioExpert_5.setCurrentText("Бухарова А.В.")
            self.FioExpert_3.addItems(["Золотаревский С.Ю."])
            self.FioExpert_3.addItems(["Гоголев Д.В."])
            self.FioExpert_3.setCurrentText("Золотаревский С.Ю.")

            self.FioExpert_12.addItems(["Бухарова А.В."])
            self.FioExpert_12.setCurrentText("Бухарова А.В.")
            self.FioExpert_11.addItems(["Золотаревский С.Ю."])
            self.FioExpert_11.addItems(["Гоголев Д.В."])
            self.FioExpert_11.setCurrentText("Золотаревский С.Ю.")

            self.FioExpert_14.addItems(["Бухарова А.В."])
            self.FioExpert_14.setCurrentText("Бухарова А.В.")
            self.FioExpert_19.addItems(["Золотаревский С.Ю."])
            self.FioExpert_19.addItems(["Гоголев Д.В."])
            self.FioExpert_19.setCurrentText("Золотаревский С.Ю.")
        elif UPR == "Управление аккредитации в сфере тяжёлой промышленности":
            self.FioExpert_5.addItems(["Макаров А.Н."])
            self.FioExpert_5.addItems(["Хазиева А.А."])
            self.FioExpert_5.setCurrentText("Хазиева А.А.")
            self.FioExpert_3.addItems(["Макаров А.Н."])
            self.FioExpert_3.addItems(["Садртдинова А.С."])
            self.FioExpert_3.setCurrentText("Макаров А.Н.")

            self.FioExpert_12.addItems(["Макаров А.Н."])
            self.FioExpert_12.addItems(["Хазиева А.А."])
            self.FioExpert_12.setCurrentText("Макаров А.Н.")
            self.FioExpert_19.addItems(["Макаров А.Н."])
            self.FioExpert_19.addItems(["Садртдинова А.С."])
            self.FioExpert_19.setCurrentText("Макаров А.Н.")

            self.FioExpert_14.addItems(["Макаров А.Н."])
            self.FioExpert_14.setCurrentText("Макаров А.Н.")
            self.FioExpert_19.addItems(["Макаров А.Н."])
            self.FioExpert_19.addItems(["Садртдинова А.С."])
            self.FioExpert_19.addItems(["Залазаев М.Б."])
            self.FioExpert_19.setCurrentText("Залазаев М.Б.")


    def VivodEZChange(self):
        if self.VivodEZ.currentText() == "Положительный":
            self.groupBox_7.setEnabled(False)
            self.groupBox_8.setEnabled(False)
            self.groupBox_9.setEnabled(False)
            self.groupBox_20.setEnabled(False)
            self.groupBox_19.setEnabled(False)
            self.LitsDoIt2GR.setEnabled(False)
        elif self.VivodEZ.currentText() == "Отрицательный":
            self.groupBox_7.setEnabled(True)
            self.groupBox_8.setEnabled(True)
            self.groupBox_9.setEnabled(True)
            self.groupBox_20.setEnabled(True)
            self.groupBox_19.setEnabled(True)
            self.LitsDoIt2GR.setEnabled(True)

    def Etap3GRChange(self):
        '''global globalNomerGU
        global UPR
        # self.StatusGU.setCurrentText("Возврат без рассмотрения")
        datenow = QDate.currentDate().toPyDate()  # сегодня
        datenow = str(datenow)
        datenow = datenow.replace("-", " ")
        datenow = datenow.split()
        datenowStr = datenow[2] + "." + datenow[1] + "." + datenow[0]

        Etap = self.Etap3GR.currentText()
        if self.Etap3GR.currentText() == "Приказ на согласовании":
            connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
            cur = connection.cursor()
            cur.execute('UPDATE GU1Group SET ДатаЗапускаПриказа3гр = ? WHERE Регистрационныйномер = ?',
                        (datenowStr, globalNomerGU))
        if self.Etap3GR.currentText() != "Приказ на согласовании" and self.Etap3GR.currentText() != "Приказ зарегистрирован":
            connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
            cur = connection.cursor()
            cur.execute('UPDATE GU1Group SET ДатаЗапускаПриказа3гр = ? WHERE Регистрационныйномер = ?',
                        ("", globalNomerGU))

        connection.commit()
        connection.close()'''






    def NomerRALChanged(self):
        NomerRALc = self.NomerRAL_6.toPlainText()
        self.NomerRAL.setPlainText(NomerRALc)

    def YrAdresChanged(self):
        YrAdresc = self.YrAdres_2.toPlainText()
        self.YrAdres.setPlainText(YrAdresc)

    def EmailRALChanged(self):
        EmailRALc = self.EmailRAL_2.toPlainText()
        self.EmailRAL.setPlainText(EmailRALc)

    def FullNaimChanged(self):
        FullNames = self.FullNaim_2.toPlainText()
        self.FullNaim.setPlainText(FullNames)

    def ShortNameChanged(self):
        ShortName = self.ShortName_3.toPlainText()
        self.ShortName.setPlainText(ShortName)
        self.ShortName_2.setPlainText(ShortName)

    def DO_Check(self):
        if self.DO.isChecked() == True:
            self.VKS.setChecked(False)
    def VO_Check(self):
        if self.VO.isChecked() == True:
            self.VKS.setChecked(True)
        else:
            self.VKS.setChecked(False)
    def DO_VO_Check(self):
        if self.DO_VO.isChecked() == True:
            self.VKS.setChecked(True)
        else:
            self.VKS.setChecked(False)


    def Itog2GrChange(self):
        if self.Itog2Gr.currentText() == "Не выбран":
            self.groupBox_4.setEnabled(False)
            self.groupBox_5.setEnabled(False)
            self.groupBox_6.setEnabled(False)
        elif self.Itog2Gr.currentText() == "Полож.":
            self.groupBox_4.setEnabled(True)
            self.groupBox_5.setEnabled(False)
            self.groupBox_6.setEnabled(False)
        elif self.Itog2Gr.currentText() == "Отриц. не отказ":
            self.groupBox_4.setEnabled(False)
            self.groupBox_5.setEnabled(True)
            self.groupBox_6.setEnabled(False)
        elif self.Itog2Gr.currentText() == "Отриц. отказ":
            self.groupBox_4.setEnabled(False)
            self.groupBox_5.setEnabled(False)
            self.groupBox_6.setEnabled(True)




    def Magic2GR(self):
        In39 =  self.TextIn39.toPlainText()
        In39 = In39.replace(",,", ",")
        In39 = In39.replace(")", ",").replace(",,", ",")
        In39 = In39.replace(" п.", " ").replace("п. ", " ")
        In39 = In39.replace("«", "").replace("»", "")
        i = 1
        for i in range(10):
            istr = str(i)
            In39 = In39.replace(istr + "а", istr + " а").replace(istr + "б", istr + " б").replace(istr + "в",                                                                                                  istr + " в").replace(
                istr + "г", istr + " г").replace(istr + "д", istr + " д").replace(istr + "е", istr + " е").replace(
                istr + "ж", istr + " ж").replace(istr + "з", istr + " з").replace(istr + "и", istr + " и").replace(
                istr + "к", istr + " к")
        In39 = In39.replace("а,", "а, ").replace("б,", "б, ").replace("в,", "в, ").replace("г,", "г, ").replace("д,",                                                                                                                "д, ").replace(
            "е,", "е, ").replace("ж,", "ж, ").replace("з,", "з, ").replace("и,", "и, ").replace("к,", "к, ")
        In39 = In39.replace(" а ", " а, ").replace(" б ", " б, ").replace(" в ", " в, ").replace(" г ", " г, ").replace(
            " д ", " д, ").replace(" е ", " е, ").replace(" ж ", " ж, ").replace(" з ", " з, ").replace(" и ",                                                                                                        " и, ").replace(
            " к ", " к, ")
        In39 = In39.replace("0,", "0, ").replace("1,", "1, ").replace("2,", "2, ").replace("3,", "3, ").replace("4,",                                                                                                                "4, ").replace(
            "5,", "5, ").replace("6,", "6, ").replace("7,", "7, ").replace("8,", "8, ").replace("9,", "9, ")
        In39 = In39.replace("  ", " ")
        In39 = In39.replace(", 1", ";1").replace(", 2", ";2").replace(", 3", ";3").replace(", 4", ";4").replace(", 5",                                                                                                                ";5").replace(
            ", 6", ";6").replace(", 7", ";7").replace(", 8", ";8").replace(", 9", ";9")
        In39 = In39.replace(", ;", ";")
        In39 = In39.split(";")
        sizeIn39 = len(In39)
        NewTextIn39 = ""
        for ii in range(sizeIn39):
            if ii == 0:
                NewTextIn39 = NewTextIn39 + In39[ii]
            else:
                NewTextIn39 = NewTextIn39 + ", "  + In39[ii]

        sizeIn39 = len(In39)
        SortAllInIn39 = []
        SravnitPunkt = ""
        for iii in range(sizeIn39):
            OneIn39 = In39[iii]
            FindProbel = OneIn39.find(" ")
            FindOnleBukva = len(OneIn39) - (FindProbel + 1)
            try:
                if FindProbel > 0:
                    OnlyPunkt = OneIn39[:FindProbel]
                if FindOnleBukva == 1:
                    OnleBukva = OneIn39[FindProbel + 1:]
                    if SravnitPunkt != OnlyPunkt:
                        if iii == 0:
                            SortAllInIn39 = OneIn39
                        else:
                            SortAllInIn39 = SortAllInIn39 + ", " + OneIn39
                    else:
                        SortAllInIn39 = SortAllInIn39 + ", " + OnleBukva
                    SravnitPunkt = OnlyPunkt
                else:
                    if iii == 0:
                        SortAllInIn39 = OneIn39
                    else:
                        SortAllInIn39 = SortAllInIn39 + ", " + OneIn39
            except:
                if iii == 0:
                    SortAllInIn39 = OneIn39
                else:
                    SortAllInIn39 = SortAllInIn39 + ", " + OneIn39

        self.TextIn39.setPlainText(SortAllInIn39)

#######################################################################################################################
        NoIn39 = self.TextNoIn39.toPlainText()
        NoIn39 = NoIn39.replace(",,", ",")
        NoIn39 = NoIn39.replace(")", ",").replace(",,", ",")
        NoIn39 = NoIn39.replace(" п.", " ").replace("п. ", " ")
        NoIn39 = NoIn39.replace("«", "").replace("»", "")
        i = 1
        for i in range(10):
            istr = str(i)
            NoIn39 = NoIn39.replace(istr + "а", istr + " а").replace(istr + "б", istr + " б").replace(istr + "в",                                                                                                  istr + " в").replace(
                istr + "г", istr + " г").replace(istr + "д", istr + " д").replace(istr + "е", istr + " е").replace(
                istr + "ж", istr + " ж").replace(istr + "з", istr + " з").replace(istr + "и", istr + " и").replace(
                istr + "к", istr + " к")
        NoIn39 = NoIn39.replace("а,", "а, ").replace("б,", "б, ").replace("в,", "в, ").replace("г,", "г, ").replace("д,",                                                                                                                "д, ").replace(
            "е,", "е, ").replace("ж,", "ж, ").replace("з,", "з, ").replace("и,", "и, ").replace("к,", "к, ")
        NoIn39 = NoIn39.replace(" а ", " а, ").replace(" б ", " б, ").replace(" в ", " в, ").replace(" г ", " г, ").replace(
            " д ", " д, ").replace(" е ", " е, ").replace(" ж ", " ж, ").replace(" з ", " з, ").replace(" и ",                                                                                                        " и, ").replace(
            " к ", " к, ")
        NoIn39 = NoIn39.replace("0,", "0, ").replace("1,", "1, ").replace("2,", "2, ").replace("3,", "3, ").replace("4,",                                                                                                                "4, ").replace(
            "5,", "5, ").replace("6,", "6, ").replace("7,", "7, ").replace("8,", "8, ").replace("9,", "9, ")
        NoIn39 = NoIn39.replace("  ", " ")
        NoIn39 = NoIn39.replace(", 1", ";1").replace(", 2", ";2").replace(", 3", ";3").replace(", 4", ";4").replace(", 5",                                                                                                                ";5").replace(
            ", 6", ";6").replace(", 7", ";7").replace(", 8", ";8").replace(", 9", ";9")
        NoIn39 = NoIn39.replace(", ;", ";")
        NoIn39 = NoIn39.split(";")
        sizeNoIn39 = len(NoIn39)
        NewTextNoIn39 = ""
        for ii in range(sizeNoIn39):
            if ii == 0:
                NewTextNoIn39 = NewTextNoIn39 + NoIn39[ii]
            else:
                NewTextNoIn39 = NewTextNoIn39 + ", " + NoIn39[ii]

        sizeIn39 = len(NoIn39)
        SortAllInNoIn39 = []
        SravnitPunkt = ""
        for iii in range(sizeIn39):
            OneIn39 = NoIn39[iii]
            FindProbel = OneIn39.find(" ")
            FindOnleBukva = len(OneIn39) - (FindProbel + 1)
            try:
                if FindProbel > 0:
                    OnlyPunkt = OneIn39[:FindProbel]
                if FindOnleBukva == 1:
                    OnleBukva = OneIn39[FindProbel + 1:]
                    print(OnlyPunkt)
                    print(OnleBukva)
                    if SravnitPunkt != OnlyPunkt:
                        if iii == 0:
                            SortAllInNoIn39 = OneIn39
                        else:
                            SortAllInNoIn39 = SortAllInNoIn39 + ", " + OneIn39
                    else:
                        SortAllInNoIn39 = SortAllInNoIn39 + ", " + OnleBukva
                    SravnitPunkt = OnlyPunkt
                else:
                    if iii == 0:
                        SortAllInNoIn39 = OneIn39
                    else:
                        SortAllInNoIn39 = SortAllInNoIn39 + ", " + OneIn39
            except:
                if iii == 0:
                    SortAllInNoIn39 = OneIn39
                else:
                    SortAllInNoIn39 = SortAllInNoIn39 + ", " + OneIn39

        self.TextNoIn39.setPlainText(SortAllInNoIn39)


        sizeNoIn39 = len(NoIn39)
        AllInNoIn39 = []

        for ii in range(sizeNoIn39):
            OneNoIn39 = NoIn39[ii]
            OneNoIn39 = OneNoIn39.replace(",", "")
            OneNoIn39 = OneNoIn39.split()
            sizeOneNoIn39Any = len(OneNoIn39)
            if len(OneNoIn39) != 1:
                for iii in range(sizeOneNoIn39Any - 1):
                    OneNoIn39Any = OneNoIn39[0] + " " + OneNoIn39[iii + 1]
                    AllInNoIn39.append(OneNoIn39Any)
            else:
                AllInNoIn39.append(NoIn39[ii])


        if "4.1" in AllInNoIn39 or "4.2" in AllInNoIn39 or "4.2 абзац 1" in AllInNoIn39 or "4.2 абзац 2" in AllInNoIn39 or "4.3" in AllInNoIn39 \
                or "4.4" in AllInNoIn39 or "5.1" in AllInNoIn39 or "5.2" in AllInNoIn39 or "5.3" in AllInNoIn39 or "5.4" in AllInNoIn39 or "9.2" in AllInNoIn39 or "9.3" in AllInNoIn39 \
                or "24.1" in AllInNoIn39 or "24.2" in AllInNoIn39 or "24.2 абзац 1" in AllInNoIn39 or "24.2 абзац 2" in AllInNoIn39 or "24.3" in AllInNoIn39 or "24.7.1" in AllInNoIn39 or "24.7.2.1" in AllInNoIn39 \
                or "24.7.2.2" in AllInNoIn39 or "24.8" in AllInNoIn39 or "29.1" in AllInNoIn39 or "29.2" in AllInNoIn39 or "29.2 абзац 1" in AllInNoIn39 \
                or "29.2 абзац 2" in AllInNoIn39 or "29.3" in AllInNoIn39 or "34.1" in AllInNoIn39 or "34.3" in AllInNoIn39 or "41" in AllInNoIn39 or "42" in AllInNoIn39 or "55 а" in AllInNoIn39 \
                or "55 е" in AllInNoIn39:
            P1a2 = "абзацу 2 пункта 1, "
        else:
            P1a2 = ""

        if "4.2 абзац 3" in AllInNoIn39 or "24.2 абзац 3" in AllInNoIn39 or "24.7.2.3" in AllInNoIn39 or "29.2 абзац 3" in AllInNoIn39 or "34.2" in AllInNoIn39:
            P1a3 = "абзацу 3 пункта 1, "
        else:
            P1a3 = ""

        if "4.5" in AllInNoIn39 or "24.4" in AllInNoIn39 or "29.5" in AllInNoIn39 or "34.4" in AllInNoIn39 or "43" in AllInNoIn39 or "44" in AllInNoIn39:
            P2 = "2, "
        else:
            P2 = ""

        if "4.6" in AllInNoIn39 or "6" in AllInNoIn39 or "7" in AllInNoIn39 or "8" in AllInNoIn39 or "9.1" in AllInNoIn39 or "10" in AllInNoIn39 \
                or "11" in AllInNoIn39 or "12" in AllInNoIn39 or "13" in AllInNoIn39 or "14" in AllInNoIn39 or "15" in AllInNoIn39 or "16" in AllInNoIn39 \
                or "17" in AllInNoIn39 or "18" in AllInNoIn39 or "23.1" in AllInNoIn39 or "23.2" in AllInNoIn39 or "23.3" in AllInNoIn39 or "24.5" in AllInNoIn39 \
                or "25" in AllInNoIn39 or "29.4" in AllInNoIn39 or "32" in AllInNoIn39 or "33" in AllInNoIn39 or "34.5" in AllInNoIn39 or "40" in AllInNoIn39:
            P3 = "3, "
        else:
            P3 = ""

        if "4.7" in AllInNoIn39 or "4.7 а" in AllInNoIn39 or "4.7 б" in AllInNoIn39 or "4.7 в" in AllInNoIn39 or "4.7 г" in AllInNoIn39 \
                or "4.7 д" in AllInNoIn39 or "4.7 е" in AllInNoIn39 or "4.7 ж" in AllInNoIn39 or "4.7 з" in AllInNoIn39 or "4.7 и" in AllInNoIn39 \
                or "29.6" in AllInNoIn39 or "29.6 а" in AllInNoIn39 or "29.6 б" in AllInNoIn39 or "29.6 в" in AllInNoIn39 or "29.6 г" in AllInNoIn39 \
                or "29.6 д" in AllInNoIn39 or "29.6 е" in AllInNoIn39 or "34.6" in AllInNoIn39 or "34.6 а" in AllInNoIn39 or "34.6 б" in AllInNoIn39 \
                or "34.6 в" in AllInNoIn39 or "34.6 г" in AllInNoIn39 or "34.6 д" in AllInNoIn39 or "34.6 е" in AllInNoIn39 or "34.6 з" in AllInNoIn39 or "34.6 и" in AllInNoIn39:
            P4 = "4, "
        else:
            P4 = ""

        if "4.8" in AllInNoIn39 or "4.8 а" in AllInNoIn39 or "4.8 б" in AllInNoIn39 or "4.8 в" in AllInNoIn39 or "24.6" in AllInNoIn39 or "24.6 а" in AllInNoIn39 or "24.6 б" in AllInNoIn39 or "24.6 в" in AllInNoIn39 or "29.8" in AllInNoIn39 or "29.8 а" in AllInNoIn39 or "29.8 б" in AllInNoIn39 or "29.8 в" in AllInNoIn39 or "46.7 г" in AllInNoIn39 or "46.7 и" in AllInNoIn39 or "46.7 к" in AllInNoIn39:
            print('пункту 5')
            P5 = "5, "
        else:
            P5 = ""

        if "49" in AllInNoIn39:
            print('абзацу 2 пункта 6')
            P6a2 = "абзацу 2 пункта 6, "
        else:
            P6a2 = ""

        if "54" in AllInNoIn39:
            print('абзацу 3 пункта 6')
            P6a3 = "абзацу 3 пункта 6, "
        else:
            P6a3 = ""
        if P2 != "":
            P2Check = 1
        else:
            P2Check = 0
        if P3 != "":
            P3Check = 1
        else:
            P3Check = 0
        if P4 != "":
            P4Check = 1
        else:
            P4Check = 0
        if P5 != "":
            P5Check = 1
        else:
            P5Check = 0

        ALLCheck = P2Check + P3Check + P4Check + P5Check
        if ALLCheck > 1:
            Punkt = "пунктам "
        else:
            Punkt = "пункту "

        P2345 = Punkt + P2 + P3 + P4 + P5
        AllPunkts = P1a2 + P1a3 + P2345 + P6a2 + P6a3
        sizeAllPunkts = len(AllPunkts)
        AllPunkts = AllPunkts[:sizeAllPunkts - 2]
        # AllPunkts = AllPunkts.replace(" , ", " ").replace(", , ", ", ").replace(", , , ", ", ").replace(", , , , ", ", ")
        if AllPunkts != "пункт":
            self.ItogVivod.setPlainText("относятся к " + AllPunkts)
        else:
            self.ItogVivod.setPlainText("Несоотвествия не относятся к 34 приказу")
        if NewTextIn39 != "":
            NewTextIn39 = ("39 (в части " + SortAllInIn39 + ")")
            #NewTextIn39 = ("Аккредитованное лицо не соответствует пунктам 39 (в части " + SortAllInIn39 + "), ")
        self.ItogText2GR.setPlainText(NewTextIn39 + SortAllInNoIn39)


        if self.ItogVivod.toPlainText() == "Несоотвествия не относятся к 34 приказу":
            self.Itog2Gr.setCurrentIndex(2)
        else:
            self.Itog2Gr.setCurrentIndex(3)
        if self.ItogText2GR.toPlainText() == "":
            self.ItogVivod.setPlainText("")
            self.Itog2Gr.setCurrentIndex(0)



    def OtcazCheck_Check(self):
        global globalNomerGU
        global UPR
        # self.StatusGU.setCurrentText("Возврат без рассмотрения")
        Ispolnitel = self.Ispolnitel2.currentText()
        datenow = QDate.currentDate().toPyDate()  # сегодня
        connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
        cur = connection.cursor()
        cur.execute('UPDATE GU1Group SET ВозвратРешение = ? WHERE Регистрационныйномер = ?',
                    ("На проверке", globalNomerGU))
        cur.execute('UPDATE GU1Group SET СтатусГУ = ? WHERE Регистрационныйномер = ?',
                    ("Отказ ГУ (договор)", globalNomerGU))
        if UPR == "Отдел аккредитации в отдельных сферах":
            cur.execute('UPDATE GU1Group SET ВозвратПроверка = ? WHERE Регистрационныйномер = ?',
                        ("Шкабура Владимир Владимирович", globalNomerGU))
        elif UPR == "Отдел аккредитации испытательных лабораторий":
            cur.execute('UPDATE GU1Group SET ВозвратПроверка = ? WHERE Регистрационныйномер = ?',
                        ("Бухарова Анастасия Владимировна", globalNomerGU))
        connection.commit()
        connection.close()
        self.StatusGU.setCurrentText("Отказ ГУ (договор)")
        self.OtcazCheck.hide()
        self.OtzivCheck.hide()
        self.PcazObOtcaze.hide()
        self.label_32.hide()
        self.FioExpert2.hide()
        self.label_36.hide()
        self.DataRospExp.hide()
        self.label_35.hide()
        self.Perevibor.hide()
        self.label_33.hide()
        self.DatePredlog.hide()
        self.label_34.hide()
        self.DataPricaza.hide()
        self.label_31.hide()
        self.DeadLine.hide()
        self.Priastanovka.hide()
        self.VozvratCheck_2.hide()
        self.VozvratCheck_3.hide()
        self.VozvratCheck.hide()

    def PcazObOtcaze_Check(self):
        global globalNomerGU
        global UPR
        # self.StatusGU.setCurrentText("Возврат без рассмотрения")
        Ispolnitel = self.Ispolnitel2.currentText()
        datenow = QDate.currentDate().toPyDate()  # сегодня
        connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
        cur = connection.cursor()
        cur.execute('UPDATE GU1Group SET Приказоботказе = ? WHERE Регистрационныйномер = ?',
                    ("Есть приказ об отказе", globalNomerGU))
        cur.execute('UPDATE GU1Group SET СтатусГУ = ? WHERE Регистрационныйномер = ?',
                    ("Приказ об отказе", globalNomerGU))
        '''if UPR == "Отдел аккредитации в отдельных сферах":
            cur.execute('UPDATE GU1Group SET ВозвратПроверка = ? WHERE Регистрационныйномер = ?',
                        ("Шкабура Владимир Владимирович", globalNomerGU))
        elif UPR == "Отдел аккредитации испытательных лабораторий":
            cur.execute('UPDATE GU1Group SET ВозвратПроверка = ? WHERE Регистрационныйномер = ?',
                        ("Бухарова Анастасия Владимировна", globalNomerGU))'''
        connection.commit()
        connection.close()
        self.StatusGU.setCurrentText("Приказ об отказе")
        self.OtcazCheck.hide()
        self.OtzivCheck.hide()
        self.PcazObOtcaze.hide()
        self.label_32.hide()
        self.FioExpert2.hide()
        self.label_36.hide()
        self.DataRospExp.hide()
        self.label_35.hide()
        self.Perevibor.hide()
        self.label_33.hide()
        self.DatePredlog.hide()
        self.label_34.hide()
        self.DataPricaza.hide()
        self.label_31.hide()
        self.DeadLine.hide()
        self.Priastanovka.hide()
        self.VozvratCheck_2.hide()
        self.VozvratCheck_3.hide()
        self.VozvratCheck.hide()

    def OtzivCheck_Check(self):
        global globalNomerGU
        global UPR
        # self.StatusGU.setCurrentText("Возврат без рассмотрения")
        Ispolnitel = self.Ispolnitel2.currentText()
        datenow = QDate.currentDate().toPyDate()  # сегодня
        connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
        cur = connection.cursor()
        cur.execute('UPDATE GU1Group SET ВозвратРешение = ? WHERE Регистрационныйномер = ?',
                    ("На проверке", globalNomerGU))
        cur.execute('UPDATE GU1Group SET СтатусГУ = ? WHERE Регистрационныйномер = ?',
                    ("Отзыв ГУ", globalNomerGU))
        if UPR == "Отдел аккредитации в отдельных сферах":
            cur.execute('UPDATE GU1Group SET ВозвратПроверка = ? WHERE Регистрационныйномер = ?',
                        ("Шкабура Владимир Владимирович", globalNomerGU))
        elif UPR == "Отдел аккредитации испытательных лабораторий":
            cur.execute('UPDATE GU1Group SET ВозвратПроверка = ? WHERE Регистрационныйномер = ?',
                        ("Бухарова Анастасия Владимировна", globalNomerGU))
        connection.commit()
        connection.close()
        self.StatusGU.setCurrentText("Отзыв ГУ")
        self.OtcazCheck.hide()
        self.OtzivCheck.hide()
        self.PcazObOtcaze.hide()
        self.label_32.hide()
        self.FioExpert2.hide()
        self.label_36.hide()
        self.DataRospExp.hide()
        self.label_35.hide()
        self.Perevibor.hide()
        self.label_33.hide()
        self.DatePredlog.hide()
        self.label_34.hide()
        self.DataPricaza.hide()
        self.label_31.hide()
        self.DeadLine.hide()
        self.Priastanovka.hide()
        self.VozvratCheck_2.hide()
        self.VozvratCheck_3.hide()
        self.VozvratCheck.hide()


    def ObshObl_Changed(self):
        Obl = self.ObshObl.currentText()
        if Obl == "ИЛ":
            self.Oblast_Button.setText("")
            self.ObshObl_text.setPlainText("ИЛ")
            self.Obl_text.setPlainText("ИЛ")
        else:
            self.ObshObl_text.setPlainText("Метрология")
            self.Obl_text.setPlainText("")

    def Priastanovka_Check(self):
        global globalNomerGU
        global UPR
        connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
        cur = connection.cursor()
        cur.execute('UPDATE GU1Group SET Приостановка = ? WHERE Регистрационныйномер = ?',
                    ("Да", globalNomerGU))
        cur.execute('UPDATE GU1Group SET СтатусГУ = ? WHERE Регистрационныйномер = ?',
                    ("Приостановка", globalNomerGU))
        self.StatusGU.setCurrentText("Приостановка")
        connection.commit()
        connection.close()

        self.Priastanovka.hide()
        self.PereviborDate.hide()
        self.label_PereviborDate.hide()
        self.VozvratCheck.hide()
        self.label_32.hide()
        self.FioExpert2.hide()
        self.label_36.hide()
        self.DataRospExp.hide()
        self.label_35.hide()
        self.Perevibor.hide()
        self.label_33.hide()
        self.DatePredlog.hide()
        self.label_34.hide()
        self.DataPricaza.hide()
        self.label_31.hide()
        self.DeadLine.hide()
        self.DogovorCheck.hide()

        self.VozvratCheck_2.show()
        self.VozvratCheck_2.setText("Возобновить ГУ")




    def Vibor_TU_Changed(self):
        TU = self.Vibor_TU.currentText()
        if TU == "Центральный федеральный округ":
            self.Vibor_TU_2.setCurrentText("ЦФО")
            self.Vibor_TU_text.setPlainText("ЦФО")
            self.SoprovodItog.setCurrentText("")
        elif TU == "Приволжский федеральный округ":
            self.Vibor_TU_2.setCurrentText("ПФО")
            self.Vibor_TU_text.setPlainText("ПФО")
            self.SoprovodItog.setCurrentText("Не нужно")
        elif TU == "Сибирский федеральный округ":
            self.Vibor_TU_2.setCurrentText("СФО")
            self.Vibor_TU_text.setPlainText("СФО")
            self.SoprovodItog.setCurrentText("Не нужно")
        elif TU == "Северо-Западный федеральный округ":
            self.Vibor_TU_2.setCurrentText("СЗФО")
            self.Vibor_TU_text.setPlainText("СЗФО")
            self.SoprovodItog.setCurrentText("Не нужно")
        elif TU == "Дальневосточный федеральный округ":
            self.Vibor_TU_2.setCurrentText("ДФО")
            self.Vibor_TU_text.setPlainText("ДФО")
            self.SoprovodItog.setCurrentText("Не нужно")
        elif TU == "Уральский федеральный округ":
            self.Vibor_TU_2.setCurrentText("УФО")
            self.Vibor_TU_text.setPlainText("УФО")
            self.SoprovodItog.setCurrentText("Не нужно")
        elif TU == "Южный и Северо-Кавказский федеральные округа":
            self.Vibor_TU_2.setCurrentText("ЮСФО")
            self.Vibor_TU_text.setPlainText("ЮСФО")
            self.SoprovodItog.setCurrentText("Не нужно")
        elif TU == "Выбрать территориальное Управление":
            self.Vibor_TU_2.setCurrentText("Тер. Управление")
            self.Vibor_TU_text.setPlainText("Не выбрано")

    def Vibor_TU_Changed2(self):
        TU = self.Vibor_TU_2.currentText()
        if TU == "ЦФО":
            self.Vibor_TU.setCurrentText("Центральный федеральный округ")
            self.Vibor_TU_text.setPlainText("ЦФО")
            self.SoprovodItog.setCurrentText("")
        elif TU == "ПФО":
            self.Vibor_TU.setCurrentText("Приволжский федеральный округ")
            self.Vibor_TU_text.setPlainText("ПФО")
            self.SoprovodItog.setCurrentText("Не нужно")
        elif TU == "СФО":
            self.Vibor_TU.setCurrentText("Сибирский федеральный округ")
            self.Vibor_TU_text.setPlainText("СФО")
            self.SoprovodItog.setCurrentText("Не нужно")
        elif TU == "СЗФО":
            self.Vibor_TU.setCurrentText("Северо-Западный федеральный округ")
            self.Vibor_TU_text.setPlainText("СЗФО")
            self.SoprovodItog.setCurrentText("Не нужно")
        elif TU == "ДФО":
            self.Vibor_TU.setCurrentText("Дальневосточный федеральный округ")
            self.Vibor_TU_text.setPlainText("ДФО")
            self.SoprovodItog.setCurrentText("Не нужно")
        elif TU == "УФО":
            self.Vibor_TU.setCurrentText("Уральский федеральный округ")
            self.Vibor_TU_text.setPlainText("УФО")
            self.SoprovodItog.setCurrentText("Не нужно")
        elif TU == "ЮСФО":
            self.Vibor_TU.setCurrentText("Южный и Северо-Кавказский федеральные округа")
            self.Vibor_TU_text.setPlainText("ЮСФО")
            self.SoprovodItog.setCurrentText("Не нужно")
        elif TU == "Тер. Управление":
            self.Vibor_TU.setCurrentText("Выбрать территориальное Управление")


    def SendMailYved(self):
        global globalNomerGU
        EmailExp = self.EmailExpert.toPlainText()
        EmailEO = self.MailREO.toPlainText()
        FullGU = self.NomerGU.toPlainText()
        datenow = QDate.currentDate().toPyDate()
        datenow = str(datenow)
        datenow = datenow.replace("-", " ")
        datenow = datenow.split()
        datenowStr = datenow[2] + "." + datenow[1] + "." + datenow[0]
        self.DataYvedT.setText(datenowStr)
        #self.DataYvedT.setAlignment(Qt.AlignCenter)
        outlook = win32.Dispatch('outlook.application')
        mail = outlook.CreateItem(0)
        mail.To = EmailExp + ";" + EmailEO
        mail.Subject = 'Уведомление о выборе ЭА (' + FullGU + ")"
        if self.Oblast_Button.text() == "ПСИ":
            mail.HTMLBody = '<p><strong>Здравствуйте!</strong></p><p>Вам была назначена работа в системе ФГИС 2. Прошу ' \
                            'Вас зайти в свой личный кабинет и сформировать согласие (либо отказ) и предложение по составу ' \
                            'экспертной группы (по образцу) в карточке ГУ. Все документы должны быть подписаны ЭЦП. ' \
                            '<strong><u>В случае возникновения проблем с работой информационной системы ФГИС 2 прошу ' \
                            'Вас направлять документы через электронную приемную ФСА (включая срок регистрации 3 дня) ' \
                            'и обращаться в службу поддержки ФГИС.</u></strong><o:p></o:p></p><p>В соответствии с ' \
                            'пунктом 8 статьи 17 Федерального закона от 28 декабря 2013 г. № 412-ФЗ &laquo;Об аккредитации ' \
                            'в национальной системе аккредитации&raquo; состав экспертной группы определяется национальным ' \
                            'органом по аккредитации на основании предложений эксперта по аккредитации о привлечении ' \
                            'технических экспертов, необходимых для проведения экспертизы представленных заявителем ' \
                            'документов и сведений, выездной экспертизы соответствия заявителя критериям аккредитации, ' \
                            'из числа технических экспертов, включенных в реестр технических экспертов. Такие предложения ' \
                            'должны быть направлены в национальный орган по аккредитации в течение пяти рабочих дней со ' \
                            'дня отбора эксперта по аккредитации и содержать сведения о согласии технических экспертов ' \
                            'на участие в проведении экспертизы представленных заявителем документов и сведений и ' \
                            'проведении выездной экспертизы соответствия заявителя критериям аккредитации.</p>' \
                            '<p>Также прошу учесть что в соответствии с Постановлением Правительства РФ № 2050 ' \
                            'от 26.11.2021 в состав экспертной группы включается<strong> ТЕХНИЧЕСКИЙ ЭКСПЕРТ,' \
                            '&nbsp;</strong><strong><u>который является работником государственного научного ' \
                            'метрологического института (ГНМИ).</u></strong></p>' \
                            '<p>Таким образом, прошу Вас направить уведомление в сроки установленные законодательством Российской ' \
                            'Федерации.</p><p>Спасибо!</p><p><strong>&nbsp;</strong></p>'  # this field is optional
        else:
            mail.HTMLBody = '<p><strong>Здравствуйте!</strong></p><p>Вам была назначена работа в системе ФГИС 2. Прошу ' \
                            'Вас зайти в свой личный кабинет и сформировать согласие (либо отказ) и предложение по составу ' \
                            'экспертной группы (по образцу) в карточке ГУ. Все документы должны быть подписаны ЭЦП. ' \
                            '<strong><u>В случае возникновения проблем с работой информационной системы ФГИС 2 прошу ' \
                            'Вас направлять документы через электронную приемную ФСА (включая срок регистрации 3 дня) ' \
                            'и обращаться в службу поддержки ФГИС.</u></strong><o:p></o:p></p><p>В соответствии с ' \
                            'пунктом 8 статьи 17 Федерального закона от 28 декабря 2013 г. № 412-ФЗ &laquo;Об аккредитации ' \
                            'в национальной системе аккредитации&raquo; состав экспертной группы определяется национальным ' \
                            'органом по аккредитации на основании предложений эксперта по аккредитации о привлечении ' \
                            'технических экспертов, необходимых для проведения экспертизы представленных заявителем ' \
                            'документов и сведений, выездной экспертизы соответствия заявителя критериям аккредитации, ' \
                            'из числа технических экспертов, включенных в реестр технических экспертов. Такие предложения ' \
                            'должны быть направлены в национальный орган по аккредитации в течение пяти рабочих дней со ' \
                            'дня отбора эксперта по аккредитации и содержать сведения о согласии технических экспертов ' \
                            'на участие в проведении экспертизы представленных заявителем документов и сведений и ' \
                            'проведении выездной экспертизы соответствия заявителя критериям аккредитации.</p><p>Таким ' \
                            'образом, прошу Вас направить уведомление в сроки установленные законодательством Российской ' \
                            'Федерации.</p><p>Спасибо!</p><p><strong>&nbsp;</strong></p>'  # this field is optional


        mail.Display(True)

    def SendMailYvedNapom(self):
        global globalNomerGU
        EmailExp = self.EmailExpert.toPlainText()
        FullGU = self.NomerGU.toPlainText()
        EmailEO = self.MailREO.toPlainText()
        datenow = QDate.currentDate().toPyDate()
        datenow = str(datenow)
        datenow = datenow.replace("-", " ")
        datenow = datenow.split()
        datenowStr = datenow[2] + "." + datenow[1] + "." + datenow[0]
        self.DataYvedT.setText(datenowStr)
        outlook = win32.Dispatch('outlook.application')
        mail = outlook.CreateItem(0)
        mail.To = EmailExp + ";" + EmailEO
        mail.Subject = 'Уведомление о выборе ЭА (' + FullGU + ")"
        mail.HTMLBody = '<p><strong>Здравствуйте!</strong></p><p> Напоминаю, что <span style="color:#ff0000;"><strong>3 и более дней назад</strong></span> Вам была назначена работа в системе ФГИС 2.Вам была назначена работа в системе ФГИС 2. Прошу Вас зайти в свой личный кабинет и сформировать согласие (либо отказ) и предложение по составу экспертной группы (по образцу) в карточке ГУ. Все документы должны быть подписаны ЭЦП. <strong><u>В случае возникновения проблем с работой информационной системы ФГИС 2 прошу Вас направлять документы через электронную приемную ФСА (включая срок регистрации 3 дня) и обращаться в службу поддержки ФГИС.</u></strong><o:p></o:p></p><p>В соответствии с пунктом 8 статьи 17 Федерального закона от 28 декабря 2013 г. № 412-ФЗ &laquo;Об аккредитации в национальной системе аккредитации&raquo; состав экспертной группы определяется национальным органом по аккредитации на основании предложений эксперта по аккредитации о привлечении технических экспертов, необходимых для проведения экспертизы представленных заявителем документов и сведений, выездной экспертизы соответствия заявителя критериям аккредитации, из числа технических экспертов, включенных в реестр технических экспертов. Такие предложения должны быть направлены в национальный орган по аккредитации в течение пяти рабочих дней со дня отбора эксперта по аккредитации и содержать сведения о согласии технических экспертов на участие в проведении экспертизы представленных заявителем документов и сведений и проведении выездной экспертизы соответствия заявителя критериям аккредитации.</p><p>Таким образом, прошу Вас направить уведомление в сроки установленные законодательством Российской Федерации.</p><p>Спасибо!</p><p><strong>&nbsp;</strong></p>>'
        mail.Display(True)

    def Vibor_TUChanged(self):
        TU = self.Vibor_TU.currentText()
        if TU == "Центральный федеральный округ":
            self.Vibor_TU_text.setPlainText("ЦФО")
            self.Rucovod_TU.setPlainText("Скрыпник Н.В.")
            self.SoprovodItog.setCurrentText("")
        elif TU == "Приволжский федеральный округ":
            self.Vibor_TU_text.setPlainText("ПФО")
            self.Rucovod_TU.setPlainText("Данилина Ю.В.")
            self.SoprovodItog.setCurrentText("Не нужно")
        elif TU == "Сибирский федеральный округ":
            self.Vibor_TU_text.setPlainText("СФО")
            self.Rucovod_TU.setPlainText("Логинов А.И.")
            self.SoprovodItog.setCurrentText("Не нужно")
        elif TU == "Северо-Западный федеральный округ":
            self.Vibor_TU_text.setPlainText("СЗФО")
            self.Rucovod_TU.setPlainText("Зайцев А.В.")
            self.SoprovodItog.setCurrentText("Не нужно")
        elif TU == "Дальневосточный федеральный округ":
            self.Vibor_TU_text.setPlainText("ДФО")
            self.Rucovod_TU.setPlainText("Куценко Т.В.")
            self.SoprovodItog.setCurrentText("Не нужно")
        elif TU == "Уральский федеральный округ":
            self.Vibor_TU_text.setPlainText("УФО")
            self.Rucovod_TU.setPlainText("Лавелина И.В.")
            self.SoprovodItog.setCurrentText("Не нужно")
        elif TU == "Южный и Северо-Кавказский федеральные округа":
            self.Vibor_TU_text.setPlainText("ЮСФО")
            self.Rucovod_TU.setPlainText("Сергеев Д.Н.")
            self.SoprovodItog.setCurrentText("Не нужно")

    def TipChanged1(self):
        TipGU = self.TipGU_3.currentText()
        self.TipGU.setCurrentText(TipGU)
        self.TipGU_text.setPlainText(TipGU)
        if TipGU == "ПК1":
            self.DO.setEnabled(True)
            self.VO.setEnabled(True)
            self.DO_VO.setEnabled(False)
            self.DO.setChecked(True)
        elif TipGU == "ПК2":
            self.DO.setEnabled(False)
            self.VO.setEnabled(True)
            self.DO_VO.setEnabled(False)
            self.VO.setChecked(True)
        else:
            self.DO.setEnabled(False)
            self.VO.setEnabled(False)
            self.DO_VO.setEnabled(True)
            self.DO_VO.setChecked(True)
        if TipGU == "ПК1" or TipGU == "ПК1+ИМОД":
            self.SrokD.setPlainText("20")
        elif TipGU == "ПК2" or TipGU == "ПК2+ИМОД":
            self.SrokD.setPlainText("20")
        elif TipGU == "АК" or TipGU == "РОА":
            self.SrokD.setPlainText("33")
        else:
            self.SrokD.setPlainText("28")


    def TipChanged2(self):
        TipGU = self.TipGU.currentText()
        self.TipGU_3.setCurrentText(TipGU)
        self.TipGU_text.setPlainText(TipGU)
        if TipGU == "ПК1":
            self.DO.setEnabled(True)
            self.VO.setEnabled(True)
            self.DO_VO.setEnabled(False)
            self.DO.setChecked(True)
        elif TipGU == "ПК2":
            self.DO.setEnabled(False)
            self.VO.setEnabled(True)
            self.DO_VO.setEnabled(False)
            self.VO.setChecked(True)
        else:
            self.DO.setEnabled(False)
            self.VO.setEnabled(False)
            self.DO_VO.setEnabled(True)
            self.DO_VO.setChecked(True)
        if TipGU == "ПК1" or TipGU == "ПК1+ИМОД":
            self.SrokD.setPlainText("20")
        elif TipGU == "ПК2" or TipGU == "ПК2+ИМОД":
            self.SrokD.setPlainText("20")
        elif TipGU == "АК" or TipGU == "РОА":
            self.SrokD.setPlainText("33")
        else:
            self.SrokD.setPlainText("28")

    def VozvratDenide(self):
        global globalNomerGU
        connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
        cur = connection.cursor()
        cur.execute('UPDATE GU1Group SET ВозвратРешение = ? WHERE Регистрационныйномер = ?',
                    (None, globalNomerGU))
        cur.execute('UPDATE GU1Group SET ДатаВозврата = ? WHERE Регистрационныйномер = ?',
                    (None, globalNomerGU))
        cur.execute('UPDATE GU1Group SET ВозвратПроверка = ? WHERE Регистрационныйномер = ?',
                    (None, globalNomerGU))
        cur.execute('UPDATE GU1Group SET СтатусГУ = ? WHERE Регистрационныйномер = ?',
                    ("В работе", globalNomerGU))
        connection.commit()
        connection.close()

        self.StatusGU.setCurrentText("В работе")

        self.label_PereviborDate.hide()
        self.PereviborDate.hide()
        self.VozvratCheck_2.hide()
        self.VozvratCheck_3.hide()
        self.label_32.show()
        self.FioExpert2.show()
        self.label_36.show()
        self.DataRospExp.show()
        self.label_35.show()
        self.Perevibor.show()
        self.label_33.show()
        self.DatePredlog.show()
        self.label_34.show()
        self.DataPricaza.show()
        self.label_31.show()
        self.DeadLine.show()

    def VozvratAccept(self):
        global globalNomerGU
        if self.VozvratCheck_2.text() == "Подтвердить возврат":
            connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
            cur = connection.cursor()
            self.StatusGU.setCurrentText("Возврат без рассмотрения")
            cur.execute('UPDATE GU1Group SET СтатусГУ = ? WHERE Регистрационныйномер = ?',
                        ("Возврат без рассмотрения", globalNomerGU))
            cur.execute('UPDATE GU1Group SET ВозвратРешение = ? WHERE Регистрационныйномер = ?',
                        ("Да", globalNomerGU))
            connection.commit()
            connection.close()
            self.VozvratCheck_3.hide()
        elif self.VozvratCheck_2.text() == "Возобновить ГУ":
            connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
            cur = connection.cursor()
            self.StatusGU.setCurrentText("В работе")
            cur.execute('UPDATE GU1Group SET СтатусГУ = ? WHERE Регистрационныйномер = ?',
                        ("В работе", globalNomerGU))
            cur.execute('UPDATE GU1Group SET Приостановка = NULL WHERE Регистрационныйномер = ?',
                        (globalNomerGU, ))
            connection.commit()
            connection.close()
        elif self.VozvratCheck_2.text() == "Подтвердить Отзыв":
            connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
            cur = connection.cursor()
            self.StatusGU.setCurrentText("Отзыв ГУ")
            cur.execute('UPDATE GU1Group SET СтатусГУ = ? WHERE Регистрационныйномер = ?',
                        ("Отзыв ГУ", globalNomerGU))
            cur.execute('UPDATE GU1Group SET ВозвратРешение = ? WHERE Регистрационныйномер = ?',
                        ("Да", globalNomerGU))
            connection.commit()
            connection.close()
            self.VozvratCheck_3.hide()

        elif self.VozvratCheck_2.text() == "Подтвердить Отказ":
            connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
            cur = connection.cursor()
            self.StatusGU.setCurrentText("Отказ ГУ (договор)")
            cur.execute('UPDATE GU1Group SET СтатусГУ = ? WHERE Регистрационныйномер = ?',
                        ("Отказ ГУ (договор)", globalNomerGU))
            cur.execute('UPDATE GU1Group SET ВозвратРешение = ? WHERE Регистрационныйномер = ?',
                        ("Да", globalNomerGU))
            connection.commit()
            connection.close()
            self.VozvratCheck_3.hide()





    def DoVozvrat(self):
        global globalNomerGU
        global UPR
        # self.StatusGU.setCurrentText("Возврат без рассмотрения")
        Ispolnitel = self.Ispolnitel2.currentText()
        datenow = QDate.currentDate().toPyDate()  # сегодня
        connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
        cur = connection.cursor()
        cur.execute('UPDATE GU1Group SET ДатаВозврата = ? WHERE Регистрационныйномер = ?',
                    (datenow, globalNomerGU))
        cur.execute('UPDATE GU1Group SET ВозвратРешение = ? WHERE Регистрационныйномер = ?',
                    ("На проверке", globalNomerGU))
        cur.execute('UPDATE GU1Group SET СтатусГУ = ? WHERE Регистрационныйномер = ?',
                    ("Возврат без рассмотрения", globalNomerGU))
        if UPR == "Отдел аккредитации в отдельных сферах":
            cur.execute('UPDATE GU1Group SET ВозвратПроверка = ? WHERE Регистрационныйномер = ?',
                        ("Шкабура Владимир Владимирович", globalNomerGU))
        elif UPR == "Отдел аккредитации испытательных лабораторий":
            cur.execute('UPDATE GU1Group SET ВозвратПроверка = ? WHERE Регистрационныйномер = ?',
                        ("Бухарова Анастасия Владимировна", globalNomerGU))

        self.PereviborDate.show()
        self.label_PereviborDate.show()
        self.label_PereviborDate.setText("Дата Возврата")
        # datenow2 = datetime.strptime(datenow, '%d.%m.%Y')
        self.PereviborDate.setDate(datenow)
        connection.commit()
        connection.close()

        self.label_32.hide()
        self.FioExpert2.hide()
        self.label_36.hide()
        self.DataRospExp.hide()
        self.label_35.hide()
        self.Perevibor.hide()
        self.label_33.hide()
        self.DatePredlog.hide()
        self.label_34.hide()
        self.DataPricaza.hide()
        self.label_31.hide()
        self.DeadLine.hide()
        self.Priastanovka.hide()

    def TakeOblast(self):
        self.Oblasti.show()
        oblast = "Не выбрана"
        self.Oblast_1.clicked.connect(self.PSI)
        self.Oblast_2.clicked.connect(self.KSI)
        self.Oblast_3.clicked.connect(self.ISO)
        self.Oblast_4.clicked.connect(self.ISI)
        self.Oblast_5.clicked.connect(self.AM)
        self.Oblast_6.clicked.connect(self.ME)
        self.Oblast_7.clicked.connect(self.Gost17020)
        self.Oblast_8.clicked.connect(self.Gost17021)
        self.Oblast_9.clicked.connect(self.Gost17024)
        self.Oblast_10.clicked.connect(self.AMiME)
        self.Oblast_11.clicked.connect(self.Gost17043)
        self.Oblast_12.clicked.connect(self.Gost17065)
        self.Oblast_13.clicked.connect(self.Gost15189)
        self.Oblast_14.clicked.connect(self.Gost14065)
        self.Oblast_Close.clicked.connect(self.OblastClose)

    def PSI(self):
        self.Oblasti.hide()
        self.Oblast_Button.setText("ПСИ")
        self.Obl_text.setPlainText("ПСИ")

    def KSI(self):
        self.Oblasti.hide()
        self.Oblast_Button.setText("КСИ")
        self.Obl_text.setPlainText("КСИ")

    def ISO(self):
        self.Oblasti.hide()
        self.Oblast_Button.setText("ИСО")
        self.Obl_text.setPlainText("ИСО")

    def ISI(self):
        self.Oblasti.hide()
        self.Oblast_Button.setText("ИСИ")
        self.Obl_text.setPlainText("ИСИ")

    def AM(self):
        self.Oblasti.hide()
        self.Oblast_Button.setText("АМ")
        self.Obl_text.setPlainText("АМ")

    def ME(self):
        self.Oblasti.hide()
        self.Oblast_Button.setText("МЭ")
        self.Obl_text.setPlainText("МЭ")

    def AMiME(self):
        self.Oblasti.hide()
        self.Oblast_Button.setText("АМ и МЭ")
        self.Obl_text.setPlainText("АМ и МЭ")

    def Gost17020(self):
        self.Oblasti.hide()
        self.Oblast_Button.setText("17020")
        self.Obl_text.setPlainText("17020")

    def Gost17021(self):
        self.Oblasti.hide()
        self.Oblast_Button.setText("17021")
        self.Obl_text.setPlainText("17021")

    def Gost17024(self):
        self.Oblasti.hide()
        self.Oblast_Button.setText("17024")
        self.Obl_text.setPlainText("17024")

    def Gost17043(self):
        self.Oblasti.hide()
        self.Oblast_Button.setText("17043")
        self.Obl_text.setPlainText("17043")

    def Gost17065(self):
        self.Oblasti.hide()
        self.Oblast_Button.setText("17065")
        self.Obl_text.setPlainText("17065")

    def Gost15189(self):
        self.Oblasti.hide()
        self.Oblast_Button.setText("15189")
        self.Obl_text.setPlainText("15189")

    def Gost14065(self):
        self.Oblasti.hide()
        self.Oblast_Button.setText("14065")
        self.Obl_text.setPlainText("14065")

    def OblastClose(self):
        self.Oblasti.hide()
        self.Oblast_Button.setText("")
        self.Obl_text.setPlainText("")

    def Open1CURL(self):
        Url1C = self.Url1CGU.toPlainText()
        QDesktopServices.openUrl(QUrl(Url1C, QUrl.TolerantMode))

    def OpenCloudURL(self):
        UrlCl = self.URLCloud.toPlainText()
        QDesktopServices.openUrl(QUrl(UrlCl, QUrl.TolerantMode))

    def OpenWikiURL(self):
        QDesktopServices.openUrl(QUrl(
            "https://ru.wikipedia.org/wiki/%D0%A4%D0%B5%D0%B4%D0%B5%D1%80%D0%B0%D0%BB%D1%8C%D0%BD%D1%8B%D0%B5_%D0%BE%D0%BA%D1%80%D1%83%D0%B3%D0%B0_%D0%A0%D0%BE%D1%81%D1%81%D0%B8%D0%B9%D1%81%D0%BA%D0%BE%D0%B9_%D0%A4%D0%B5%D0%B4%D0%B5%D1%80%D0%B0%D1%86%D0%B8%D0%B8",
            QUrl.TolerantMode))

    def OpenRAL(self):
        NomerRAL = self.NomerRAL.toPlainText()
        connection = sqlite3.connect('/Portal/UseFile/BaseFrom1C.db')
        cur = connection.cursor()
        cur.execute("SELECT * FROM RalReestr WHERE Номераттестатааккредитации = ?", (NomerRAL,))
        result = cur.fetchone()

        if result is not None:
            nomer = str(result[0])
            QDesktopServices.openUrl(QUrl("http://10.250.74.17/ral/view/" + nomer, QUrl.TolerantMode))

    def DataDogChanged(self):
        datenow = self.DataDogD.text()
        datenow = str(datenow)
        self.DataDogT.setText(datenow)

    def DataPricaz3GRChanged(self):
        datenow = self.DataPricaz3GR.text()
        datenow = str(datenow)
        self.DataPricaz3GRT.setText(datenow)


    def DataPricaz2GRChanged(self):
        datenow = self.DataPricaz2GR.text()
        datenow = str(datenow)
        self.DataPricaz2GRT.setText(datenow)


    def DataEZChanged(self):
        datenow = self.DataEZ.text()
        datenow = str(datenow)
        self.DataEZT.setText(datenow)

    def DataActChanged(self):
        datenow = self.DataActa.text()
        datenow = str(datenow)
        self.DataActaT.setText(datenow)

    def DogovorCheckB(self):
        global globalNomerGU
        connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
        cur = connection.cursor()
        DogovorText = self.DogovorCheck.text()





        if DogovorText == "Указать дату договора":
            datenow = QDate.currentDate().toPyDate()
            self.DataDogT.show()
            self.DataDogD.show()
            datenow = str(datenow)
            datenow = datenow.replace("-", " ")
            datenow = datenow.split()
            datenowStr = datenow[2] + "." + datenow[1] + "." + datenow[0]
            self.DataDogT.setText(datenowStr)
            cur.execute('UPDATE GU1Group SET Договор1Г = ? WHERE Регистрационныйномер = ?',
                        (datenowStr, globalNomerGU))
            self.DogovorCheck.setText("Договор подписан")
            self.DogovorCheck.setStyleSheet('QPushButton:hover { background-color: rgb(159, 211, 140); border-radius:15px; font: 10pt "MS Shell Dlg 2"; }QPushButton:!hover { background-color: rgb(152, 202, 134); border-radius:15px; font: 10pt "MS Shell Dlg 2";  }QPushButton:pressed { background-color: rgb(128, 171, 113); border-radius:15px; font: 10pt "MS Shell Dlg 2";  }')
        elif DogovorText == "Договор подписан":
            self.DataDogT.hide()
            self.DataDogD.hide()
            cur.execute('UPDATE GU1Group SET Договор1Г = ? WHERE Регистрационныйномер = ?',
                        ("", globalNomerGU))
            self.DataDogT.setText("")
            self.DogovorCheck.setText("Указать дату договора")
            self.DogovorCheck.setStyleSheet('QPushButton:hover { background-color: rgb(249, 69, 75); border-radius:5px; font: 10pt "MS Shell Dlg 2"; } QPushButton:!hover { background-color:rgb(230, 64, 67); border-radius:5px; font: 10pt "MS Shell Dlg 2";  } QPushButton:pressed { background-color: rgb(209, 58, 63); border-radius:5px; font: 10pt "MS Shell Dlg 2";  }')
        connection.commit()
        connection.close()

    def FixFullNaim(self):
        global FullNaim
        FullNaim = self.FullNaim.toPlainText()
        FullNaim = FullNaim.replace("ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ",
                                    "Общество с ограниченной ответственностью")
        FullNaim = FullNaim.replace("ПУБЛИЧНОЕ АКЦИОНЕРНОЕ ОБЩЕСТВО", "Публичное акционерное общество")
        FullNaim = FullNaim.replace("ЗАКРЫТОЕ АКЦИОНЕРНОЕ ОБЩЕСТВО", "Закрытое акционерное общество")
        FullNaim = FullNaim.replace("ОТКРЫТОЕ АКЦИОНЕРНОЕ ОБЩЕСТВО", "Открытого акционерного общества")
        FullNaim = FullNaim.replace("АКЦИОНЕРНОЕ ОБЩЕСТВО", "Акционерное общество")
        FullNaim = FullNaim.replace("ФЕДЕРАЛЬНОЕ БЮДЖЕТНОЕ УЧРЕЖДЕНИЕ ЗДРАВООХРАНЕНИЯ",
                                    "Федеральное бюджетное учреждение здравоохранения")
        FullNaim = FullNaim.replace("ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ВОДОХОЗЯЙСТВЕННОЕ УЧРЕЖДЕНИЕ",
                                    "Федеральное государственное бюджетное водохозяйственное учреждение")
        FullNaim = FullNaim.replace("ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ УЧРЕЖДЕНИЕ",
                                    "Федеральное государственное бюджетное учреждение")
        FullNaim = FullNaim.replace("НАУЧНО-ТЕХНИЧЕСКИЙ ЦЕНТР", "Научно-технический центр")
        FullNaim = FullNaim.replace("ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ УЧРЕЖДЕНИЕ", "Научно-технический центр")
        FullNaim = FullNaim.replace("ФЕДЕРАЛЬНОЕ БЮДЖЕТНОЕ УЧРЕЖДЕНИЕ", "Федеральное бюджетное учреждение")
        FullNaim = FullNaim.replace("САНКТ-ПЕТЕРБУРГСКОЕ", "Санкт-Петербургское")
        FullNaim = FullNaim.replace("КАЗЕННОЕ ПРЕДПРИЯТИЕ", "Казенное предприятие")
        FullNaim = FullNaim.replace("АССОЦИАЦИЯ", "Ассоциация")
        FullNaim = FullNaim.replace("НАУЧНО-ПРОИЗВОДСТВЕННОЕ ОБЪЕДИНЕНИЕ", "Научно-производственное объединение")
        FullNaim = FullNaim.replace("ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ УНИТАРНОЕ ПРЕДПРИЯТИЕ",
                                    "Федеральное государственное унитарное предприятие")

        self.FullNaim.setPlainText(FullNaim)

    def FixFullNaimKogo(self):
        global FullNaimKogo
        FullNaimKogo = self.FullNaim.toPlainText()
        FullNaimKogo = FullNaimKogo.replace("Общество с ограниченной ответственностью",
                                            "Общества с ограниченной ответственностью")
        FullNaimKogo = FullNaimKogo.replace("Публичное акционерное общество", "Публичного акционерного общества")
        FullNaimKogo = FullNaimKogo.replace("Закрытое акционерное общество", "Закрытого акционерного общества")
        FullNaimKogo = FullNaimKogo.replace("Открытое акционерное общество", "Открытого акционерного общества")
        FullNaimKogo = FullNaimKogo.replace("Акционерное общество", "Акционерного общества")
        FullNaimKogo = FullNaimKogo.replace("Федеральное бюджетное учреждение здравоохранения",
                                            "Федерального бюджетного учреждения здравоохранения")
        FullNaimKogo = FullNaimKogo.replace("Федеральное государственное бюджетное водохозяйственное учреждение",
                                            "Федерального государственного бюджетного водохозяйственного учреждения")
        FullNaimKogo = FullNaimKogo.replace("Федеральное государственное бюджетное учреждение",
                                            "Федерального государственного бюджетного учреждения")
        FullNaimKogo = FullNaimKogo.replace("Государственное бюджетное учреждение",
                                            "Государственного бюджетного учреждения")
        FullNaimKogo = FullNaimKogo.replace(" Научно-технический центр", " Научно-техническиого центра")
        FullNaimKogo = FullNaimKogo.replace("Государственное бюджетное учреждение",
                                            "Государственного бюджетного учреждения")
        FullNaimKogo = FullNaimKogo.replace("Федеральное бюджетное учреждение", "Федерального бюджетного учреждения")
        FullNaimKogo = FullNaimKogo.replace("Бюджетное учреждение здравоохранения",
                                            "Бюджетного учреждения здравоохранения")
        FullNaimKogo = FullNaimKogo.replace("Федеральное казенное учреждение", "Федерального казенного учреждения")
        FullNaimKogo = FullNaimKogo.replace("Казенное предприятие", "Казенного предприятия")
        FullNaimKogo = FullNaimKogo.replace("Ассоциация", "Ассоциации")
        FullNaimKogo = FullNaimKogo.replace("Федеральное государственное унитарное предприятие",
                                            "Федерального государственного унитарного предприятия")

        #self.FullNaim.setPlainText(FullNaimKogo)

# Создание Word по 3 группе
    def CreatePricaz3GR(self):
        global FullNaimKogo, ALLTeh
        global FullFioLogin
        global MsgError

        if self.Vibor_TU.currentText() == "Выбрать территориальное Управление":
            MsgError = "Выберете ТУ во вкладке (Данные ГУ)"
            self.gotoAlarm()
            return
        if self.Vibor_TU.currentText() == "Центральный федеральный округ" and (self.EmailRAL.toPlainText() == "" or self.YrAdres.toPlainText() == ""):
            MsgError = "ЦФО? Не повезло\n Нужно заполнить Email или Почтовый адрес\nдля Сопровода"
            self.gotoAlarm()
            return
        if self.DataAct_2.text() == "":
            MsgError = "Заполните дату Акта (котора в самом акте)\n мне тоже это не нравится"
            self.gotoAlarm()
            return
        if (self.TipGU_3.currentText() == "АК" or self.TipGU_3.currentText() == "РОА") and self.DataEZ_2.text() == "":
            MsgError = "Поздравляю! У Вас РОА или АК\nЗаполните дату ЭЗ (котора в самом ЭЗ)\n#Давайте отменим эту запись"
            self.gotoAlarm()
            return
        if self.Obl_text.toPlainText() == "Не выбрано" or self.Obl_text.toPlainText() == "":
            MsgError = "Это Вам не ИЛ. \nНужно выбрать точную область \n(вкладка 1 Группа)"
            self.gotoAlarm()
            return

        #############


        FullNaimWho = self.FullNaim_2.toPlainText()
        self.FixFullNaimKogo()  # испраить ООО на падеж

        FullGU = self.NomerGU.toPlainText()
        FullGUSplit = FullGU.split()
        DataPricaza = FullGUSplit[2]
        DataPricazaSplit = DataPricaza.replace('.', ' ')
        DataPricazaSplit = DataPricazaSplit.split()
        DataPricaza0 = DataPricazaSplit[0]
        DataPricaza1 = DataPricazaSplit[1]
        DataPricaza2 = DataPricazaSplit[2]
        DataPricaza0 = DataPricaza0.replace('01', '1').replace('02', '2').replace('03', '3').replace('04', '4').replace(
            '05', '5').replace('06', '6').replace('07', '7').replace('08', '8').replace('09', '9')
        DataPricaza1 = DataPricaza1.replace('01', 'января').replace('02', 'февралья').replace('03', 'марта').replace(
            '04', 'апреля').replace('05', 'мая').replace('06', 'июня').replace('07', 'июля').replace('08',
                                                                                                     'августа').replace(
            '09', 'сентября').replace('10', 'октября').replace('11', 'ноября').replace('12', 'декабря')
        DataPricaza = DataPricaza0 + " " + DataPricaza1 + " " + DataPricaza2 + " г."
        NomerGU = FullGUSplit[0]

        NomerRAL = self.NomerRAL.toPlainText()
        DataAct = self.DataAct_2.text()
        DatActSplit = DataAct.replace('.', ' ')
        DatActSplit = DatActSplit.split()
        DatAct0 = DatActSplit[0]
        DatAct1 = DatActSplit[1]
        DatAct2 = DatActSplit[2]
        DatAct0 = DatAct0.replace('01', '1').replace('02', '2').replace('03', '3').replace('04', '4').replace(
            '05', '5').replace('06', '6').replace('07', '7').replace('08', '8').replace('09', '9')
        DatAct1 = DatAct1.replace('01', 'января').replace('02', 'февралья').replace('03', 'марта').replace(
            '04', 'апреля').replace('05', 'мая').replace('06', 'июня').replace('07', 'июля').replace('08',
                                                                                                     'августа').replace(
            '09', 'сентября').replace('10', 'октября').replace('11', 'ноября').replace('12', 'декабря')
        DataAct = DatAct0 + " " + DatAct1 + " " + DatAct2 + " г."
        DataEZ = ""
        if (self.TipGU_3.currentText() == "АК" or self.TipGU_3.currentText() == "РОА"):
            DataEZ = self.DataEZ_2.text()
            DataEZSplit = DataEZ.replace('.', ' ')
            DataEZSplit = DataEZSplit.split()
            DataEZ0 = DataEZSplit[0]
            DataEZ1 = DataEZSplit[1]
            DataEZ2 = DataEZSplit[2]
            DataEZ0 = DataEZ0.replace('01', '1').replace('02', '2').replace('03', '3').replace('04', '4').replace(
                '05', '5').replace('06', '6').replace('07', '7').replace('08', '8').replace('09', '9')
            DataEZ1 = DataEZ1.replace('01', 'января').replace('02', 'февралья').replace('03', 'марта').replace(
                '04', 'апреля').replace('05', 'мая').replace('06', 'июня').replace('07', 'июля').replace('08',
                                                                                                         'августа').replace(
                '09', 'сентября').replace('10', 'октября').replace('11', 'ноября').replace('12', 'декабря')
            DataEZ = DataEZ0 + " " + DataEZ1 + " " + DataEZ2 + " г."

        Podpisant = self.FioExpert_19.currentText()
        if Podpisant == "Золотаревский С.Ю.":
            Podpisant = "С.Ю. Золотаревский"
            Doljnost = "Начальник Управления \nаккредитации в сфере \nдобровольного подтверждения \nсоответствия, метрологии \nи иных сферах деятельности"
            Kontrol = "Контроль за исполнением настоящего приказа оставляю за собой."
        elif Podpisant == "Шкабура В.В.":
            Podpisant = "В.В. Шкабура"
            Doljnost = "Начальник отдела аккредитации \nв отдельных сферах \nУправления аккредитации в сфере \nдобровольного подтверждения \nсоответствия, метрологии и иных \nсферах деятельности"
        elif Podpisant == "Бухарова А.В.":
            Podpisant = "А.В. Бухарова"
            Doljnost = "Начальник отдела аккредитации \nиспытательных лабораторий \nУправления аккредитации \nв сфере добровольного подтверждения \nсоответствия, метрологии \nи иных сферах деятельности"
        elif Podpisant == "Гоголев Д.В.":
            Podpisant = "Д.В. Гоголев"
            Doljnost = "Заместитель руководителя"
            Kontrol = "Контроль за исполнением настоящего приказа возложить на начальника Управления аккредитации в сфере добровольного подтверждения соответствия, метрологии и иных сферах деятельности С.Ю. Золотаревского."
        elif Podpisant == "Белогуров С.И.":
            Podpisant = "С.И. Белогуров"
            Doljnost = "Заместитель начальника отдела \nаккредитации в отдельных сферах \nУправления аккредитации в сфере \nдобровольного подтверждения \nсоответствия, метрологии и иных \nсферах деятельности"
        elif Podpisant == "Макаров А.Н.":
            Podpisant = "А.Н. Макаров"
            Doljnost = "И.о. начальника Управления аккредитации \nв сфере подтверждения соответствия \nпродукции машиностроения, \nэлектротехнической продукции, \nстроительных материалов \nи пожарной безопасности"
            Kontrol = "Контроль за исполнением настоящего приказа оставляю за собой."
        elif Podpisant == "Хазиева А.А.":
            Podpisant = "А.А. Хазиева"
            Doljnost = "Начальник отдела аккредитации в сфере \nподтверждения соответствия продукции машиностроения \nУправления аккредитации в сфере подтверждения \nсоответствия продукции машиностроения, \nэлектротехнической продукции, строительных \nматериалов и пожарной безопасности"

        TipGU_VO_DO = self.TipGU.currentText()
        if TipGU_VO_DO == "ПК1" or TipGU_VO_DO == "ПК2" or TipGU_VO_DO == "ПК5":
            path = "SourceGitHub/Shablon/Group3/ОЕИ/Приказ/1. ПК (положительный).docx"
        elif TipGU_VO_DO == "ПК1+РОА" or TipGU_VO_DO == "ПК2+РОА" or TipGU_VO_DO == "ПК5+РОА":
            path = "SourceGitHub/Shablon/Group3/ОЕИ/Приказ/1. ПК+РОА (положительный).docx"
        elif TipGU_VO_DO == "ПК1+РОА+ИМОД" or TipGU_VO_DO == "ПК2+РОА+ИМОД" or TipGU_VO_DO == "ПК5+РОА+ИМОД":
            path = "SourceGitHub/Shablon/Group3/ОЕИ/Приказ/1. ПК+РОА+ИМОД (положительный).docx"
        elif TipGU_VO_DO == "РОА":
            path = "SourceGitHub/Shablon/Group3/ОЕИ/Приказ/1. РОА (положительный).docx"
        elif TipGU_VO_DO == "АК":
            path = "SourceGitHub/Shablon/Group3/ОЕИ/Приказ/1. АК (положительный).docx"
        else:
            path = "SourceGitHub/Shablon/Group3/ОЕИ/Приказ/1. ПК+ИМОД (положительный).docx"


        if self.Sokr.isChecked() == True:
            SokrNeSokr = ""
        else:
            SokrNeSokr = ""

        if self.Sogl.isChecked() == True:
            SoglNeSogl = ", в том числе по результатам заседания комиссии, проведенной в соответствии с Порядком рассмотрения экспертного заключения, " \
                       "акта выездной экспертизы, акта экспертизы на предмет соответствия требованиям законодательства Российской Федерации об аккредитации " \
                       "в национальной системе аккредитации, утвержденным приказом Министерства экономического развития Российской Федерации " \
                       "от 29 октября 2021 г. № 657 «Об установлении порядка проведения проверки экспертного заключения, акта выездной экспертизы, " \
                       "акта экспертизы на предмет соответствия требованиям законодательства Российской Федерации об аккредитации в национальной системе аккредитации»,"
        else:
            SoglNeSogl = ""
        punk_imod = ""
        punk_roa = ""
        punkt2or3 = ""
        punkt3or4 = ""
        if self.Voz.isChecked() == True:
            VozNeVoz = "\a2.\xa0Возобновить действие аккредитации Аккредитованного лица в связи с предоставлением в Федеральную службу по аккредитации " \
                       "заявления о проведении процедуры подтверждения компетентности Аккредитованного лица в соответствии \nс частью 4 статьи 24 " \
                       "Федерального закона от 28 декабря 2013 г. № 412-ФЗ \n«Об аккредитации в национальной системе аккредитации» и прохождением " \
                       "процедуры подтверждения компетентности Аккредитованного лица."
            VozNeVoz2 = " и возобновлении действия аккредитации"
            if TipGU_VO_DO == "ПК1" or TipGU_VO_DO == "ПК2" or TipGU_VO_DO == "ПК5":
                punk_imod = ""
                punk_roa = ""
                punkt2or3 = "3"
                punkt3or4 = "4"
            elif TipGU_VO_DO == "ПК1+РОА" or TipGU_VO_DO == "ПК2+РОА" or TipGU_VO_DO == "ПК5+РОА":
                punk_imod = ""
                punk_roa = "3"
                punkt2or3 = "4"
                punkt3or4 = "5"
            elif TipGU_VO_DO == "ПК1+РОА+ИМОД" or TipGU_VO_DO == "ПК2+РОА+ИМОД" or TipGU_VO_DO == "ПК5+РОА+ИМОД":
                punk_imod = "3"
                punk_roa = "4"
                punkt2or3 = "5"
                punkt3or4 = "6"
            else:
                punk_imod = "3"
                punk_roa = ""
                punkt2or3 = "4"
                punkt3or4 = "5"
        else:
            VozNeVoz = ""
            VozNeVoz2 = ""
            if TipGU_VO_DO == "ПК1" or TipGU_VO_DO == "ПК2" or TipGU_VO_DO == "ПК5":
                punk_imod = ""
                punk_roa = ""
                punkt2or3 = "2"
                punkt3or4 = "3"
            elif TipGU_VO_DO == "ПК1+РОА" or TipGU_VO_DO == "ПК2+РОА" or TipGU_VO_DO == "ПК5+РОА":
                punk_imod = ""
                punk_roa = "2"
                punkt2or3 = "3"
                punkt3or4 = "4"
            elif TipGU_VO_DO == "ПК1+РОА+ИМОД" or TipGU_VO_DO == "ПК2+РОА+ИМОД" or TipGU_VO_DO == "ПК5+РОА+ИМОД":
                punk_imod = "2"
                punk_roa = "3"
                punkt2or3 = "4"
                punkt3or4 = "5"
            else:
                punk_imod = "2"
                punk_roa = ""
                punkt2or3 = "3"
                punkt3or4 = "4"


        Gost = ""
        Gost2 = ""
        ObshObl = self.ObshObl.currentText()
        Obl = self.Oblast_Button.text()
        if ObshObl == "ИЛ":
            accObl = "в качестве испытательной лаборатории"
            Gost = " и требованиям ГОСТ ISO/IEC 17025-2019 «Общие требования к компетентности испытательных и калибровочных лабораторий»"
            Gost2 = " и требованиям ГОСТ ISO/IEC 17025-2019"
        elif ObshObl == "Метрология" and Obl == "ПСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по поверке средств измерений"
        elif ObshObl == "Метрология" and Obl == "КСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по калибровке средств измерений"
            Gost = " и требованиям ГОСТ ISO/IEC 17025-2019 «Общие требования к компетентности испытательных и калибровочных лабораторий»"
            Gost2 = " и требованиям ГОСТ ISO/IEC 17025-2019"
        elif ObshObl == "Метрология" and Obl == "ИСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по испытанию средств измерений в целях утверждения типа"
        elif ObshObl == "Метрология" and Obl == "ИСО":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по испытанию стандартных образцов в целях утверждения типа"
        elif ObshObl == "Метрология" and Obl == "АМ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по аттестации методик (методов) измерений"
        elif ObshObl == "Метрология" and Obl == "МЭ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по метрологической экспертизе"
        elif ObshObl == "Метрология" and Obl == "АМ и МЭ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по аттестации методик (методов) измерений и метрологической экспертизе"
        elif ObshObl == "Метрология" and Obl == "17020":
            accObl = "в качестве органа инспекции"
            Gost = " и требованиям ГОСТ Р ИСО/МЭК 17020-2012 «Оценка соответствия. Требования к работе различных типов органов инспекции»"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 17020-2012"
        elif ObshObl == "Метрология" and Obl == "17021":
            accObl = "в качестве органа по сертификации систем менеджмента"
            Gost = " и требованиям ГОСТ Р ИСО/МЭК 17021-1-2017 «Оценка соответствия. Требования к органам, проводящим аудит и сертификацию систем менеджмента»"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 17021-1-2017"
        elif ObshObl == "Метрология" and Obl == "17024":
            accObl = "в качестве органа по сертификации персонала"
            Gost = " и требованиям ГОСТ Р ИСО/МЭК 17024-1-2017 «Оценка соответствия. Общие требования к органам, проводящим сертификацию персонала»"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 17024-1-2017"
        elif ObshObl == "Метрология" and Obl == "17043":
            accObl = "в качестве провайдера межлабораторных сличительных испытаний"  # Не уверен
            Gost = " и требованиям ГОСТ Р ИСО/МЭК 17043-1-2017 «Основные требования к проведению проверки квалификации»"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 17043-1-2017"
        elif ObshObl == "Метрология" and Obl == "17065":
            accObl = "в качестве органа по сертификации продукции, услуг"
            Gost = " и требованиям ГОСТ Р ИСО/МЭК 17065-1-2012 «Оценка соответствия. Требования к органам по сертификации продукции, процессов и услуг»"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 17065-1-2012"
        elif ObshObl == "Метрология" and Obl == "15189":
            accObl = "в качестве испытательной лаборатории"
        elif ObshObl == "Метрология" and Obl == "14065":
            accObl = "в качестве органа по валидации и верификации парниковых газов"
            Gost = " и требованиям ГОСТ Р ИСО/МЭК 14065-2014 «Газы парниковые. Требования к органам по валидации и верификации парниковых газов для их применения при аккредитации или других формах признания»"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 14065-2014"
        elif ObshObl == "Метрология" and Obl == "Не выбрано":
            accObl = "в качестве испытательной лаборатории"  # Потом доработать исключение

        ShortFIO = FullFioLogin.split()
        ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."

        if self.Vibor_TU_2.currentText() == "ЦФО":
            VMerop1 = "Управлению аккредитации в сфере добровольного подтверждения соответствия, метрологии и иных сферах деятельности"
        if self.Vibor_TU_2.currentText() == "ПФО":
            VMerop1 = "Управлению Федеральной службы по аккредитации Приволжскому федеральному округу"
        if self.Vibor_TU_2.currentText() == "СФО":
            VMerop1 = "Управлению Федеральной службы по аккредитации по Сибирскому федеральному округу"
        if self.Vibor_TU_2.currentText() == "СЗФО":
            VMerop1 = "Управлению Федеральной службы по аккредитации по Северо-Западному федеральному округу"
        if self.Vibor_TU_2.currentText() == "ДФО":
            VMerop1 = "Управлению Федеральной службы по аккредитации по Дальневосточному федеральному округу"
        if self.Vibor_TU_2.currentText() == "УФО":
            VMerop1 = "Управлению Федеральной службы по аккредитации по Уральскому федеральному округу"
        if self.Vibor_TU_2.currentText() == "ЮСФО":
            VMerop1 = "Управлению Федеральной службы по аккредитации по Южному и Северо-Кавказскому федеральным округам"

        doc = DocxTemplate(path)

        context = {"nomer_gu": NomerGU, "data_gu": DataPricaza, "full_name_org": FullNaimKogo.replace(' "', ' «').replace('"', '»'),
                   "full_name_who": FullNaimKogo.replace(' "', ' «').replace('"', '»'), "SokrNeSokr": SokrNeSokr,
                   "acc_obl": accObl, "nomer_ral": NomerRAL, "gost": Gost, "dolj_podpisant": Doljnost, "fio_podpisant": Podpisant,
                   "Kontrol": Kontrol, "VozNeVoz": VozNeVoz, "VozNeVoz2": VozNeVoz2, "DataAct": DataAct, "VMerop1": VMerop1,
                   "sogl_ne_sog": SoglNeSogl, "sogl_ne_sog": SoglNeSogl, "punkt2or3": punkt2or3, "punkt3or4": punkt3or4,
                   "punk_imod": punk_imod, "punk_roa": punk_roa, "DataEZ": DataEZ, "short_isp": ShortFIO}

        doc.render(context)
        try:
            doc.save(
                "SourceGitHub/ГУ/1 Группа/" + FullGU + "/5. Итог Приказ " + TipGU_VO_DO + " " + FullGU + ".docx")
            os.startfile(
                "SourceGitHub/ГУ/1 Группа/" + FullGU + "/5. Итог Приказ " + TipGU_VO_DO + " " + FullGU + ".docx")
        except:
            print("Ошибка при создании уведомления")
    def CreateSzPolojGroup3(self):
        global FullNaimKogo, ALLTeh
        global FullFioLogin
        global MsgError

        if FullFioLogin == "Приостановка":
            MsgError = "Нельзя использовать с приостановленной ГУ"
            self.gotoAlarm()
            return
        if self.Vibor_TU.currentText() == "Выбрать территориальное Управление":
            MsgError = "Выберете ТУ во вкладке (Данные ГУ)"
            self.gotoAlarm()
            return
        if self.Vibor_TU.currentText() == "Центральный федеральный округ" and (
                self.EmailRAL.toPlainText() == "" or self.YrAdres.toPlainText() == ""):
            MsgError = "ЦФО? Не повезло\n Нужно заполнить Email или Почтовый адрес\nдля Сопровода"
            self.gotoAlarm()
            return
        if self.DataAct_2.text() == "" or self.NomerAct.toPlainText() == "":
            MsgError = "Заполните дату или номер Акта"
            self.gotoAlarm()
            return
        if self.NomerEG.toPlainText() == "" or self.NomerEG.toPlainText() == "":
            MsgError = "Заполните дату или номер Приказа по ЭГ"
            self.gotoAlarm()
            return
        if self.Obl_text.toPlainText() == "Не выбрано" or self.Obl_text.toPlainText() == "":
            MsgError = "Это Вам не ИЛ. \nНужно выбрать точную область \n(вкладка 1 Группа)"
            self.gotoAlarm()
            return
        #############

        self.FixFullNaimKogo()  # испраить ООО на падеж
        FullGU = self.NomerGU.toPlainText()
        FullGUSplit = FullGU.split()
        DataPricaza = FullGUSplit[2]
        DataPricazaSplit = DataPricaza.replace('.', ' ')
        DataPricazaSplit = DataPricazaSplit.split()
        DataPricaza0 = DataPricazaSplit[0]
        DataPricaza1 = DataPricazaSplit[1]
        DataPricaza2 = DataPricazaSplit[2]
        DataPricaza0 = DataPricaza0.replace('01', '1').replace('02', '2').replace('03', '3').replace('04', '4').replace(
            '05', '5').replace('06', '6').replace('07', '7').replace('08', '8').replace('09', '9')
        DataPricaza1 = DataPricaza1.replace('01', 'января').replace('02', 'февралья').replace('03', 'марта').replace(
            '04', 'апреля').replace('05', 'мая').replace('06', 'июня').replace('07', 'июля').replace('08',
                                                                                                     'августа').replace(
            '09', 'сентября').replace('10', 'октября').replace('11', 'ноября').replace('12', 'декабря')
        DataPricaza = DataPricaza0 + " " + DataPricaza1 + " " + DataPricaza2 + " г."
        NomerGU = FullGUSplit[0]

        NomerRAL = self.NomerRAL.toPlainText()
        NomerEG = self.NomerEG.toPlainText()
        DataEG = self.DataEG.text()
        DataEGSplit = DataEG.replace('.', ' ')
        DataEGSplit = DataEGSplit.split()
        DataEG0 = DataEGSplit[0]
        DataEG1 = DataEGSplit[1]
        DataEG2 = DataEGSplit[2]
        DataEG0 = DataEG0.replace('01', '1').replace('02', '2').replace('03', '3').replace('04', '4').replace(
            '05', '5').replace('06', '6').replace('07', '7').replace('08', '8').replace('09', '9')
        DataEG1 = DataEG1.replace('01', 'января').replace('02', 'февралья').replace('03', 'марта').replace(
            '04', 'апреля').replace('05', 'мая').replace('06', 'июня').replace('07', 'июля').replace('08',
                                                                                                     'августа').replace(
            '09', 'сентября').replace('10', 'октября').replace('11', 'ноября').replace('12', 'декабря')
        DataEG = DataEG0 + " " + DataEG1 + " " + DataEG2 + " г."

        NomerAct = self.NomerAct.toPlainText()
        DataAct = self.DataAct.text()
        DataActSplit = DataAct.replace('.', ' ')
        DataActSplit = DataActSplit.split()
        DataAct0 = DataActSplit[0]
        DataAct1 = DataActSplit[1]
        DataAct2 = DataActSplit[2]
        DataAct0 = DataAct0.replace('01', '1').replace('02', '2').replace('03', '3').replace('04', '4').replace(
            '05', '5').replace('06', '6').replace('07', '7').replace('08', '8').replace('09', '9')
        DataAct1 = DataAct1.replace('01', 'января').replace('02', 'февралья').replace('03', 'марта').replace(
            '04', 'апреля').replace('05', 'мая').replace('06', 'июня').replace('07', 'июля').replace('08',
                                                                                                     'августа').replace(
            '09', 'сентября').replace('10', 'октября').replace('11', 'ноября').replace('12', 'декабря')
        DataAct = DataAct0 + " " + DataAct1 + " " + DataAct2 + " г."



        Podpisant = self.FioExpert_14.currentText()
        if Podpisant == "Шкабура В.В.":
            Podpisant = "В.В. Шкабура"
            Doljnost = "Начальник отдела аккредитации \nв отдельных сферах \nУправления аккредитации в сфере \nдобровольного подтверждения \nсоответствия, метрологии и иных \nсферах деятельности"
        elif Podpisant == "Бухарова А.В.":
            Podpisant = "А.В. Бухарова"
            Doljnost = "Начальник отдела аккредитации \nиспытательных лабораторий \nУправления аккредитации \nв сфере добровольного подтверждения \nсоответствия, метрологии \nи иных сферах деятельности"
        elif Podpisant == "Гоголев Д.В.":
            Podpisant = "Д.В. Гоголев"
            Doljnost = "Заместитель руководителя"
            DoljnostBIG = "Заместителю руководителя\n Федеральной службы \n по аккредитации"
            PodpisantBIG = "Д.В. Гоголеву"
            FULLPodpisant = "Уважаемый Дмитрий Владимирович!"
        elif Podpisant == "Белогуров С.И.":
            Podpisant = "С.И. Белогуров"
            Doljnost = "Заместитель начальника отдела \nаккредитации в отдельных сферах \nУправления аккредитации в сфере \nдобровольного подтверждения \nсоответствия, метрологии и иных \nсферах деятельности"
        elif Podpisant == "Золотаревский С.Ю.":
            DoljnostBIG = "Начальнику Управления аккредитации \nв сфере добровольного \nподтверждения соответствия, \nметрологии и иных сферах деятельности"
            PodpisantBIG = "С.Ю. Золотаревскому"
            FULLPodpisant = "Уважаемый Сергей Юрьевич!"
        elif Podpisant == "Макаров А.Н.":
            Podpisant = "А.Н. Макаров"
            PodpisantBIG = "А.Н. Макарову"
            FULLPodpisant = "Уважаемый Андрей Николаевич!"
            Doljnost = "И.о. начальника Управления аккредитации \nв сфере подтверждения соответствия \nпродукции машиностроения, \nэлектротехнической продукции, \nстроительных материалов \nи пожарной безопасности"
        elif Podpisant == "Хазиева А.А.":
            Podpisant = "А.А. Хазиева"
            Doljnost = "Начальник отдела аккредитации в сфере \nподтверждения соответствия продукции машиностроения \nУправления аккредитации в сфере подтверждения \nсоответствия продукции машиностроения, \nэлектротехнической продукции, строительных \nматериалов и пожарной безопасности"


        Komy = self.FioExpert_19.currentText()
        if Komy == "Гоголев Д.В.":
            DoljnostBIG = "Заместителю руководителя\n Федеральной службы \n по аккредитации"
            PodpisantBIG = "Д.В. Гоголеву"
            FULLPodpisant = "Уважаемый Дмитрий Владимирович!"
        elif Komy == "Золотаревский С.Ю.":
            DoljnostBIG = "Начальнику Управления \aаккредитации \aв сфере добровольного \aподтверждения соответствия, \aметрологии и иных сферах \aдеятельности"
            PodpisantBIG = "С.Ю. Золотаревскому"
            FULLPodpisant = "Уважаемый Сергей Юрьевич!"
        elif Podpisant == "Макаров А.Н.":
            Podpisant = "А.Н. Макаров"
            PodpisantBIG = "А.Н. Макарову"
            FULLPodpisant = "Уважаемый Андрей Николаевич!"
            Doljnost = "И.о. начальника Управления аккредитации \nв сфере подтверждения соответствия \nпродукции машиностроения, \nэлектротехнической продукции, \nстроительных материалов \nи пожарной безопасности"
        elif Podpisant == "Хазиева А.А.":
            Podpisant = "А.А. Хазиева"
            Doljnost = "Начальник отдела аккредитации в сфере \nподтверждения соответствия продукции машиностроения \nУправления аккредитации в сфере подтверждения \nсоответствия продукции машиностроения, \nэлектротехнической продукции, строительных \nматериалов и пожарной безопасности"

        Gost = ""
        Gost2 = ""
        ObshObl = self.ObshObl.currentText()
        Obl = self.Oblast_Button.text()
        if ObshObl == "ИЛ":
            accObl = "в качестве испытательной лаборатории"
            Gost = ", а также ГОСТ ISO/IEC 17025-2019 «Общие требования к компетентности испытательных и калибровочных лабораторий»"
        elif ObshObl == "Метрология" and Obl == "ПСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по поверке средств измерений"
        elif ObshObl == "Метрология" and Obl == "КСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по калибровке средств измерений"
            Gost = ", а также ГОСТ ISO/IEC 17025-2019 «Общие требования к компетентности испытательных и калибровочных лабораторий»"
        elif ObshObl == "Метрология" and Obl == "ИСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по испытанию средств измерений в целях утверждения типа"
        elif ObshObl == "Метрология" and Obl == "ИСО":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по испытанию стандартных образцов в целях утверждения типа"
        elif ObshObl == "Метрология" and Obl == "АМ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по аттестации методик (методов) измерений"
        elif ObshObl == "Метрология" and Obl == "МЭ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по метрологической экспертизе"
        elif ObshObl == "Метрология" and Obl == "АМ и МЭ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по аттестации методик (методов) измерений и метрологической экспертизе"
        elif ObshObl == "Метрология" and Obl == "17020":
            accObl = "в качестве органа инспекции"
            Gost = ", а также ГОСТ Р ИСО/МЭК 17020-2012 «Оценка соответствия. Требования к работе различных типов органов инспекции»"
        elif ObshObl == "Метрология" and Obl == "17021":
            accObl = "в качестве органа по сертификации систем менеджмента"
            Gost = ", а также ГОСТ Р ИСО/МЭК 17021-1-2017 «Оценка соответствия. Требования к органам, проводящим аудит и сертификацию систем менеджмента»"
        elif ObshObl == "Метрология" and Obl == "17024":
            accObl = "в качестве органа по сертификации персонала"
            Gost = ", а также ГОСТ Р ИСО/МЭК 17024-1-2017 «Оценка соответствия. Общие требования к органам, проводящим сертификацию персонала»"
        elif ObshObl == "Метрология" and Obl == "17043":
            accObl = "в качестве провайдера межлабораторных сличительных испытаний"  # Не уверен
            Gost = ", а также ГОСТ Р ИСО/МЭК 17043-1-2017 «Основные требования к проведению проверки квалификации»"
        elif ObshObl == "Метрология" and Obl == "17065":
            accObl = "в качестве органа по сертификации продукции, услуг"
            Gost = ", а также ГОСТ Р ИСО/МЭК 17065-1-2012 «Оценка соответствия. Требования к органам по сертификации продукции, процессов и услуг»"
        elif ObshObl == "Метрология" and Obl == "15189":
            accObl = "в качестве испытательной лаборатории"
        elif ObshObl == "Метрология" and Obl == "14065":
            accObl = "в качестве органа по валидации и верификации парниковых газов"
            Gost = ", а также ГОСТ Р ИСО/МЭК 14065-2014 «Газы парниковые. Требования к органам по валидации и верификации парниковых газов для их применения при аккредитации или других формах признания»"
        elif ObshObl == "Метрология" and Obl == "Не выбрано":
            accObl = "в качестве испытательной лаборатории"  # Потом доработать исключение

        ShortFIO = FullFioLogin.split()
        ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."

        TipGU_VO_DO = self.TipGU.currentText()
        if TipGU_VO_DO == "ПК1":
            TipGU = "подтверждения компетентности"
            TipGU2 = "утвержденной области"
        elif TipGU_VO_DO == "ПК2":
            TipGU = "подтверждения компетентности"
            TipGU2 = "утвержденной области"
        elif TipGU_VO_DO == "ПК5":
            TipGU = "подтверждения компетентности"
            TipGU2 = "утвержденной области"
        elif TipGU_VO_DO == "АК":
            TipGU = "аккредитации"
            TipGU2 = "заявленной области"
        elif TipGU_VO_DO == "РОА":
            TipGU = "расширения области аккредитации"
            TipGU2 = "расширяемой области"
        elif TipGU_VO_DO == "ПК1+РОА" or TipGU_VO_DO == "ПК2+РОА" or TipGU_VO_DO == "ПК5+РОА":
            TipGU = "подтверждения компетентности и расширения области аккредитации"
            TipGU2 = "утвержденной и расширяемой областях"
        elif TipGU_VO_DO == "ПК1+ИМОД" or TipGU_VO_DO == "ПК2+ИМОД":
            TipGU = "подтверждения компетентности и изменения места (мест) осуществления деятельности"
            TipGU2 = "утвержденной области"
        elif TipGU_VO_DO == "ПК5+ИМОД":
            TipGU = "подтверждения компетентности и изменения места (мест) осуществления деятельности"
            TipGU2 = "утвержденной области"
        elif TipGU_VO_DO == "РОА+ИМОД":
            TipGU = "расширения области аккредитации и изменения места (мест) осуществления деятельности"
            TipGU2 = "расширяемой области"
        elif TipGU_VO_DO == "ПК1+РОА+ИМОД" or TipGU_VO_DO == "ПК2+РОА+ИМОД" or TipGU_VO_DO == "ПК5+РОА+ИМОД":
            TipGU = "подтверждения компетентности, расширения области аккредитации и изменения места (мест) осуществления деятельности"
            TipGU2 = "утвержденной и расширяемой областях"


        if TipGU_VO_DO != "АК" and TipGU_VO_DO != "РОА":
            path = "SourceGitHub/Shablon/Group3/ОЕИ/СЗ/1. СЗ ПК (положительный).docx"
        elif TipGU_VO_DO == "РОА":
            path = "SourceGitHub/Shablon/Group3/ОЕИ/СЗ/1. СЗ РОА (положительный).docx"
        elif TipGU_VO_DO == "АК":
            path = "SourceGitHub/Shablon/Group3/ОЕИ/СЗ/1. СЗ АК (положительный).docx"

        doc = DocxTemplate(path)

        context = { "nomer_gu": NomerGU, "data_gu": DataPricaza, "full_name_org": FullNaimKogo.replace(' "', ' «').replace('"', '»'),
                   "acc_obl": accObl, "nomer_ral": NomerRAL, "DataEG": DataEG, "NomerEG": NomerEG, "DataAct": DataAct, "NomerAct": NomerAct,
                   "dolj_podpisant": Doljnost, "fio_podpisant": Podpisant, "short_isp": ShortFIO,
                   "gost": Gost, "TipGU2": TipGU2, "TipGU": TipGU,
                   "big_dolj": DoljnostBIG, "big_fio": PodpisantBIG, "full_big_fio": FULLPodpisant,}

        doc.render(context)
        try:
            doc.save(
                "SourceGitHub/ГУ/1 Группа/" + FullGU + "/4. СЗ Итог " + TipGU_VO_DO + " " + FullGU + ".docx")
            os.startfile(
                "SourceGitHub/ГУ/1 Группа/" + FullGU + "/4. СЗ Итог " + TipGU_VO_DO + " " + FullGU + ".docx")
        except:
            print("Ошибка при создании уведомления")
    def CreateSoprPolojGroup3(self):
        global FullNaimKogo, ALLTeh
        global FullFioLogin
        global MsgError

        if self.TipGU.currentText() == "ПК" or self.TipGU.currentText() == "ПК+РОА" or self.TipGU.currentText() == "ПК+ИМОД" or self.TipGU.currentText() == "ПК+РОА+ИМОД":
            MsgError = "Уточните Тип ГУ"
            self.gotoAlarm()
            return

        if self.EmailRAL.toPlainText() == "":
            MsgError = "Введите эл. почту"
            self.gotoAlarm()
            return
        if self.YrAdres_2.toPlainText() == "":
            MsgError = "Введите Почтовый адрес"
            self.gotoAlarm()
            return
        if self.FioExpert.currentText() == "Метрология" and (
                self.Oblast_Button.text() == "Не выбрано" or self.Oblast_Button.text() == ""):
            MsgError = "Уточните конкретную область"
            self.gotoAlarm()
            return
        #############

        self.FixFullNaimKogo()  # испраить ООО на падеж
        FullGU = self.NomerGU.toPlainText()
        FullGUSplit = FullGU.split()
        DataPricaza = FullGUSplit[2]
        DataPricazaSplit = DataPricaza.replace('.', ' ')
        DataPricazaSplit = DataPricazaSplit.split()
        DataPricaza0 = DataPricazaSplit[0]
        DataPricaza1 = DataPricazaSplit[1]
        DataPricaza2 = DataPricazaSplit[2]
        DataPricaza0 = DataPricaza0.replace('01', '1').replace('02', '2').replace('03', '3').replace('04', '4').replace(
            '05', '5').replace('06', '6').replace('07', '7').replace('08', '8').replace('09', '9')
        DataPricaza1 = DataPricaza1.replace('01', 'января').replace('02', 'февралья').replace('03', 'марта').replace(
            '04', 'апреля').replace('05', 'мая').replace('06', 'июня').replace('07', 'июля').replace('08',
                                                                                                     'августа').replace(
            '09', 'сентября').replace('10', 'октября').replace('11', 'ноября').replace('12', 'декабря')
        DataPricaza = DataPricaza0 + " " + DataPricaza1 + " " + DataPricaza2 + " г."
        NomerGU = FullGUSplit[0]

        NomerRAL = self.NomerRAL.toPlainText()
        ShortNaim = self.ShortName_2.toPlainText()
        EmailOrg = self.EmailRAL.toPlainText()
        AdresOrg = self.YrAdres_2.toPlainText()

        Podpisant = self.FioExpert_14.currentText()
        if Podpisant == "Шкабура В.В.":
            Podpisant = "В.В. Шкабура"
            Doljnost = "Начальник отдела аккредитации \nв отдельных сферах \nУправления аккредитации в сфере \nдобровольного подтверждения \nсоответствия, метрологии и иных \nсферах деятельности"
        elif Podpisant == "Бухарова А.В.":
            Podpisant = "А.В. Бухарова"
            Doljnost = "Начальник отдела аккредитации \nиспытательных лабораторий \nУправления аккредитации \nв сфере добровольного подтверждения \nсоответствия, метрологии \nи иных сферах деятельности"
        elif Podpisant == "Гоголев Д.В.":
            Podpisant = "Д.В. Гоголев"
            Doljnost = "Заместитель руководителя"
        elif Podpisant == "Белогуров С.И.":
            Podpisant = "С.И. Белогуров"
            Doljnost = "Заместитель начальника отдела \nаккредитации в отдельных сферах \nУправления аккредитации в сфере \nдобровольного подтверждения \nсоответствия, метрологии и иных \nсферах деятельности"
        elif Podpisant == "Макаров А.Н.":
            Podpisant = "А.Н. Макаров"
            Doljnost = "И.о. начальника Управления аккредитации \nв сфере подтверждения соответствия \nпродукции машиностроения, \nэлектротехнической продукции, \nстроительных материалов \nи пожарной безопасности"
        elif Podpisant == "Хазиева А.А.":
            Podpisant = "А.А. Хазиева"
            Doljnost = "Начальник отдела аккредитации в сфере \nподтверждения соответствия продукции машиностроения \nУправления аккредитации в сфере подтверждения \nсоответствия продукции машиностроения, \nэлектротехнической продукции, строительных \nматериалов и пожарной безопасности"


        ShortFIO = FullFioLogin.split()
        ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."

        TipGU_VO_DO = self.TipGU.currentText()
        if TipGU_VO_DO == "ПК1":
            TipGU = "о подтверждении компетентности"
            TipGU2 = "утвержденная"
        elif TipGU_VO_DO == "ПК2":
            TipGU = "о подтверждении компетентности"
            TipGU2 = "утвержденная"
        elif TipGU_VO_DO == "ПК5":
            TipGU = "о подтверждении компетентности"
            TipGU2 = "утвержденная"
        elif TipGU_VO_DO == "АК":
            TipGU = "об аккредитации"
            TipGU2 = "аккредитации "
        elif TipGU_VO_DO == "РОА":
            TipGU = "о расширении области аккредитации"
            TipGU2 = "расширяемая"
        elif TipGU_VO_DO == "ПК1+РОА" or TipGU_VO_DO == "ПК2+РОА" or TipGU_VO_DO == "ПК5+РОА":
            TipGU = "о подтверждении компетентности и расширении области аккредитации"
            TipGU2 = "утвержденная и расширяемая"
        elif TipGU_VO_DO == "ПК1+ИМОД" or TipGU_VO_DO == "ПК2+ИМОД":
            TipGU = "о подтверждении компетентности и изменении места (мест) осуществления деятельности"
            TipGU2 = "утвержденная"
        elif TipGU_VO_DO == "ПК5+ИМОД":
            TipGU = "о подтверждении компетентности и изменении места (мест) осуществления деятельности"
            TipGU2 = "утвержденная"
        elif TipGU_VO_DO == "РОА+ИМОД":
            TipGU = "о расширении области аккредитации и изменении места (мест) осуществления деятельности"
            TipGU2 = "расширяемая"
        elif TipGU_VO_DO == "ПК1+РОА+ИМОД" or TipGU_VO_DO == "ПК2+РОА+ИМОД" or TipGU_VO_DO == "ПК5+РОА+ИМОД":
            TipGU = "о подтверждении компетентности, расширении области аккредитации и изменении места (мест) осуществления деятельности"
            TipGU2 = "утвержденная и расширяемая"

        if TipGU_VO_DO == "АК":
            path = "SourceGitHub/Shablon/Group3/ОЕИ/Уведомление/1. Сопровод АК (положительный).docx"
        elif TipGU_VO_DO == "РОА" or TipGU_VO_DO == "РОА+ИМОД":
            path = "SourceGitHub/Shablon/Group3/ОЕИ/Уведомление/1. Сопровод РОА (положительный).docx"
        else:
            path = "SourceGitHub/Shablon/Group3/ОЕИ/Уведомление/1. Сопровод ПК (положительный).docx"


        doc = DocxTemplate(path)

        context = {"short_name_org": ShortNaim.replace(' "', ' «').replace('"', '»'), "email_org": EmailOrg,
                   "AdresOrg": AdresOrg, "nomer_gu": NomerGU,
                   "data_gu": DataPricaza, "full_name_org": FullNaimKogo.replace(' "', ' «').replace('"', '»'),
                   "nomer_ral": NomerRAL, "nomer_gu2": FullGU, "TipGU2": TipGU2,
                   "dolj_podpisant": Doljnost, "fio_podpisant": Podpisant, "short_isp": ShortFIO, "tip_gu": TipGU}

        doc.render(context)
        try:
            doc.save(
                "SourceGitHub/ГУ/1 Группа/" + FullGU + "/6. Сопроводительной письмо Итог " + FullGU + ".docx")
            os.startfile(
                "SourceGitHub/ГУ/1 Группа/" + FullGU + "/6. Сопроводительной письмо Итог " + FullGU + ".docx")
        except:
            print("Ошибка при создании уведомления")


# Создание Word по 2 группе
    def CreatePolojGroup2(self):
        global FullNaimKogo, ALLTeh
        global FullFioLogin
        global MsgError

        if self.FioExpert.currentText() == "Эксперт не найден":
            MsgError = "Выберете эксперта"
            self.gotoAlarm()
            return
        if self.EmailRAL.toPlainText() == "":
            MsgError = "Введите эл. почту в разделе \n(Данные по организации)\n(Не обязательно, но если Вам не трудно, заполните ещё почтовый адрес (если он пустой))"
            self.gotoAlarm()
            return
        if FullFioLogin == "Приостановка":
            MsgError = "Нельзя использовать с приостановленной ГУ"
            self.gotoAlarm()
            return



        #############

        self.FixFullNaimKogo()  # испраить ООО на падеж
        FullGU = self.NomerGU.toPlainText()
        FullGUSplit = FullGU.split()
        DataPricaza = FullGUSplit[2]
        DataPricazaSplit = DataPricaza.replace('.', ' ')
        DataPricazaSplit = DataPricazaSplit.split()
        DataPricaza0 = DataPricazaSplit[0]
        DataPricaza1 = DataPricazaSplit[1]
        DataPricaza2 = DataPricazaSplit[2]
        DataPricaza0 = DataPricaza0.replace('01', '1').replace('02', '2').replace('03', '3').replace('04', '4').replace(
            '05', '5').replace('06', '6').replace('07', '7').replace('08', '8').replace('09', '9')
        DataPricaza1 = DataPricaza1.replace('01', 'января').replace('02', 'февралья').replace('03', 'марта').replace(
            '04', 'апреля').replace('05', 'мая').replace('06', 'июня').replace('07', 'июля').replace('08',
                                                                                                     'августа').replace(
            '09', 'сентября').replace('10', 'октября').replace('11', 'ноября').replace('12', 'декабря')
        DataPricaza = DataPricaza0 + " " + DataPricaza1 + " " + DataPricaza2 + " г."
        NomerGU = FullGUSplit[0]

        NomerRAL = self.NomerRAL.toPlainText()
        ShortNaim = self.ShortName_2.toPlainText()
        EmailOrg = self.EmailRAL.toPlainText()
        FIOExp = self.FioExpert.currentText()
        EmailExp = self.EmailExpert.toPlainText()
        NameEO = self.REO.currentText()
        EmailEO = self.MailREO.toPlainText()
        DataEZ = self.DataEZT.text()
        DatEZSplit = DataEZ.replace('.', ' ')
        DatEZSplit = DatEZSplit.split()
        DatEZ0 = DatEZSplit[0]
        DatEZ1 = DatEZSplit[1]
        DatEZ2 = DatEZSplit[2]
        DatEZ0 = DatEZ0.replace('01', '1').replace('02', '2').replace('03', '3').replace('04', '4').replace(
            '05', '5').replace('06', '6').replace('07', '7').replace('08', '8').replace('09', '9')
        DatEZ1 = DatEZ1.replace('01', 'января').replace('02', 'февралья').replace('03', 'марта').replace(
            '04', 'апреля').replace('05', 'мая').replace('06', 'июня').replace('07', 'июля').replace('08',
                                                                                                     'августа').replace(
            '09', 'сентября').replace('10', 'октября').replace('11', 'ноября').replace('12', 'декабря')
        DatEZ = DatEZ0 + " " + DatEZ1 + " " + DatEZ2 + " г."
        NomEG = self.NomerEG.toPlainText()
        DatEG = self.DataEG.text()
        DatEGSplit = DatEG.replace('.', ' ')
        DatEGSplit = DatEGSplit.split()
        DatEG0 = DatEGSplit[0]
        DatEG1 = DatEGSplit[1]
        DatEG2 = DatEGSplit[2]
        DatEG0 = DatEG0.replace('01', '1').replace('02', '2').replace('03', '3').replace('04', '4').replace(
            '05', '5').replace('06', '6').replace('07', '7').replace('08', '8').replace('09', '9')
        DatEG1 = DatEG1.replace('01', 'января').replace('02', 'февралья').replace('03', 'марта').replace(
            '04', 'апреля').replace('05', 'мая').replace('06', 'июня').replace('07', 'июля').replace('08',
                                                                                                     'августа').replace(
            '09', 'сентября').replace('10', 'октября').replace('11', 'ноября').replace('12', 'декабря')
        DatEG = DatEG0 + " " + DatEG1 + " " + DatEG2 + " г."

        Podpisant = self.FioExpert_12.currentText()
        if Podpisant == "Шкабура В.В.":
            Podpisant = "В.В. Шкабура"
            Doljnost = "Начальник отдела аккредитации \nв отдельных сферах \nУправления аккредитации в сфере \nдобровольного подтверждения \nсоответствия, метрологии и иных \nсферах деятельности"
        elif Podpisant == "Бухарова А.В.":
            Podpisant = "А.В. Бухарова"
            Doljnost = "Начальник отдела аккредитации \nиспытательных лабораторий \nУправления аккредитации \nв сфере добровольного подтверждения \nсоответствия, метрологии \nи иных сферах деятельности"
        elif Podpisant == "Гоголев Д.В.":
            Podpisant = "Д.В. Гоголев"
            Doljnost = "Заместитель руководителя"
        elif Podpisant == "Белогуров С.И.":
            Podpisant = "С.И. Белогуров"
            Doljnost = "Заместитель начальника отдела \nаккредитации в отдельных сферах \nУправления аккредитации в сфере \nдобровольного подтверждения \nсоответствия, метрологии и иных \nсферах деятельности"
        elif Podpisant == "Макаров А.Н.":
            Podpisant = "А.Н. Макаров"
            Doljnost = "И.о. начальника Управления аккредитации \nв сфере подтверждения соответствия \nпродукции машиностроения, \nэлектротехнической продукции, \nстроительных материалов \nи пожарной безопасности"
        elif Podpisant == "Хазиева А.А.":
            Podpisant = "А.А. Хазиева"
            Doljnost = "Начальник отдела аккредитации в сфере \nподтверждения соответствия продукции машиностроения \nУправления аккредитации в сфере подтверждения \nсоответствия продукции машиностроения, \nэлектротехнической продукции, строительных \nматериалов и пожарной безопасности"

        Komy = self.FioExpert_11.currentText()
        if Komy == "Гоголев Д.В.":
            DoljnostBIG = "Заместителю руководителя\n Федеральной службы \n по аккредитации"
            PodpisantBIG = "Д.В. Гоголеву"
            FULLPodpisant = "Уважаемый Дмитрий Владимирович!"
        elif Komy == "Золотаревский С.Ю.":
            DoljnostBIG = "Начальнику Управления \aаккредитации \aв сфере добровольного \aподтверждения соответствия, \aметрологии и иных сферах \aдеятельности"
            PodpisantBIG = "С.Ю. Золотаревскому"
            FULLPodpisant = "Уважаемый Сергей Юрьевич!"
        elif Podpisant == "Макаров А.Н.":
            Podpisant = "А.Н. Макаров"
            PodpisantBIG = "А.Н. Макарову"
            FULLPodpisant = "Уважаемый Андрей Николаевич!"
            Doljnost = "И.о. начальника Управления аккредитации \nв сфере подтверждения соответствия \nпродукции машиностроения, \nэлектротехнической продукции, \nстроительных материалов \nи пожарной безопасности"
        elif Podpisant == "Хазиева А.А.":
            Podpisant = "А.А. Хазиева"
            Doljnost = "Начальник отдела аккредитации в сфере \nподтверждения соответствия продукции машиностроения \nУправления аккредитации в сфере подтверждения \nсоответствия продукции машиностроения, \nэлектротехнической продукции, строительных \nматериалов и пожарной безопасности"

        Gost = ""
        Gost2 = ""
        ObshObl = self.ObshObl.currentText()
        Obl = self.Oblast_Button.text()
        if ObshObl == "ИЛ":
            accObl = "в качестве испытательной лаборатории"
            Gost = " ГОСТ ISO/IEC 17025-2019 «Общие требования к компетентности испытательных и калибровочных лабораторий» (далее – ГОСТ ISO/IEC 17025-2019)"
            Gost2 = " и требованиям ГОСТ ISO/IEC 17025-2019"
        elif ObshObl == "Метрология" and Obl == "ПСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по поверке средств измерений"
        elif ObshObl == "Метрология" and Obl == "КСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по калибровке средств измерений"
            Gost = " ГОСТ ISO/IEC 17025-2019 «Общие требования к компетентности испытательных и калибровочных лабораторий» (далее – ГОСТ ISO/IEC 17025-2019)"
            Gost2 = " и требованиям ГОСТ ISO/IEC 17025-2019"
        elif ObshObl == "Метрология" and Obl == "ИСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по испытанию средств измерений в целях утверждения типа"
        elif ObshObl == "Метрология" and Obl == "ИСО":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по испытанию стандартных образцов в целях утверждения типа"
        elif ObshObl == "Метрология" and Obl == "АМ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по аттестации методик (методов) измерений"
        elif ObshObl == "Метрология" and Obl == "МЭ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по метрологической экспертизе"
        elif ObshObl == "Метрология" and Obl == "АМ и МЭ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по аттестации методик (методов) измерений и метрологической экспертизе"
        elif ObshObl == "Метрология" and Obl == "17020":
            accObl = "в качестве органа инспекции"
            Gost = " ГОСТ Р ИСО/МЭК 17020-2012 «Оценка соответствия. Требования к работе различных типов органов инспекции» (далее – ГОСТ Р ИСО/МЭК 17020-2012)"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 17020-2012"
        elif ObshObl == "Метрология" and Obl == "17021":
            accObl = "в качестве органа по сертификации систем менеджмента"
            Gost = " ГОСТ Р ИСО/МЭК 17021-1-2017 «Оценка соответствия. Требования к органам, проводящим аудит и сертификацию систем менеджмента» (далее – ГОСТ Р ИСО/МЭК 17021-1-2017)"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 17021-1-2017"
        elif ObshObl == "Метрология" and Obl == "17024":
            accObl = "в качестве органа по сертификации персонала"
            Gost = " ГОСТ Р ИСО/МЭК 17024-1-2017 «Оценка соответствия. Общие требования к органам, проводящим сертификацию персонала» (далее – ГОСТ Р ИСО/МЭК 17024-1-2017)"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 17024-1-2017"
        elif ObshObl == "Метрология" and Obl == "17043":
            accObl = "в качестве провайдера межлабораторных сличительных испытаний"  # Не уверен
            Gost = " ГОСТ Р ИСО/МЭК 17043-1-2017 «Основные требования к проведению проверки квалификации» (далее – ГОСТ Р ИСО/МЭК 17043-1-2017)"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 17043-1-2017"
        elif ObshObl == "Метрология" and Obl == "17065":
            accObl = "в качестве органа по сертификации продукции, услуг"
            Gost = " ГОСТ Р ИСО/МЭК 17065-1-2012 «Оценка соответствия. Требования к органам по сертификации продукции, процессов и услуг» (далее – ГОСТ Р ИСО/МЭК 17065-1-2012)"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 17065-1-2012"
        elif ObshObl == "Метрология" and Obl == "15189":
            accObl = "в качестве испытательной лаборатории"
        elif ObshObl == "Метрология" and Obl == "14065":
            accObl = "в качестве органа по валидации и верификации парниковых газов"
            Gost = " ГОСТ Р ИСО/МЭК 14065-2014 «Газы парниковые. Требования к органам по валидации и верификации парниковых газов для их применения при аккредитации или других формах признания» (далее – ГОСТ Р ИСО/МЭК 14065-2014)"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 14065-2014"
        elif ObshObl == "Метрология" and Obl == "Не выбрано":
            accObl = "в качестве испытательной лаборатории"  # Потом доработать исключение

        In39TextPunkts = self.TextIn39.toPlainText()
        NoIn39TextPunkts = self.TextNoIn39.toPlainText()
        ItogTextPunkts = self.ItogText2GR.toPlainText()
        if (len(NoIn39TextPunkts) > 5) or (In39TextPunkts != "" and NoIn39TextPunkts != ""):
            ItogTextPunkts = "пунктам " + ItogTextPunkts
        else:
            ItogTextPunkts = "пункту " + ItogTextPunkts

        ItogTextGOST = self.TextGost.toPlainText()
        if ItogTextGOST != "":
            ItogTextGOST = ", а также пунктам " + ItogTextGOST
        else:
            ItogTextGOST = ""
            Gost = ""
            Gost2 = ""

        ShortFIO = FullFioLogin.split()
        ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."

        TipGU_VO_DO = self.TipGU.currentText()
        ItogVivod = self.Itog2Gr.currentText()


        if TipGU_VO_DO == "АК" and ItogVivod == "Полож.":
            path = "SourceGitHub/Shablon/Group2/Подожительный/Увед. полож. АКК.docx"
        elif (TipGU_VO_DO == "РОА" or TipGU_VO_DO == "РОА+ИМОД") and ItogVivod == "Полож.":
            path = "SourceGitHub/Shablon/Group2/Подожительный/Увед. полож. РОА.docx"
        elif TipGU_VO_DO == "АК" and ItogVivod == "Отриц. не отказ":
            path = "SourceGitHub/Shablon/Group2/ОтрицательныйНе34/Уведомление АКК выезд с перечнем.docx"
        elif (TipGU_VO_DO == "РОА" or TipGU_VO_DO == "РОА+ИМОД") and ItogVivod == "Отриц. не отказ":
            path = "SourceGitHub/Shablon/Group2/ОтрицательныйНе34/Уведомление РОА выезд с перечнем.docx"

        doc = DocxTemplate(path)

        context = {"short_name_org": ShortNaim.replace(' "', ' «').replace('"', '»'), "email_org": EmailOrg,
                   "fio_exp": FIOExp, "email_exp": EmailExp, "name_eo": NameEO.replace(' "', ' «').replace('"', '»'),
                   "email_eo": EmailEO, "nomer_gu": NomerGU, "data_eg": DatEG, "nomer_eg": NomEG,
                   "data_gu": DataPricaza, "full_name_org": FullNaimKogo.replace(' "', ' «').replace('"', '»'),
                   "acc_obl": accObl, "nomer_ral": NomerRAL, "data_ez": DatEZ, "real_punkt": ItogTextPunkts,
                   "gost_punkt": ItogTextGOST,"gost": Gost,"gost2": Gost2,
                   "dolj_podpisant": Doljnost, "fio_podpisant": Podpisant, "short_isp": ShortFIO}

        doc.render(context)
        try:
            doc.save(
                "SourceGitHub/ГУ/1 Группа/" + FullGU + "/3. Уведомление по ВО (полож.) " + TipGU_VO_DO + " " + FullGU + ".docx")
            os.startfile(
                "SourceGitHub/ГУ/1 Группа/" + FullGU + "/3. Уведомление по ВО (полож.) " + TipGU_VO_DO + " " + FullGU + ".docx")
        except:
            print("Ошибка при создании уведомления")
    def CreateSzOtricGroup2(self):
        global FullNaimKogo, ALLTeh
        global FullFioLogin
        global MsgError

        if self.FioExpert.currentText() == "Эксперт не найден":
            MsgError = "Выберете эксперта"
            self.gotoAlarm()
            return
        if self.EmailRAL.toPlainText() == "":
            MsgError = "Введите эл. почту в разделе \n(Данные по организации)\n(Не обязательно, но если Вам не трудно, заполните ещё почтовый адрес (если он пустой))"
            self.gotoAlarm()
            return
        if FullFioLogin == "Приостановка":
            MsgError = "Нельзя использовать с приостановленной ГУ"
            self.gotoAlarm()
            return
        if self.Ispolnitel2Gr.currentText() == "Исполнитель не выбран":
            MsgError = "Выберете исполнителя по 2 Группе"
            self.gotoAlarm()
            return
        if self.DataEZT.text() == "":
            MsgError = "Введите даты ЭЗ"
            self.gotoAlarm()
            return
        if self.NomerEZ.text() == "":
            MsgError = "Введите номер ЭЗ"
            self.gotoAlarm()
            return
        if self.VivodEZ.currentText() == "":
            MsgError = "Введите Вывод по ЭЗ"
            self.gotoAlarm()
            return
        if self.Itog2Gr.currentText() == "":
            MsgError = "Введите Итог по ЭЗ"
            self.gotoAlarm()
            return
        #############

        self.FixFullNaimKogo()  # испраить ООО на падеж
        FullGU = self.NomerGU.toPlainText()
        FullGUSplit = FullGU.split()
        DataPricaza = FullGUSplit[2]
        DataPricazaSplit = DataPricaza.replace('.', ' ')
        DataPricazaSplit = DataPricazaSplit.split()
        DataPricaza0 = DataPricazaSplit[0]
        DataPricaza1 = DataPricazaSplit[1]
        DataPricaza2 = DataPricazaSplit[2]
        DataPricaza0 = DataPricaza0.replace('01', '1').replace('02', '2').replace('03', '3').replace('04', '4').replace(
            '05', '5').replace('06', '6').replace('07', '7').replace('08', '8').replace('09', '9')
        DataPricaza1 = DataPricaza1.replace('01', 'января').replace('02', 'февралья').replace('03', 'марта').replace(
            '04', 'апреля').replace('05', 'мая').replace('06', 'июня').replace('07', 'июля').replace('08',
                                                                                                     'августа').replace(
            '09', 'сентября').replace('10', 'октября').replace('11', 'ноября').replace('12', 'декабря')
        DataPricaza = DataPricaza0 + " " + DataPricaza1 + " " + DataPricaza2 + " г."
        NomerGU = FullGUSplit[0]

        NomerRAL = self.NomerRAL.toPlainText()
        ShortNaim = self.ShortName_2.toPlainText()
        EmailOrg = self.EmailRAL.toPlainText()
        FIOExp = self.FioExpert.currentText()
        EmailExp = self.EmailExpert.toPlainText()
        NameEO = self.REO.currentText()
        EmailEO = self.MailREO.toPlainText()
        NomEZ = self.NomerEZ.toPlainText()
        DataEZ = self.DataEZT.text()
        DatEZSplit = DataEZ.replace('.', ' ')
        DatEZSplit = DatEZSplit.split()
        DatEZ0 = DatEZSplit[0]
        DatEZ1 = DatEZSplit[1]
        DatEZ2 = DatEZSplit[2]
        DatEZ0 = DatEZ0.replace('01', '1').replace('02', '2').replace('03', '3').replace('04', '4').replace(
            '05', '5').replace('06', '6').replace('07', '7').replace('08', '8').replace('09', '9')
        DatEZ1 = DatEZ1.replace('01', 'января').replace('02', 'февралья').replace('03', 'марта').replace(
            '04', 'апреля').replace('05', 'мая').replace('06', 'июня').replace('07', 'июля').replace('08',
                                                                                                     'августа').replace(
            '09', 'сентября').replace('10', 'октября').replace('11', 'ноября').replace('12', 'декабря')
        DatEZ = DatEZ0 + " " + DatEZ1 + " " + DatEZ2 + " г."



        Podpisant = self.FioExpert_12.currentText()
        if Podpisant == "Шкабура В.В.":
            Podpisant = "В.В. Шкабура"
            Doljnost = "Начальник отдела аккредитации \nв отдельных сферах \nУправления аккредитации в сфере \nдобровольного подтверждения \nсоответствия, метрологии и иных \nсферах деятельности"
        elif Podpisant == "Бухарова А.В.":
            Podpisant = "А.В. Бухарова"
            Doljnost = "Начальник отдела аккредитации \nиспытательных лабораторий \nУправления аккредитации \nв сфере добровольного подтверждения \nсоответствия, метрологии \nи иных сферах деятельности"
        elif Podpisant == "Гоголев Д.В.":
            Podpisant = "Д.В. Гоголев"
            Doljnost = "Заместитель руководителя"
            DoljnostBIG = "Заместителю руководителя\n Федеральной службы \n по аккредитации"
            PodpisantBIG = "Д.В. Гоголеву"
            FULLPodpisant = "Уважаемый Дмитрий Владимирович!"
        elif Podpisant == "Белогуров С.И.":
            Podpisant = "С.И. Белогуров"
            Doljnost = "Заместитель начальника отдела \nаккредитации в отдельных сферах \nУправления аккредитации в сфере \nдобровольного подтверждения \nсоответствия, метрологии и иных \nсферах деятельности"
        elif Podpisant == "Золотаревский С.Ю.":
            DoljnostBIG = "Начальнику Управления аккредитации \nв сфере добровольного \nподтверждения соответствия, \nметрологии и иных сферах деятельности"
            PodpisantBIG = "С.Ю. Золотаревскому"
            FULLPodpisant = "Уважаемый Сергей Юрьевич!"
        elif Podpisant == "Макаров А.Н.":
            Podpisant = "А.Н. Макаров"
            Doljnost = "И.о. начальника Управления аккредитации \nв сфере подтверждения соответствия \nпродукции машиностроения, \nэлектротехнической продукции, \nстроительных материалов \nи пожарной безопасности"
        elif Podpisant == "Хазиева А.А.":
            Podpisant = "А.А. Хазиева"
            Doljnost = "Начальник отдела аккредитации в сфере \nподтверждения соответствия продукции машиностроения \nУправления аккредитации в сфере подтверждения \nсоответствия продукции машиностроения, \nэлектротехнической продукции, строительных \nматериалов и пожарной безопасности"

        Komy = self.FioExpert_11.currentText()
        if Komy == "Гоголев Д.В.":
            DoljnostBIG = "Заместителю руководителя\n Федеральной службы \n по аккредитации"
            PodpisantBIG = "Д.В. Гоголеву"
            FULLPodpisant = "Уважаемый Дмитрий Владимирович!"
        elif Komy == "Золотаревский С.Ю.":
            DoljnostBIG = "Начальнику Управления \aаккредитации \aв сфере добровольного \aподтверждения соответствия, \aметрологии и иных сферах \aдеятельности"
            PodpisantBIG = "С.Ю. Золотаревскому"
            FULLPodpisant = "Уважаемый Сергей Юрьевич!"
        elif Podpisant == "Макаров А.Н.":
            Podpisant = "А.Н. Макаров"
            PodpisantBIG = "А.Н. Макарову"
            FULLPodpisant = "Уважаемый Андрей Николаевич!"
            Doljnost = "И.о. начальника Управления аккредитации \nв сфере подтверждения соответствия \nпродукции машиностроения, \nэлектротехнической продукции, \nстроительных материалов \nи пожарной безопасности"
        elif Podpisant == "Хазиева А.А.":
            Podpisant = "А.А. Хазиева"
            Doljnost = "Начальник отдела аккредитации в сфере \nподтверждения соответствия продукции машиностроения \nУправления аккредитации в сфере подтверждения \nсоответствия продукции машиностроения, \nэлектротехнической продукции, строительных \nматериалов и пожарной безопасности"

        Gost = ""
        Gost2 = ""
        ObshObl = self.ObshObl.currentText()
        Obl = self.Oblast_Button.text()
        if ObshObl == "ИЛ":
            accObl = "в качестве испытательной лаборатории"
            Gost = " ГОСТ ISO/IEC 17025-2019 «Общие требования к компетентности испытательных и калибровочных лабораторий» (далее – ГОСТ ISO/IEC 17025-2019)"
            Gost2 = " и требованиям ГОСТ ISO/IEC 17025-2019"
        elif ObshObl == "Метрология" and Obl == "ПСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по поверке средств измерений"
        elif ObshObl == "Метрология" and Obl == "КСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по калибровке средств измерений"
            Gost = " ГОСТ ISO/IEC 17025-2019 «Общие требования к компетентности испытательных и калибровочных лабораторий» (далее – ГОСТ ISO/IEC 17025-2019)"
            Gost2 = " и требованиям ГОСТ ISO/IEC 17025-2019"
        elif ObshObl == "Метрология" and Obl == "ИСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по испытанию средств измерений в целях утверждения типа"
        elif ObshObl == "Метрология" and Obl == "ИСО":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по испытанию стандартных образцов в целях утверждения типа"
        elif ObshObl == "Метрология" and Obl == "АМ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по аттестации методик (методов) измерений"
        elif ObshObl == "Метрология" and Obl == "МЭ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по метрологической экспертизе"
        elif ObshObl == "Метрология" and Obl == "АМ и МЭ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по аттестации методик (методов) измерений и метрологической экспертизе"
        elif ObshObl == "Метрология" and Obl == "17020":
            accObl = "в качестве органа инспекции"
            Gost = " ГОСТ Р ИСО/МЭК 17020-2012 «Оценка соответствия. Требования к работе различных типов органов инспекции» (далее – ГОСТ Р ИСО/МЭК 17020-2012)"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 17020-2012"
        elif ObshObl == "Метрология" and Obl == "17021":
            accObl = "в качестве органа по сертификации систем менеджмента"
            Gost = " ГОСТ Р ИСО/МЭК 17021-1-2017 «Оценка соответствия. Требования к органам, проводящим аудит и сертификацию систем менеджмента» (далее – ГОСТ Р ИСО/МЭК 17021-1-2017)"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 17021-1-2017"
        elif ObshObl == "Метрология" and Obl == "17024":
            accObl = "в качестве органа по сертификации персонала"
            Gost = " ГОСТ Р ИСО/МЭК 17024-1-2017 «Оценка соответствия. Общие требования к органам, проводящим сертификацию персонала» (далее – ГОСТ Р ИСО/МЭК 17024-1-2017)"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 17024-1-2017"
        elif ObshObl == "Метрология" and Obl == "17043":
            accObl = "в качестве провайдера межлабораторных сличительных испытаний"  # Не уверен
            Gost = " ГОСТ Р ИСО/МЭК 17043-1-2017 «Основные требования к проведению проверки квалификации» (далее – ГОСТ Р ИСО/МЭК 17043-1-2017)"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 17043-1-2017"
        elif ObshObl == "Метрология" and Obl == "17065":
            accObl = "в качестве органа по сертификации продукции, услуг"
            Gost = " ГОСТ Р ИСО/МЭК 17065-1-2012 «Оценка соответствия. Требования к органам по сертификации продукции, процессов и услуг» (далее – ГОСТ Р ИСО/МЭК 17065-1-2012)"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 17065-1-2012"
        elif ObshObl == "Метрология" and Obl == "15189":
            accObl = "в качестве испытательной лаборатории"
        elif ObshObl == "Метрология" and Obl == "14065":
            accObl = "в качестве органа по валидации и верификации парниковых газов"
            Gost = " ГОСТ Р ИСО/МЭК 14065-2014 «Газы парниковые. Требования к органам по валидации и верификации парниковых газов для их применения при аккредитации или других формах признания» (далее – ГОСТ Р ИСО/МЭК 14065-2014)"
            Gost2 = " и требованиям ГОСТ Р ИСО/МЭК 14065-2014"
        elif ObshObl == "Метрология" and Obl == "Не выбрано":
            accObl = "в качестве испытательной лаборатории"  # Потом доработать исключение

        In39TextPunkts = self.TextIn39.toPlainText()
        NoIn39TextPunkts = self.TextNoIn39.toPlainText()
        ItogTextPunkts = self.ItogText2GR.toPlainText()
        if (len(NoIn39TextPunkts) > 5) or (In39TextPunkts != "" and NoIn39TextPunkts != ""):
            ItogTextPunkts = "пунктам " + ItogTextPunkts
        else:
            ItogTextPunkts = "пункту " + ItogTextPunkts


        ItogTextGOST = self.TextGost.toPlainText()
        if ItogTextGOST != "":
            ItogTextGOST = " и пунктам " + ItogTextGOST
        else:
            ItogTextGOST = ""
            Gost = ""
            Gost2 = ""
        FullFioLoginFix = self.Ispolnitel2Gr.currentText()
        ShortFIO = FullFioLoginFix.split()
        ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."

        TipGU_VO_DO = self.TipGU.currentText()


        if TipGU_VO_DO == "АК":
            path = "SourceGitHub/Shablon/Group2/ОтрицательныйНе34/СЗ АКК выезд с перечнем.docx"
        elif TipGU_VO_DO == "РОА" or TipGU_VO_DO == "РОА+ИМОД":
            path = "SourceGitHub/Shablon/Group2/ОтрицательныйНе34/СЗ РОА выезд с перечнем.docx"

        doc = DocxTemplate(path)

        context = {"short_name_org": ShortNaim.replace(' "', ' «').replace('"', '»'), "email_org": EmailOrg,
                   "fio_exp": FIOExp, "email_exp": EmailExp, "name_eo": NameEO.replace(' "', ' «').replace('"', '»'),
                   "email_eo": EmailEO, "nomer_gu": NomerGU,
                   "data_gu": DataPricaza, "full_name_org": FullNaimKogo.replace(' "', ' «').replace('"', '»'),
                   "acc_obl": accObl, "nomer_ral": NomerRAL, "data_ez": DatEZ, "nomer_ez": DatEZ,
                   "dolj_podpisant": Doljnost, "fio_podpisant": Podpisant, "short_isp": ShortFIO,
                   "real_punkt": ItogTextPunkts, "gost_punkt": ItogTextGOST,"gost": Gost,"gost2": Gost2,
                   "big_dolj": DoljnostBIG, "big_fio": PodpisantBIG, "full_big_fio": FULLPodpisant,}

        doc.render(context)
        try:
            doc.save(
                "SourceGitHub/ГУ/1 Группа/" + FullGU + "/4. СЗ по ВО (отриц. не отказ) " + TipGU_VO_DO + " " + FullGU + ".docx")
            os.startfile(
                "SourceGitHub/ГУ/1 Группа/" + FullGU + "/4. СЗ по ВО (отриц. не отказ) " + TipGU_VO_DO + " " + FullGU + ".docx")
        except:
            print("Ошибка при создании уведомления")
#---------------------------------
    def CreateYvedEG(self):
        global FullNaimKogo, ALLTeh
        global FullFioLogin
        global MsgError

        if self.TipGU.currentText() == "ПК" or self.TipGU.currentText() == "ПК+РОА" or self.TipGU.currentText() == "ПК+ИМОД" or self.TipGU.currentText() == "ПК+РОА+ИМОД":
            MsgError = "Уточните Тип ГУ"
            self.gotoAlarm()
            return
        if self.FioExpert.currentText() == "Эксперт не найден":
            MsgError = "Выберете эксперта"
            self.gotoAlarm()
            return
        if self.EmailRAL.toPlainText() == "":
            MsgError = "Введите эл. почту в разделе \n(Данные по организации)\n(Не обязательно, но если Вам не трудно, заполните ещё почтовый адрес (если он пустой))"
            self.gotoAlarm()
            return
        if self.FioExpert.currentText() == "Метрология" and (
                self.Oblast_Button.text() == "Не выбрано" or self.Oblast_Button.text() == ""):
            MsgError = "Уточните конкретную область"
            self.gotoAlarm()
            return
        if self.Teh_1.currentText() == "":
            MsgError = "Тех. Эксперт не выбран"
            self.gotoAlarm()
            return
        if self.VMeropr.isChecked() == True and self.Vibor_TU_2.currentText() == "Тер. Управление":
            MsgError = "Выбрано поле (Выездное мероприятие по ОС), но не вбрано ТУ"
            self.gotoAlarm()
            return
        if FullFioLogin == "Приостановка":
            MsgError = "Нельзя использовать с приостановленной ГУ"
            self.gotoAlarm()
            return



        cal = Russia()
        datenow = QDate.currentDate().toPyDate()
        datta = cal.add_working_days(datenow, 2)  # оОчень важно по датам
        # datta = cal.get_working_days_delta(date(2018, 4, 2), date(2018, 6, 17)) # считает даты в рабочем
        datenow = str(datenow)
        datenow = datenow.replace("-", " ")
        datenow = datenow.split()
        datenowStr = datenow[2] + "." + datenow[1] + "." + datenow[0]
        self.DataYvedEGT.setText(datenowStr)
        #############

        self.FixFullNaimKogo()  # испраить ООО на падеж
        FullGU = self.NomerGU.toPlainText()
        FullGUSplit = FullGU.split()
        DataPricaza = FullGUSplit[2]
        DataPricazaSplit = DataPricaza.replace('.', ' ')
        DataPricazaSplit = DataPricazaSplit.split()
        DataPricaza0 = DataPricazaSplit[0]
        DataPricaza1 = DataPricazaSplit[1]
        DataPricaza2 = DataPricazaSplit[2]
        DataPricaza0 = DataPricaza0.replace('01', '1').replace('02', '2').replace('03', '3').replace('04', '4').replace(
            '05', '5').replace('06', '6').replace('07', '7').replace('08', '8').replace('09', '9')
        DataPricaza1 = DataPricaza1.replace('01', 'января').replace('02', 'февралья').replace('03', 'марта').replace(
            '04', 'апреля').replace('05', 'мая').replace('06', 'июня').replace('07', 'июля').replace('08',
                                                                                                     'августа').replace(
            '09', 'сентября').replace('10', 'октября').replace('11', 'ноября').replace('12', 'декабря')
        DataPricaza = DataPricaza0 + " " + DataPricaza1 + " " + DataPricaza2 + " г."
        NomerGU = FullGUSplit[0]

        NomerRAL = self.NomerRAL.toPlainText()
        ShortNaim = self.ShortName_2.toPlainText()
        EmailOrg = self.EmailRAL.toPlainText()
        FIOExp = self.FioExpert.currentText()
        EmailExp = self.EmailExpert.toPlainText()
        NameEO = self.REO.currentText()
        EmailEO = self.MailREO.toPlainText()

        if self.Teh_1.currentText() != "":
            ALLTeh = self.Teh_1.currentText()
            textAllTeh = "член экспертной группы (технический эксперт)"
        if self.Teh_2.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_2.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"
        if self.Teh_3.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_3.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"
        if self.Teh_4.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_4.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"
        if self.Teh_5.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_5.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"
        if self.Teh_6.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_6.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"
        if self.Teh_7.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_7.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"
        if self.Teh_8.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_8.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"
        if self.Teh_9.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_9.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"
        if self.Teh_10.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_10.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"

        ObshObl = self.ObshObl.currentText()
        Obl = self.Oblast_Button.text()
        if ObshObl == "ИЛ":
            accObl = "в качестве испытательной лаборатории"
        elif ObshObl == "Метрология" and Obl == "ПСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по поверке средств измерений"
        elif ObshObl == "Метрология" and Obl == "КСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по калибровке средств измерений"
        elif ObshObl == "Метрология" and Obl == "ИСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по испытанию средств измерений в целях утверждения типа"
        elif ObshObl == "Метрология" and Obl == "ИСО":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по испытанию стандартных образцов в целях утверждения типа"
        elif ObshObl == "Метрология" and Obl == "АМ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по аттестации методик (методов) измерений"
        elif ObshObl == "Метрология" and Obl == "МЭ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по метрологической экспертизе"
        elif ObshObl == "Метрология" and Obl == "АМ и МЭ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по аттестации методик (методов) измерений и метрологической экспертизе"
        elif ObshObl == "Метрология" and Obl == "17020":
            accObl = "в качестве органа инспекции"
        elif ObshObl == "Метрология" and Obl == "17021":
            accObl = "в качестве органа по сертификации систем менеджмента"
        elif ObshObl == "Метрология" and Obl == "17024":
            accObl = "в качестве органа по сертификации персонала"
        elif ObshObl == "Метрология" and Obl == "17043":
            accObl = "в качестве провайдера межлабораторных сличительных испытаний"  # Не уверен
        elif ObshObl == "Метрология" and Obl == "17065":
            accObl = "в качестве органа по сертификации продукции, услуг"
        elif ObshObl == "Метрология" and Obl == "15189":
            accObl = "в качестве испытательной лаборатории"
        elif ObshObl == "Метрология" and Obl == "14065":
            accObl = "в качестве органа по валидации и верификации парниковых газов"
        elif ObshObl == "Метрология" and Obl == "Не выбрано":
            accObl = "в качестве испытательной лаборатории"  # Потом доработать исключение

        Podpisant = self.FioExpert_5.currentText()
        if Podpisant == "Шкабура В.В.":
            Podpisant = "В.В. Шкабура"
            Doljnost = "Начальник отдела аккредитации \nв отдельных сферах \nУправления аккредитации в сфере \nдобровольного подтверждения \nсоответствия, метрологии и иных \nсферах деятельности"
        elif Podpisant == "Бухарова А.В.":
            Podpisant = "А.В. Бухарова"
            Doljnost = "Начальник отдела аккредитации \nиспытательных лабораторий \nУправления аккредитации \nв сфере добровольного подтверждения \nсоответствия, метрологии \nи иных сферах деятельности"
        elif Podpisant == "Гоголев Д.В.":
            Podpisant = "Д.В. Гоголев"
            Doljnost = "Заместитель руководителя"
        elif Podpisant == "Белогуров С.И.":
            Podpisant = "С.И. Белогуров"
            Doljnost = "Заместитель начальника отдела \nаккредитации в отдельных сферах \nУправления аккредитации в сфере \nдобровольного подтверждения \nсоответствия, метрологии и иных \nсферах деятельности"
        elif Podpisant == "Макаров А.Н.":
            Podpisant = "А.Н. Макаров"
            Doljnost = "И.о. начальника Управления аккредитации \nв сфере подтверждения соответствия \nпродукции машиностроения, \nэлектротехнической продукции, \nстроительных материалов \nи пожарной безопасности"
        elif Podpisant == "Хазиева А.А.":
            Podpisant = "А.А. Хазиева"
            Doljnost = "Начальник отдела аккредитации в сфере \nподтверждения соответствия продукции машиностроения \nУправления аккредитации в сфере подтверждения \nсоответствия продукции машиностроения, \nэлектротехнической продукции, строительных \nматериалов и пожарной безопасности"

        ShortFIO = FullFioLogin.split()
        ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."

        TipGU_VO_DO = self.TipGU.currentText()
        if TipGU_VO_DO == "ПК1":
            if self.VO.isChecked() == True:
                TipGU = "подтверждения компетентности"
                VO_DO = "выездной экспертизы соответствия критериям аккредитации"
            elif self.DO.isChecked() == True:
                TipGU = "подтверждения компетентности"
                VO_DO = "экспертизы представленных документов и сведений"
        elif TipGU_VO_DO == "ПК2":
            TipGU = "подтверждения компетентности"
            VO_DO = "выездной экспертизы соответствия критериям аккредитации"
        elif TipGU_VO_DO == "ПК5":
            TipGU = "подтверждения компетентности"
            VO_DO = "экспертизы представленных документов и сведений и выездной экспертизы соответствия критериям аккредитации"
        elif TipGU_VO_DO == "АК":
            TipGU = "аккредитации"
            VO_DO = "экспертизы представленных документов и сведений и выездной экспертизы соответствия критериям аккредитации"
        elif TipGU_VO_DO == "РОА":
            TipGU = "расширения области аккредитации"
            VO_DO = "экспертизы представленных документов и сведений и выездной экспертизы соответствия критериям аккредитации"
        elif TipGU_VO_DO == "ПК1+РОА" or TipGU_VO_DO == "ПК2+РОА" or TipGU_VO_DO == "ПК5+РОА":
            TipGU = "подтверждения компетентности и расширения области аккредитации"
            VO_DO = "экспертизы представленных документов и сведений и выездной экспертизы соответствия критериям аккредитации"
        elif TipGU_VO_DO == "ПК1+ИМОД" or TipGU_VO_DO == "ПК2+ИМОД":
            TipGU = "подтверждения компетентности и изменения места (мест) осуществления деятельности"
            VO_DO = "выездной экспертизы соответствия критериям аккредитации"
        elif TipGU_VO_DO == "ПК5+ИМОД":
            TipGU = "подтверждения компетентности и изменения места (мест) осуществления деятельности"
            VO_DO = "экспертизы представленных документов и сведений и выездной экспертизы соответствия критериям аккредитации"
        elif TipGU_VO_DO == "РОА+ИМОД":
            TipGU = "расширения области аккредитации и изменения места (мест) осуществления деятельности"
            VO_DO = "экспертизы представленных документов и сведений и выездной экспертизы соответствия критериям аккредитации"
        elif TipGU_VO_DO == "ПК1+РОА+ИМОД" or TipGU_VO_DO == "ПК2+РОА+ИМОД" or TipGU_VO_DO == "ПК5+РОА+ИМОД":
            TipGU = "подтверждения компетентности, расширения области аккредитации и изменения места (мест) осуществления деятельности"
            VO_DO = "экспертизы представленных документов и сведений и выездной экспертизы соответствия критериям аккредитации"

        if TipGU_VO_DO == "АК":
            path = "SourceGitHub/Shablon/Group1/YvedOSostaveEG/2. Уведомление по ЭГ АКК.docx"
        elif TipGU_VO_DO == "РОА" or TipGU_VO_DO == "РОА+ИМОД":
            path = "SourceGitHub/Shablon/Group1/YvedOSostaveEG/3. Уведомление по ЭГ РОА.docx"
        else:
            path = "SourceGitHub/Shablon/Group1/YvedOSostaveEG/1. Уведомление по ЭГ.docx"


        doc = DocxTemplate(path)

        context = {"short_name_org": ShortNaim.replace(' "', ' «').replace('"', '»'), "email_org": EmailOrg,
                   "fio_exp": FIOExp, "email_exp": EmailExp, "name_eo": NameEO.replace(' "', ' «').replace('"', '»'),
                   "email_eo": EmailEO, "nomer_gu": NomerGU,
                   "data_gu": DataPricaza, "full_name_org": FullNaimKogo.replace(' "', ' «').replace('"', '»'),
                   "acc_obl": accObl, "nomer_ral": NomerRAL, "nomer_gu2": FullGU, "all_teh": ALLTeh,
                   "text_teh": textAllTeh,
                   "dolj_podpisant": Doljnost, "fio_podpisant": Podpisant, "short_isp": ShortFIO, "vo_do": VO_DO,
                   "tip_gu": TipGU}

        doc.render(context)
        try:
            doc.save(
                "SourceGitHub/ГУ/1 Группа/" + FullGU + "/1. Уведомление по ЭГ " + FullGU + ".docx")
            os.startfile(
                "SourceGitHub/ГУ/1 Группа/" + FullGU + "/1. Уведомление по ЭГ " + FullGU + ".docx")
        except:
            print("Ошибка при создании уведомления")

    def likeanim(self):
        self.label_save.show()
        self.movie = QMovie("SourceGitHub/Pictures/Darg2.gif")
        self.label_save.setMovie(self.movie)
        self.movie.start()
        sleep(2)
        self.label_save.hide()

    def SaveALLS(self, result, ):
        try:
            connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
            cur = connection.cursor()
            global z
            global globalNomerGU
            global FindNumGU
            global Deadline
            global DeadlineDay
            global MsgError

            try:
                if self.DataPricaza.text() != "":
                    DateCheck = datetime.strptime(self.DataPricaza.text(), '%d.%m.%Y')
            except:
                MsgError = "Ошибка в дате Решения по 1 группе"
                self.DataPricaza.setStyleSheet('background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(243, 34, 6)')
                self.gotoAlarm()
                self.SaveGif.hide()
                return
            try:
                if self.DataPricaz2GRT.text() != "":
                    DateCheck = datetime.strptime(self.DataPricaz2GRT.text(), '%d.%m.%Y')
            except:
                MsgError = "Ошибка в дате Решения по 2 группе"
                self.DataPricaz2GRT.setStyleSheet('background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(243, 34, 6)')
                self.gotoAlarm()
                self.SaveGif.hide()
                return
            try:
                if self.DataYvedT.text() != "":
                    DateCheck = datetime.strptime(self.DataYvedT.text(), '%d.%m.%Y')
            except:
                MsgError = "Ошибка в дате Уведомления эксперта"
                self.DataYvedT.setStyleSheet('background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(243, 34, 6)')
                self.gotoAlarm()
                self.SaveGif.hide()
                return
            try:
                if self.DataYvedEGT.text() != "":
                    DateCheck = datetime.strptime(self.DataYvedEGT.text(), '%d.%m.%Y')
            except:
                MsgError = "Ошибка в дате Уведомления о составе ЭГ"
                self.DataYvedEGT.setStyleSheet('background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(243, 34, 6)')
                self.gotoAlarm()
                self.SaveGif.hide()
                return
            try:
                if self.DataPricaz3GRT.text() != "":
                    DateCheck = datetime.strptime(self.DataPricaz3GRT.text(), '%d.%m.%Y')
            except:
                MsgError = "Ошибка в дате Решения по 3 группе"
                self.DataPricaz3GRT.setStyleSheet('background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(243, 34, 6)')
                self.gotoAlarm()
                self.SaveGif.hide()
                return

            try:
                if self.DataDogT.text() != "":
                    DateCheck = datetime.strptime(self.DataDogT.text(), '%d.%m.%Y')
            except:
                MsgError = "Ошибка в дате Догвора"
                self.DataDogT.setStyleSheet('background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(243, 34, 6)')
                self.gotoAlarm()
                self.SaveGif.hide()
                return
            try:
                if self.DataEZT.text() != "":
                    DateCheck = datetime.strptime(self.DataEZT.text(), '%d.%m.%Y')
            except:
                MsgError = "Ошибка в дате ЭЗ"
                self.DataEZT.setStyleSheet('background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(243, 34, 6)')
                self.gotoAlarm()
                self.SaveGif.hide()
                return
            try:
                if self.DataActaT.text() != "":
                    DateCheck = datetime.strptime(self.DataActaT.text(), '%d.%m.%Y')
            except:
                MsgError = "Ошибка в дате Акта"
                self.DataActaT.setStyleSheet('background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(243, 34, 6)')
                self.gotoAlarm()
                self.SaveGif.hide()
                return
            try:
                if self.DataEG.text() != "":
                    DateCheck = datetime.strptime(self.DataEG.text(), '%d.%m.%Y')
            except:
                MsgError = "Ошибка в дате приказа ЭГ"
                self.DataEG.setStyleSheet('background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(243, 34, 6)')
                self.gotoAlarm()
                self.SaveGif.hide()
                return
            try:
                if self.DataAct.text() != "":
                    DateCheck = datetime.strptime(self.DataAct.text(), '%d.%m.%Y')
            except:
                MsgError = "Ошибка в дате Акта"
                self.DataAct.setStyleSheet('background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(243, 34, 6)')
                self.gotoAlarm()
                self.SaveGif.hide()
                return
            try:
                if self.DataAct_2.text() != "":
                    DateCheck = datetime.strptime(self.DataAct_2.text(), '%d.%m.%Y')
            except:
                MsgError = "Ошибка в дате Акта"
                self.DataAct_2.setStyleSheet('background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(243, 34, 6)')
                self.gotoAlarm()
                self.SaveGif.hide()
                return
            try:
                if self.DataEZ_2.text() != "":
                    DateCheck = datetime.strptime(self.DataEZ_2.text(), '%d.%m.%Y')
            except:
                MsgError = "Ошибка в дате ЭЗ"
                self.DataEZ_2.setStyleSheet('background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(243, 34, 6)')
                self.gotoAlarm()
                self.SaveGif.hide()
                return

            self.DataPricaza.setStyleSheet(
                'background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(0, 0, 0)')
            self.DataPricaz2GRT.setStyleSheet(
                'background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(0, 0, 0)')
            self.DataPricaz3GRT.setStyleSheet(
                'background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(0, 0, 0)')
            self.DataDogT.setStyleSheet(
                'background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(0, 0, 0)')
            self.DataEZT.setStyleSheet(
                'background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(0, 0, 0)')
            self.DataActaT.setStyleSheet(
                'background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(0, 0, 0)')
            self.DataEG.setStyleSheet(
                'background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(0, 0, 0)')
            self.DataAct.setStyleSheet(
                'background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(0, 0, 0)')
            self.DataAct_2.setStyleSheet(
                'background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(0, 0, 0)')
            self.DataEZ_2.setStyleSheet(
                'background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(0, 0, 0)')
            self.DataYvedT.setStyleSheet(
                'background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(0, 0, 0)')
            self.DataYvedEGT.setStyleSheet(
                'background-color:rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; border-radius: 10px; color: rgb(0, 0, 0)')



            self.SaveGif.show()
            self.movie = QMovie("SourceGitHub/Pictures/well-done.gif")
            self.SaveGif.setMovie(self.movie)
            self.movie.start()
            self.animation1 = QPropertyAnimation(self.SaveGif, b"geometry")  # Animate minimumWidht
            self.animation1.setStartValue(QRect(731, -100, 70, 70))  # Start value is the current menu width
            self.animation1.setEndValue(QRect(731, -16, 70, 70))  # end value is the new menu width
            self.animation1.setDuration(500)
            self.animation1.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation1.start()






            FindGU = globalNomerGU
            NewShortName = self.ShortName_2.toPlainText()
            NewNomerRAL = self.NomerRAL.toPlainText()
            NewFullNaim = self.FullNaim.toPlainText()
            NewEmailRAL = self.EmailRAL.toPlainText()
            NewYrAdres = self.YrAdres.toPlainText()
            NewTipGU = self.TipGU_3.currentText()
            self.TipGU.setCurrentText(NewTipGU)
            NewFioExpert = self.FioExpert.currentText()
            NewEmailExpert = self.EmailExpert.toPlainText()
            NewREO = self.REO.currentText()
            NewMailREO = self.MailREO.toPlainText()
            NewNumberTeh = self.NumberTeh.currentText()
            NewTeh_1 = self.Teh_1.currentText()
            NewEoTeh_1 = self.EoTeh_1.toPlainText()
            NewTeh_2 = self.Teh_2.currentText()
            NewEoTeh_2 = self.EoTeh_2.toPlainText()
            NewTeh_3 = self.Teh_3.currentText()
            NewEoTeh_3 = self.EoTeh_3.toPlainText()
            NewTeh_4 = self.Teh_4.currentText()
            NewEoTeh_4 = self.EoTeh_4.toPlainText()
            NewTeh_5 = self.Teh_5.currentText()
            NewEoTeh_5 = self.EoTeh_5.toPlainText()
            NewTeh_6 = self.Teh_6.currentText()
            NewEoTeh_6 = self.EoTeh_6.toPlainText()
            NewTeh_7 = self.Teh_7.currentText()
            NewEoTeh_7 = self.EoTeh_7.toPlainText()
            NewTeh_8 = self.Teh_8.currentText()
            NewEoTeh_8 = self.EoTeh_8.toPlainText()
            NewTeh_9 = self.Teh_9.currentText()
            NewEoTeh_9 = self.EoTeh_9.toPlainText()
            NewTeh_10 = self.Teh_10.currentText()
            NewEoTeh_10 = self.EoTeh_10.toPlainText()
            NewObshObl = self.ObshObl.currentText()
            NewDataYvedT = self.DataYvedT.text()
            NewSrokD = self.SrokD.toPlainText()
            NewDataPricaz1GR = self.DataPricaz1GR.text()
            if NewDataPricaz1GR == "01.01.2000":
                NewDataPricaz1GR = None
            NewNomerPricaz1Gr = self.NomerPricaz1Gr.toPlainText()
            NewPriznakSK = self.PriznakSK.currentText()
            NewComment = self.Comment.toPlainText()
            NewPerevibor = self.Perevibor.currentText()
            NewPereviborDate = self.PereviborDate.text()
            if NewPereviborDate == "01.01.2000":
                NewPereviborDate = None
            NewStatusGU = self.StatusGU.currentText()
            NewDataYvedEGT = self.DataYvedEGT.text()
            NewDataRospExp = self.DataRospExp.text()
            if NewDataRospExp == "01.01.2000":
                NewDataRospExp = None
            NewDatePredlog = self.DatePredlog.text()
            if NewDatePredlog == "01.01.2000":
                NewDatePredlog = None
            NewDataPricaza = self.DataPricaza.text()
            if NewDataPricaza == "01.01.2000":
                NewDataPricaza = None
            if self.NomerEG.toPlainText() != "":
                NewDataPricaza = "№"+ self.NomerEG.toPlainText() + " от " + self.DataEG.text() + "г."

            NewTochnObl = self.Oblast_Button.text()
            NewUrl1C = self.Url1CGU.toPlainText()
            NewUrlCloud = self.URLCloud.toPlainText()
            NewVibor_TU = self.Vibor_TU.currentText()
            NewRucovod_TU = self.Rucovod_TU.toPlainText()
            if self.VKS.isChecked() == True:
                newVKS = "Да"
            else:
                newVKS = "Нет"
            if self.VMeropr.isChecked() == True:
                newVMeropr = "Да"
            else:
                newVMeropr = "Нет"
            NewDataDogT = self.DataDogT.text()
            NewDataEZT = self.DataEZT.text()
            NewVivodEZ = self.VivodEZ.currentText()
            NewItog2Gr = self.Itog2Gr.currentText()
            NewDataPricaz2GR = self.DataPricaz2GRT.text()
            NewTextIn39 = self.TextIn39.toPlainText()
            NewTextNoIn39 = self.TextNoIn39.toPlainText()
            NewTextGost = self.TextGost.toPlainText()
            NewItogText2GR = self.ItogText2GR.toPlainText()
            NewItogVivod = self.ItogVivod.toPlainText()
            NewIspolnitel3Gr = self.Ispolnitel3Gr.currentText()
            NewDataActaT = self.DataActaT.text()
            NewVivodActa = self.VivodActa.currentText()
            NewRisk = self.Risk.currentText()
            NewIspolnitel2Gr = self.Ispolnitel2Gr.currentText()
            NewComment_2 = self.Comment_2.toPlainText()
            NewComment_3 = self.Comment_3.toPlainText()
            NewEtap3GR = self.Etap3GR.currentText()
            NewNomerAct = self.NomerAct.toPlainText()
            NewDataPricaz3GRT = self.DataPricaz3GRT.text()
            if self.DopZapr.isChecked() == True:
                newDopZapr = "Да"
            else:
                newDopZapr = "Нет"
            if self.Voz.isChecked() == True:
                newVoz = "Да"
            else:
                newVoz = "Нет"
            if self.Sogl.isChecked() == True:
                newSogl = "Да"
            else:
                newSogl = "Нет"
            if self.Sokr.isChecked() == True:
                newSokr = "Да"
            else:
                newSokr = "Нет"
            NewDataAct_2 = self.DataAct_2.text()
            NewDataEZ_2 = self.DataEZ_2.text()
            NewSoprovodItog = self.SoprovodItog.currentText()


            cur.execute(
                "UPDATE GU1Group SET Наименованиезаявителя = ?, РегистрационныйномерАЛвРАЛ = ?, "
                "Полноенаименованиезаявителя = ?, ЭлектроннаяПочтаЗаявителя = ?, ЮрАдресЗаявителя = ?, ТипГУ = ?, "
                "ФИОЭкспертапоаккредитации = ?, ЭлПочтаЭксперта = ?, ЭоЭксперта = ?, ЭлПочтаЭоЭксперта = ?, КолВоТех = ?, "
                " ФИОТех1 = ?, ЭлТех1 = ?, ФИОТех2 = ?, ЭлТех2 = ?, ФИОТех3 = ?, ЭлТех3 = ?, ФИОТех4 = ?, ЭлТех4 = ?, "
                "ФИОТех5 = ?, ЭлТех5 = ?, ФИОТех6 = ?, ЭлТех6 = ?, ФИОТех7 = ?, ЭлТех7 = ?, ФИОТех8 = ?, ЭлТех8 = ?, "
                "ФИОТех9 = ?, ЭлТех9 = ?, ФИОТех10 = ?, ЭлТех10 = ?, ОбщаяОбласть = ?, ДатаУведомления = ?, СрокДней = ?, "
                "ДатаПриказа = ?, НомерПриказа = ?, ПризнакСК =?, Комментарий = ?, Перевыбор = ?, ДатаПеревыбора = ?, "
                "СтатусГУ = ?, ДатаУведЭГ = ?, Датавыбораэксперта = ?, ПредложениепосоставуЭГ = ?, ДатаприказаосоставеЭГ "
                "= ?, ТочнаяОбласть = ?, СсылкаНа1С = ?, СсылкаНаОблако = ?, ТУ = ?, РукТУ = ?, ВКС = ?, "
                "ВМероприятиеОС = ?, Договор1Г = ?, ДатаЭЗ = ?, ВыводЭЗ = ?, Итог2Гр = ?, ДатаПриказа2гр = ?, "
                "В34Группа2 = ?, НеВ34Группа2 = ?, ГОСТГруппа2 = ?, ИтогТекстГруппа2 = ?, ВыводГруппа2 = ?,"
                "Исполнитель3Гр = ?, ВыводАкта = ?, РискАкта = ?, ДатаАкта = ?, Исполнитель2Гр = ?, Комментарий2Гр = ?, "
                "Комментарий3Гр = ?, Этап3гр = ?, НомерАкта = ?, ДатаПриказа3гр = ? , ДополнительныйЗапрос = ?, "
                "Возобновление3гр = ?, СогласительнаяКомиссия = ?, ПриостановкаАКилиРОА = ?, ДатаАктаДок = ?, ДатаЭзДок = ?,"
                "СопроводИтог = ? WHERE Регистрационныйномер = ?",
                (NewShortName, NewNomerRAL, NewFullNaim, NewEmailRAL, NewYrAdres, NewTipGU, NewFioExpert, NewEmailExpert,
                 NewREO, NewMailREO, NewNumberTeh, NewTeh_1, NewEoTeh_1, NewTeh_2, NewEoTeh_2, NewTeh_3, NewEoTeh_3,
                 NewTeh_4, NewEoTeh_4, NewTeh_5, NewEoTeh_5, NewTeh_6, NewEoTeh_6, NewTeh_7, NewEoTeh_7, NewTeh_8,
                 NewEoTeh_8, NewTeh_9, NewEoTeh_9, NewTeh_10, NewEoTeh_10, NewObshObl, NewDataYvedT, NewSrokD,
                 NewDataPricaz1GR, NewNomerPricaz1Gr, NewPriznakSK, NewComment, NewPerevibor, NewPereviborDate,
                 NewStatusGU, NewDataYvedEGT, NewDataRospExp, NewDatePredlog, NewDataPricaza, NewTochnObl, NewUrl1C,
                 NewUrlCloud, NewVibor_TU, NewRucovod_TU, newVKS, newVMeropr, NewDataDogT, NewDataEZT, NewVivodEZ,
                 NewItog2Gr, NewDataPricaz2GR, NewTextIn39, NewTextNoIn39, NewTextGost, NewItogText2GR, NewItogVivod,
                 NewIspolnitel3Gr, NewVivodActa, NewRisk, NewDataActaT,NewIspolnitel2Gr, NewComment_2, NewComment_3,
                 NewEtap3GR, NewNomerAct, NewDataPricaz3GRT, newDopZapr, newVoz, newSogl, newSokr, NewDataAct_2,
                 NewDataEZ_2, NewSoprovodItog, FindGU,))
            result = cur.fetchall()
            connection.commit()
            connection.close()
        except:
            connection.close()
            self.SaveALLS()

        # self.close()
        # self.setFocus()
        # MainWindow()
        # MainWindow().DoTable()
        # self.DoTable()
        # MainWindow().FindButton.click()

    def TehNumberChanged(self):
        Number = self.NumberTeh.currentText()
        Number = int(Number) - 1
        self.TehSlider.setMaximum(Number)

    def TehSliderChanged(self):
        Number = self.TehSlider.value() + 1
        Number = str(Number)
        self.label_Teh.setText("ФИО тех. Эксперта " + Number)
        self.label_EOTeh.setText("ЭО тех. Эксперта " + Number)
        if Number == "1":
            self.Teh_1.show()
            self.Teh_2.hide()
            self.Teh_3.hide()
            self.Teh_4.hide()
            self.Teh_5.hide()
            self.Teh_6.hide()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.hide()
            self.Teh_10.hide()
            self.EoTeh_1.show()
            self.EoTeh_2.hide()
            self.EoTeh_3.hide()
            self.EoTeh_4.hide()
            self.EoTeh_5.hide()
            self.EoTeh_6.hide()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.hide()
            self.EoTeh_10.hide()
        elif Number == "2":
            self.Teh_1.hide()
            self.Teh_2.show()
            self.Teh_3.hide()
            self.Teh_4.hide()
            self.Teh_5.hide()
            self.Teh_6.hide()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.hide()
            self.Teh_10.hide()
            self.EoTeh_1.hide()
            self.EoTeh_2.show()
            self.EoTeh_3.hide()
            self.EoTeh_4.hide()
            self.EoTeh_5.hide()
            self.EoTeh_6.hide()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.hide()
            self.EoTeh_10.hide()
        elif Number == "3":
            self.Teh_1.hide()
            self.Teh_2.hide()
            self.Teh_3.show()
            self.Teh_4.hide()
            self.Teh_5.hide()
            self.Teh_6.hide()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.hide()
            self.Teh_10.hide()
            self.EoTeh_1.hide()
            self.EoTeh_2.hide()
            self.EoTeh_3.show()
            self.EoTeh_4.hide()
            self.EoTeh_5.hide()
            self.EoTeh_6.hide()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.hide()
            self.EoTeh_10.hide()
        elif Number == "4":
            self.Teh_1.hide()
            self.Teh_2.hide()
            self.Teh_3.hide()
            self.Teh_4.show()
            self.Teh_5.hide()
            self.Teh_6.hide()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.hide()
            self.Teh_10.hide()
            self.EoTeh_1.hide()
            self.EoTeh_2.hide()
            self.EoTeh_3.hide()
            self.EoTeh_4.show()
            self.EoTeh_5.hide()
            self.EoTeh_6.hide()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.hide()
            self.EoTeh_10.hide()
        elif Number == "5":
            self.Teh_1.hide()
            self.Teh_2.hide()
            self.Teh_3.hide()
            self.Teh_4.hide()
            self.Teh_5.show()
            self.Teh_6.hide()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.hide()
            self.Teh_10.hide()
            self.EoTeh_1.hide()
            self.EoTeh_2.hide()
            self.EoTeh_3.hide()
            self.EoTeh_4.hide()
            self.EoTeh_5.show()
            self.EoTeh_6.hide()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.hide()
            self.EoTeh_10.hide()
        elif Number == "6":
            self.Teh_1.hide()
            self.Teh_2.hide()
            self.Teh_3.hide()
            self.Teh_4.hide()
            self.Teh_5.hide()
            self.Teh_6.show()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.hide()
            self.Teh_10.hide()
            self.EoTeh_1.hide()
            self.EoTeh_2.hide()
            self.EoTeh_3.hide()
            self.EoTeh_4.hide()
            self.EoTeh_5.hide()
            self.EoTeh_6.show()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.hide()
            self.EoTeh_10.hide()
        elif Number == "7":
            self.Teh_1.hide()
            self.Teh_2.hide()
            self.Teh_3.hide()
            self.Teh_4.hide()
            self.Teh_5.hide()
            self.Teh_6.show()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.hide()
            self.Teh_10.hide()
            self.EoTeh_1.hide()
            self.EoTeh_2.hide()
            self.EoTeh_3.hide()
            self.EoTeh_4.hide()
            self.EoTeh_5.hide()
            self.EoTeh_6.show()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.hide()
            self.EoTeh_10.hide()
        elif Number == "8":
            self.Teh_1.hide()
            self.Teh_2.hide()
            self.Teh_3.hide()
            self.Teh_4.hide()
            self.Teh_5.hide()
            self.Teh_6.hide()
            self.Teh_7.hide()
            self.Teh_8.show()
            self.Teh_9.hide()
            self.Teh_10.hide()
            self.EoTeh_1.hide()
            self.EoTeh_2.hide()
            self.EoTeh_3.hide()
            self.EoTeh_4.hide()
            self.EoTeh_5.hide()
            self.EoTeh_6.hide()
            self.EoTeh_7.hide()
            self.EoTeh_8.show()
            self.EoTeh_9.hide()
            self.EoTeh_10.hide()
        elif Number == "9":
            self.Teh_1.hide()
            self.Teh_2.hide()
            self.Teh_3.hide()
            self.Teh_4.hide()
            self.Teh_5.hide()
            self.Teh_6.hide()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.show()
            self.Teh_10.hide()
            self.EoTeh_1.hide()
            self.EoTeh_2.hide()
            self.EoTeh_3.hide()
            self.EoTeh_4.hide()
            self.EoTeh_5.hide()
            self.EoTeh_6.hide()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.show()
            self.EoTeh_10.hide()
        elif Number == "10":
            self.Teh_1.hide()
            self.Teh_2.hide()
            self.Teh_3.hide()
            self.Teh_4.hide()
            self.Teh_5.hide()
            self.Teh_6.hide()
            self.Teh_7.hide()
            self.Teh_8.hide()
            self.Teh_9.hide()
            self.Teh_10.show()
            self.EoTeh_1.hide()
            self.EoTeh_2.hide()
            self.EoTeh_3.hide()
            self.EoTeh_4.hide()
            self.EoTeh_5.hide()
            self.EoTeh_6.hide()
            self.EoTeh_7.hide()
            self.EoTeh_8.hide()
            self.EoTeh_9.hide()
            self.EoTeh_10.show()

    def ChangedTeh(self):
        global Teh
        Number = self.TehSlider.value() + 1
        if Number == 1:
            Teh = self.Teh_1.currentText()
        elif Number == 2:
            Teh = self.Teh_2.currentText()
        elif Number == 3:
            Teh = self.Teh_3.currentText()
        elif Number == 4:
            Teh = self.Teh_4.currentText()
        elif Number == 5:
            Teh = self.Teh_5.currentText()
        elif Number == 6:
            Teh = self.Teh_6.currentText()
        elif Number == 7:
            Teh = self.Teh_7.currentText()
        elif Number == 8:
            Teh = self.Teh_8.currentText()
        elif Number == 9:
            Teh = self.Teh_9.currentText()
        elif Number == 10:
            Teh = self.Teh_10.currentText()

        connection = sqlite3.connect('/Portal/UseFile/BaseFrom1C.db')
        reo = connection.cursor()
        reo.execute("SELECT * FROM ReestTeh WHERE ФИОтехническогоэксперта = ?", (Teh,))
        resultREO = reo.fetchall()
        connection.commit()
        connection.close()

        if Number == 1:
            self.EoTeh_1.setPlainText(resultREO[0][30])
        elif Number == 2:
            self.EoTeh_2.setPlainText(resultREO[0][30])
        elif Number == 3:
            self.EoTeh_3.setPlainText(resultREO[0][30])
        elif Number == 4:
            self.EoTeh_4.setPlainText(resultREO[0][30])
        elif Number == 5:
            self.EoTeh_5.setPlainText(resultREO[0][30])
        elif Number == 6:
            self.EoTeh_6.setPlainText(resultREO[0][30])
        elif Number == 7:
            self.EoTeh_7.setPlainText(resultREO[0][30])
        elif Number == 8:
            self.EoTeh_8.setPlainText(resultREO[0][30])
        elif Number == 9:
            self.EoTeh_9.setPlainText(resultREO[0][30])
        elif Number == 10:
            self.EoTeh_10.setPlainText(resultREO[0][30])

    def ChangedREO(self):
        REO = self.REO.currentText()
        connection = sqlite3.connect('/Portal/UseFile/BaseFrom1C.db')
        reo = connection.cursor()
        reo.execute("SELECT * FROM ReestrREO WHERE Сокращенноенаименование = ?", (REO,))
        resultREO = reo.fetchall()
        connection.commit()
        connection.close()
        self.MailREO.setPlainText(resultREO[0][12])

    def ChangedExp(self):
        connection = sqlite3.connect('/Portal/UseFile/BaseFrom1C.db')
        exp = connection.cursor()
        FioExpert = self.FioExpert.currentText()
        self.FioExpert2.setCurrentText(FioExpert)
        exp.execute("SELECT * FROM Experts WHERE ФИОэкспертапоаккредитации = ?", (FioExpert,))
        OneExpert = exp.fetchall()
        EmailExpert = OneExpert[0][8]
        self.EmailExpert.setPlainText(EmailExpert)
        connection.commit()
        connection.close()

        connection = sqlite3.connect('/Portal/UseFile/BaseFrom1C.db')
        reo = connection.cursor()
        reo.execute("SELECT * FROM ReestrREO WHERE Адресэлектроннойпочты = ?", (OneExpert[0][7],))
        resultREO = reo.fetchall()
        connection.commit()
        connection.close()
        self.REO.setCurrentText(resultREO[0][9])
        self.MailREO.setPlainText(resultREO[0][12])

    def open_file(self):
        index = self.treeView.currentIndex()
        file_path = self.model.filePath(index)
        os.startfile(file_path)

    def onClicked(self, index):
        path = self.sender().model().filePath(index)
        os.startfile(path)

    def DataPricazaChanged(self):
        self.DataPricaza.setStyleSheet(
            'background-color: rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; color: rgb(0, 0, 0);')
        if self.DataPricaza.text() == "01.01.2000":
            self.DogovorCheck.hide()
        else:
            self.DogovorCheck.show()

    def DatePredlogChanged(self):
        self.DatePredlog.setStyleSheet(
            'background-color: rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; color: rgb(0, 0, 0);')

    def DataRospExpChanged(self):
        self.DataRospExp.setStyleSheet(
            'background-color: rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; color: rgb(0, 0, 0);')

    def DataYvedEGTChanged(self):
        datenow = self.DataYvedEG.text()
        datenow = str(datenow)
        self.DataYvedEGT.setText(datenow)

    def PereviborDateChanged(self):
        self.PereviborDate.setStyleSheet(
            'background-color: rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; color: rgb(0, 0, 0);')

    def DataPricaz1GRChanged(self):
        self.DataPricaz1GR.setStyleSheet(
            'background-color: rgb(255, 255, 255); font: 14pt "MS Shell Dlg 2"; color: rgb(0, 0, 0);')

    def DataYvedTChanged(self):
        datenow = self.DataYved.text()
        datenow = str(datenow)
        self.DataYvedT.setText(datenow)


    # _________________________________________
    def Create_Pricaz_1Group(self):  # Создание приказ 1 группа
        global FullNaimKogo
        global FullFioLogin
        global MsgError

        if self.TipGU.currentText() == "ПК" or self.TipGU.currentText() == "ПК+РОА" or self.TipGU.currentText() == "ПК+ИМОД" or self.TipGU.currentText() == "ПК+РОА+ИМОД":
            MsgError = "Уточните Тип ГУ"
            self.gotoAlarm()
            return
        if self.FioExpert.currentText() == "Эксперт не найден":
            MsgError = "Выберете эксперта"
            self.gotoAlarm()
            return
        if self.FioExpert.currentText() == "Метрология" and (
                self.Oblast_Button.text() == "Не выбрано" or self.Oblast_Button.text() == ""):
            MsgError = "Уточните конкретную область"
            self.gotoAlarm()
            return
        if self.Teh_1.currentText() == "":
            MsgError = "Тех. Эксперт не выбран"
            self.gotoAlarm()
            return
        if self.VMeropr.isChecked() == True and self.Vibor_TU_2.currentText() == "Тер. Управление":
            MsgError = "Выбрано поле (Выездное мероприятие по ОС), но не вбрано ТУ"
            self.gotoAlarm()
            return
        if FullFioLogin == "Приостановка":
            MsgError = "Нельзя использовать с приостановленной ГУ"
            self.gotoAlarm()
            return
        if self.FioExpert.currentText() == "" or self.FioExpert.currentText() is None:
            MsgError = "Не выбрана общая область"
            self.gotoAlarm()
            return

        TipGU_VO_DO = self.TipGU.currentText()
        ObshObl = self.ObshObl.currentText()
        Obl = self.Oblast_Button.text()
        acc = ""
        if ObshObl == "ИЛ":
            accObl = "в качестве испытательной лаборатории"
            acc = "\n(испытательная лаборатория)"
            if TipGU_VO_DO == "АК":
                doc2 = Document("SourceGitHub/Shablon/Group1/"
                                     "Programma_IL/ИЛ АКК.docx")
            elif TipGU_VO_DO == "РОА":
                doc2 = Document("SourceGitHub/Shablon/Group1/"
                                     "Programma_IL/ИЛ РОА.docx")
            else:
                doc2 = Document("SourceGitHub/Shablon/Group1/"
                                     "Programma_IL/ИЛ ПК.docx")
            shema = "со схемой аккредитации испытательных лабораторий (центров) в национальной системе аккредитации, утвержденной национальным органом по аккредитации 11.03.2022 (СМ № 03.1-9.0013)"
        elif ObshObl == "Метрология" and Obl == "ПСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по поверке средств измерений"
            if TipGU_VO_DO == "АК":
                doc2 = Document("SourceGitHub/Shablon/Group1/Programma_OEI/ПСИ (АК).docx")
            else:
                doc2 = Document("SourceGitHub/Shablon/Group1/Programma_OEI/ПСИ.docx")
            shema = "со схемой аккредитации юридических лиц и индивидуальных предпринимателей, выполняющих работы и (или) оказывающих услуги по поверке средств измерений, в национальной системе аккредитации 01.03.2022 (СМ № 03.1-9.0005)"
        elif ObshObl == "Метрология" and Obl == "КСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по калибровке средств измерений"
            if TipGU_VO_DO == "АК":
                doc2 = Document("SourceGitHub/Shablon/Group1/Programma_OEI/КСИ (АК).docx")
            else:
                doc2 = Document("SourceGitHub/Shablon/Group1/Programma_OEI/КСИ.docx")
            shema = "со схемой аккредитации юридических лиц и индивидуальных предпринимателей, выполняющих работы и (или) оказывающих услуги по калибровке средств измерений, в национальной системе аккредитации 01.03.2022 (СМ № 03.1-9.0007)"
        elif ObshObl == "Метрология" and Obl == "ИСИ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по испытанию средств измерений в целях утверждения типа"
            doc2 = Document("SourceGitHub/Shablon/Group1/Programma_OEI/ИСИ.docx")
            shema = "со схемой аккредитации юридических лиц и индивидуальных предпринимателей, выполняющих работы и (или) оказывающих услуги по проведению испытаний средств измерений в целях утверждения типа, в национальной системе аккредитации 01.03.2022 (СМ № 03.1-9.0015)"
        elif ObshObl == "Метрология" and Obl == "ИСО":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по испытанию стандартных образцов в целях утверждения типа"
            doc2 = Document(
                "SourceGitHub/Shablon/Group1/Programma_OEI/ИСО.docx")  # Ещё нет!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
            shema = "со схемой аккредитации юридических лиц и индивидуальных предпринимателей, выполняющих работы и (или) оказывающих услуги по проведению испытаний стандартных образцов в целях утверждения типа, в национальной системе аккредитации 01.03.2022 (СМ № 03.1-9.0016)"
        elif ObshObl == "Метрология" and Obl == "АМ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по аттестации методик (методов) измерений"
            if TipGU_VO_DO == "АК":
                doc2 = Document("SourceGitHub/Shablon/Group1/Programma_OEI/АМ (АК).docx")
            else:
                doc2 = Document("SourceGitHub/Shablon/Group1/Programma_OEI/АМ.docx")
            shema = "со схемой аккредитации юридических лиц и индивидуальных предпринимателей, выполняющих работы и (или) оказывающих услуги по аттестации методик (методов) измерений и (или) метрологической экспертизе, в национальной системе аккредитации 01.03.2022 (СМ № 03.1-9.0018)"
        elif ObshObl == "Метрология" and Obl == "МЭ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по метрологической экспертизе"
            doc2 = Document("SourceGitHub/Shablon/Group1/Programma_OEI/МЭ.docx")
            shema = "со схемой аккредитации юридических лиц и индивидуальных предпринимателей, выполняющих работы и (или) оказывающих услуги по аттестации методик (методов) измерений и (или) метрологической экспертизе, в национальной системе аккредитации 01.03.2022 (СМ № 03.1-9.0018)"
        elif ObshObl == "Метрология" and Obl == "АМ и МЭ":
            accObl = "в области обеспечения единства измерений для выполнения работ и (или) оказания услуг по аттестации методик (методов) измерений и метрологической экспертизе"
            if TipGU_VO_DO == "АК":
                doc2 = Document("SourceGitHub/Shablon/Group1/Programma_OEI/АМ МЭ (АК).docx")
            else:
                doc2 = Document("SourceGitHub/Shablon/Group1/Programma_OEI/АМ МЭ.docx")
            shema = "со схемой аккредитации юридических лиц и индивидуальных предпринимателей, выполняющих работы и (или) оказывающих услуги по аттестации методик (методов) измерений и (или) метрологической экспертизе, в национальной системе аккредитации 01.03.2022 (СМ № 03.1-9.0018)"
        elif ObshObl == "Метрология" and Obl == "17020":
            accObl = "в качестве органа инспекции"
            acc = "(орган инспекции)"
            doc2 = Document(
                "SourceGitHub/Shablon/Group1/Programma_OEI/ОИ 17020.docx")  # Ещё нет!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
            shema = "со схемой аккредитации органов инспекций в национальной системе аккредитации 11.03.2022 (СМ № 03.1-9.0011)"
        elif ObshObl == "Метрология" and Obl == "17021":
            accObl = "в качестве органа по сертификации систем менеджмента"
            doc2 = Document(
                "SourceGitHub/Shablon/Group1/Programma_OEI/ОС 17021 СМК 9000.docx")  # ХЗ как выбирать (надо спросить)!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
            doc2 = Document(
                "SourceGitHub/Shablon/Group1/Programma_OEI/ОС 17021 СМК 22000.docx")  # ХЗ как выбирать (надо спросить)!!!!!!!!!!!!!!!!!!!!!!!!!!
            shema = "со схемой аккредитации органов по сертификации систем менеджмента в национальной системе аккредитации 01.03.2022 (СМ № 03.1-9.0004)"
        elif ObshObl == "Метрология" and Obl == "17024":
            accObl = "в качестве органа по сертификации персонала"
            doc2 = Document(
                "SourceGitHub/Shablon/Group1/Programma_OEI/ОС 17024.docx")
        elif ObshObl == "Метрология" and Obl == "17043":
            accObl = "в качестве провайдера межлабораторных сличительных испытаний"  # Не уверен
            doc2 = Document(
                "SourceGitHub/Shablon/Group1/Programma_OEI/МСИ 17043.docx")
            shema = "со схемой аккредитации испытательных лабораторий (центров) в национальной системе аккредитации 11.03.2022 (СМ № 03.1-9.0013)"
        elif ObshObl == "Метрология" and Obl == "17065":
            accObl = "в качестве органа по сертификации продукции, услуг"
            doc2 = Document(
                "SourceGitHub/Shablon/Group1/Programma_OEI/ОС 17065.docx")
            shema = "со схемой аккредитации органов по сертификации продукции в национальной системе аккредитации 11.03.2022 (СМ № 03.1-9.0017)"
        elif ObshObl == "Метрология" and Obl == "15189":
            accObl = "в качестве испытательной лаборатории"
            doc2 = Document(
                "SourceGitHub/Shablon/Group1/Programma_OEI/15189.docx")  # Ещё нет!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
            shema = "со схемой аккредитации испытательных лабораторий (центров) в национальной системе аккредитации, утвержденной национальным органом по аккредитации 11.03.2022 (СМ № 03.1-9.0013)"  # Не уверен
        elif ObshObl == "Метрология" and Obl == "14065":
            accObl = "в качестве органа по валидации и верификации парниковых газов"
            doc2 = Document(
                "SourceGitHub/Shablon/Group1/Programma_OEI/14065.docx")  # Ещё нет!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
            shema = "со схемой аккредитации органов по валидации и верификации парниковых газов в национальной системе аккредитации 01.03.2022 (СМ № 03.1-9.0017)"
        elif ObshObl == "Метрология" and Obl == "Не выбрано":
            accObl = "в качестве испытательной лаборатории"  # Потом доработать исключение

        FullGU = self.NomerGU.toPlainText()
        if TipGU_VO_DO == "АК":
            master = Document(
                "SourceGitHub/Shablon/Group1/Pricaz/2. Приказ АК или РОА.docx")
        elif TipGU_VO_DO == "РОА":
            master = Document(
                "SourceGitHub/Shablon/Group1/Pricaz/2. Приказ АК или РОА.docx")
        else:
            master = Document("SourceGitHub/Shablon/Group1/Pricaz/1. Приказ.docx")
        composer = Composer(master)

        if self.VO.isChecked() == True:
            VO_DO = "ВО"
        elif self.DO.isChecked() == True:
            VO_DO = "ДО"
        elif self.DO_VO.isChecked() == True:
            VO_DO = "ДО+ВО"
        if VO_DO == "ДО" and TipGU_VO_DO != "АК" and TipGU_VO_DO != "РОА":
            path = "SourceGitHub/Shablon/Group1/Pricaz/1. Приказ.docx"
            doc = DocxTemplate(path)
            next_str = ""
        else:
            composer.append(doc2)
            composer.save(
                "SourceGitHub/ГУ/1 Группа/" + FullGU + "/2. Приказ " + Obl + " " + FullGU + " " + VO_DO +".docx")
            #############
            path = "SourceGitHub/ГУ/1 Группа/" + FullGU + "/2. Приказ " + Obl + " " + FullGU + " " + VO_DO +".docx"
            doc = DocxTemplate(path)
            next_str = "\f"

        # sd = DocxTemplate("SourceGitHub/Shablon/Group1/Programma_OEI/ПСИ.docx")
        self.FixFullNaimKogo()  # испраить ООО на падеж
        FullGU = self.NomerGU.toPlainText()
        FullGUSplit = FullGU.split()
        DataPricaza = FullGUSplit[2]
        DataPricazaSplit = DataPricaza.replace('.', ' ')
        DataPricazaSplit = DataPricazaSplit.split()
        DataPricaza0 = DataPricazaSplit[0]
        DataPricaza1 = DataPricazaSplit[1]
        DataPricaza2 = DataPricazaSplit[2]
        DataPricaza0 = DataPricaza0.replace('01', '1').replace('02', '2').replace('03', '3').replace('04', '4').replace(
            '05', '5').replace('06', '6').replace('07', '7').replace('08', '8').replace('09', '9')
        DataPricaza1 = DataPricaza1.replace('01', 'января').replace('02', 'февралья').replace('03', 'марта').replace(
            '04', 'апреля').replace('05', 'мая').replace('06', 'июня').replace('07', 'июля').replace('08',
                                                                                                     'августа').replace(
            '09', 'сентября').replace('10', 'октября').replace('11', 'ноября').replace('12', 'декабря')
        DataPricaza = DataPricaza0 + " " + DataPricaza1 + " " + DataPricaza2 + " г."
        NomerGU = FullGUSplit[0]

        NomerRAL = self.NomerRAL.toPlainText()
        ShortNaim = self.ShortName_2.toPlainText()
        EmailOrg = self.EmailRAL.toPlainText()
        FIOExp = self.FioExpert.currentText()
        EmailExp = self.EmailExpert.toPlainText()
        NameEO = self.REO.currentText()
        EmailEO = self.MailREO.toPlainText()

        if self.Teh_1.currentText() != "":
            ALLTeh = self.Teh_1.currentText()
            textAllTeh = "член экспертной группы (технический эксперт)"
        if self.Teh_2.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_2.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"
        if self.Teh_3.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_3.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"
        if self.Teh_4.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_4.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"
        if self.Teh_5.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_5.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"
        if self.Teh_6.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_6.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"
        if self.Teh_7.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_7.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"
        if self.Teh_8.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_8.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"
        if self.Teh_9.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_9.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"
        if self.Teh_10.currentText() != "":
            ALLTeh = ALLTeh + ", " + self.Teh_10.currentText()
            textAllTeh = "члены экспертной группы (технические эксперты)"

        Kontrol = ""
        Podpisant = self.FioExpert_3.currentText()
        if Podpisant == "Золотаревский С.Ю.":
            Podpisant = "С.Ю. Золотаревский"
            Doljnost = "Начальник Управления \nаккредитации в сфере \nдобровольного подтверждения \nсоответствия, метрологии \nи иных сферах деятельности"
            Kontrol = "Контроль за исполнением настоящего приказа оставляю за собой."
        elif Podpisant == "Шкабура В.В.":
            Podpisant = "В.В. Шкабура"
            Doljnost = "Начальник отдела аккредитации \nв отдельных сферах \nУправления аккредитации в сфере \nдобровольного подтверждения \nсоответствия, метрологии и иных \nсферах деятельности"
        elif Podpisant == "Бухарова А.В.":
            Podpisant = "А.В. Бухарова"
            Doljnost = "Начальник отдела аккредитации \nиспытательных лабораторий \nУправления аккредитации \nв сфере добровольного подтверждения \nсоответствия, метрологии \nи иных сферах деятельности"
        elif Podpisant == "Гоголев Д.В.":
            Podpisant = "Д.В. Гоголев"
            Doljnost = "Заместитель руководителя"
            Kontrol = "Контроль за исполнением настоящего приказа возложить на начальника Управления аккредитации в сфере добровольного подтверждения соответствия, метрологии и иных сферах деятельности С.Ю. Золотаревского."
        elif Podpisant == "Белогуров С.И.":
            Podpisant = "С.И. Белогуров"
            Doljnost = "Заместитель начальника отдела \nаккредитации в отдельных сферах \nУправления аккредитации в сфере \nдобровольного подтверждения \nсоответствия, метрологии и иных \nсферах деятельности"
        elif Podpisant == "Макаров А.Н.":
            Podpisant = "А.Н. Макаров"
            Doljnost = "И.о. начальника Управления аккредитации \nв сфере подтверждения соответствия \nпродукции машиностроения, \nэлектротехнической продукции, \nстроительных материалов \nи пожарной безопасности"
            Kontrol = "Контроль за исполнением настоящего приказа оставляю за собой."
        elif Podpisant == "Гоголев Д.В.":
            Podpisant = "Д.В. Гоголев"
            Doljnost = "Заместитель руководителя"
            Kontrol = "Контроль за исполнением настоящего приказа возложить на начальника Управления аккредитации в сфере добровольного подтверждения соответствия, метрологии и иных сферах деятельности С.Ю. Золотаревского."
        elif Podpisant == "Хазиева А.А.":
            Podpisant = "А.А. Хазиева"
            Doljnost = "Начальник отдела аккредитации в сфере \nподтверждения соответствия продукции машиностроения \nУправления аккредитации в сфере подтверждения \nсоответствия продукции машиностроения, \nэлектротехнической продукции, строительных \nматериалов и пожарной безопасности"

        ShortFIO = FullFioLogin.split()
        ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."


        VKS1 = ""
        VKS2 = ""
        VMerop1 = ""
        VMerop2 = ""
        VMerop3 = ""
        WTF1 = ""
        WTF2 = ""
        WTF3 = ""
        WhoIsWho = ""
        WhoIsWhoo = ""
        WhoIsWhooo = ""
        TipGU2 = ""
        VO_DO2 = ""

        if self.VKS.isChecked() == True:
            if TipGU_VO_DO == "АК":
                VKS1 = "\aВыездную оценку соответствия Заявителя критериям аккредитации провести " \
                   "с использованием средств дистанционного взаимодействия, в том числе " \
                   "посредством аудио- и видео-конференц-связи."
            else:
                VKS1 = "\aВыездную оценку соответствия Аккредитованного лица критериям аккредитации провести " \
                       "с использованием средств дистанционного взаимодействия, в том числе " \
                       "посредством аудио- и видео-конференц-связи."
            if TipGU_VO_DO == "АК":
                VKS2 = "Подходы к организации проведения выездной оценки заявителя критериям аккредитации в форме удаленной " \
                       "оценки, а также методы, применяемые экспертной группой, должностным лицом Росаккредитации в рамках " \
                       "проведения такой оценки определены в СМ № 03.1-1.0007 «Руководство по проведению удаленной оценки».\a"
            else:
                VKS2 = "Выездная оценка проводится в форме удаленной оценки в соответствии \nс документом " \
                       "системы менеджмента качества Росаккредитации «Руководство \nпо проведению " \
                       "удаленной оценки», утвержденным руководителем Росаккредитации 30 декабря " \
                       "2020 г. (далее - Руководство).\aДля выполнения каждого мероприятия программы " \
                       "выездной оценки в форме удаленной оценки применяются методы (техники), " \
                       "позволяющие обеспечить результативность удаленной оценки, определенные в разделе 7 Руководства.\a"
        if TipGU_VO_DO == "ПК1":
            TipGU = "подтверждения компетентности"
            if self.VO.isChecked() == True:
                VO_DO = "выездной оценки соответствия аккредитованного лица критериям аккредитации"  # ВО
                VO_DO2 = "выездной экспертизы соответствия Аккредитованного лица критериям аккредитации"
                Punkt3 = "Утвердить программу выездной оценки соответствия аккредитованного лица критериям аккредитации согласно приложению к настоящему приказу."
                NomerP1 = "4."
                NomerP2 = "5."
                NomerP3 = "\a3."
            elif self.DO.isChecked() == True:
                VO_DO = "документарной оценки соответствия аккредитованного лица критериям аккредитации"  # ДО
                VO_DO2 = "экспертизы представленных Аккредитованным лицом документов и сведений"  # ДО
                Punkt3 = ""
                NomerP1 = "3."
                NomerP2 = "4."
                NomerP3 = ""
            ProgTip1 = "аккредитованного"
            ProgTip2 = "(уникальный номер записи об аккредитации в реестре аккредитованных лиц "
            ProgTip3 = ", "
        elif TipGU_VO_DO == "ПК2":
            TipGU = "подтверждения компетентности"
            VO_DO = "выездной оценки соответствия аккредитованного лица критериям аккредитации"  # ВО
            VO_DO2 = "выездной экспертизы соответствия Аккредитованного лица критериям аккредитации"
            Punkt3 = "Утвердить программу выездной оценки соответствия аккредитованного лица критериям аккредитации согласно приложению к настоящему приказу."
            NomerP1 = "4."
            NomerP2 = "5."
            NomerP3 = "\a3."
            ProgTip1 = "аккредитованного"
            ProgTip2 = "(уникальный номер записи об аккредитации в реестре аккредитованных лиц "
            ProgTip3 = ", "
        elif TipGU_VO_DO == "ПК5":
            TipGU = "подтверждения компетентности"
            VO_DO = "документарной оценки соответствия аккредитованного лица критериям аккредитации и выездной оценки соответствия аккредитованного лица критериям аккредитации"  # ВО + ДО
            VO_DO2 = "экспертизы представленных Аккредитованным лицом документов и сведений и выездной экспертизы соответствия Аккредитованного лица критериям аккредитации"
            Punkt3 = "Утвердить программу выездной оценки соответствия аккредитованного лица критериям аккредитации согласно приложению к настоящему приказу."
            NomerP1 = "4."
            NomerP2 = "5."
            NomerP3 = "\a3."
            ProgTip1 = "аккредитованного"
            ProgTip2 = "(уникальный номер записи об аккредитации в реестре аккредитованных лиц "
            ProgTip3 = ", "
        elif TipGU_VO_DO == "АК":
            TipGU = "аккредитации"
            TipGU2 = "об аккредитации"
            WhoIsWho = "заявителя"
            WhoIsWhoo = "заявителем"
            WhoIsWhooo = "заявитель"
            NomerRAL = ""
            VO_DO = "экспертизы представленных документов и сведений и выездной экспертизы соответствия критериям аккредитации"
            Punkt3 = "Утвердить программу выездной оценки соответствия заявителя критериям аккредитации согласно приложению к настоящему приказу."
            NomerP1 = "4."
            NomerP2 = "5."
            NomerP3 = "\a3."
            ProgTip1 = ""
            ProgTip2 = ""
            ProgTip3 = "( "
        elif TipGU_VO_DO == "РОА":
            TipGU = "расширения области аккредитации"
            TipGU2 = "о расширении области аккредитации"
            WTF1 = "(аккредитованного "
            WTF2 = ", уникальный номер записи об аккредитации в реестре аккредитованных лиц "
            WTF3 = ")"
            WhoIsWho = "аккредитованного лица"
            WhoIsWhoo = "аккредитованным лицом"
            WhoIsWhooo = "аккредитованное лицо"
            VO_DO = "экспертизы представленных документов и сведений и выездной экспертизы соответствия критериям аккредитации"
            Punkt3 = "Утвердить программу выездной оценки соответствия аккредитованного лица критериям аккредитации согласно приложению к настоящему приказу."
            NomerP1 = "4."
            NomerP2 = "5."
            NomerP3 = "\a3."
            ProgTip1 = ""
            ProgTip2 = ""
            ProgTip3 = "( "
        elif TipGU_VO_DO == "ПК1+РОА" or TipGU_VO_DO == "ПК2+РОА" or TipGU_VO_DO == "ПК5+РОА":
            TipGU = "подтверждения компетентности и расширения области аккредитации"
            VO_DO = "экспертизы представленных документов и сведений и выездной экспертизы соответствия критериям аккредитации"
            VO_DO2 = "экспертизы представленных Аккредитованным лицом документов и сведений и выездной экспертизы соответствия Аккредитованного лица критериям аккредитации"
            Punkt3 = "Утвердить программу выездной оценки соответствия аккредитованного лица критериям аккредитации согласно приложению к настоящему приказу."
            NomerP1 = "4."
            NomerP2 = "5."
            NomerP3 = "\a3."
            ProgTip1 = "аккредитованного"
            ProgTip2 = "(уникальный номер записи об аккредитации в реестре аккредитованных лиц "
            ProgTip3 = ", "
        elif TipGU_VO_DO == "ПК1+ИМОД" or TipGU_VO_DO == "ПК2+ИМОД": # ВО
            TipGU = "подтверждения компетентности и изменения места (мест) осуществления деятельности"
            VO_DO = "выездной оценки соответствия аккредитованного лица критериям аккредитации"
            VO_DO2 = "выездной экспертизы соответствия Аккредитованного лица критериям аккредитации"
            Punkt3 = "Утвердить программу выездной оценки соответствия аккредитованного лица критериям аккредитации согласно приложению к настоящему приказу."
            NomerP1 = "4."
            NomerP2 = "5."
            NomerP3 = "\a3."
            ProgTip1 = "аккредитованного"
            ProgTip2 = "(уникальный номер записи об аккредитации в реестре аккредитованных лиц "
            ProgTip3 = ", "
        elif TipGU_VO_DO == "ПК5+ИМОД":
            TipGU = "подтверждения компетентности и изменения места (мест) осуществления деятельности"
            VO_DO = "экспертизы представленных документов и сведений и выездной экспертизы соответствия критериям аккредитации"
            VO_DO2 = "экспертизы представленных Аккредитованным лицом документов и сведений и выездной экспертизы соответствия Аккредитованного лица критериям аккредитации"
            Punkt3 = "Утвердить программу выездной оценки соответствия аккредитованного лица критериям аккредитации согласно приложению к настоящему приказу."
            NomerP1 = "4."
            NomerP2 = "5."
            NomerP3 = "\a3."
            ProgTip1 = "аккредитованного"
            ProgTip2 = "(уникальный номер записи об аккредитации в реестре аккредитованных лиц "
            ProgTip3 = ", "
        elif TipGU_VO_DO == "РОА+ИМОД":
            TipGU = "расширения области аккредитации и изменения места (мест) осуществления деятельности"
            VO_DO = "экспертизы представленных документов и сведений и выездной экспертизы соответствия критериям аккредитации"
            VO_DO2 = "экспертизы представленных Аккредитованным лицом документов и сведений и выездной экспертизы соответствия Аккредитованного лица критериям аккредитации"
            Punkt3 = "Утвердить программу выездной оценки соответствия аккредитованного лица критериям аккредитации согласно приложению к настоящему приказу."
            NomerP1 = "4."
            NomerP2 = "5."
            NomerP3 = "\a3."
            ProgTip1 = "аккредитованного"
            ProgTip2 = "(уникальный номер записи об аккредитации в реестре аккредитованных лиц "
            ProgTip3 = ", "
        elif TipGU_VO_DO == "ПК1+РОА+ИМОД" or TipGU_VO_DO == "ПК2+РОА+ИМОД" or TipGU_VO_DO == "ПК5+РОА+ИМОД":
            TipGU = "подтверждения компетентности, расширения области аккредитации и изменения места (мест) осуществления деятельности"
            VO_DO = "экспертизы представленных документов и сведений и выездной экспертизы соответствия критериям аккредитации"
            VO_DO2 = "экспертизы представленных Аккредитованным лицом документов и сведений и выездной экспертизы соответствия Аккредитованного лица критериям аккредитации"
            Punkt3 = "Утвердить программу выездной оценки соответствия аккредитованного лица критериям аккредитации согласно приложению к настоящему приказу."
            NomerP1 = "4."
            NomerP2 = "5."
            NomerP3 = "\a3."
            ProgTip1 = "аккредитованного"
            ProgTip2 = "(уникальный номер записи об аккредитации в реестре аккредитованных лиц "
            ProgTip3 = ", "

        VMeroprBool = self.VMeropr.isChecked()
        NomerP4 = ""
        if VMeroprBool == True:
            VMerop3 = "определить должностное лицо для участия в выездных мероприятиях по оценке соответствия " \
                      "аккредитованного лица критериям аккредитации согласно программе, указанной в пункте 3 настоящего приказа."
            NomerP1 = "4."
            NomerP2 = "6."
            NomerP4 = "\a5."
            if self.Vibor_TU_2.currentText() == "ЦФО":
                VMerop1 = "Управлению аккредитации в сфере добровольного подтверждения соответствия, метрологии и иных сферах деятельности"
            if self.Vibor_TU_2.currentText() == "ПФО":
                VMerop1 = "Управлению Федеральной службы по аккредитации по "
                VMerop2 = "Приволжскому федеральному округу"
            if self.Vibor_TU_2.currentText() == "СФО":
                VMerop1 = "Управлению Федеральной службы по аккредитации по "
                VMerop2 = "Сибирскому федеральному округу"
            if self.Vibor_TU_2.currentText() == "СЗФО":
                VMerop1 = "Управлению Федеральной службы по аккредитации по "
                VMerop2 = "Северо-Западному федеральному округу"
            if self.Vibor_TU_2.currentText() == "ДФО":
                VMerop1 = "Управлению Федеральной службы по аккредитации по "
                VMerop2 = "Дальневосточному федеральному округу"
            if self.Vibor_TU_2.currentText() == "УФО":
                VMerop1 = "Управлению Федеральной службы по аккредитации по "
                VMerop2 = "Уральскому федеральному округу"
            if self.Vibor_TU_2.currentText() == "ЮСФО":
                VMerop1 = "Управлению Федеральной службы по аккредитации по "
                VMerop2 = "Южному и Северо-Кавказскому федеральным округам"

        context = {"short_name_org": ShortNaim.replace(' "', ' «').replace('"', '»'), "email_org": EmailOrg,
                   "fio_exp": FIOExp, "email_exp": EmailExp, "name_eo": NameEO.replace(' "', ' «').replace('"', '»'),
                   "email_eo": EmailEO, "nomer_gu": NomerGU, "wtf1": WTF1, "wtf2": WTF2, "wtf3": WTF3, "hish": WhoIsWho,
                   "hisho": WhoIsWhoo, "shema": shema, "prog_tip1": ProgTip1, "prog_tip2": ProgTip2,
                   "prog_tip3": ProgTip3,
                   "data_gu": DataPricaza, "full_name_org": FullNaimKogo.replace(' "', ' «').replace('"', '»'),
                   "acc_obl": accObl, "nomer_ral": NomerRAL, "nomer_gu2": FullGU, "all_teh": ALLTeh,
                   "text_teh": textAllTeh, "tip_gu2": TipGU2, "kontrol": Kontrol,
                   "dolj_podpisant": Doljnost, "fio_podpisant": Podpisant, "short_isp": ShortFIO, "vo_do": VO_DO,
                   "tip_gu": TipGU, "acc": acc, "vks": VKS1, "vmerop1": VMerop1, "vmerop2": VMerop2, "vmerop3": VMerop3,
                   "punkt3": Punkt3, "vks2": VKS2, "next_str": next_str, "nomerp1": NomerP1, "nomerp2": NomerP2,
                   "nomerp3": NomerP3, "nomerp4": NomerP4, "vo_do2":VO_DO2, "hishasha":WhoIsWhooo}
        doc.render(context)
        if self.VO.isChecked() == True:
            VO_DO = "ВО"
        elif self.DO.isChecked() == True:
            VO_DO = "ДО"
        elif self.DO_VO.isChecked() == True:
            VO_DO = "ДО+ВО"

        try:
            doc.save(
                "SourceGitHub/ГУ/1 Группа/" + FullGU + "/2. Приказ " + Obl + " " + FullGU + " " + VO_DO +".docx")
            os.startfile(
                "SourceGitHub/ГУ/1 Группа/" + FullGU + "/2. Приказ " + Obl + " " + FullGU + " " + VO_DO +".docx")
        except:
            print("Что-то пошло не так")

    def gotoWhat(self):
        create = CreateWhat()
        create.exec_()

    def gotoAlarm(self):
        create = CreateAlarm()
        create.exec_()

    def OpenGUFolder(self):  # Проверить и создать папку
        FullGU = self.NomerGU.toPlainText()
        GUFolder = "SourceGitHub/ГУ/1 Группа/" + FullGU
        if os.path.exists(GUFolder):
            os.startfile(GUFolder)
        else:
            os.startfile(GUFolder)

    def HidePerevibor(self):
        if self.Perevibor.currentText() == "Нет":
            self.PereviborDate.hide()
            self.label_PereviborDate.hide()

        else:
            self.PereviborDate.show()
            self.label_PereviborDate.show()


class PatchNote(QDialog):
    def __init__(self):
        global ZY
        super(PatchNote, self).__init__()
        loadUi("SourceGitHub/UI/PatchNote.ui", self)
        self.Closeer.clicked.connect(self.ClosePatchNote)
        self.PatchW.move(2420, 300)
        self.PatchW.clicked = False

    def ClosePatchNote(self):
        create = MainWindow()
        widget.addWidget(create)
        widget.setCurrentIndex(widget.currentIndex() + 1)
        widget.setGeometry(-1920, 0, 5760, 1080)
        widget.show()

    def mousePressEvent(self, event):
        self.old_pos = event.screenPos()

    def mouseMoveEvent(self, event):
        if self.PatchW.clicked:
            dx = self.old_pos.x() - event.screenPos().x()
            dy = self.old_pos.y() - event.screenPos().y()
            self.move(self.pos().x() - dx, self.pos().y() - dy)
        self.old_pos = event.screenPos()
        self.PatchW.clicked = True



class CreateUpdate(QDialog):
    def __init__(self):
        global FileUpdate
        super(CreateUpdate, self).__init__()
        loadUi("SourceGitHub/UI/Updater.ui", self)

        self.FIndFile.clicked.connect(self.getfile)
        self.Update.clicked.connect(self.DoUpdate)
        self.Update.setEnabled(False)

        self.movie = QMovie("SourceGitHub/Pictures/peach-cat.gif")
        self.ErrorDrag1.setMovie(self.movie)
        self.movie.start()
        self.movie = QMovie("SourceGitHub/Pictures/mochi-cat-chibi-cat.gif")
        self.ErrorDrag2.setMovie(self.movie)
        self.movie.start()



    def DoUpdate(self):
        global FileUpdate
        #FileUpdate = FileUpdate.split()
        df = pd.DataFrame(pd.read_excel('SourceGitHub/DB Excel/xlsx/Список.xlsx'))

        connection = sqlite3.connect('SourceGitHub/DB/BaseFrom1C.db')
        cursor = connection.cursor()
        cursor.execute('DELETE FROM TableFrom1C WHERE Датарегистрации != ""')
        connection.commit()
        connection.close()

        connection = sqlite3.connect('SourceGitHub/DB/BaseFrom1C.db')
        cursor = connection.cursor()
        for i in range(len(df)):
            cursor.execute(
                'INSERT INTO TableFrom1C (ТипГУ, Регистрационныйномер, Датарегистрации, РегистрационныйномерАЛвРАЛ, Наименованиезаявителя, Полноенаименованиезаявителя, ИННЗаявителя'
                ', Ответственныйисполнитель, Управлениеответственногоисполнителя, ЭтапГУ, СтатусГУ, Информацияотекущейактивнойзадаче, Датаподачизаявления, ЗаявлениенаотзывГУ, Датаназначенияответственногоисполнителя, ДатапроверкизаявленияОИ'
                ', СЗовозвратевОДК, Приказоботказе, Датавыбораэксперта, УведомлениеоботбореЭА, ФИОЭкспертапоаккредитации, НаименованиеЭО, ВремямеждурегистрациейивыборомЭАврабочихднях, Согласиеилиотказэкспертапоаккредитации, ПредложениепосоставуЭГ'
                ', ДатаприказаосоставеЭГ, Экспертноезаключениепорезультатамдокументарнойоценки, ПисьмоовозвратеЭЗ, ПриказнаприостановкупорезультатамДО, Письмообустранениинарушенийвнутреннийдокументобустранениинарушений, Приказнавыезднуюоценку, Приказоботказеваккредитации, Актвыезднойэкспертизы, Письмоовозвратеактавыезднойэкспертизы'
                ', СЗсотчетомоВО, ПриказнаприостановкупорезультатамВО, Письмоорассмотренииотчетазаявителяэкспертомпоаккредитации, СЗосогласительнойкомиссии, Протоколсогласительнойкомиссии, Запросдополненийкактувыезднойэкспертизыотэкспертнойгруппыи,'
                ' Запроснадополнениякакту, Заключениеобоценкеустранениязаявителемвыявленныхнесоответствий, ПриказозавершенииГУ, ПодготовилприказозавершенииГУ, ПриказоботказевпредоставленииГУ)'
                ' VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
                (df.iat[i, 0], df.iat[i, 1], df.iat[i, 2], df.iat[i, 3], df.iat[i, 4], df.iat[i, 5], df.iat[i, 6],
                 df.iat[i, 7]
                 , df.iat[i, 8], df.iat[i, 9], df.iat[i, 10], df.iat[i, 11], df.iat[i, 12], df.iat[i, 13],
                 df.iat[i, 14], df.iat[i, 15]
                 , df.iat[i, 16], df.iat[i, 17], df.iat[i, 18], df.iat[i, 19], df.iat[i, 20], df.iat[i, 21],
                 df.iat[i, 22], df.iat[i, 23], df.iat[i, 24]
                 , df.iat[i, 25], df.iat[i, 26], df.iat[i, 27], df.iat[i, 28], df.iat[i, 29], df.iat[i, 30],
                 df.iat[i, 31], df.iat[i, 32], df.iat[i, 33]
                 , df.iat[i, 34], df.iat[i, 35], df.iat[i, 36], df.iat[i, 37], df.iat[i, 38], df.iat[i, 39],
                 df.iat[i, 40], df.iat[i, 41], df.iat[i, 42]
                 , df.iat[i, 43], df.iat[i, 44]))

        connection.commit()
        connection.close()
        self.Main = MainWindow()
        self.Main.UpdateGU1()
        self.close()



    def getfile(self):
        global FileUpdate
        FileUpdate = QFileDialog.getOpenFileName(self, 'Выбрать',
                                            'c:\\', "Файл выгрузки из 1С (*.xlsx)")
        self.MsgFile.setText(FileUpdate[0])
        self.Update.setEnabled(True)





class CreateAleshaChoise(QDialog):
    def __init__(self):
        global dfIspMail
        super(CreateAleshaChoise, self).__init__()
        #self.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.CustomizeWindowHint)
        cal = Russia()
        cal.holidays(2022)
        pd.set_option("display.max_rows", None, "display.max_columns", None)
        loadUi("SourceGitHub/UI/AleshaChoise.ui", self)

        conn = sqlite3.connect('SourceGitHub/DB/shop_data.db')
        cursor = conn.cursor()
        cursor.execute('SELECT myname, firstname, lastname, email FROM Login')
        IspMail = cursor.fetchall()
        cursor.execute('SELECT COUNT(*) FROM Login')
        Longe = cursor.fetchone()
        dfIspMail = pd.DataFrame({'ФИО Исполнителя': [], 'Почта': []})
        dfIspMail.rename(columns={0: 'ФИО Исполнителя', 1: 'Почта'}, inplace=True);
        for i in range(Longe[0]):
            FIO = IspMail[i][1] + " " + IspMail[i][0] + " " + IspMail[i][2]
            dfIspMail.loc[len(dfIspMail.index)] = [FIO, IspMail[i][3]]

        self.movie = QMovie("SourceGitHub/Pictures/parrot.gif")
        self.AleshaGif1.setMovie(self.movie)
        self.movie.start()
        self.movie = QMovie("SourceGitHub/Pictures/car.gif")
        self.AleshaGif2.setMovie(self.movie)
        self.movie.start()

        self.PredlALL.clicked.connect(self.CreateAleshaPredlALL)
        self.PredlOEI.clicked.connect(self.CreateAleshaPredlOEI)
        self.PredlIL.clicked.connect(self.CreateAleshaPredlIL)

        self.DogovorlALL.clicked.connect(self.CreateAleshaDogovorlALL)
        self.DogovorlOEI.clicked.connect(self.CreateAleshaDogovorlOEI)
        self.DogovorlIL.clicked.connect(self.CreateAleshaDogovorlIL)

        self.EZALL.clicked.connect(self.CreateAleshaEZALL)
        self.EZOEI.clicked.connect(self.CreateAleshaEZOEI)
        self.EZIL.clicked.connect(self.CreateAleshaEZIL)

        self.ActALL.clicked.connect(self.CreateAleshaActALL)
        self.ActOEI.clicked.connect(self.CreateAleshaActOEI)
        self.ActIL.clicked.connect(self.CreateAleshaActIL)


    def CreateAleshaPredlALL(self):
        global AleshaOtdeal1
        global AleshaOtdeal2
        AleshaOtdeal1 = "Отдел аккредитации в отдельных сферах"
        AleshaOtdeal2 = "Отдел аккредитации испытательных лабораторий"
        create = CreateAleshaPredl()
        create.exec_()
    def CreateAleshaPredlOEI(self):
        global AleshaOtdeal1
        global AleshaOtdeal2
        AleshaOtdeal1 = "Отдел аккредитации в отдельных сферах"
        AleshaOtdeal2 = "Отдел аккредитации в отдельных сферах"
        create = CreateAleshaPredl()
        create.exec_()
    def CreateAleshaPredlIL(self):
        global AleshaOtdeal1
        global AleshaOtdeal2
        AleshaOtdeal1 = "Отдел аккредитации испытательных лабораторий"
        AleshaOtdeal2 = "Отдел аккредитации испытательных лабораторий"
        create = CreateAleshaPredl()
        create.exec_()

    def CreateAleshaDogovorlALL(self):
        global AleshaOtdeal1
        global AleshaOtdeal2
        AleshaOtdeal1 = "Отдел аккредитации в отдельных сферах"
        AleshaOtdeal2 = "Отдел аккредитации испытательных лабораторий"
        create = CreateAleshaDogovor()
        create.exec_()
    def CreateAleshaDogovorlOEI(self):
        global AleshaOtdeal1
        global AleshaOtdeal2
        AleshaOtdeal1 = "Отдел аккредитации в отдельных сферах"
        AleshaOtdeal2 = "Отдел аккредитации в отдельных сферах"
        create = CreateAleshaDogovor()
        create.exec_()
    def CreateAleshaDogovorlIL(self):
        global AleshaOtdeal1
        global AleshaOtdeal2
        AleshaOtdeal1 = "Отдел аккредитации испытательных лабораторий"
        AleshaOtdeal2 = "Отдел аккредитации испытательных лабораторий"
        create = CreateAleshaDogovor()
        create.exec_()

    def CreateAleshaEZALL(self):
        global AleshaOtdeal1
        global AleshaOtdeal2
        AleshaOtdeal1 = "Отдел аккредитации в отдельных сферах"
        AleshaOtdeal2 = "Отдел аккредитации испытательных лабораторий"
        create = CreateAleshaEZ()
        create.exec_()
    def CreateAleshaEZOEI(self):
        global AleshaOtdeal1
        global AleshaOtdeal2
        AleshaOtdeal1 = "Отдел аккредитации в отдельных сферах"
        AleshaOtdeal2 = "Отдел аккредитации в отдельных сферах"
        create = CreateAleshaEZ()
        create.exec_()
    def CreateAleshaEZIL(self):
        global AleshaOtdeal1
        global AleshaOtdeal2
        AleshaOtdeal1 = "Отдел аккредитации испытательных лабораторий"
        AleshaOtdeal2 = "Отдел аккредитации испытательных лабораторий"
        create = CreateAleshaEZ()
        create.exec_()

    def CreateAleshaActALL(self):
        global AleshaOtdeal1
        global AleshaOtdeal2
        AleshaOtdeal1 = "Отдел аккредитации в отдельных сферах"
        AleshaOtdeal2 = "Отдел аккредитации испытательных лабораторий"
        create = CreateAleshaAct()
        create.exec_()
    def CreateAleshaActOEI(self):
        global AleshaOtdeal1
        global AleshaOtdeal2
        AleshaOtdeal1 = "Отдел аккредитации в отдельных сферах"
        AleshaOtdeal2 = "Отдел аккредитации в отдельных сферах"
        create = CreateAleshaAct()
        create.exec_()
    def CreateAleshaActIL(self):
        global AleshaOtdeal1
        global AleshaOtdeal2
        AleshaOtdeal1 = "Отдел аккредитации испытательных лабораторий"
        AleshaOtdeal2 = "Отдел аккредитации испытательных лабораторий"
        create = CreateAleshaAct()
        create.exec_()

    def AleshaToExcel(self):
        global AleshaDF
        global AleshaDF18
        global AleshaDF3
        global AleshaWho
        try:
            df = AleshaDF
        except:
            pass
        datenow = QDate.currentDate().toPyDate()
        try:
            if AleshaWho == "Предложения":
                VigFile = 'SourceGitHub/Выгрузки/Просрочки по Предложениям'
            elif AleshaWho == "Договор":
                VigFile = 'SourceGitHub/Выгрузки/Просрочки по Договору'
            if os.path.exists(VigFile):
                What = "Папка уже создана"
            else:
                os.mkdir(VigFile)
            VigFolder = VigFile
            if AleshaWho == "Предложения":
                VigFile = VigFile + "/" + "Просрочка по Предложениям на " + str(datenow) + ".xlsx"
                writer = pd.ExcelWriter(VigFile)
            elif AleshaWho == "Договор":
                VigFile18 = VigFile + "/" + "Просрочка по Договору 18 дней на " + str(datenow) + ".xlsx"
                VigFile3 = VigFile + "/" + "Просрочка по Договору 3 дня на " + str(datenow) + ".xlsx"
                writer18 = pd.ExcelWriter(VigFile18)
                writer3 = pd.ExcelWriter(VigFile3)
            if AleshaWho == "Предложения":
                df.to_excel(writer, sheet_name='All_GU', index=False)
                for column in df:
                    column_width = max(df[column].astype(str).map(len).max(), len(column))
                    col_idx = df.columns.get_loc(column)
                    writer.sheets['All_GU'].set_column(col_idx, col_idx, column_width)
                writer.save()
                os.startfile(VigFolder)
            elif AleshaWho == "Договор":
                AleshaDF18.to_excel(writer18, sheet_name='All_GU', index=False)
                for column in AleshaDF18:
                    column_width = max(AleshaDF18[column].astype(str).map(len).max(), len(column))
                    col_idx = AleshaDF18.columns.get_loc(column)
                    writer18.sheets['All_GU'].set_column(col_idx, col_idx, column_width)
                writer18.save()
                os.startfile(VigFolder)
                AleshaDF3.to_excel(writer3, sheet_name='All_GU', index=False)
                for column in AleshaDF3:
                    column_width = max(AleshaDF3[column].astype(str).map(len).max(), len(column))
                    col_idx = AleshaDF3.columns.get_loc(column)
                    writer3.sheets['All_GU'].set_column(col_idx, col_idx, column_width)
                writer3.save()
                os.startfile(VigFolder)
        except:
            if AleshaWho == "Предложения":
                VigFile = 'SourceGitHub/Выгрузки/Просрочки по Предложениям'
            elif AleshaWho == "Договор":
                VigFile = 'SourceGitHub/Выгрузки/Просрочки по Договору'
            if os.path.exists(VigFile):
                What = "Папка уже создана"
            else:
                os.mkdir(VigFile)
            VigFolder = VigFile
            if AleshaWho == "Предложения":
                VigFile = VigFile + "/" + "Просрочка по Предложениям на " + str(datenow) + ".xlsx"
                writer = pd.ExcelWriter(VigFile)
            elif AleshaWho == "Договор":
                VigFile18 = VigFile + "/" + "Просрочка по Договору 18 дней на " + str(datenow) + " копия " + str(random.uniform(0, 20)) + ".xlsx"
                VigFile3 = VigFile + "/" + "Просрочка по Договору 3 дня на " + str(datenow) + " копия " + str(random.uniform(0, 20)) + ".xlsx"
                writer18 = pd.ExcelWriter(VigFile18)
                writer3 = pd.ExcelWriter(VigFile3)
            if AleshaWho == "Предложения":
                df.to_excel(writer, sheet_name='All_GU', index=False)
                for column in df:
                    column_width = max(df[column].astype(str).map(len).max(), len(column))
                    col_idx = df.columns.get_loc(column)
                    writer.sheets['All_GU'].set_column(col_idx, col_idx, column_width)
                writer.save()
                os.startfile(VigFolder)
            elif AleshaWho == "Договор":
                AleshaDF18.to_excel(writer18, sheet_name='All_GU', index=False)
                for column in AleshaDF18:
                    column_width = max(AleshaDF18[column].astype(str).map(len).max(), len(column))
                    col_idx = AleshaDF18.columns.get_loc(column)
                    writer18.sheets['All_GU'].set_column(col_idx, col_idx, column_width)
                writer18.save()
                os.startfile(VigFolder)
                AleshaDF3.to_excel(writer3, sheet_name='All_GU', index=False)
                for column in AleshaDF3:
                    column_width = max(AleshaDF3[column].astype(str).map(len).max(), len(column))
                    col_idx = AleshaDF3.columns.get_loc(column)
                    writer3.sheets['All_GU'].set_column(col_idx, col_idx, column_width)
                writer3.save()
                os.startfile(VigFolder)


class CreateAleshaPredl(QDialog):
    def __init__(self):
        global AleshaOtdeal1
        global AleshaOtdeal2
        global AleshaDF
        global AleshaWho
        AleshaWho = "Предложения"
        super(CreateAleshaPredl, self).__init__()
        cal = Russia()
        cal.holidays(2022)
        pd.set_option("display.max_rows", None, "display.max_columns", None)
        loadUi("SourceGitHub/UI/AleshaPredl.ui", self)
        self.AleshaTable.setColumnCount(8)
        self.AleshaTable.setColumnWidth(0, 70)
        self.AleshaTable.setColumnWidth(1, 70)
        self.AleshaTable.setColumnWidth(2, 120)
        self.AleshaTable.setColumnWidth(3, 120)
        self.AleshaTable.setColumnWidth(4, 150)
        self.AleshaTable.setColumnWidth(5, 150)
        self.AleshaTable.setColumnWidth(6, 90)
        self.AleshaTable.setColumnWidth(7, 70)
        self.SendMail.clicked.connect(self.AleshaMailPredl)
        conn = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
        cursor = conn.cursor()
        cursor.execute(
            'SELECT Регистрационныйномер, Датарегистрации, Ответственныйисполнитель, '
            'ФИОЭкспертапоаккредитации, ЭлПочтаЭксперта, НаименованиеЭО, Датавыбораэксперта, ТочнаяОбласть FROM GU1Group WHERE (Датавыбораэксперта IS NOT NULL AND '
            'Датавыбораэксперта != "") AND '
            '(ДатаприказаосоставеЭГ IS NULL or ДатаприказаосоставеЭГ = "Не зарегистрирован" or ДатаприказаосоставеЭГ = "") AND '
            '(ОтделУП = "Отдел аккредитации в отдельных сферах" or ОтделУП = "Отдел аккредитации испытательных лабораторий") AND '
            'СтатусГУ = "В работе" AND (ПредложениепосоставуЭГ IS NULL or ПредложениепосоставуЭГ = "") AND (ОтделУП = ? or ОтделУП = ?)', (AleshaOtdeal1, AleshaOtdeal2 ))  # Находит кол-во ГУ на выборе эксперта
        rows = cursor.fetchall()
        cursor.execute(
            'SELECT COUNT(*) FROM GU1Group WHERE (Датавыбораэксперта IS NOT NULL AND '
            'Датавыбораэксперта != "") AND '
            '(ДатаприказаосоставеЭГ IS NULL or ДатаприказаосоставеЭГ = "Не зарегистрирован" or ДатаприказаосоставеЭГ = "") AND '
            '(ОтделУП = "Отдел аккредитации в отдельных сферах" or ОтделУП = "Отдел аккредитации испытательных лабораторий") AND '
            'СтатусГУ = "В работе" AND (ПредложениепосоставуЭГ IS NULL or ПредложениепосоставуЭГ = "") AND (ОтделУП = ? or ОтделУП = ?)', (AleshaOtdeal1, AleshaOtdeal2 ))
        Longe = cursor.fetchone()
        df = pd.DataFrame({'Номер ГУ': [], 'Дата регистрации': [], 'Исполнитель': [], 'Эксперт': [], 'Э.п. Эксперт': [], 'ЭО': [], 'Дата в. эксп.': [], 'Прошло': [], 'Точная Область': []})
        df.rename(columns={0: 'Номер ГУ', 1: 'Дата регистрации', 2: 'Исполнитель', 3: 'Эксперт', 4: 'Э.п. Эксперт', 5: 'ЭО', 6: 'Дата в. эксп.', 7: 'Прошло', 8: 'Точная Область'}, inplace=True);
        for i in range(Longe[0]):
            Vib_Exp = rows[i][6]
            Date_Vib_Exp_Date = datetime.strptime(Vib_Exp, '%d.%m.%Y')
            Now = datetime.today()
            Days = cal.get_working_days_delta(Date_Vib_Exp_Date, Now)
            if int(Days) >= 2:
                df.loc[len(df.index)] = [rows[i][0], rows[i][1], rows[i][2], rows[i][3], rows[i][4], rows[i][5], rows[i][6], Days, rows[i][7]]
        AleshaDF = df
        print(df)
        self.AleshaTable.setRowCount(len(df))
        for i in range(len(df)):
            self.AleshaTable.setItem(i, 0, QtWidgets.QTableWidgetItem(df.iat[i, 0]))
            self.AleshaTable.setItem(i, 1, QtWidgets.QTableWidgetItem(df.iat[i, 1]))
            ShortFIO = df.iat[i, 2]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTable.setItem(i, 2, QtWidgets.QTableWidgetItem(ShortFIO))
            ShortFIO = df.iat[i, 3]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTable.setItem(i, 3, QtWidgets.QTableWidgetItem(ShortFIO))
            self.AleshaTable.setItem(i, 4, QtWidgets.QTableWidgetItem(df.iat[i, 4]))
            self.AleshaTable.setItem(i, 5, QtWidgets.QTableWidgetItem(df.iat[i, 5]))
            self.AleshaTable.setItem(i, 6, QtWidgets.QTableWidgetItem(df.iat[i, 6]))
            days = df.iat[i, 7]
            self.AleshaTable.setItem(i, 7, QtWidgets.QTableWidgetItem(str(days)))
        dfEO = df.groupby(['ЭО']).size().reset_index(name='count')
        self.AleshaTableEO.setColumnCount(2)
        self.AleshaTableEO.setColumnWidth(0, 160)
        self.AleshaTableEO.setColumnWidth(1, 65)
        self.AleshaTableEO.setRowCount(len(dfEO))
        print(dfEO)
        for i in range(len(dfEO)):
            self.AleshaTableEO.setItem(i, 0, QtWidgets.QTableWidgetItem(dfEO.iat[i, 0]))
            self.AleshaTableEO.setItem(i, 1, QtWidgets.QTableWidgetItem(str(dfEO.iat[i, 1])))

        dfExp = df.groupby(['Эксперт']).size().reset_index(name='count')
        self.AleshaTableExp.setColumnCount(2)
        self.AleshaTableExp.setColumnWidth(0, 160)
        self.AleshaTableExp.setColumnWidth(1, 65)
        self.AleshaTableExp.setRowCount(len(dfExp))
        print(dfEO)
        for i in range(len(dfExp)):
            ShortFIO = dfExp.iat[i, 0]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTableExp.setItem(i, 0, QtWidgets.QTableWidgetItem(ShortFIO))
            self.AleshaTableExp.setItem(i, 1, QtWidgets.QTableWidgetItem(str(dfExp.iat[i, 1])))

        dfIsp = df.groupby(['Исполнитель']).size().reset_index(name='count')
        self.AleshaTableIsp.setColumnCount(2)
        self.AleshaTableIsp.setColumnWidth(0, 160)
        self.AleshaTableIsp.setColumnWidth(1, 65)
        self.AleshaTableIsp.setRowCount(len(dfIsp))
        print(dfIsp)
        for i in range(len(dfIsp)):
            ShortFIO = dfIsp.iat[i, 0]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTableIsp.setItem(i, 0, QtWidgets.QTableWidgetItem(ShortFIO))
            self.AleshaTableIsp.setItem(i, 1, QtWidgets.QTableWidgetItem(str(dfIsp.iat[i, 1])))

        self.ButAlesha = CreateAleshaChoise()
        self.ButtonExcel.clicked.connect(self.ButAlesha.AleshaToExcel)

    def AleshaMailPredl(self):
        global globalNomerGU
        global AleshaDF
        global dfIspMail
        df = AleshaDF
        LongDF = len((df.index))
        LongdfIspMail = len((dfIspMail.index))
        for i in range(LongDF):
            FullGU = "№ " + df.iat[i, 0] + " от " + df.iat[i, 1]
            FioExp = df.iat[i, 4]
            EmailExp = df.iat[i, 4]
            ShortFIOIsp = df.iat[i, 2]
            for ii in range(LongdfIspMail):
                if ShortFIOIsp == dfIspMail.iat[ii, 0]:
                    MailIsp = dfIspMail.iat[ii, 1]


            datenow = QDate.currentDate().toPyDate()
            datenow = str(datenow)
            datenow = datenow.replace("-", " ")
            datenow = datenow.split()
            datenowStr = datenow[2] + "." + datenow[1] + "." + datenow[0]
            outlook = win32.Dispatch('outlook.application')
            mail = outlook.CreateItem(0)
            #mail.To = EmailExp
            mail.To = "IvanovAA@fsa.gov.ru"
            mail.Subject = 'Направить согласие по ' + FullGU
            Tip = df.iat[i, 8]
            if Tip != "ПСИ":
                mail.HTMLBody = '<p>Добрый день, Вам назначена государственная услуга ' + FullGU + '.</p>' \
                                '<p>Обращаем Ваше внимание, что в соответствии с пунктом 11 Порядка формирования экспертной ' \
                                'группы, утверждённого приказом Минэкономразвития России от 29.10.2021 № 657, установлено, ' \
                                'что эксперт по аккредитации, назначенный руководителем экспертной группы, <strong>в ' \
                                'течение 2 рабочих </strong>дней <strong>со дня</strong> его отбора осуществляет формирование ' \
                                'и направление в Федеральную службу по аккредитации предложения о привлечении технических экспертов.</p>' \
                                '<p>Сообщаем о необходимости загрузить во ФГИС &nbsp;указанные документы по назначенной ' \
                                'вам государственной услуге ' + FullGU + ' в срок не позднее ' + datenowStr +';.</p>' \
                                '<p>В случае отсутствия технической возможности подгрузки в личном кабинете Вам необходимо направить ' \
                                'соответствующее обращение через интерактивный помощник ФГИС (с описанием проблемы), &nbsp;а также &nbsp;' \
                                'продублировать документы куратору по государственной услуге.</p>' \
                                '<p><strong>Куратор по данной Государственной услуге</strong><br /><strong>' + ShortFIOIsp + '<br />' \
                                '<a href="' + MailIsp + '">' + MailIsp + '</a></strong></p> ' \
                                '<p><sub>Уведомления генерируются автоматически. Пожалуйста, не отвечайте на это сообщение.</sub><o:p></o:p></p>' \
                                '<p><strong><img src="' + ':\Управление Аккредитации\Программа\PortalFSA\Pictures\logo2.jpg" width="240" height="100" alt="" /></strong></p>'
            else:
                mail.HTMLBody = '<p>Добрый день, Вам назначена государственная услуга ' + FullGU + '.</p>' \
                                '<p>Обращаем Ваше внимание, что в соответствии с пунктом 11 Порядка формирования экспертной ' \
                                'группы, утверждённого приказом Минэкономразвития России от 29.10.2021 № 657, установлено, ' \
                                'что эксперт по аккредитации, назначенный руководителем экспертной группы, <strong>в ' \
                                'течение 2 рабочих </strong>дней <strong>со дня</strong> его отбора осуществляет формирование ' \
                                'и направление в Федеральную службу по аккредитации предложения о привлечении технических экспертов.</p>' \
                                '<p>Сообщаем о необходимости загрузить во ФГИС &nbsp;указанные документы по назначенной ' \
                                'вам государственной услуге ' + FullGU + ' в срок не позднее ' + datenowStr + ';.</p>' \
                                '<p>В случае отсутствия технической возможности подгрузки в личном кабинете Вам необходимо направить ' \
                                'соответствующее обращение через интерактивный помощник ФГИС (с описанием проблемы), &nbsp;а также &nbsp;' \
                                'продублировать документы куратору по государственной услуге.</p>' \
                                '<p><strong>Куратор по данной Государственной услуге</strong><br /><strong>' + ShortFIOIsp + '<br />' \
                                '<a href="' + MailIsp + '">' + MailIsp + '</a></strong></p> ' \
                                '<p>Также прошу учесть что в соответствии с Постановлением Правительства РФ № 2050 ' \
                                'от 26.11.2021 в состав экспертной группы включается<strong> ТЕХНИЧЕСКИЙ ЭКСПЕРТ,' \
                                '&nbsp;</strong><strong><u>который является работником государственного научного ' \
                                'метрологического института (ГНМИ).</u></strong></p>' \
                                '<p><sub>Уведомления генерируются автоматически. Пожалуйста, не отвечайте на это сообщение.</sub><o:p></o:p></p>' \
                                '<p><strong><img src="' + ':\Управление Аккредитации\Программа\PortalFSA\Pictures\logo2.jpg" width="240" height="100" alt="" /></strong></p>'

            #mail.Display(True)
            mail.Send()

class CreateAleshaDogovor(QDialog):
    def __init__(self):
        global AleshaOtdeal1
        global AleshaOtdeal2
        global AleshaDF18
        global AleshaDF3
        global AleshaWho
        AleshaWho = "Договор"
        #from CreateAleshaChoise import AleshaToExcel
        super(CreateAleshaDogovor, self).__init__()
        cal = Russia()
        cal.holidays(2022)
        pd.set_option("display.max_rows", None, "display.max_columns", None)
        loadUi("SourceGitHub/UI/AleshaDogovor.ui", self)

        self.AleshaTable.setColumnCount(9)
        self.AleshaTable.setColumnWidth(0, 50)
        self.AleshaTable.setColumnWidth(1, 70)
        self.AleshaTable.setColumnWidth(2, 70)
        self.AleshaTable.setColumnWidth(3, 200)
        self.AleshaTable.setColumnWidth(4, 120)
        self.AleshaTable.setColumnWidth(5, 120)
        self.AleshaTable.setColumnWidth(6, 150)
        self.AleshaTable.setColumnWidth(7, 150)
        self.AleshaTable.setColumnWidth(8, 65)
        self.AleshaTable_2.setColumnCount(9)
        self.AleshaTable_2.setColumnWidth(0, 50)
        self.AleshaTable_2.setColumnWidth(1, 70)
        self.AleshaTable_2.setColumnWidth(2, 70)
        self.AleshaTable_2.setColumnWidth(3, 200)
        self.AleshaTable_2.setColumnWidth(4, 120)
        self.AleshaTable_2.setColumnWidth(5, 120)
        self.AleshaTable_2.setColumnWidth(6, 150)
        self.AleshaTable_2.setColumnWidth(7, 150)
        self.AleshaTable_2.setColumnWidth(8, 65)



        conn = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
        cursor = conn.cursor()
        cursor.execute(
            'SELECT ТипГУ, Регистрационныйномер, Датарегистрации, Наименованиезаявителя, ЭлектроннаяПочтаЗаявителя, Ответственныйисполнитель, '
            'ФИОЭкспертапоаккредитации, ЭлПочтаЭксперта, НаименованиеЭО, ЭлПочтаЭоЭксперта, ДатаприказаосоставеЭГ FROM GU1Group WHERE '
            '(ДатаприказаосоставеЭГ IS NOT NULL AND ДатаприказаосоставеЭГ != "Не зарегистрирован" AND ДатаприказаосоставеЭГ != "") AND '
            'СтатусГУ = "В работе" AND (Договор1Г IS NULL or Договор1Г = "")AND (ОтделУП = ? or ОтделУП = ?)', (AleshaOtdeal1, AleshaOtdeal2 ))  # Находит кол-во ГУ на выборе эксперта
        rows = cursor.fetchall()
        cursor.execute(
            'SELECT COUNT(*) FROM GU1Group WHERE (ДатаприказаосоставеЭГ IS NOT NULL AND '
            'ДатаприказаосоставеЭГ != "Не зарегистрирован" AND ДатаприказаосоставеЭГ != "") AND '
            'СтатусГУ = "В работе" AND (Договор1Г IS NULL or Договор1Г = "")AND (ОтделУП = ? or ОтделУП = ?)', (AleshaOtdeal1, AleshaOtdeal2 ))
        Longe = cursor.fetchone()
        df = pd.DataFrame({'Тип ГУ': [], 'Номер ГУ': [], 'Дата регистрации': [], 'Заявитель': [], 'Эл.п. Заявителя': [],
                           'Исполнитель': [], 'Эксперт': [], 'Эл.п. Эксперта': [], 'ЭО': [], 'Эл.п. ЭО': [], 'Дата Приказа': [], 'Прошло': []})
        df.rename(columns={0: 'Тип ГУ', 1: 'Номер ГУ', 2: 'Дата регистрации', 3: 'Заявитель', 4: 'Эл.п. Заявителя' ,
                           5: 'Исполнитель', 6: 'Эксперт', 7: 'Эл.п. Эксперта', 8: 'ЭО', 9: 'Эл.п. ЭО', 10: 'Дата Приказа', 11: 'Прошло'}, inplace=True);
        for i in range(Longe[0]):
            Vib_Exp = rows[i][10]
            Vib_Exp = Vib_Exp.split()
            Len_Vib_Exp = len(Vib_Exp)
            if Len_Vib_Exp > 1:
                Vib_Exp = Vib_Exp[2].replace("г.", "")
            else:
                Vib_Exp = Vib_Exp[0]
            Date_Vib_Exp_Date = datetime.strptime(Vib_Exp, '%d.%m.%Y')
            Now = datetime.today()
            Days = cal.get_working_days_delta(Date_Vib_Exp_Date, Now)
            if int(Days) >= 18:
                df.loc[len(df.index)] = [rows[i][0], rows[i][1], rows[i][2], rows[i][3], rows[i][4], rows[i][5], rows[i][6], rows[i][7], rows[i][8], rows[i][9], rows[i][10], Days]
        print(df)
        AleshaDF18 = df
        self.AleshaTable.setRowCount(len(df))
        for i in range(len(df)):
            self.AleshaTable.setItem(i, 0, QtWidgets.QTableWidgetItem(df.iat[i, 0]))
            self.AleshaTable.setItem(i, 1, QtWidgets.QTableWidgetItem(df.iat[i, 1]))
            self.AleshaTable.setItem(i, 2, QtWidgets.QTableWidgetItem(df.iat[i, 2]))
            self.AleshaTable.setItem(i, 3, QtWidgets.QTableWidgetItem(df.iat[i, 3]))
            ShortFIO = df.iat[i, 5]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTable.setItem(i, 4, QtWidgets.QTableWidgetItem(ShortFIO))
            ShortFIO = df.iat[i, 6]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTable.setItem(i, 5, QtWidgets.QTableWidgetItem(ShortFIO))
            self.AleshaTable.setItem(i, 6, QtWidgets.QTableWidgetItem(df.iat[i, 8]))
            self.AleshaTable.setItem(i, 7, QtWidgets.QTableWidgetItem(df.iat[i,10]))
            days = df.iat[i, 11]
            self.AleshaTable.setItem(i, 8, QtWidgets.QTableWidgetItem(str(days)))

        dfEO = df.groupby(['ЭО']).size().reset_index(name='count')
        self.AleshaTableEO.setColumnCount(2)
        self.AleshaTableEO.setColumnWidth(0, 160)
        self.AleshaTableEO.setColumnWidth(1, 65)
        self.AleshaTableEO.setRowCount(len(dfEO))
        print(dfEO)
        for i in range(len(dfEO)):
            self.AleshaTableEO.setItem(i, 0, QtWidgets.QTableWidgetItem(dfEO.iat[i, 0]))
            self.AleshaTableEO.setItem(i, 1, QtWidgets.QTableWidgetItem(str(dfEO.iat[i, 1])))

        dfExp = df.groupby(['Эксперт']).size().reset_index(name='count')
        self.AleshaTableExp.setColumnCount(2)
        self.AleshaTableExp.setColumnWidth(0, 160)
        self.AleshaTableExp.setColumnWidth(1, 65)
        self.AleshaTableExp.setRowCount(len(dfExp))
        print(dfEO)
        for i in range(len(dfExp)):
            ShortFIO = dfExp.iat[i, 0]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTableExp.setItem(i, 0, QtWidgets.QTableWidgetItem(ShortFIO))
            self.AleshaTableExp.setItem(i, 1, QtWidgets.QTableWidgetItem(str(dfExp.iat[i, 1])))

        dfIsp = df.groupby(['Исполнитель']).size().reset_index(name='count')
        self.AleshaTableIsp.setColumnCount(2)
        self.AleshaTableIsp.setColumnWidth(0, 160)
        self.AleshaTableIsp.setColumnWidth(1, 65)
        self.AleshaTableIsp.setRowCount(len(dfIsp))
        print(dfIsp)
        for i in range(len(dfIsp)):
            ShortFIO = dfIsp.iat[i, 0]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTableIsp.setItem(i, 0, QtWidgets.QTableWidgetItem(ShortFIO))
            self.AleshaTableIsp.setItem(i, 1, QtWidgets.QTableWidgetItem(str(dfIsp.iat[i, 1])))

#-----------------------------------------------------------------------------
        for i in range(Longe[0]):
            Vib_Exp = rows[i][10]
            Vib_Exp = Vib_Exp.split()
            Len_Vib_Exp = len(Vib_Exp)
            if Len_Vib_Exp > 1:
                Vib_Exp = Vib_Exp[2].replace("г.", "")
            else:
                Vib_Exp = Vib_Exp[0]
            Date_Vib_Exp_Date = datetime.strptime(Vib_Exp, '%d.%m.%Y')
            Now = datetime.today()
            Days = cal.get_working_days_delta(Date_Vib_Exp_Date, Now)
            if int(Days) > 3:
                df.loc[len(df.index)] = [rows[i][0], rows[i][1], rows[i][2], rows[i][3], rows[i][4], rows[i][5], rows[i][6], rows[i][7], rows[i][8], rows[i][9], rows[i][10], Days]
        print(df)
        AleshaDF3 = df
        self.AleshaTable_2.setRowCount(len(df))
        for i in range(len(df)):
            self.AleshaTable_2.setItem(i, 0, QtWidgets.QTableWidgetItem(df.iat[i, 0]))
            self.AleshaTable_2.setItem(i, 1, QtWidgets.QTableWidgetItem(df.iat[i, 1]))
            self.AleshaTable_2.setItem(i, 2, QtWidgets.QTableWidgetItem(df.iat[i, 2]))
            self.AleshaTable_2.setItem(i, 3, QtWidgets.QTableWidgetItem(df.iat[i, 3]))
            ShortFIO = df.iat[i, 5]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTable_2.setItem(i, 4, QtWidgets.QTableWidgetItem(ShortFIO))
            ShortFIO = df.iat[i, 6]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTable_2.setItem(i, 5, QtWidgets.QTableWidgetItem(ShortFIO))
            self.AleshaTable_2.setItem(i, 6, QtWidgets.QTableWidgetItem(df.iat[i, 8]))
            self.AleshaTable_2.setItem(i, 7, QtWidgets.QTableWidgetItem(df.iat[i,10]))
            days = df.iat[i, 11]
            self.AleshaTable_2.setItem(i, 8, QtWidgets.QTableWidgetItem(str(days)))

        dfEO = df.groupby(['ЭО']).size().reset_index(name='count')
        self.AleshaTableEO_2.setColumnCount(2)
        self.AleshaTableEO_2.setColumnWidth(0, 160)
        self.AleshaTableEO_2.setColumnWidth(1, 65)
        self.AleshaTableEO_2.setRowCount(len(dfEO))
        print(dfEO)
        for i in range(len(dfEO)):
            self.AleshaTableEO_2.setItem(i, 0, QtWidgets.QTableWidgetItem(dfEO.iat[i, 0]))
            self.AleshaTableEO_2.setItem(i, 1, QtWidgets.QTableWidgetItem(str(dfEO.iat[i, 1])))

        dfExp = df.groupby(['Эксперт']).size().reset_index(name='count')
        self.AleshaTableExp_2.setColumnCount(2)
        self.AleshaTableExp_2.setColumnWidth(0, 160)
        self.AleshaTableExp_2.setColumnWidth(1, 65)
        self.AleshaTableExp_2.setRowCount(len(dfExp))
        print(dfEO)
        for i in range(len(dfExp)):
            ShortFIO = dfExp.iat[i, 0]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTableExp_2.setItem(i, 0, QtWidgets.QTableWidgetItem(ShortFIO))
            self.AleshaTableExp_2.setItem(i, 1, QtWidgets.QTableWidgetItem(str(dfExp.iat[i, 1])))

        dfIsp = df.groupby(['Исполнитель']).size().reset_index(name='count')
        self.AleshaTableIsp_2.setColumnCount(2)
        self.AleshaTableIsp_2.setColumnWidth(0, 160)
        self.AleshaTableIsp_2.setColumnWidth(1, 65)
        self.AleshaTableIsp_2.setRowCount(len(dfIsp))
        print(dfIsp)
        for i in range(len(dfIsp)):
            ShortFIO = dfIsp.iat[i, 0]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTableIsp_2.setItem(i, 0, QtWidgets.QTableWidgetItem(ShortFIO))
            self.AleshaTableIsp_2.setItem(i, 1, QtWidgets.QTableWidgetItem(str(dfIsp.iat[i, 1])))

        self.ButAlesha = CreateAleshaChoise()
        self.ButtonExcel.clicked.connect(self.ButAlesha.AleshaToExcel)
        self.ButtonExcel_2.clicked.connect(self.ButAlesha.AleshaToExcel)
        self.ButtonMailEO3.clicked.connect(self.AleshaMailDogovorEO3)
        self.ButtonMailZav18.clicked.connect(self.AleshaMailDogovorZav18)

    def AleshaMailDogovorEO3(self):
        global globalNomerGU
        global AleshaDF3
        df = AleshaDF3
        dfEO = df.groupby(['ЭО']).size().reset_index(name='count')


        LongdfEO = len((dfEO.index))
        for i in range(LongdfEO):
            QWE = dfEO.iat[i, 0]
            dfEOFilter = df[df.ЭО == QWE]
            dfEOFilter = dfEOFilter.drop(columns=['Тип ГУ', 'Эл.п. Заявителя', 'Исполнитель', 'Эл.п. Эксперта', 'ЭО', 'Эл.п. ЭО'], axis=1)
            dfEOFilter = dfEOFilter.reset_index()
            dfEOFilter = dfEOFilter.drop(columns=['index'], axis=1)
            print(dfEOFilter)
            outlook = win32.Dispatch('outlook.application')
            mail = outlook.CreateItem(0)
            #mail.To = EmailExp
            mail.To = "IvanovAA@fsa.gov.ru"
            mail.Subject = 'Список ГУ по которым нет информации о заключении Договора '
            #mail.HTMLBody = '<html><body>' + dfEOFilter.to_html() + '</body></html>'
            mail.HTMLBody = '<p>Добрый день.</p><p>Пунктом 31 Правил осуществления аккредитации в национальной системе аккредитации, а ' \
                            'также пунктом 30 Правил проведения процедуры подтверждения компетентности аккредитованного лица, ' \
                            'утвержденных постановлением Правительства Российской Федерации от 26.11.2021 № 2050, установлено, ' \
                            'что срок заключения договора об оказании услуг между заявителем (аккредитованным лицом) &nbsp;' \
                            'и экспертной организацией, составляет <strong>3 рабочих дня </strong>со дня направления заявителю ' \
                            '(аккредитованному лицу), экспертной организации и эксперту по аккредитации приказа национального ' \
                            'органа по аккредитации, предусмотренного пунктом 30 Правил осуществления аккредитации в национальной ' \
                            'системе аккредитации и пунктом 29 Правил проведения процедуры подтверждения компетентности ' \
                            'аккредитованного лица.</p><p>Напоминаем вам о необходимости проинформировать национальный орган ' \
                            'по аккредитации о заключении (незаключении) договора с заявителем (аккредитованным лицом) не ' \
                            'позднее одного рабочего дня со дня истечения вышеуказанного срока по следующим государственным услугам</p>' \
                            '<p><html><body>' + dfEOFilter.to_html() + '</body></html></p><p>В случае заключения указанного договора, &nbsp;в представляемой информации просим указывать <strong>' \
                            'дату заключения договора</strong>.</p><p><sub>Уведомления генерируются автоматически. Пожалуйста, ' \
                            'не отвечайте на это сообщение.</sub><o:p></o:p><sub></sub></p>' \
                            '<p><strong><img src="' + ':\Управление Аккредитации\Программа\PortalFSA\Pictures\logo2.jpg" width="240" height="100" alt="" /></strong></p>'

            mail.Display(True)
            mail.Send()


    def AleshaMailDogovorZav18(self):
        global globalNomerGU
        global AleshaDF18
        df = AleshaDF18
        LongDF = len((df.index))
        for i in range(LongDF):
            FullGU = "№ " + df.iat[i, 0] + " от " + df.iat[i, 1]
            EmailExp = df.iat[i, 4]
            datenow = QDate.currentDate().toPyDate()
            datenow = str(datenow)
            datenow = datenow.replace("-", " ")
            datenow = datenow.split()
            datenowStr = datenow[2] + "." + datenow[1] + "." + datenow[0]
            outlook = win32.Dispatch('outlook.application')
            mail = outlook.CreateItem(0)
            #mail.To = EmailExp
            mail.To = "IvanovAA@fsa.gov.ru"
            mail.Subject = 'О заключении договора по ' + FullGU
            mail.HTMLBody = '<p>Добрый день.</p><p>Пунктом 31 Правил осуществления аккредитации в национальной системе аккредитации, ' \
                            'а также пунктом 30 Правил проведения процедуры подтверждения компетентности аккредитованного лица, ' \
                            'утвержденных постановлением Правительства Российской Федерации от 26.11.2021 № 2050, установлено, ' \
                            'что срок заключения договора об оказании услуг между заявителем (аккредитованным лицом) и экспертной организацией, ' \
                            'составляет 3 рабочих дня со дня направления заявителю (аккредитованному лицу), ' \
                            'экспертной организации и эксперту по аккредитации приказа национального органа по аккредитации, ' \
                            'предусмотренного пунктом 30 Правил осуществления аккредитации в национальной системе аккредитации и ' \
                            'пунктом 29 Правил проведения процедуры подтверждения компетентности аккредитованного лица.</p>' \
                            '<p>Обращаем внимание, что в случае непредставления в национальный орган по аккредитации договора в ' \
                            'течение 21 рабочего дня со дня направления заявителю (аккредитованному лицу) вышеуказанного приказа ' \
                            'национальный орган по аккредитации принимает решение <strong>об отказе в аккредитации</strong> ' \
                            'при предоставлении государственной услуги по аккредитации или <strong>прекращении предоставления ' \
                            'государственной услуги по подтверждению компетентности аккредитованного лица и приостановлении ' \
                            'действия аккредитации.</strong></p><p><strong>На текущую дату в Росаккредитации отсутствует ' \
                            'информация о заключении вышеуказанного договора по государственной услуге' +  FullGU + '.</strong></p>' \
                            '<p>В случае, если договор заключен, просим представить информацию (с приложением копии договора) ' \
                            'в адрес Exproblem@fsa.gov.ru.</p><p>В случае наличия объективных причин, не позволяющих в указанный ' \
                            'срок заключить договор между заявителем (аккредитованным лицом) и экспертной организацией просим ' \
                            'проинформировать национальный орган, направив соответствующее официальное обращение &nbsp;в ' \
                            'Росаккредитацию за подписью уполномоченного представителя заявителя (аккредитованного лица), ' \
                            'до истечения вышеуказанного срока.</p><p>&nbsp;<sub>Уведомления генерируются автоматически. ' \
                            'Пожалуйста, не отвечайте на это сообщение.</sub><o:p>' \
                            '</o:p><sub></sub></p><p><strong><img src="' + ':\Управление Аккредитации\Программа\PortalFSA\Pictures\logo2.jpg" width="240" height="100" alt="" /></strong></p>'

            mail.Display(True)
            mail.Send()

class CreateAleshaEZ(QDialog):
    def __init__(self):
        global AleshaOtdeal1
        global AleshaOtdeal2
        super(CreateAleshaEZ, self).__init__()
        cal = Russia()
        cal.holidays(2022)
        pd.set_option("display.max_rows", None, "display.max_columns", None)
        loadUi("SourceGitHub/UI/AleshaEZ.ui", self)
        self.AleshaTable.setColumnCount(8)
        self.AleshaTable.setColumnWidth(0, 50)
        self.AleshaTable.setColumnWidth(1, 70)
        self.AleshaTable.setColumnWidth(2, 70)
        self.AleshaTable.setColumnWidth(3, 120)
        self.AleshaTable.setColumnWidth(4, 120)
        self.AleshaTable.setColumnWidth(5, 150)
        self.AleshaTable.setColumnWidth(6, 90)
        self.AleshaTable.setColumnWidth(7, 65)


        conn = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
        cursor = conn.cursor()
        cursor.execute(
            'SELECT ТипГУ, Регистрационныйномер, Датарегистрации, Ответственныйисполнитель, '
            'ФИОЭкспертапоаккредитации, НаименованиеЭО, Договор1Г FROM GU1Group WHERE '
            '(ДатаприказаосоставеЭГ IS NOT NULL AND ДатаприказаосоставеЭГ != "Не зарегистрирован" AND ДатаприказаосоставеЭГ != "") AND '
            '(ОтделУП = "Отдел аккредитации в отдельных сферах" or ОтделУП = "Отдел аккредитации испытательных лабораторий") AND '
            'СтатусГУ = "В работе" AND (ДатаЭЗ IS NULL or ДатаЭЗ = "") AND'
            '(Договор1Г IS NOT NULL AND Договор1Г != "") AND (ТипГУ = "АК" OR ТипГУ = "РОА") AND (ОтделУП = ? or ОтделУП = ?)', (AleshaOtdeal1, AleshaOtdeal2 ))  # Находит кол-во ГУ на выборе эксперта
        rows = cursor.fetchall()
        cursor.execute(
            'SELECT COUNT(*) FROM GU1Group WHERE (ДатаприказаосоставеЭГ IS NOT NULL AND ДатаприказаосоставеЭГ != "Не зарегистрирован" AND ДатаприказаосоставеЭГ != "") AND '
            '(ОтделУП = "Отдел аккредитации в отдельных сферах" or ОтделУП = "Отдел аккредитации испытательных лабораторий") AND '
            'СтатусГУ = "В работе" AND (ДатаЭЗ IS NULL or ДатаЭЗ = "") AND'
            '(Договор1Г IS NOT NULL AND Договор1Г != "") AND (ТипГУ = "АК" OR ТипГУ = "РОА") AND (ОтделУП = ? or ОтделУП = ?)', (AleshaOtdeal1, AleshaOtdeal2 ))
        Longe = cursor.fetchone()
        df = pd.DataFrame({'Тип ГУ': [], 'Номер ГУ': [], 'Дата регистрации': [], 'Исполнитель': [], 'Эксперт': [], 'ЭО': [], 'Договор': [], 'Прошло': []})
        df.rename(columns={0: 'Тип ГУ', 1: 'Номер ГУ', 2: 'Дата регистрации', 3: 'Исполнитель', 4: 'Эксперт', 5: 'ЭО', 6: 'Договор', 7: 'Прошло'}, inplace=True);
        for i in range(Longe[0]):
            Vib_Exp = rows[i][6]
            Vib_Exp = Vib_Exp.split()
            Len_Vib_Exp = len(Vib_Exp)
            if Len_Vib_Exp > 1:
                Vib_Exp = Vib_Exp[2].replace("г.", "")
            else:
                Vib_Exp = Vib_Exp[0]
            Date_Vib_Exp_Date = datetime.strptime(Vib_Exp, '%d.%m.%Y')
            Now = datetime.today()
            Days = cal.get_working_days_delta(Date_Vib_Exp_Date, Now)
            if int(Days) >= 10:
                df.loc[len(df.index)] = [rows[i][0], rows[i][1], rows[i][2], rows[i][3], rows[i][4], rows[i][5], Vib_Exp, Days]
        print(df)
        self.AleshaTable.setRowCount(len(df))
        for i in range(len(df)):
            self.AleshaTable.setItem(i, 0, QtWidgets.QTableWidgetItem(df.iat[i, 0]))
            self.AleshaTable.setItem(i, 1, QtWidgets.QTableWidgetItem(df.iat[i, 1]))
            self.AleshaTable.setItem(i, 2, QtWidgets.QTableWidgetItem(df.iat[i, 2]))
            ShortFIO = df.iat[i, 3]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTable.setItem(i, 3, QtWidgets.QTableWidgetItem(ShortFIO))
            ShortFIO = df.iat[i, 4]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTable.setItem(i, 4, QtWidgets.QTableWidgetItem(ShortFIO))
            self.AleshaTable.setItem(i, 5, QtWidgets.QTableWidgetItem(df.iat[i, 5]))
            self.AleshaTable.setItem(i, 6, QtWidgets.QTableWidgetItem(df.iat[i, 6]))
            days = df.iat[i, 7]
            self.AleshaTable.setItem(i, 7, QtWidgets.QTableWidgetItem(str(days)))
        dfEO = df.groupby(['ЭО']).size().reset_index(name='count')
        self.AleshaTableEO.setColumnCount(2)
        self.AleshaTableEO.setColumnWidth(0, 160)
        self.AleshaTableEO.setColumnWidth(1, 65)
        self.AleshaTableEO.setRowCount(len(dfEO))
        print(dfEO)
        for i in range(len(dfEO)):
            self.AleshaTableEO.setItem(i, 0, QtWidgets.QTableWidgetItem(dfEO.iat[i, 0]))
            self.AleshaTableEO.setItem(i, 1, QtWidgets.QTableWidgetItem(str(dfEO.iat[i, 1])))

        dfExp = df.groupby(['Эксперт']).size().reset_index(name='count')
        self.AleshaTableExp.setColumnCount(2)
        self.AleshaTableExp.setColumnWidth(0, 160)
        self.AleshaTableExp.setColumnWidth(1, 65)
        self.AleshaTableExp.setRowCount(len(dfExp))
        print(dfEO)
        for i in range(len(dfExp)):
            ShortFIO = dfExp.iat[i, 0]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTableExp.setItem(i, 0, QtWidgets.QTableWidgetItem(ShortFIO))
            self.AleshaTableExp.setItem(i, 1, QtWidgets.QTableWidgetItem(str(dfExp.iat[i, 1])))

        dfIsp = df.groupby(['Исполнитель']).size().reset_index(name='count')
        self.AleshaTableIsp.setColumnCount(2)
        self.AleshaTableIsp.setColumnWidth(0, 160)
        self.AleshaTableIsp.setColumnWidth(1, 65)
        self.AleshaTableIsp.setRowCount(len(dfIsp))
        print(dfIsp)
        for i in range(len(dfIsp)):
            ShortFIO = dfIsp.iat[i, 0]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTableIsp.setItem(i, 0, QtWidgets.QTableWidgetItem(ShortFIO))
            self.AleshaTableIsp.setItem(i, 1, QtWidgets.QTableWidgetItem(str(dfIsp.iat[i, 1])))

class CreateAleshaAct(QDialog):
    def __init__(self):
        global AleshaOtdeal1
        global AleshaOtdeal2
        super(CreateAleshaAct, self).__init__()
        cal = Russia()
        cal.holidays(2022)
        pd.set_option("display.max_rows", None, "display.max_columns", None)
        loadUi("SourceGitHub/UI/AleshaAct.ui", self)
        self.AleshaTable.setColumnCount(8)
        self.AleshaTable.setColumnWidth(0, 50)
        self.AleshaTable.setColumnWidth(1, 70)
        self.AleshaTable.setColumnWidth(2, 70)
        self.AleshaTable.setColumnWidth(3, 120)
        self.AleshaTable.setColumnWidth(4, 120)
        self.AleshaTable.setColumnWidth(5, 150)
        self.AleshaTable.setColumnWidth(6, 90)
        self.AleshaTable.setColumnWidth(7, 65)


        conn = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
        cursor = conn.cursor()
        cursor.execute(
            "SELECT ТипГУ, Регистрационныйномер, Датарегистрации, Ответственныйисполнитель, ФИОЭкспертапоаккредитации, НаименованиеЭО, Договор1Г FROM GU1Group WHERE "
            "(((((ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND "
            "(ТипГУ = 'РОА' OR ТипГУ = 'АК') AND (ДатаПриказа2гр != '' and ДатаПриказа2гр is NOT NULL) AND "
            "Итог2Гр != 'Отриц. отказ') or ((ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ != '' AND "
            "ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND "
            "(ТипГУ != 'РОА' AND ТипГУ != 'АК'))) AND (Исполнитель3Гр = '' OR Исполнитель3Гр is NULL)) AND "
            "(ОтделУП = ? or ОтделУП = ?)) AND (ДатаПриказа3гр = '' OR ДатаПриказа3гр is NULL) AND "
            "(Договор1Г != '' AND Договор1Г is NOT NULL)", (AleshaOtdeal1, AleshaOtdeal2))  # Находит кол-во ГУ на выборе эксперта
        rows = cursor.fetchall()
        cursor.execute(
            "SELECT COUNT(*) FROM GU1Group WHERE (((((ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND "
            "(ТипГУ = 'РОА' OR ТипГУ = 'АК') AND (ДатаПриказа2гр != '' and ДатаПриказа2гр is NOT NULL) AND "
            "Итог2Гр != 'Отриц. отказ') or ((ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ != '' AND "
            "ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND "
            "(ТипГУ != 'РОА' AND ТипГУ != 'АК'))) AND (Исполнитель3Гр = '' OR Исполнитель3Гр is NULL)) AND "
            "(ОтделУП = ? or ОтделУП = ?)) AND (ДатаПриказа3гр = '' OR ДатаПриказа3гр is NULL) AND "
            "(Договор1Г != '' AND Договор1Г is NOT NULL)", (AleshaOtdeal1, AleshaOtdeal2))
        Longe = cursor.fetchone()
        df = pd.DataFrame({'Тип ГУ': [], 'Номер ГУ': [], 'Дата регистрации': [], 'Исполнитель': [], 'Эксперт': [], 'ЭО': [], 'Договор': [], 'Прошло': []})
        df.rename(columns={0: 'Тип ГУ', 1: 'Номер ГУ', 2: 'Дата регистрации', 3: 'Исполнитель', 4: 'Эксперт', 5: 'ЭО', 6: 'Договор', 7: 'Прошло'}, inplace=True);
        for i in range(Longe[0]):
            Vib_Exp = rows[i][6]
            Vib_Exp = Vib_Exp.split()
            Len_Vib_Exp = len(Vib_Exp)
            if Len_Vib_Exp > 1:
                Vib_Exp = Vib_Exp[2].replace("г.", "")
            else:
                Vib_Exp = Vib_Exp[0]
            Date_Vib_Exp_Date = datetime.strptime(Vib_Exp, '%d.%m.%Y')
            Now = datetime.today()
            Days = cal.get_working_days_delta(Date_Vib_Exp_Date, Now)
            if (rows[i][0] == "АК" or rows[i][0] == "РОА") and int(Days) >= 33:
                df.loc[len(df.index)] = [rows[i][0], rows[i][1], rows[i][2], rows[i][3], rows[i][4], rows[i][5], Vib_Exp, Days]
            elif (rows[i][0] == "ПК1" or rows[i][0] == "ПК2" or rows[i][0] == "ПК2+ИМОД" or rows[i][0] == "ПК2+ИМОД") and int(Days) >= 20:
                df.loc[len(df.index)] = [rows[i][0], rows[i][1], rows[i][2], rows[i][3], rows[i][4], rows[i][5], Vib_Exp, Days]
            elif (rows[i][0] != "АК" and rows[i][0] != "РОА" and rows[i][0] != "ПК1" and rows[i][0] != "ПК2" and rows[i][0] != "ПК2+ИМОД" and rows[i][0] != "ПК2+ИМОД")  and int(Days) >= 28:
                df.loc[len(df.index)] = [rows[i][0], rows[i][1], rows[i][2], rows[i][3], rows[i][4], rows[i][5], Vib_Exp, Days]
        print(df)
        self.AleshaTable.setRowCount(len(df))
        for i in range(len(df)):
            self.AleshaTable.setItem(i, 0, QtWidgets.QTableWidgetItem(df.iat[i, 0]))
            self.AleshaTable.setItem(i, 1, QtWidgets.QTableWidgetItem(df.iat[i, 1]))
            self.AleshaTable.setItem(i, 2, QtWidgets.QTableWidgetItem(df.iat[i, 2]))
            ShortFIO = df.iat[i, 3]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTable.setItem(i, 3, QtWidgets.QTableWidgetItem(ShortFIO))
            ShortFIO = df.iat[i, 4]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTable.setItem(i, 4, QtWidgets.QTableWidgetItem(ShortFIO))
            self.AleshaTable.setItem(i, 5, QtWidgets.QTableWidgetItem(df.iat[i, 5]))
            self.AleshaTable.setItem(i, 6, QtWidgets.QTableWidgetItem(df.iat[i, 6]))
            days = df.iat[i, 7]
            self.AleshaTable.setItem(i, 7, QtWidgets.QTableWidgetItem(str(days)))
        dfEO = df.groupby(['ЭО']).size().reset_index(name='count')
        self.AleshaTableEO.setColumnCount(2)
        self.AleshaTableEO.setColumnWidth(0, 160)
        self.AleshaTableEO.setColumnWidth(1, 65)
        self.AleshaTableEO.setRowCount(len(dfEO))
        print(dfEO)
        for i in range(len(dfEO)):
            self.AleshaTableEO.setItem(i, 0, QtWidgets.QTableWidgetItem(dfEO.iat[i, 0]))
            self.AleshaTableEO.setItem(i, 1, QtWidgets.QTableWidgetItem(str(dfEO.iat[i, 1])))

        dfExp = df.groupby(['Эксперт']).size().reset_index(name='count')
        self.AleshaTableExp.setColumnCount(2)
        self.AleshaTableExp.setColumnWidth(0, 160)
        self.AleshaTableExp.setColumnWidth(1, 65)
        self.AleshaTableExp.setRowCount(len(dfExp))
        print(dfEO)
        for i in range(len(dfExp)):
            ShortFIO = dfExp.iat[i, 0]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTableExp.setItem(i, 0, QtWidgets.QTableWidgetItem(ShortFIO))
            self.AleshaTableExp.setItem(i, 1, QtWidgets.QTableWidgetItem(str(dfExp.iat[i, 1])))

        dfIsp = df.groupby(['Исполнитель']).size().reset_index(name='count')
        self.AleshaTableIsp.setColumnCount(2)
        self.AleshaTableIsp.setColumnWidth(0, 160)
        self.AleshaTableIsp.setColumnWidth(1, 65)
        self.AleshaTableIsp.setRowCount(len(dfIsp))
        print(dfIsp)
        for i in range(len(dfIsp)):
            ShortFIO = dfIsp.iat[i, 0]
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.AleshaTableIsp.setItem(i, 0, QtWidgets.QTableWidgetItem(ShortFIO))
            self.AleshaTableIsp.setItem(i, 1, QtWidgets.QTableWidgetItem(str(dfIsp.iat[i, 1])))


class SendMailRass(QDialog):
    def __init__(self):
        global MsgError
        super(SendMailRass, self).__init__()
        self.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.CustomizeWindowHint)
        loadUi("SourceGitHub/UI/Mailer.ui", self)

        self.CopySpisok.clicked.connect(self.CopyToTable)


    def CopyToTable(self):
        Spisok = load_workbook("SourceGitHub/Shablon/Список рассылки/Список рассылки.xlsx")
        c = Spisok['Лист1']
        for i in range(0, 3000):
            self.MailTable.setItem(i, 0, QtWidgets.QTableWidgetItem(c.cell(row=i + 2, column=1).value))
            self.MailTable.setItem(i, 1, QtWidgets.QTableWidgetItem(c.cell(row=i + 2, column=2).value))
            self.MailTable.setItem(i, 2, QtWidgets.QTableWidgetItem(c.cell(row=i + 2, column=3).value))
            self.MailTable.setItem(i, 3, QtWidgets.QTableWidgetItem(c.cell(row=i + 2, column=4).value))
            self.MailTable.setItem(i, 4, QtWidgets.QTableWidgetItem(c.cell(row=i + 2, column=5).value))





    def SendMailRassStart(self):
        global globalNomerGU
        EmailExp = self.EmailExpert.toPlainText()
        EmailEO = self.MailREO.toPlainText()
        FullGU = self.NomerGU.toPlainText()
        datenow = QDate.currentDate().toPyDate()
        datenow = str(datenow)
        datenow = datenow.replace("-", " ")
        datenow = datenow.split()
        datenowStr = datenow[2] + "." + datenow[1] + "." + datenow[0]
        self.DataYvedT.setText(datenowStr)
        #self.DataYvedT.setAlignment(Qt.AlignCenter)
        outlook = win32.Dispatch('outlook.application')
        mail = outlook.CreateItem(0)
        mail.To = EmailExp + ";" + EmailEO
        mail.Subject = 'Уведомление о выборе ЭА (' + FullGU + ")"
        mail.HTMLBody = '<p><strong>Здравствуйте!</strong></p><p>Вам была назначена работа в системе ФГИС 2. Прошу ' \
                        'Вас зайти в свой личный кабинет и сформировать согласие (либо отказ) и предложение по составу ' \
                        'экспертной группы (по образцу) в карточке ГУ. Все документы должны быть подписаны ЭЦП. ' \
                        '<strong><u>В случае возникновения проблем с работой информационной системы ФГИС 2 прошу ' \
                        'Вас направлять документы через электронную приемную ФСА (включая срок регистрации 3 дня) ' \
                        'и обращаться в службу поддержки ФГИС.</u></strong><o:p></o:p></p><p>В соответствии с ' \
                        'пунктом 8 статьи 17 Федерального закона от 28 декабря 2013 г. № 412-ФЗ &laquo;Об аккредитации ' \
                        'в национальной системе аккредитации&raquo; состав экспертной группы определяется национальным ' \
                        'органом по аккредитации на основании предложений эксперта по аккредитации о привлечении ' \
                        'технических экспертов, необходимых для проведения экспертизы представленных заявителем ' \
                        'документов и сведений, выездной экспертизы соответствия заявителя критериям аккредитации, ' \
                        'из числа технических экспертов, включенных в реестр технических экспертов. Такие предложения ' \
                        'должны быть направлены в национальный орган по аккредитации в течение пяти рабочих дней со ' \
                        'дня отбора эксперта по аккредитации и содержать сведения о согласии технических экспертов ' \
                        'на участие в проведении экспертизы представленных заявителем документов и сведений и ' \
                        'проведении выездной экспертизы соответствия заявителя критериям аккредитации.</p><p>Таким ' \
                        'образом, прошу Вас направить уведомление в сроки установленные законодательством Российской ' \
                        'Федерации.</p><p>Спасибо!</p><p><strong>&nbsp;</strong></p>'  # this field is optional
        mail.Display(True)




class MainWindow(QDialog):
    def __init__(self):
        global FindNumGU
        global z, NomerGUforDeadline
        global Deadline
        global DeadlineDay
        global FullFioLogin
        global TipDop
        global UPR
        global user
        global UseTableGU
        super(MainWindow, self).__init__()
        loadUi("SourceGitHub/UI/mainForm.ui", self)

        # self.loading_screen = LoadScreen()
        UseTableGU = "1 Группа"

        self.bgwidget.clicked = False
        self.bgwidget.move(2120, 158)
        self.left_menu.move(2120, 158)
        self.tableWidget.setColumnWidth(0, 70)
        self.tableWidget.setColumnWidth(1, 150)
        self.tableWidget.setColumnWidth(2, 50)
        self.tableWidget.setColumnWidth(3, 70)
        self.tableWidget.setColumnWidth(4, 70)
        self.tableWidget.setColumnWidth(5, 130)
        self.tableWidget.setColumnWidth(6, 120)
        self.tableWidget.setColumnWidth(7, 75)
        self.tableWidget.setColumnWidth(8, 156)
        self.tableWidget.setColumnWidth(9, 120)
        self.tableWidget.setColumnWidth(10, 150)
        self.tableWidget.setColumnWidth(11, 70)
        self.tableWidget.setColumnWidth(12, 44)
        self.tableWidget.setColumnWidth(13, 85)
        self.tableWidget.setColumnWidth(14, 81)
        self.Button_GR1_1.hide()
        self.Button_GR1_2.hide()
        self.Button_GR1_3.hide()
        self.Button_GR2_1.hide()
        self.Button_GR2_2.hide()
        self.Button_GR2_3.hide()
        self.Button_GR3_1.hide()
        self.Button_GR3_2.hide()
        self.Button_GR3_3.hide()

        self.Button_GR1.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Button_GR1_0.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Button_GR1_1.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Button_GR1_2.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Button_GR1_3.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Button_GR2.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Button_GR2_1.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Button_GR2_2.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Button_GR2_3.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Button_GR3.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Button_GR3_1.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Button_GR3_2.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Button_GR3_3.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.FindButton.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Ispolnitel.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.FilterGU.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Mini.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Closeer.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.FilterStatus.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Back.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))
        self.Next.setGraphicsEffect(QGraphicsDropShadowEffect(blurRadius=5, xOffset=3, yOffset=3))





        self.movie = QMovie("SourceGitHub/Pictures/Book.gif")
        self.Book.setMovie(self.movie)
        self.movie.start()
        self.Book_button.clicked.connect(self.OpenBook)
        #self.movie = QMovie("SourceGitHub/Pictures/lady-noir.gif")
        #self.Alesha_l.setMovie(self.movie)
        #self.movie.start()
        self.Alesha_b.clicked.connect(self.goToCreateAleshaChoise)
        self.Update_b.clicked.connect(self.goToCreateUpdate)

        #self.movie = QMovie("SourceGitHub/Pictures/giphy.gif")
        #self.Download_Excel.setMovie(self.movie)
        #self.movie.start()
        self.Download_Excel_button.clicked.connect(self.gotoWhat)
        #self.movie = QMovie("SourceGitHub/Pictures/sharingan-item.gif")
        #self.Portal_Ku_l.setMovie(self.movie)
        #self.movie.start()
        self.Portal_Ku_b.clicked.connect(self.DoPortalKu)
        #self.movie = QMovie("SourceGitHub/Pictures/Mailer.gif")
        #self.Mailer_l.setMovie(self.movie)
        #self.movie.start()
        #self.Mailer_b.clicked.connect(self.DoPortalKu)
        #self.movie = QMovie("SourceGitHub/Pictures/DashIcon.gif")
        #self.DashL.setMovie(self.movie)
        #self.movie.start()
        # self.DashB.clicked.connect(self.DoPortalKu)

        self.FilterStatus.hide()
        self.Back.hide()
        self.Next.hide()
        self.Long_Pages.hide()

        self.tableWidget.hide()

        TipDop = int(TipDop)
        if TipDop != 10:
            self.Update_b.setEnabled(False)


        self.Closeer.clicked.connect(QCoreApplication.instance().quit)
        self.Mini.clicked.connect(self.Minamal)
        self.Mailer_b.clicked.connect(self.gotoSendMailRass)

        self.Left_minimizer.clicked.connect(lambda: self.slideLeftMenu())

        self.Button_GR1.clicked.connect(lambda: self.Menu1G())
        self.Button_GR1_0.clicked.connect(self.GR0_NeRospis)
        self.Button_GR2.clicked.connect(lambda: self.Menu2G())
        self.Button_GR2_1.clicked.connect(self.DoTable2G)
        self.Button_GR2_3.clicked.connect(self.DoTableEndGU2GR)
        self.Button_GR3.clicked.connect(lambda: self.Menu3G())
        self.Button_GR3_1.clicked.connect(self.AllIn3Group)
        self.Button_GR3_2.clicked.connect(self.In3Group)
        self.Button_GR3_3.clicked.connect(self.DoTableEndGU)
        self.loaddata()

        self.Button_GR1_1.clicked.connect(self.GR1_NeRospis)
        #if TipDop == 10:
        self.Button_GR1_2.clicked.connect(self.GR2_NeRospis)
        self.Button_GR1_3.clicked.connect(self.GR3_NeRospis)
        self.Back.clicked.connect(self.BackPage)
        self.Next.clicked.connect(self.NextPage)
        connection777 = sqlite3.connect("SourceGitHub/DB/shop_data.db")
        isp = connection777.cursor()
        isp.execute("SELECT DownloadDB FROM Login WHERE username = ?", (user,))
        DownloadDB = isp.fetchone()
        connection777.close()
        if DownloadDB[0] != "Обновил":
            self.UpdateGU1()
        #else:
            #self.DoTable()
    def In3Group(self):
        global Allchek
        global curPage
        curPage = 1
        Allchek = "Расписанная 3"
        self.DoTable3G()


    def AllIn3Group(self):
        global Allchek
        global NextPage
        global curPage
        curPage = 1
        NextPage = "Нерасписанныая 3"
        Allchek = "Нерасписанныая 3"
        self.DoTable3G()


    def gotoWhat(self):
        create = CreateWhat()
        create.exec_()

############################################
    def UpdateGU1(self):
        try:
            connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
            connection.commit()
            connection.close()

            shutil.copy(r'SourceGitHub/DB/BaseFrom1C.db',
                        r'C:/Portal/UseFile/BaseFrom1C.db')
            FindSettings = shutil.which(r'C:/Portal/UseFile/Settings.db')
            connection = sqlite3.connect('/Portal/UseFile/BaseFrom1C.db')
            cur1C = connection.cursor()
            cur1C.execute("SELECT COUNT(0) FROM TableFrom1C")
            sqllong = cur1C.fetchone()  # Находим длинну базы данных в кортеже
            sqllong2 = sqllong[0]  # Вытаскиваем длину баззы данных из первого картежа
            cur1C.execute("SELECT * FROM TableFrom1C")
            allGU1Group = cur1C.fetchall()
            connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
            curTo1g = connection.cursor()

            connection777 = sqlite3.connect("SourceGitHub/DB/shop_data.db")
            isp = connection777.cursor()
            isp.execute("SELECT DownloadDB FROM Login WHERE username = ?", (user, ))
            DownloadDB = isp.fetchone()

            if  True == True:
                curTo1g.execute("UPDATE GU1Group SET ОтделУП = NULL")
                isp.execute("SELECT * FROM Login")
                allIsp = isp.fetchall()
                isp.execute("SELECT COUNT(*) as count FROM Login")
                HowMuchIsp = isp.fetchone()
                HowMuchIsp = int(HowMuchIsp[0])
                i = 0
                for i in range(HowMuchIsp):
                    FullIspList = allIsp[i][3] + " " + allIsp[i][2] + " " + allIsp[i][4]
                    UPRst = allIsp[i][9]
                    curTo1g.execute("SELECT COUNT(*) FROM GU1Group  WHERE Ответственныйисполнитель = ? and (ОтделУП is NULL or ОтделУП != ?)", (FullIspList, UPRst))
                    Nado = curTo1g.fetchone()
                    if Nado[0] != 0:
                        try:
                            curTo1g.execute("UPDATE GU1Group SET ОтделУП = ? WHERE Ответственныйисполнитель = ? and (ОтделУП is NULL or ОтделУП != ? or ОтделУП = '')",
                                        (UPRst, FullIspList, UPRst))
                        except:
                            print("WTF")
                connection777.commit()
                connection777.close()
                try:
                    curTo1g.execute("UPDATE GU1Group SET ТипГУ = 'ПК' WHERE ТипГУ = 'Подтверждение компетентности'")
                except:
                    print("Null change ПК")
                try:
                    curTo1g.execute("UPDATE GU1Group SET ТипГУ = 'АК' WHERE ТипГУ = 'Аккредитация'")
                except:
                    print("Null change РОА")
                try:
                    curTo1g.execute("UPDATE GU1Group SET ТипГУ = 'РОА' WHERE ТипГУ = 'Расширение области аккредитации'")
                except:
                    print("Null change АК")

                for i in range(sqllong2):
                    curTo1g.execute("SELECT * FROM GU1Group WHERE Регистрационныйномер = ?", (allGU1Group[i][1],))
                    FindAllGU1Group = curTo1g.fetchall()
                    if len(FindAllGU1Group) == 0:
                        curTo1g.execute(
                            "INSERT INTO GU1Group (ТипГУ, Регистрационныйномер, Датарегистрации, РегистрационныйномерАЛвРАЛ, "
                            "Наименованиезаявителя, Полноенаименованиезаявителя, ИННЗаявителя, Ответственныйисполнитель, "
                            "Управлениеответственногоисполнителя, ЭтапГУ, СтатусГУ, Информацияотекущейактивнойзадаче, "
                            "Датаподачизаявления, ЗаявлениенаотзывГУ, Датаназначенияответственногоисполнителя, "
                            "ДатапроверкизаявленияОИ, СЗовозвратевОДК, Приказоботказе, Датавыбораэксперта, "
                            "УведомлениеоботбореЭА, ФИОЭкспертапоаккредитации, НаименованиеЭО, "
                            "ВремямеждурегистрациейивыборомЭАврабочихднях, Согласиеилиотказэкспертапоаккредитации, "
                            "ПредложениепосоставуЭГ, ДатаприказаосоставеЭГ, "
                            "Экспертноезаключениепорезультатамдокументарнойоценки, ПисьмоовозвратеЭЗ, "
                            "ПриказнаприостановкупорезультатамДО, Письмообустранениинарушений, Приказнавыезднуюоценку, "
                            "Приказоботказеваккредитации, Актвыезднойэкспертизы, Письмоовозвратеактавыезднойэкспертизы, "
                            "СЗсотчетомоВО, ПриказнаприостановкупорезультатамВО, "
                            "Письмоорассмотренииотчетазаявителяэкспертомпоаккредитации, СЗосогласительнойкомиссии, "
                            "Протоколсогласительнойкомиссии, "
                            "ЗапросдополненийкактувыезднойэкспертизыотэкспертнойгруппыиилизапросвФНСРоссиииПФР, "
                            "Запроснадополнениякакту, Заключениеобоценкеустранениязаявителемвыявленныхнесоответствий, "
                            "ПриказозавершенииГУ, ПодготовилприказозавершенииГУ, ПриказоботказевпредоставленииГУ) VALUES (?,"
                            "?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                            (allGU1Group[i][0], allGU1Group[i][1], allGU1Group[i][2], allGU1Group[i][3], allGU1Group[i][4],
                             allGU1Group[i][5], allGU1Group[i][6], allGU1Group[i][7], allGU1Group[i][8], allGU1Group[i][9],
                             allGU1Group[i][10], allGU1Group[i][11], allGU1Group[i][12], allGU1Group[i][13], allGU1Group[i][14],
                             allGU1Group[i][15], allGU1Group[i][16], allGU1Group[i][17], allGU1Group[i][18], allGU1Group[i][19],
                             allGU1Group[i][20], allGU1Group[i][21], allGU1Group[i][22], allGU1Group[i][23], allGU1Group[i][24],
                             allGU1Group[i][25], allGU1Group[i][26], allGU1Group[i][27], allGU1Group[i][28], allGU1Group[i][29],
                             allGU1Group[i][30], allGU1Group[i][31], allGU1Group[i][32], allGU1Group[i][33], allGU1Group[i][34],
                             allGU1Group[i][35], allGU1Group[i][36], allGU1Group[i][37], allGU1Group[i][38], allGU1Group[i][39],
                             allGU1Group[i][40], allGU1Group[i][41], allGU1Group[i][42], allGU1Group[i][43],
                             allGU1Group[i][44]))  # работает

                    NomerGUforDeadline = allGU1Group[i][1]
                    print(i)
                    if FindAllGU1Group != []:
                        Stat = allGU1Group[i][7]
                        Din = FindAllGU1Group[0][7]
                        if Stat != Din:
                            curTo1g.execute('UPDATE GU1Group SET Ответственныйисполнитель = ? WHERE Регистрационныйномер = ?',
                                            (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][10]
                        Din = FindAllGU1Group[0][10]
                        if Din is None:
                            curTo1g.execute(
                                'UPDATE GU1Group SET СтатусГУ = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))
                        if Din == "":
                            curTo1g.execute(
                                'UPDATE GU1Group SET СтатусГУ = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))
                        if Din == "На регистрации":
                            curTo1g.execute(
                                'UPDATE GU1Group SET СтатусГУ = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))
                        if Stat == "Возврат без рассмотрения":
                            curTo1g.execute(
                                'UPDATE GU1Group SET СтатусГУ = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][13]
                        Din = FindAllGU1Group[0][13]
                        if Stat != Din:
                            curTo1g.execute('UPDATE GU1Group SET ЗаявлениенаотзывГУ = ? WHERE Регистрационныйномер = ?',
                                            (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][14]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][14]
                        if Stat != Din:
                            curTo1g.execute(
                                'UPDATE GU1Group SET Датаназначенияответственногоисполнителя = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][15]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][15]
                        if Stat != Din:
                            curTo1g.execute('UPDATE GU1Group SET ДатапроверкизаявленияОИ = ? WHERE Регистрационныйномер = ?',
                                            (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][16]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][16]
                        if Stat != Din:
                            curTo1g.execute('UPDATE GU1Group SET СЗовозвратевОДК = ? WHERE Регистрационныйномер = ?',
                                            (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][17]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][17]
                        if Stat != Din:
                            curTo1g.execute('UPDATE GU1Group SET Приказоботказе = ? WHERE Регистрационныйномер = ?',
                                            (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][18]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][18]
                        if Stat != Din:
                            curTo1g.execute('UPDATE GU1Group SET Датавыбораэксперта = ? WHERE Регистрационныйномер = ?',
                                            (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][20]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][20]
                        if Stat != Din:
                            curTo1g.execute('UPDATE GU1Group SET ФИОЭкспертапоаккредитации = ? WHERE Регистрационныйномер = ?',
                                            (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][21]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][21]
                        if Stat != Din:
                            curTo1g.execute('UPDATE GU1Group SET НаименованиеЭО = ? WHERE Регистрационныйномер = ?',
                                            (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][22]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][22]
                        if Stat != Din:
                            curTo1g.execute(
                                'UPDATE GU1Group SET ВремямеждурегистрациейивыборомЭАврабочихднях = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][24]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][24]
                        if Stat != Din:
                            curTo1g.execute('UPDATE GU1Group SET ПредложениепосоставуЭГ = ? WHERE Регистрационныйномер = ?',
                                            (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][25]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][25]
                        if Stat != Din :
                            curTo1g.execute(
                                'UPDATE GU1Group SET ДатаприказаосоставеЭГ = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][26]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][26]
                        if Stat != Din:
                            curTo1g.execute(
                                'UPDATE GU1Group SET Экспертноезаключениепорезультатамдокументарнойоценки = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][27]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][27]
                        if Stat != Din:
                            curTo1g.execute('UPDATE GU1Group SET ПисьмоовозвратеЭЗ = ? WHERE Регистрационныйномер = ?',
                                            (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][28]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][28]
                        if Stat != Din:
                            curTo1g.execute(
                                'UPDATE GU1Group SET ПриказнаприостановкупорезультатамДО = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][29]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][29]
                        if Stat != Din:
                            curTo1g.execute(
                                'UPDATE GU1Group SET Письмообустранениинарушений = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][30]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][30]
                        if Stat != Din:
                            curTo1g.execute('UPDATE GU1Group SET Приказнавыезднуюоценку = ? WHERE Регистрационныйномер = ?',
                                            (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][31]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][31]
                        if Stat != Din:
                            curTo1g.execute(
                                'UPDATE GU1Group SET Приказоботказеваккредитации = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][32]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][32]
                        if Stat != Din:
                            curTo1g.execute('UPDATE GU1Group SET Актвыезднойэкспертизы = ? WHERE Регистрационныйномер = ?',
                                            (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][33]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][33]
                        if Stat != Din:
                            curTo1g.execute(
                                'UPDATE GU1Group SET Письмоовозвратеактавыезднойэкспертизы = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][34]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][34]
                        if Stat != Din:
                            curTo1g.execute('UPDATE GU1Group SET СЗсотчетомоВО = ? WHERE Регистрационныйномер = ?',
                                            (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][35]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][35]
                        if Stat != Din:
                            curTo1g.execute(
                                'UPDATE GU1Group SET ПриказнаприостановкупорезультатамВО = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][36]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][36]
                        if Stat != Din:
                            curTo1g.execute(
                                'UPDATE GU1Group SET Письмоорассмотренииотчетазаявителяэкспертомпоаккредитации = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][37]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][37]
                        if Stat != Din:
                            curTo1g.execute('UPDATE GU1Group SET СЗосогласительнойкомиссии = ? WHERE Регистрационныйномер = ?',
                                            (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][38]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][38]
                        if Stat != Din:
                            curTo1g.execute(
                                'UPDATE GU1Group SET Протоколсогласительнойкомиссии = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][39]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][39]
                        if Stat != Din:
                            curTo1g.execute(
                                'UPDATE GU1Group SET ЗапросдополненийкактувыезднойэкспертизыотэкспертнойгруппыиилизапросвФНСРоссиииПФР = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][40]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][40]
                        if Stat != Din:
                            curTo1g.execute(
                                'UPDATE GU1Group SET Запроснадополнениякакту = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][41]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][41]
                        if Stat != Din:
                            curTo1g.execute(
                                'UPDATE GU1Group SET Заключениеобоценкеустранениязаявителемвыявленныхнесоответствий = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][42]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][42]
                        if Stat != Din:
                            curTo1g.execute('UPDATE GU1Group SET ПриказозавершенииГУ = ? WHERE Регистрационныйномер = ?',
                                            (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][43]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][43]
                        if Stat != Din:
                            curTo1g.execute(
                                'UPDATE GU1Group SET ПодготовилприказозавершенииГУ = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][44]
                        if Stat == "		Не зарегистрирован ":
                            Stat = "Не зарегистрирован"
                        elif Stat == "		Не зарегистрировано ":
                            Stat = "Не зарегистрировано"
                        Din = FindAllGU1Group[0][44]
                        if Stat != Din:
                            curTo1g.execute(
                                'UPDATE GU1Group SET ПриказоботказевпредоставленииГУ = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

                        Stat = allGU1Group[i][7]
                        Din = FindAllGU1Group[0][112]
                        Obl = FindAllGU1Group[0][0]
                        Pricaz = FindAllGU1Group[0][25]
                        if (Din == "" or Din is None or Din == []) and (Obl == "РОА" or Obl == "АК") and (Pricaz != "" and Pricaz is not None and Pricaz != []):
                            curTo1g.execute(
                                'UPDATE GU1Group SET Исполнитель2Гр = ? WHERE Регистрационныйномер = ?',
                                (Stat, NomerGUforDeadline))

            curTo1g.execute('DELETE FROM GU1Group WHERE Регистрационныйномер is Null')
            #isp.execute("UPDATE Login SET DownloadDB = ? WHERE username = ?", ("Обновил", user,))

            connection.commit()
            connection.close()

        except:
            connection777.close()
            connection.close()
            self.UpdateGU1()

        ############################################

    def goToCreateAleshaChoise(self):
        create = CreateAleshaChoise()
        create.exec_()

    def goToCreateUpdate(self):
        create = CreateUpdate()
        create.exec_()

    def gotoSendMailRass(self):
        create = SendMailRass()
        create.exec_()

    def OpenBook(self):
        GUFolder = "SourceGitHub/Book"
        if os.path.exists(GUFolder):
            os.startfile(GUFolder)
        else:
            # os.mkdir(GUFolder)
            os.startfile(GUFolder)

    def GR1_NeRospis(self):
        global FullFioLogin
        FullFioLogin = "NULL"
        self.DoTable()

    def GR0_NeRospis(self):
        global FullFioLogin
        global FindNumGU
        FullFioLogin = self.Ispolnitel.currentText()
        FindNumGU = "*"
        self.DoTable()

    def GR2_NeRospis(self):
        global FullFioLogin
        global FindNumGU
        FindNumGU = "*"
        FullFioLogin = "Приостановка"
        self.DoTable()

    def GR3_NeRospis(self):
        global FullFioLogin
        FullFioLogin = "Group1End"
        self.DoTable()

    def Minamal(self):
        widget.showMinimized()

    def Menu1G(self):
        self.Button_GR1_0.show()
        self.Button_GR1_1.show()
        self.Button_GR1_2.show()
        self.Button_GR1_3.show()
        self.Button_GR2_1.hide()
        self.Button_GR2_2.hide()
        self.Button_GR2_3.hide()
        self.Button_GR3_1.hide()
        self.Button_GR3_2.hide()
        self.Button_GR3_3.hide()
        if self.left_menu.width() == 60:
            self.slideLeftMenu()
        zxc = self.Button_GR1_0.y()
        if self.left_menu.width() != 60 and (
            self.Button_GR1_0.y() == 70 or self.Button_GR1_0.y() == 164 or self.Button_GR1_0.y() == 140):
            self.animation015 = QPropertyAnimation(self.Button_GR1_0, b"geometry")  # Animate minimumWidht
            self.animation015.setDuration(500)
            self.animation015.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation015.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation015.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation015.start()
            self.animation15 = QPropertyAnimation(self.Button_GR1_1, b"geometry")  # Animate minimumWidht
            self.animation15.setDuration(500)
            self.animation15.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation15.setEndValue(QRect(7, 240, 216, 44))  # end value is the new menu width
            self.animation15.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation15.start()
            self.animation16 = QPropertyAnimation(self.Button_GR1_2, b"geometry")  # Animate minimumWidht
            self.animation16.setDuration(500)
            self.animation16.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation16.setEndValue(QRect(7, 290, 216, 44))  # end value is the new menu width
            self.animation16.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation16.start()
            self.animation17 = QPropertyAnimation(self.Button_GR1_3, b"geometry")  # Animate minimumWidht
            self.animation17.setDuration(500)
            self.animation17.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation17.setEndValue(QRect(7, 340, 216, 44))  # end value is the new menu width
            self.animation17.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation17.start()
            self.animation18 = QPropertyAnimation(self.Button_GR2, b"geometry")  # Animate minimumWidht
            self.animation18.setDuration(500)
            self.animation18.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation18.setEndValue(QRect(7, 390, 216, 44))  # end value is the new menu width
            self.animation18.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation18.start()
            self.animation19 = QPropertyAnimation(self.Button_GR3, b"geometry")  # Animate minimumWidht
            self.animation19.setDuration(500)
            self.animation19.setStartValue(QRect(7, 240, 216, 44))  # Start value is the current menu width
            self.animation19.setEndValue(QRect(7, 440, 216, 44))  # end value is the new menu width
            self.animation19.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation19.start()
            self.animation20 = QPropertyAnimation(self.Button_GR2_1, b"geometry")  # Animate minimumWidht
            self.animation20.setDuration(500)
            self.animation20.setStartValue(QRect(7, 390, 216, 44))  # Start value is the current menu width
            self.animation20.setEndValue(QRect(7, 390, 216, 44))  # end value is the new menu width
            self.animation20.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation20.start()
            self.animation21 = QPropertyAnimation(self.Button_GR2_2, b"geometry")  # Animate minimumWidht
            self.animation21.setDuration(500)
            self.animation21.setStartValue(QRect(7, 390, 216, 44))  # Start value is the current menu width
            self.animation21.setEndValue(QRect(7, 390, 216, 44))  # end value is the new menu width
            self.animation21.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation21.start()
            self.animation22 = QPropertyAnimation(self.Button_GR2_3, b"geometry")  # Animate minimumWidht
            self.animation22.setDuration(500)
            self.animation22.setStartValue(QRect(7, 390, 216, 44))  # Start value is the current menu width
            self.animation22.setEndValue(QRect(7, 390, 216, 44))  # end value is the new menu width
            self.animation22.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation22.start()
            self.animation23 = QPropertyAnimation(self.Button_GR3_1, b"geometry")  # Animate minimumWidht
            self.animation23.setDuration(500)
            self.animation23.setStartValue(QRect(7, 440, 216, 44))  # Start value is the current menu width
            self.animation23.setEndValue(QRect(7, 440, 216, 44))  # end value is the new menu width
            self.animation23.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation23.start()
            self.animation24 = QPropertyAnimation(self.Button_GR3_2, b"geometry")  # Animate minimumWidht
            self.animation24.setDuration(500)
            self.animation24.setStartValue(QRect(7, 440, 216, 44))  # Start value is the current menu width
            self.animation24.setEndValue(QRect(7, 440, 216, 44))  # end value is the new menu width
            self.animation24.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation24.start()
            self.animation25 = QPropertyAnimation(self.Button_GR3_3, b"geometry")  # Animate minimumWidht
            self.animation25.setDuration(500)
            self.animation25.setStartValue(QRect(7, 440, 216, 44))  # Start value is the current menu width
            self.animation25.setEndValue(QRect(7, 440, 216, 44))  # end value is the new menu width
            self.animation25.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation25.start()

        elif self.Button_GR1_0.y() == 190:
            self.animation015 = QPropertyAnimation(self.Button_GR1_0, b"geometry")  # Animate minimumWidht
            self.animation015.setDuration(500)
            self.animation015.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation015.setEndValue(QRect(7, 164, 216, 0))  # end value is the new menu width
            self.animation015.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation015.start()
            self.animation15 = QPropertyAnimation(self.Button_GR1_1, b"geometry")  # Animate minimumWidht
            self.animation15.setDuration(500)
            self.animation15.setStartValue(QRect(7, 240, 216, 44))  # Start value is the current menu width
            self.animation15.setEndValue(QRect(7, 164, 216, 0))  # end value is the new menu width
            self.animation15.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation15.start()
            self.animation16 = QPropertyAnimation(self.Button_GR1_2, b"geometry")  # Animate minimumWidht
            self.animation16.setDuration(500)
            self.animation16.setStartValue(QRect(7, 290, 216, 44))  # Start value is the current menu width
            self.animation16.setEndValue(QRect(7, 164, 216, 0))  # end value is the new menu width
            self.animation16.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation16.start()
            self.animation17 = QPropertyAnimation(self.Button_GR1_3, b"geometry")  # Animate minimumWidht
            self.animation17.setDuration(500)
            self.animation17.setStartValue(QRect(7, 340, 216, 44))  # Start value is the current menu width
            self.animation17.setEndValue(QRect(7, 164, 216, 0))  # end value is the new menu width
            self.animation17.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation17.start()
            self.animation18 = QPropertyAnimation(self.Button_GR2, b"geometry")  # Animate minimumWidht
            self.animation18.setDuration(500)
            self.animation18.setStartValue(QRect(7, 390, 216, 44))  # Start value is the current menu width
            self.animation18.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation18.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation18.start()
            self.animation19 = QPropertyAnimation(self.Button_GR3, b"geometry")  # Animate minimumWidht
            self.animation19.setDuration(500)
            self.animation19.setStartValue(QRect(7, 440, 216, 44))  # Start value is the current menu width
            self.animation19.setEndValue(QRect(7, 240, 216, 44))  # end value is the new menu width
            self.animation19.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation19.start()
            self.animation20 = QPropertyAnimation(self.Button_GR2_1, b"geometry")  # Animate minimumWidht
            self.animation20.setDuration(500)
            self.animation20.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation20.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation20.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation20.start()
            self.animation21 = QPropertyAnimation(self.Button_GR2_2, b"geometry")  # Animate minimumWidht
            self.animation21.setDuration(500)
            self.animation21.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation21.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation21.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation21.start()
            self.animation22 = QPropertyAnimation(self.Button_GR2_3, b"geometry")  # Animate minimumWidht
            self.animation22.setDuration(500)
            self.animation22.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation22.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation22.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation22.start()
            self.animation23 = QPropertyAnimation(self.Button_GR3_1, b"geometry")  # Animate minimumWidht
            self.animation23.setDuration(500)
            self.animation23.setStartValue(QRect(7, 240, 216, 44))  # Start value is the current menu width
            self.animation23.setEndValue(QRect(7, 240, 216, 44))  # end value is the new menu width
            self.animation23.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation23.start()
            self.animation24 = QPropertyAnimation(self.Button_GR3_2, b"geometry")  # Animate minimumWidht
            self.animation24.setDuration(500)
            self.animation24.setStartValue(QRect(7, 240, 216, 44))  # Start value is the current menu width
            self.animation24.setEndValue(QRect(7, 240, 216, 44))  # end value is the new menu width
            self.animation24.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation24.start()
            self.animation25 = QPropertyAnimation(self.Button_GR3_3, b"geometry")  # Animate minimumWidht
            self.animation25.setDuration(500)
            self.animation25.setStartValue(QRect(7, 240, 216, 44))  # Start value is the current menu width
            self.animation25.setEndValue(QRect(7, 240, 216, 44))  # end value is the new menu width
            self.animation25.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation25.start()

    def Menu2G(self):
        self.Button_GR1_0.hide()
        self.Button_GR1_1.hide()
        self.Button_GR1_2.hide()
        self.Button_GR1_3.hide()
        self.Button_GR2_1.show()
        self.Button_GR2_2.show()
        self.Button_GR2_3.show()
        self.Button_GR3_1.hide()
        self.Button_GR3_2.hide()
        self.Button_GR3_3.hide()
        if self.left_menu.width() == 60:
            self.slideLeftMenu()
        zxc = self.Button_GR2_1.y()
        if self.left_menu.width() != 60 and (
                self.Button_GR2_1.y() == 120 or self.Button_GR2_1.y() == 212 or self.Button_GR2_1.y() == 190 or self.Button_GR2_1.y() == 390):
            self.animation015 = QPropertyAnimation(self.Button_GR1_0, b"geometry")  # Animate minimumWidht
            self.animation015.setDuration(500)
            self.animation015.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation015.setEndValue(QRect(7, 140, 216, 44))  # end value is the new menu width
            self.animation015.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation015.start()
            self.animation15 = QPropertyAnimation(self.Button_GR1_1, b"geometry")  # Animate minimumWidht
            self.animation15.setDuration(500)
            self.animation15.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation15.setEndValue(QRect(7, 140, 216, 44))  # end value is the new menu width
            self.animation15.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation15.start()
            self.animation16 = QPropertyAnimation(self.Button_GR1_2, b"geometry")  # Animate minimumWidht
            self.animation16.setDuration(500)
            self.animation16.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation16.setEndValue(QRect(7, 140, 216, 44))  # end value is the new menu width
            self.animation16.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation16.start()
            self.animation17 = QPropertyAnimation(self.Button_GR1_3, b"geometry")  # Animate minimumWidht
            self.animation17.setDuration(500)
            self.animation17.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation17.setEndValue(QRect(7, 140, 216, 44))  # end value is the new menu width
            self.animation17.setEndValue(QRect(7, 140, 216, 44))  # end value is the new menu width
            self.animation17.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation17.start()
            self.animation18 = QPropertyAnimation(self.Button_GR2, b"geometry")  # Animate minimumWidht
            self.animation18.setDuration(500)
            self.animation18.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation18.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation18.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation18.start()
            self.animation19 = QPropertyAnimation(self.Button_GR3, b"geometry")  # Animate minimumWidht
            self.animation19.setDuration(500)
            self.animation19.setStartValue(QRect(7, 240, 216, 44))  # Start value is the current menu width
            self.animation19.setEndValue(QRect(7, 390, 216, 44))  # end value is the new menu width
            self.animation19.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation19.start()
            self.animation20 = QPropertyAnimation(self.Button_GR2_1, b"geometry")  # Animate minimumWidht
            self.animation20.setDuration(500)
            self.animation20.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation20.setEndValue(QRect(7, 240, 216, 44))  # end value is the new menu width
            self.animation20.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation20.start()
            self.animation21 = QPropertyAnimation(self.Button_GR2_2, b"geometry")  # Animate minimumWidht
            self.animation21.setDuration(500)
            self.animation21.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation21.setEndValue(QRect(7, 290, 216, 44))  # end value is the new menu width
            self.animation21.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation21.start()
            self.animation22 = QPropertyAnimation(self.Button_GR2_3, b"geometry")  # Animate minimumWidht
            self.animation22.setDuration(500)
            self.animation22.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation22.setEndValue(QRect(7, 340, 216, 44))  # end value is the new menu width
            self.animation22.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation22.start()
            self.animation23 = QPropertyAnimation(self.Button_GR3_1, b"geometry")  # Animate minimumWidht
            self.animation23.setDuration(500)
            self.animation23.setStartValue(QRect(7, 390, 216, 44))  # Start value is the current menu width
            self.animation23.setEndValue(QRect(7, 390, 216, 44))  # end value is the new menu width
            self.animation23.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation23.start()
            self.animation24 = QPropertyAnimation(self.Button_GR3_2, b"geometry")  # Animate minimumWidht
            self.animation24.setDuration(500)
            self.animation24.setStartValue(QRect(7, 390, 216, 44))  # Start value is the current menu width
            self.animation24.setEndValue(QRect(7, 390, 216, 44))  # end value is the new menu width
            self.animation24.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation24.start()
            self.animation25 = QPropertyAnimation(self.Button_GR3_3, b"geometry")  # Animate minimumWidht
            self.animation25.setDuration(500)
            self.animation25.setStartValue(QRect(7, 390, 216, 44))  # Start value is the current menu width
            self.animation25.setEndValue(QRect(7, 390, 216, 44))  # end value is the new menu width
            self.animation25.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation25.start()

        elif self.Button_GR2_1.y() == 240:
            self.animation015 = QPropertyAnimation(self.Button_GR1_0, b"geometry")  # Animate minimumWidht
            self.animation015.setDuration(500)
            self.animation015.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation015.setEndValue(QRect(7, 140, 216, 0))  # end value is the new menu width
            self.animation015.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation015.start()
            self.animation15 = QPropertyAnimation(self.Button_GR1_1, b"geometry")  # Animate minimumWidht
            self.animation15.setDuration(500)
            self.animation15.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation15.setEndValue(QRect(7, 140, 216, 0))  # end value is the new menu width
            self.animation15.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation15.start()
            self.animation16 = QPropertyAnimation(self.Button_GR1_2, b"geometry")  # Animate minimumWidht
            self.animation16.setDuration(500)
            self.animation16.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation16.setEndValue(QRect(7, 140, 216, 0))  # end value is the new menu width
            self.animation16.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation16.start()
            self.animation17 = QPropertyAnimation(self.Button_GR1_3, b"geometry")  # Animate minimumWidht
            self.animation17.setDuration(500)
            self.animation17.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation17.setEndValue(QRect(7, 140, 216, 0))  # end value is the new menu width
            self.animation17.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation17.start()
            self.animation18 = QPropertyAnimation(self.Button_GR2, b"geometry")  # Animate minimumWidht
            self.animation18.setDuration(500)
            self.animation18.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation18.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation18.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation18.start()
            self.animation19 = QPropertyAnimation(self.Button_GR3, b"geometry")  # Animate minimumWidht
            self.animation19.setDuration(500)
            self.animation19.setStartValue(QRect(7, 390, 216, 44))  # Start value is the current menu width
            self.animation19.setEndValue(QRect(7, 240, 216, 44))  # end value is the new menu width
            self.animation19.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation19.start()
            self.animation20 = QPropertyAnimation(self.Button_GR2_1, b"geometry")  # Animate minimumWidht
            self.animation20.setDuration(500)
            self.animation20.setStartValue(QRect(7, 240, 216, 44))  # Start value is the current menu width
            self.animation20.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation20.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation20.start()
            self.animation21 = QPropertyAnimation(self.Button_GR2_2, b"geometry")  # Animate minimumWidht
            self.animation21.setDuration(500)
            self.animation21.setStartValue(QRect(7, 290, 216, 44))  # Start value is the current menu width
            self.animation21.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation21.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation21.start()
            self.animation22 = QPropertyAnimation(self.Button_GR2_3, b"geometry")  # Animate minimumWidht
            self.animation22.setDuration(500)
            self.animation22.setStartValue(QRect(7, 340, 216, 44))  # Start value is the current menu width
            self.animation22.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation22.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation22.start()
            self.animation23 = QPropertyAnimation(self.Button_GR3_1, b"geometry")  # Animate minimumWidht
            self.animation23.setDuration(500)
            self.animation23.setStartValue(QRect(7, 240, 216, 44))  # Start value is the current menu width
            self.animation23.setEndValue(QRect(7, 240, 216, 44))  # end value is the new menu width
            self.animation23.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation23.start()
            self.animation24 = QPropertyAnimation(self.Button_GR3_2, b"geometry")  # Animate minimumWidht
            self.animation24.setDuration(500)
            self.animation24.setStartValue(QRect(7, 240, 216, 44))  # Start value is the current menu width
            self.animation24.setEndValue(QRect(7, 240, 216, 44))  # end value is the new menu width
            self.animation24.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation24.start()
            self.animation25 = QPropertyAnimation(self.Button_GR3_3, b"geometry")  # Animate minimumWidht
            self.animation25.setDuration(500)
            self.animation25.setStartValue(QRect(7, 240, 216, 44))  # Start value is the current menu width
            self.animation25.setEndValue(QRect(7, 240, 216, 44))  # end value is the new menu width
            self.animation25.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation25.start()

    def Menu3G(self):
        self.Button_GR1_0.hide()
        self.Button_GR1_1.hide()
        self.Button_GR1_2.hide()
        self.Button_GR1_3.hide()
        self.Button_GR2_1.hide()
        self.Button_GR2_2.hide()
        self.Button_GR2_3.hide()
        self.Button_GR3_1.show()
        self.Button_GR3_2.show()
        self.Button_GR3_3.show()

        if self.left_menu.width() == 60:
            self.slideLeftMenu()
        zxc = self.Button_GR3_1.y()
        if self.left_menu.width() != 60 and (
                self.Button_GR3_1.y() == 170 or self.Button_GR3_1.y() == 252 or self.Button_GR3_1.y() == 240 or self.Button_GR3_1.y() == 390 or self.Button_GR3_1.y() == 440):
            self.animation015 = QPropertyAnimation(self.Button_GR1_0, b"geometry")  # Animate minimumWidht
            self.animation015.setDuration(500)
            self.animation015.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation015.setEndValue(QRect(7, 140, 216, 44))  # end value is the new menu width
            self.animation015.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation015.start()
            self.animation15 = QPropertyAnimation(self.Button_GR1_1, b"geometry")  # Animate minimumWidht
            self.animation15.setDuration(500)
            self.animation15.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation15.setEndValue(QRect(7, 140, 216, 44))  # end value is the new menu width
            self.animation15.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation15.start()
            self.animation16 = QPropertyAnimation(self.Button_GR1_2, b"geometry")  # Animate minimumWidht
            self.animation16.setDuration(500)
            self.animation16.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation16.setEndValue(QRect(7, 140, 216, 44))  # end value is the new menu width
            self.animation16.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation16.start()
            self.animation17 = QPropertyAnimation(self.Button_GR1_3, b"geometry")  # Animate minimumWidht
            self.animation17.setDuration(500)
            self.animation17.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation17.setEndValue(QRect(7, 140, 216, 44))  # end value is the new menu width
            self.animation17.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation17.start()
            self.animation18 = QPropertyAnimation(self.Button_GR2, b"geometry")  # Animate minimumWidht
            self.animation18.setDuration(500)
            self.animation18.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation18.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation18.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation18.start()
            self.animation19 = QPropertyAnimation(self.Button_GR3, b"geometry")  # Animate minimumWidht
            self.animation19.setDuration(500)
            self.animation19.setStartValue(QRect(7, 240, 216, 44))  # Start value is the current menu width
            self.animation19.setEndValue(QRect(7, 240, 216, 44))  # end value is the new menu width
            self.animation19.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation19.start()
            self.animation20 = QPropertyAnimation(self.Button_GR2_1, b"geometry")  # Animate minimumWidht
            self.animation20.setDuration(500)
            self.animation20.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation20.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation20.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation20.start()
            self.animation21 = QPropertyAnimation(self.Button_GR2_2, b"geometry")  # Animate minimumWidht
            self.animation21.setDuration(500)
            self.animation21.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation21.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation21.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation21.start()
            self.animation22 = QPropertyAnimation(self.Button_GR2_3, b"geometry")  # Animate minimumWidht
            self.animation22.setDuration(500)
            self.animation22.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation22.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation22.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation22.start()
            self.animation23 = QPropertyAnimation(self.Button_GR3_1, b"geometry")  # Animate minimumWidht
            self.animation23.setDuration(500)
            self.animation23.setStartValue(QRect(7, 240, 216, 44))  # Start value is the current menu width
            self.animation23.setEndValue(QRect(7, 290, 216, 44))  # end value is the new menu width
            self.animation23.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation23.start()
            self.animation24 = QPropertyAnimation(self.Button_GR3_2, b"geometry")  # Animate minimumWidht
            self.animation24.setDuration(500)
            self.animation24.setStartValue(QRect(7, 240, 216, 44))  # Start value is the current menu width
            self.animation24.setEndValue(QRect(7, 340, 216, 44))  # end value is the new menu width
            self.animation24.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation24.start()
            self.animation25 = QPropertyAnimation(self.Button_GR3_3, b"geometry")  # Animate minimumWidht
            self.animation25.setDuration(500)
            self.animation25.setStartValue(QRect(7, 240, 216, 44))  # Start value is the current menu width
            self.animation25.setEndValue(QRect(7, 390, 216, 44))  # end value is the new menu width
            self.animation25.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation25.start()

        elif self.Button_GR3_1.y() == 290:
            self.animation015 = QPropertyAnimation(self.Button_GR1_0, b"geometry")  # Animate minimumWidht
            self.animation015.setDuration(500)
            self.animation015.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation015.setEndValue(QRect(7, 140, 216, 0))  # end value is the new menu width
            self.animation015.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation015.start()
            self.animation15 = QPropertyAnimation(self.Button_GR1_1, b"geometry")  # Animate minimumWidht
            self.animation15.setDuration(500)
            self.animation15.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation15.setEndValue(QRect(7, 140, 216, 0))  # end value is the new menu width
            self.animation15.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation15.start()
            self.animation16 = QPropertyAnimation(self.Button_GR1_2, b"geometry")  # Animate minimumWidht
            self.animation16.setDuration(500)
            self.animation16.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation16.setEndValue(QRect(7, 140, 216, 0))  # end value is the new menu width
            self.animation16.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation16.start()
            self.animation17 = QPropertyAnimation(self.Button_GR1_3, b"geometry")  # Animate minimumWidht
            self.animation17.setDuration(500)
            self.animation17.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation17.setEndValue(QRect(7, 140, 216, 0))  # end value is the new menu width
            self.animation17.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation17.start()
            self.animation18 = QPropertyAnimation(self.Button_GR2, b"geometry")  # Animate minimumWidht
            self.animation18.setDuration(500)
            self.animation18.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation18.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation18.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation18.start()
            self.animation19 = QPropertyAnimation(self.Button_GR3, b"geometry")  # Animate minimumWidht
            self.animation19.setDuration(500)
            self.animation19.setStartValue(QRect(7, 240, 216, 44))  # Start value is the current menu width
            self.animation19.setEndValue(QRect(7, 240, 216, 44))  # end value is the new menu width
            self.animation19.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation19.start()
            self.animation20 = QPropertyAnimation(self.Button_GR2_1, b"geometry")  # Animate minimumWidht
            self.animation20.setDuration(500)
            self.animation20.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation20.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation20.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation20.start()
            self.animation21 = QPropertyAnimation(self.Button_GR2_2, b"geometry")  # Animate minimumWidht
            self.animation21.setDuration(500)
            self.animation21.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation21.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation21.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation21.start()
            self.animation22 = QPropertyAnimation(self.Button_GR2_3, b"geometry")  # Animate minimumWidht
            self.animation22.setDuration(500)
            self.animation22.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation22.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation22.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation22.start()
            self.animation23 = QPropertyAnimation(self.Button_GR3_1, b"geometry")  # Animate minimumWidht
            self.animation23.setDuration(500)
            self.animation23.setStartValue(QRect(7, 290, 216, 44))  # Start value is the current menu width
            self.animation23.setEndValue(QRect(7, 240, 216, 44))  # end value is the new menu width
            self.animation23.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation23.start()
            self.animation24 = QPropertyAnimation(self.Button_GR3_2, b"geometry")  # Animate minimumWidht
            self.animation24.setDuration(500)
            self.animation24.setStartValue(QRect(7, 340, 216, 44))  # Start value is the current menu width
            self.animation24.setEndValue(QRect(7, 240, 216, 44))  # end value is the new menu width
            self.animation24.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation24.start()
            self.animation25 = QPropertyAnimation(self.Button_GR3_3, b"geometry")  # Animate minimumWidht
            self.animation25.setDuration(500)
            self.animation25.setStartValue(QRect(7, 390, 216, 44))  # Start value is the current menu width
            self.animation25.setEndValue(QRect(7, 240, 216, 44))  # end value is the new menu width
            self.animation25.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation25.start()

    ########################################################################
    # Slide left menu
    ########################################################################
    def slideLeftMenu(self):
        # Get current left menu width
        width = self.left_menu.width()

        # Animate the transition
        if width == 60:
            self.animation = QPropertyAnimation(self.left_menu, b"geometry")  # Animate minimumWidht
            self.animation.setDuration(500)
            self.animation.setStartValue(QRect(2120, 158, 60, 800))  # Start value is the current menu width
            self.animation.setEndValue(QRect(2120, 158, 231, 800))  # end value is the new menu width
            self.animation.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation.start()
            self.animation2 = QPropertyAnimation(self.Left_minimizer, b"geometry")  # Animate minimumWidht
            self.animation2.setDuration(500)
            self.animation2.setStartValue(QRect(0, 0, 60, 60))  # Start value is the current menu width
            self.animation2.setEndValue(QRect(0, 0, 230, 130))  # end value is the new menu width
            self.animation2.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation2.start()
            self.animation3 = QPropertyAnimation(self.Left_minimizer, b"iconSize")  # Animate minimumWidht
            self.animation3.setDuration(500)
            self.animation3.setStartValue(QSize(50, 50))  # Start value is the current menu width
            self.animation3.setEndValue(QSize(120, 150))  # end value is the new menu width
            self.animation3.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation3.start()
            self.animation4 = QPropertyAnimation(self.tableWidget, b"geometry")  # Animate minimumWidht
            self.animation4.setDuration(500)
            self.animation4.setStartValue(QRect(60, 60, 1484, 741))  # Start value is the current menu width
            self.animation4.setEndValue(QRect(230, 60, 1484, 741))  # end value is the new menu width
            self.animation4.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation4.start()
            self.animation5 = QPropertyAnimation(self.frame, b"geometry")  # Animate minimumWidht
            self.animation5.setDuration(500)
            self.animation5.setStartValue(QRect(30, 0, 1520, 811))  # Start value is the current menu width
            self.animation5.setEndValue(QRect(30, 0, 1690, 811))  # end value is the new menu width
            self.animation5.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation5.start()
            self.animation6 = QPropertyAnimation(self.Ispolnitel, b"geometry")  # Animate minimumWidht
            self.animation6.setDuration(500)
            self.animation6.setStartValue(QRect(50, 15, 431, 34))  # Start value is the current menu width
            self.animation6.setEndValue(QRect(220, 15, 431, 34))  # end value is the new menu width
            self.animation6.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation6.start()
            self.animation7 = QPropertyAnimation(self.FilterGU, b"geometry")  # Animate minimumWidht
            self.animation7.setDuration(500)
            self.animation7.setStartValue(QRect(490, 15, 271, 34))  # Start value is the current menu width
            self.animation7.setEndValue(QRect(660, 15, 271, 34))  # end value is the new menu width
            self.animation7.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation7.start()
            self.animation8 = QPropertyAnimation(self.FindButton, b"geometry")  # Animate minimumWidht
            self.animation8.setDuration(500)
            self.animation8.setStartValue(QRect(780, 15, 141, 34))  # Start value is the current menu width
            self.animation8.setEndValue(QRect(950, 15, 141, 34))  # end value is the new menu width
            self.animation8.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation8.start()
            self.animation9 = QPropertyAnimation(self.Refresher, b"geometry")  # Animate minimumWidht
            self.animation9.setDuration(500)
            self.animation9.setStartValue(QRect(1310, 5, 131, 51))  # Start value is the current menu width
            self.animation9.setEndValue(QRect(1480, 5, 131, 51))  # end value is the new menu width
            self.animation9.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation9.start()
            self.animation10 = QPropertyAnimation(self.Mini, b"geometry")  # Animate minimumWidht
            self.animation10.setDuration(500)
            self.animation10.setStartValue(QRect(1480, 0, 31, 31))  # Start value is the current menu width
            self.animation10.setEndValue(QRect(1650, 0, 31, 31))  # end value is the new menu width
            self.animation10.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation10.start()
            self.animation11 = QPropertyAnimation(self.Closeer, b"geometry")  # Animate minimumWidht
            self.animation11.setDuration(500)
            self.animation11.setStartValue(QRect(1510, 0, 31, 31))  # Start value is the current menu width
            self.animation11.setEndValue(QRect(1680, 0, 31, 31))  # end value is the new menu width
            self.animation11.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation11.start()
            self.animation15 = QPropertyAnimation(self.frame_2, b"geometry")  # Animate minimumWidht
            self.animation15.setDuration(500)
            self.animation15.setStartValue(QRect(930, 0, 371, 61))  # Start value is the current menu width
            self.animation15.setEndValue(QRect(1100, 0, 371, 61))  # end value is the new menu width
            self.animation15.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation15.start()

            self.animation12 = QPropertyAnimation(self.Button_GR1, b"geometry")  # Animate minimumWidht
            self.animation12.setDuration(500)
            self.animation12.setStartValue(QRect(7, 70, 44, 44))  # Start value is the current menu width
            self.animation12.setEndValue(QRect(7, 140, 216, 44))  # end value is the new menu width
            self.animation12.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation12.start()
            self.animation13 = QPropertyAnimation(self.Button_GR2, b"geometry")  # Animate minimumWidht
            self.animation13.setDuration(500)
            self.animation13.setStartValue(QRect(7, 120, 44, 44))  # Start value is the current menu width
            self.animation13.setEndValue(QRect(7, 190, 216, 44))  # end value is the new menu width
            self.animation13.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation13.start()
            self.animation14 = QPropertyAnimation(self.Button_GR3, b"geometry")  # Animate minimumWidht
            self.animation14.setDuration(500)
            self.animation14.setStartValue(QRect(7, 170, 44, 44))  # Start value is the current menu width
            self.animation14.setEndValue(QRect(7, 240, 216, 44))  # end value is the new menu width
            self.animation14.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation14.start()
            self.Button_GR3.setText("3 Группа")
            self.Button_GR1.setText("1 Группа")
            self.Button_GR2.setText("2 Группа")
            self.Button_GR1_0.hide()
            self.Button_GR1_1.hide()
            self.Button_GR1_2.hide()
            self.Button_GR1_3.hide()
            self.Button_GR2_1.hide()
            self.Button_GR2_2.hide()
            self.Button_GR2_3.hide()
            self.Button_GR3_1.hide()
            self.Button_GR3_2.hide()
            self.Button_GR3_3.hide()

        else:
            self.animation = QPropertyAnimation(self.left_menu, b"geometry")  # Animate minimumWidht
            self.animation.setDuration(500)
            self.animation.setStartValue(QRect(2120, 158, 231, 800))  # Start value is the current menu width
            self.animation.setEndValue(QRect(2120, 158, 60, 800))  # end value is the new menu width
            self.animation.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation.start()
            self.animation2 = QPropertyAnimation(self.Left_minimizer, b"geometry")  # Animate minimumWidht
            self.animation2.setDuration(500)
            self.animation2.setStartValue(QRect(0, 0, 230, 130))  # Start value is the current menu width
            self.animation2.setEndValue(QRect(0, 0, 60, 60))  # end value is the new menu width
            self.animation2.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation2.start()
            self.animation3 = QPropertyAnimation(self.Left_minimizer, b"iconSize")  # Animate minimumWidht
            self.animation3.setDuration(500)
            self.animation3.setStartValue(QSize(120, 150))  # Start value is the current menu width
            self.animation3.setEndValue(QSize(50, 50))  # end value is the new menu width
            self.animation3.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation3.start()
            self.animation4 = QPropertyAnimation(self.tableWidget, b"geometry")  # Animate minimumWidht
            self.animation4.setDuration(500)
            self.animation4.setStartValue(QRect(230, 60, 1484, 741))  # Start value is the current menu width
            self.animation4.setEndValue(QRect(60, 60, 1484, 741))  # end value is the new menu width
            self.animation4.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation4.start()
            self.animation5 = QPropertyAnimation(self.frame, b"geometry")  # Animate minimumWidht
            self.animation5.setDuration(500)
            self.animation5.setStartValue(QRect(30, 0, 1690, 811))  # Start value is the current menu width
            self.animation5.setEndValue(QRect(30, 0, 1520, 811))  # end value is the new menu width
            self.animation5.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation5.start()
            self.animation6 = QPropertyAnimation(self.Ispolnitel, b"geometry")  # Animate minimumWidht
            self.animation6.setDuration(500)
            self.animation6.setStartValue(QRect(220, 15, 431, 34))  # Start value is the current menu width
            self.animation6.setEndValue(QRect(50, 15, 431, 34))  # end value is the new menu width
            self.animation6.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation6.start()
            self.animation7 = QPropertyAnimation(self.FilterGU, b"geometry")  # Animate minimumWidht
            self.animation7.setDuration(500)
            self.animation7.setStartValue(QRect(660, 15, 271, 34))  # Start value is the current menu width
            self.animation7.setEndValue(QRect(490, 15, 271, 34))  # end value is the new menu width
            self.animation7.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation7.start()
            self.animation8 = QPropertyAnimation(self.FindButton, b"geometry")  # Animate minimumWidht
            self.animation8.setDuration(500)
            self.animation8.setStartValue(QRect(950, 15, 141, 34))  # Start value is the current menu width
            self.animation8.setEndValue(QRect(780, 15, 141, 34))  # end value is the new menu width
            self.animation8.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation8.start()
            self.animation9 = QPropertyAnimation(self.Refresher, b"geometry")  # Animate minimumWidht
            self.animation9.setDuration(500)
            self.animation9.setStartValue(QRect(1480, 5, 131, 51))  # Start value is the current menu width
            self.animation9.setEndValue(QRect(1310, 5, 131, 51))  # end value is the new menu width
            self.animation9.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation9.start()
            self.animation10 = QPropertyAnimation(self.Mini, b"geometry")  # Animate minimumWidht
            self.animation10.setDuration(500)
            self.animation10.setStartValue(QRect(1650, 0, 31, 31))  # Start value is the current menu width
            self.animation10.setEndValue(QRect(1480, 0, 31, 31))  # end value is the new menu width
            self.animation10.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation10.start()
            self.animation11 = QPropertyAnimation(self.Closeer, b"geometry")  # Animate minimumWidht
            self.animation11.setDuration(500)
            self.animation11.setStartValue(QRect(1680, 0, 31, 31))  # Start value is the current menu width
            self.animation11.setEndValue(QRect(1510, 0, 31, 31))  # end value is the new menu width
            self.animation11.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation11.start()
            self.animation15 = QPropertyAnimation(self.frame_2, b"geometry")  # Animate minimumWidht
            self.animation15.setDuration(500)
            self.animation15.setStartValue(QRect(1100, 0, 371, 61))  # Start value is the current menu width
            self.animation15.setEndValue(QRect(930, 0, 371, 61))  # end value is the new menu width
            self.animation15.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation15.start()

            self.animation12 = QPropertyAnimation(self.Button_GR1, b"geometry")  # Animate minimumWidht
            self.animation12.setDuration(500)
            self.animation12.setStartValue(QRect(7, 140, 216, 44))  # Start value is the current menu width
            self.animation12.setEndValue(QRect(7, 70, 44, 44))  # end value is the new menu width
            self.animation12.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation12.start()
            self.animation13 = QPropertyAnimation(self.Button_GR2, b"geometry")  # Animate minimumWidht
            self.animation13.setDuration(500)
            self.animation13.setStartValue(QRect(7, 190, 216, 44))  # Start value is the current menu width
            self.animation13.setEndValue(QRect(7, 120, 44, 44))  # end value is the new menu width
            self.animation13.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation13.start()
            self.animation14 = QPropertyAnimation(self.Button_GR3, b"geometry")  # Animate minimumWidht
            self.animation14.setDuration(500)
            self.animation14.setStartValue(QRect(7, 240, 216, 44))  # Start value is the current menu width
            self.animation14.setEndValue(QRect(7, 170, 44, 44))  # end value is the new menu width
            self.animation14.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
            self.animation14.start()
            self.Button_GR1.setText("1")
            self.Button_GR2.setText("2")
            self.Button_GR3.setText("3")
            self.Button_GR1_0.hide()
            self.Button_GR1_1.hide()
            self.Button_GR1_2.hide()
            self.Button_GR1_3.hide()
            self.Button_GR2_1.hide()
            self.Button_GR2_2.hide()
            self.Button_GR2_3.hide()
            self.Button_GR3_1.hide()
            self.Button_GR3_2.hide()
            self.Button_GR3_3.hide()

    # /////////////////////////////////////////////////////////////////////////

    def mousePressEvent(self, event):
        self.old_pos = event.screenPos()

    def mouseMoveEvent(self, event):
        if self.bgwidget.clicked:
            dx = self.old_pos.x() - event.screenPos().x()
            dy = self.old_pos.y() - event.screenPos().y()
            self.move(self.pos().x() - dx, self.pos().y() - dy)
        self.old_pos = event.screenPos()
        self.bgwidget.clicked = True



    def gotoGroup1GU(self):
        create = CreateGroup1GU()
        create.exec_()

    def loaddata(self):
        global loginUser
        global FullFioLogin
        global TipDop
        global UPR
        connection = sqlite3.connect("SourceGitHub/DB/shop_data.db")
        log = connection.cursor()
        log.execute("SELECT * FROM Login WHERE username = ?", (loginUser,))
        WhoLogin = log.fetchall()
        FullFioLogin = WhoLogin[0][3] + " " + WhoLogin[0][2] + " " + WhoLogin[0][4]


        isp = connection.cursor()
        isp.execute("SELECT * FROM Login WHERE otdel = ?", (UPR, ))
        allIsp = isp.fetchall()
        isp.execute("SELECT COUNT(*) as count FROM Login WHERE otdel = ?", (UPR, ))
        HowMuchIsp = isp.fetchone()
        i = 0
        for i in range(HowMuchIsp[0]):
            FullIspList = allIsp[i][3] + " " + allIsp[i][2] + " " + allIsp[i][4]
            self.Ispolnitel.addItems([FullIspList])
        self.Ispolnitel.setCurrentText(FullFioLogin)
        # self.Ispolnitel.currentIndexChanged[str].connect(self.ChangedIsp)
        # self.Ispolnitel.currentIndexChanged[str].connect(self.DoTable)
        # self.FilterGU.currentIndexChanged[str].connect(self.SearchGU)  # создание таблицы при изменения поиска

        self.FindButton.clicked.connect(self.SearchGU)  # создание таблицы при изменения поиска
        self.Refresher.clicked.connect(self.SearchGU)  # создание таблицы при изменения поиска

        TipDop = int(TipDop)
        if TipDop >= 5:
            self.Ispolnitel.setEnabled(True)


    def TakeNomerGU(self, nomerGU):
        global globalNomerGU
        globalNomerGU = nomerGU

    def TakePortalKuStatus(self, pkud):
        global PortalKuStatus777
        PortalKuStatus777 = pkud

    def ChangedIsp(self):
        global FullFioLogin
        global UseTableGU
        global curPage
        curPage = 1
        FullFioLogin = self.Ispolnitel.currentText()
        self.FilterGU.clear()
        if UseTableGU == "1 Группа":
            self.DoTable()
        elif UseTableGU == "2 Группа":
            self.DoTable2G()
        elif UseTableGU == "3 Группа":
            self.DoTable3G()
        elif UseTableGU == "Старый портал":
            self.DoPortalKu()


    def Loading2(self):
        self.movie = QMovie("SourceGitHub/Pictures/Gear.gif")
        self.loading.setMovie(self.movie)
        self.movie.start()

    def SearchGU(self):
        # Делаем поиск по номеру ГУ
        global FindNumGU
        global FullFioLogin
        global curPage

        if self.FilterGU.currentText() == "":
            FindNumGU = "*"
        elif self.FilterGU.currentText() != "":
            FindNumGU = self.FilterGU.currentText()

        if FullFioLogin != self.Ispolnitel.currentText() and self.FilterGU.itemText(0) != "Все ГУ из портала К":
            self.FilterGU.clear()
            FullFioLogin = self.Ispolnitel.currentText()
            #FindNumGU = "*"

        if UseTableGU == "1 Группа":
            curPage = 1
            self.DoTable()
        elif UseTableGU == "2 Группа":
            curPage = 1
            self.DoTable2G()
        elif UseTableGU == "3 Группа":
            curPage = 1
            self.DoTable3G()
        elif UseTableGU == "Старый портал":
            curPage = 1
            self.DoPortalKu()
        elif UseTableGU == "Завершенные":
            curPage = 1
            self.DoTableEndGU()
        elif UseTableGU == "Завершенные 2гр":
            curPage = 1
            self.DoTableEndGU2GR()

        # _____________________________________________________

    def DoTable(self):
        try:
            global z, NomerGUforDeadline
            global FindNumGU
            global Deadline
            global DeadlineDay
            global FullFioLogin
            global TipDop
            global UPR
            global UseTableGU
            if FullFioLogin != "Приостановка" and FullFioLogin != "Group1End" and FullFioLogin != "NULL":
                FullFioLogin = self.Ispolnitel.currentText()
            if self.FilterGU.currentText() == "" and FindNumGU != "":
                FindNumGU = FindNumGU
            if UseTableGU != "1 Группа":
                FindNumGU = "*"
                UseTableGU = "1 Группа"
            self.tableWidget.clear()
            self.tableWidget.setSortingEnabled(True)
            self.tableWidget.setColumnCount(15)
            self.tableWidget.setColumnWidth(0, 70)
            self.tableWidget.setColumnWidth(1, 150)
            self.tableWidget.setColumnWidth(2, 50)
            self.tableWidget.setColumnWidth(3, 70)
            self.tableWidget.setColumnWidth(4, 70)
            self.tableWidget.setColumnWidth(5, 130)
            self.tableWidget.setColumnWidth(6, 120)
            self.tableWidget.setColumnWidth(7, 156)
            self.tableWidget.setColumnWidth(8, 75)
            self.tableWidget.setColumnWidth(9, 120)
            self.tableWidget.setColumnWidth(10, 140)
            self.tableWidget.setColumnWidth(11, 70)
            self.tableWidget.setColumnWidth(12, 44)
            self.tableWidget.setColumnWidth(13, 95)
            self.tableWidget.setColumnWidth(14, 81)
            self.tableWidget.setHorizontalHeaderLabels(
                ["", "Комментарий", "Тип ГУ", "Номер ГУ", "Дата ГУ", "Номер в РАЛ", "Исполнитель", "Наименование",
                 "Д. в. эксп.", "Эксперт", "Статус ГУ", "Дедлайн", "Дней", "Дата Решения", "Договор"])

            self.FilterStatus.hide()
            self.Back.hide()
            self.Next.hide()
            self.Long_Pages.hide()
            self.tableWidget.show()
            self.Welcome_image.hide()

            self.tableWidget.setRowCount(0)
            self.tableWidget.setSortingEnabled(False)
            self.FilterGU.clear()
            self.FilterGU.addItems([""])
            connection = sqlite3.connect('SourceGitHub/DB/DinamicBaseNew.db')
            #connection.execute('PRAGMA synchronous = OFF')
            #connection.execute('PRAGMA auto_vacuum = FULL')

            cur = connection.cursor()
            # PricazGU1 = ""
            if FindNumGU == "*":
                if FullFioLogin == "Выбрать исполнителя":
                    if UPR == 'Управление аккредитации' or UPR == 'Все Управления':
                        cur.execute(
                            "SELECT COUNT(0) FROM GU1Group WHERE Ответственныйисполнитель is NOT NULL AND "
                            "СтатусГУ != 'Возврат без рассмотрения' AND СтатусГУ != 'Отказ ГУ (договор)' AND СтатусГУ != 'Отзыв ГУ' AND СтатусГУ != 'Приказ об отказе' AND ДатаВозврата is NULL AND Приостановка is NULL AND (Договор1Г = '' or Договор1Г is NULL)")  # AND ЗаявлениенаотзывГУ is NULL Запрос в sql на кол-ва ячеек в базе данных ||AND (ДатаприказаосоставеЭГ is NULL OR ДатаприказаосоставеЭГ = 'Не зарегистрирован')||
                    elif UPR != 'Управление аккредитации' or UPR != 'Все Управления':
                        cur.execute(
                            "SELECT COUNT(0) FROM GU1Group WHERE Ответственныйисполнитель is NOT NULL AND "
                            "СтатусГУ != 'Возврат без рассмотрения' AND СтатусГУ != 'Отказ ГУ (договор)' AND СтатусГУ != 'Отзыв ГУ' AND СтатусГУ != 'Приказ об отказе' AND ДатаВозврата is NULL AND Приостановка is NULL AND ОтделУП = ? AND (Договор1Г = '' or Договор1Г is NULL)", (UPR, ))  # AND ЗаявлениенаотзывГУ is NULL Запрос в sql на кол-ва ячеек в базе данных ||AND (ДатаприказаосоставеЭГ is NULL OR ДатаприказаосоставеЭГ = 'Не зарегистрирован')||
                elif FullFioLogin == "NULL":
                    cur.execute(
                        "SELECT COUNT(0) FROM GU1Group WHERE Ответственныйисполнитель is NULL AND "
                        "СтатусГУ != 'Возврат без рассмотрения' AND СтатусГУ != 'Отказ ГУ (договор)' AND СтатусГУ != 'Отзыв ГУ' AND СтатусГУ != 'Приказ об отказе' AND ДатаВозврата is NULL AND Приостановка is NULL AND (Договор1Г = '' or Договор1Г is NULL)")  # AND ЗаявлениенаотзывГУ is NULL
                elif FullFioLogin == "Group1End":
                    cur.execute(
                        "SELECT COUNT(0) FROM GU1Group WHERE Актвыезднойэкспертизы is NOT NULL OR СтатусГУ == 'Возврат без рассмотрения' OR ДатаВозврата is NOT NULL AND (Договор1Г = '' or Договор1Г is NULL)")
                elif FullFioLogin == "Шкабура Владимир Владимирович" or FullFioLogin == "Бухарова Анастасия Владимировна":
                    cur.execute(
                        "SELECT COUNT(0) FROM GU1Group WHERE (Ответственныйисполнитель = ? AND "
                        " СтатусГУ != 'Возврат без рассмотрения' AND СтатусГУ != 'Отказ ГУ (договор)' AND СтатусГУ != 'Отзыв ГУ' AND СтатусГУ != 'Приказ об отказе') OR (ВозвратПроверка = ? AND ВозвратРешение = 'На проверке') AND (Договор1Г = '' or Договор1Г is NULL)", (FullFioLogin, FullFioLogin,))  # Запрос в sql на кол-ва ячеек в базе данных || AND (ДатаприказаосоставеЭГ is NULL OR ДатаприказаосоставеЭГ = 'Не зарегистрирован')||
                elif FullFioLogin == "Приостановка":
                    if UPR == 'Управление аккредитации' or UPR == 'Все Управления':
                        cur.execute("SELECT COUNT(0) FROM GU1Group WHERE Приостановка = 'Да'")
                    else:
                        cur.execute("SELECT COUNT(0) FROM GU1Group WHERE (Приостановка = 'Да' AND ОтделУП = ?) AND (Договор1Г = '' or Договор1Г is NULL)", (UPR, ))
                else:
                    cur.execute(
                        "SELECT COUNT(0) FROM GU1Group WHERE (Ответственныйисполнитель = ? AND "
                        "СтатусГУ != 'Возврат без рассмотрения' AND СтатусГУ != 'Отказ ГУ (договор)' AND "
                        "СтатусГУ != 'Отзыв ГУ' AND СтатусГУ != 'Приказ об отказе' AND ДатаВозврата is NULL AND Приостановка is NULL AND ТипГУ != 'АК' AND ТипГУ != 'РОА' AND (Договор1Г = '' or Договор1Г is NULL)) or "
                        "(Ответственныйисполнитель = ? AND "
                        "СтатусГУ != 'Возврат без рассмотрения' AND СтатусГУ != 'Отказ ГУ (договор)' AND "
                        "СтатусГУ != 'Отзыв ГУ' AND СтатусГУ != 'Приказ об отказе' AND ДатаВозврата is NULL AND Приостановка is NULL AND (ТипГУ = 'АК' OR ТипГУ = 'РОА') AND "
                        "(ДатаприказаосоставеЭГ = '' OR ДатаприказаосоставеЭГ is NULL OR ДатаприказаосоставеЭГ = 'Не зарегистрирован'))", (FullFioLogin, FullFioLogin))
 # Запрос в sql на кол-ва ячеек в базе данных || AND (ДатаприказаосоставеЭГ is NULL OR ДатаприказаосоставеЭГ = 'Не зарегистрирован')||
            elif FindNumGU != "*":
                cur.execute('SELECT COUNT(0) FROM GU1Group WHERE Регистрационныйномер = ? AND (Договор1Г = "" OR Договор1Г is NULL)',
                            (FindNumGU,))  # Запрос в sql на кол-ва ячеек в базе данных

            sqllong = cur.fetchone()  # Находим длинну базы данных в кортеже
            sqllong2 = sqllong[0]  # Вытаскиваем длину баззы данных из первого картежа
            if FindNumGU == "*":
                # sqlstr = 'SELECT * FROM GU1Group LIMIT (SELECT COUNT(0) FROM GU1Group)'  # Запрос в sql сколько запихнуть в созданную таблицу
                if FullFioLogin == "Выбрать исполнителя":
                    if UPR == 'Управление аккредитации' or UPR == 'Все Управления':
                        results = cur.execute(
                            "SELECT * FROM GU1Group WHERE Ответственныйисполнитель is NOT NULL AND "
                            "СтатусГУ != 'Возврат без рассмотрения' AND СтатусГУ != 'Отказ ГУ (договор)' AND СтатусГУ != 'Отзыв ГУ' AND СтатусГУ != 'Приказ об отказе' AND ДатаВозврата is NULL AND Приостановка is NULL "
                            " AND (Договор1Г = '' or Договор1Г is NULL)")  # AND ЗаявлениенаотзывГУ is NULL Запрос в sql на кол-ва ячеек в базе данных ||AND (ДатаприказаосоставеЭГ is NULL OR ДатаприказаосоставеЭГ = 'Не зарегистрирован')||
                    else:
                        results = cur.execute(
                            "SELECT * FROM GU1Group WHERE Ответственныйисполнитель is NOT NULL AND "
                            "СтатусГУ != 'Возврат без рассмотрения' AND СтатусГУ != 'Отказ ГУ (договор)' AND СтатусГУ != 'Отзыв ГУ' AND СтатусГУ != 'Приказ об отказе' AND ДатаВозврата is NULL AND Приостановка is NULL "
                            "AND ОтделУП = ? AND (Договор1Г = '' or Договор1Г is NULL)", (UPR,))  # AND ЗаявлениенаотзывГУ is NULL Запрос в sql на кол-ва ячеек в базе данных ||AND (ДатаприказаосоставеЭГ is NULL OR ДатаприказаосоставеЭГ = 'Не зарегистрирован')||
                elif FullFioLogin == "NULL":
                    results = cur.execute(
                        "SELECT * FROM GU1Group WHERE Ответственныйисполнитель is NULL AND "
                        " СтатусГУ != 'Возврат без рассмотрения' AND СтатусГУ != 'Отказ ГУ (договор)' AND СтатусГУ != 'Отзыв ГУ' AND СтатусГУ != 'Приказ об отказе' AND ДатаВозврата is NULL AND Приостановка is NULL AND (Договор1Г = '' or Договор1Г is NULL)")  # Запрос в sql сколько запихнуть в созданную таблицу || AND (ДатаприказаосоставеЭГ is NULL OR ДатаприказаосоставеЭГ = 'Не зарегистрирован')||
                elif FullFioLogin == "Group1End":
                    results = cur.execute(
                        "SELECT * FROM GU1Group WHERE (Актвыезднойэкспертизы is NOT NULL OR СтатусГУ == 'Возврат без рассмотрения' OR ДатаВозврата is NOT NULL) AND (Договор1Г = '' or Договор1Г is NULL)")  # Запрос в sql сколько запихнуть в созданную таблицу || AND (ДатаприказаосоставеЭГ is NULL OR ДатаприказаосоставеЭГ = 'Не зарегистрирован')||
                elif FullFioLogin == "Шкабура Владимир Владимирович" or FullFioLogin == "Бухарова Анастасия Владимировна":
                    results = cur.execute(
                        "SELECT * FROM GU1Group WHERE (Ответственныйисполнитель = ? AND "
                        " СтатусГУ != 'Возврат без рассмотрения' AND СтатусГУ != 'Отказ ГУ (договор)' AND СтатусГУ != 'Отзыв ГУ' AND СтатусГУ != 'Приказ об отказе') OR (ВозвратПроверка = ? AND ВозвратРешение = 'На проверке') AND (Договор1Г = '' or Договор1Г is NULL)",
                        (FullFioLogin,FullFioLogin))  # Запрос в sql сколько запихнуть в созданную таблицу || AND (ДатаприказаосоставеЭГ is NULL OR ДатаприказаосоставеЭГ = 'Не зарегистрирован')||
                elif FullFioLogin == "Приостановка":
                    if UPR == 'Управление аккредитации' or UPR == 'Все Управления':
                        results = cur.execute("SELECT * FROM GU1Group WHERE Приостановка = 'Да' AND (Договор1Г = '' or Договор1Г is NULL)")
                    else:
                        results = cur.execute("SELECT * FROM GU1Group WHERE (Приостановка = 'Да' AND ОтделУП = ?) AND (Договор1Г = '' or Договор1Г is NULL)", (UPR, ))
                else:
                    results = cur.execute(
                        "SELECT * FROM GU1Group WHERE (Ответственныйисполнитель = ? AND "
                        "СтатусГУ != 'Возврат без рассмотрения' AND СтатусГУ != 'Отказ ГУ (договор)' AND "
                        "СтатусГУ != 'Отзыв ГУ' AND СтатусГУ != 'Приказ об отказе' AND ДатаВозврата is NULL AND Приостановка is NULL AND ТипГУ != 'АК' AND ТипГУ != 'РОА' AND (Договор1Г = '' or Договор1Г is NULL)) or "
                        "(Ответственныйисполнитель = ? AND "
                        "СтатусГУ != 'Возврат без рассмотрения' AND СтатусГУ != 'Отказ ГУ (договор)' AND "
                        "СтатусГУ != 'Отзыв ГУ' AND СтатусГУ != 'Приказ об отказе' AND ДатаВозврата is NULL AND Приостановка is NULL AND (ТипГУ = 'АК' OR ТипГУ = 'РОА') AND "
                        "(ДатаприказаосоставеЭГ = '' OR ДатаприказаосоставеЭГ is NULL OR ДатаприказаосоставеЭГ = 'Не зарегистрирован'))", (FullFioLogin, FullFioLogin))  # Запрос в sql сколько запихнуть в созданную таблицу || AND (ДатаприказаосоставеЭГ is NULL OR ДатаприказаосоставеЭГ = 'Не зарегистрирован')||
                tablerow = 0
            elif FindNumGU != "*":
                results = cur.execute('SELECT * FROM GU1Group WHERE Регистрационныйномер = ? AND (Договор1Г = "" or Договор1Г is NULL)', (FindNumGU,))  # Запрос в sql сколько запихнуть в созданную таблицу
                tablerow = 0

            self.tableWidget.setRowCount(sqllong2)  # sqllong2 = все ГУ
            for row in results:
                #if QtWidgets.QTableWidgetItem(row[7]).text() == FullFioLogin or FullFioLogin == "Выбрать исполнителя" or FullFioLogin == "NULL" or row[86] == FullFioLogin or FullFioLogin == "Group1End" or FullFioLogin == "Приостановка":
                if True == True:
                    if QtWidgets.QTableWidgetItem(row[0]).text() == "Подтверждение компетентности":
                        tipGU = "ПК"
                        self.tableWidget.setItem(tablerow, 2, QtWidgets.QTableWidgetItem(tipGU))
                    elif QtWidgets.QTableWidgetItem(row[0]).text() == "Аккредитация":
                        tipGU = "АК"
                        self.tableWidget.setItem(tablerow, 2, QtWidgets.QTableWidgetItem(tipGU))
                    elif QtWidgets.QTableWidgetItem(row[0]).text() == "Расширение области аккредитации":
                        tipGU = "РОА"
                        self.tableWidget.setItem(tablerow, 2, QtWidgets.QTableWidgetItem(tipGU))
                    else:
                        self.tableWidget.setItem(tablerow, 2,
                                                 QtWidgets.QTableWidgetItem(QtWidgets.QTableWidgetItem(row[0]).text()))

                    self.tableWidget.setItem(tablerow, 3, QtWidgets.QTableWidgetItem(row[1]))
                    self.tableWidget.setItem(tablerow, 4, QtWidgets.QTableWidgetItem(row[2]))
                    if FullFioLogin != "NULL" or FullFioLogin != [] or FullFioLogin != "":
                        ShortFIO = QtWidgets.QTableWidgetItem(row[7]).text()
                        ShortFIO = ShortFIO.split()
                        if ShortFIO != []:
                            ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
                        else:
                            ShortFIO = ""
                    else:
                        ShortFIO = ""
                    self.tableWidget.setItem(tablerow, 1, QtWidgets.QTableWidgetItem(row[77]))
                    self.tableWidget.setItem(tablerow, 6, QtWidgets.QTableWidgetItem(ShortFIO))
                    self.tableWidget.setItem(tablerow, 5, QtWidgets.QTableWidgetItem(row[3]))
                    self.tableWidget.setItem(tablerow, 8, QtWidgets.QTableWidgetItem(row[18]))
                    self.tableWidget.setItem(tablerow, 7, QtWidgets.QTableWidgetItem(row[4]))
                    # self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem("------------"))
                    qwers = QtWidgets.QTableWidgetItem(row[84]).text()
                    asdasd = QtWidgets.QTableWidgetItem(row[26]).text()
                    asdasdw = QtWidgets.QTableWidgetItem(row[32]).text()
                    if (QtWidgets.QTableWidgetItem(row[84]).text() is None or QtWidgets.QTableWidgetItem(row[84]).text()  == "") and (QtWidgets.QTableWidgetItem(row[26]).text() is not None and QtWidgets.QTableWidgetItem(row[26]).text() != ""):
                        self.tableWidget.setItem(tablerow, 14, QtWidgets.QTableWidgetItem("Уже зашло ЭЗ"))
                        self.tableWidget.item(tablerow, 14).setForeground(QColor("#ff0000"))
                    elif (QtWidgets.QTableWidgetItem(row[84]).text() is None or QtWidgets.QTableWidgetItem(row[84]).text()  == "") and (QtWidgets.QTableWidgetItem(row[32]).text() is not None and QtWidgets.QTableWidgetItem(row[32]).text() != ""):
                        self.tableWidget.setItem(tablerow, 14, QtWidgets.QTableWidgetItem("Уже зашел Акт"))
                        self.tableWidget.item(tablerow, 14).setForeground(QColor("#ff0000"))
                    else:
                        self.tableWidget.setItem(tablerow, 14, QtWidgets.QTableWidgetItem(row[84]))

                    if QtWidgets.QTableWidgetItem(row[20]).text() != "":
                        ShortEXP = QtWidgets.QTableWidgetItem(row[20]).text()
                        ShortEXP = ShortEXP.split()
                        ShortEXP = ShortEXP[0] + " " + ShortEXP[1][0:1] + "." + ShortEXP[2][0:1] + "."
                    else:
                        ShortEXP = ""

                    self.tableWidget.setItem(tablerow, 9, QtWidgets.QTableWidgetItem(ShortEXP))
                    pole10 = QtWidgets.QTableWidgetItem(row[10]).text()
                    self.tableWidget.setItem(tablerow, 10, QtWidgets.QTableWidgetItem(pole10))



                    if QtWidgets.QTableWidgetItem(row[10]).text() == "В работе" or QtWidgets.QTableWidgetItem(
                            row[10]).text() == "Не рассмотрено" and FullFioLogin != "Group1End":
                        StatusGUFirst = "Формирование ЭГ"
                    else:
                        StatusGUFirst = QtWidgets.QTableWidgetItem(row[10]).text()
                    self.tableWidget.setItem(tablerow, 10, QtWidgets.QTableWidgetItem(StatusGUFirst))

                    if FindNumGU == "*":
                        self.FilterGU.addItems([QtWidgets.QTableWidgetItem(row[1]).text()])  # Добавить ГУ в поиск

                    cal = Russia()
                    Perevibor = QtWidgets.QTableWidgetItem(row[78]).text()
                    if Perevibor == "Нет" or Perevibor is None or Perevibor == "":
                        Perevibor = 0
                    elif Perevibor == "Да":
                        Perevibor = 1
                    else:
                        Perevibor = int(Perevibor)
                    Data_Dead = QtWidgets.QTableWidgetItem(row[2]).text()
                    Data_EG = QtWidgets.QTableWidgetItem(row[80]).text()
                    Data_Dead = Data_Dead[:10]
                    Data_EG = Data_EG[:10]

                    if QtWidgets.QTableWidgetItem(row[24]).text() == "":
                        Data_ISP_date = datetime.strptime(Data_Dead, '%d.%m.%Y')
                        datta = cal.add_working_days(Data_ISP_date, 9 + (Perevibor * 7))
                        datta = datta.strftime("%d.%m.%Y")
                        datta10 = str(datta)
                        self.tableWidget.setItem(tablerow, 11, QtWidgets.QTableWidgetItem(datta10))
                        self.tableWidget.setItem(tablerow, 10, QtWidgets.QTableWidgetItem("Формирование ЭГ"))
                        GU1Status = 1

                        datenow = QDate.currentDate().toPyDate()  # сегодня
                        dattaFormEG = cal.get_working_days_delta(Data_ISP_date, datenow) + 1  # считает даты в рабочем
                        datta11 = str(dattaFormEG)
                        self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem(datta11))
                        if Perevibor == 1:
                            self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem(str(Perevibor) + " Перевыбор"))
                        elif Perevibor > 1 and Perevibor < 5:
                            self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem(str(Perevibor) + " Перевыбора"))
                        elif Perevibor == 5:
                            self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem(str(Perevibor) + " Перевыборов"))
                        if dattaFormEG < (9 + (Perevibor * 7)):
                            self.tableWidget.item(tablerow, 11).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                            self.tableWidget.item(tablerow, 12).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                        else:
                            self.tableWidget.item(tablerow, 11).setForeground(QColor("#ff0000"))  # цвет текста!!!!!!!!!!!
                            self.tableWidget.item(tablerow, 12).setForeground(QColor("#ff0000"))  # цвет текста!!!!!!!!!!!

                    if QtWidgets.QTableWidgetItem(row[24]).text() != "":
                        Data_ISP_date = datetime.strptime(Data_Dead, '%d.%m.%Y')
                        datta = cal.add_working_days(Data_ISP_date, 9 + (Perevibor * 7))
                        datta = datta.strftime("%d.%m.%Y")
                        datta10 = str(datta)
                        self.tableWidget.setItem(tablerow, 11, QtWidgets.QTableWidgetItem(datta10))
                        self.tableWidget.setItem(tablerow, 10, QtWidgets.QTableWidgetItem("Предложение зашло"))
                        GU1Status = 2

                        datenow = QDate.currentDate().toPyDate()  # сегодня
                        dattaFormEG = cal.get_working_days_delta(Data_ISP_date, datenow) + 1  # считает даты в рабочем
                        datta11 = str(dattaFormEG)
                        self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem(datta11))
                        if Perevibor == 1:
                            self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem(str(Perevibor) + " Перевыбор"))
                        elif Perevibor > 1 and Perevibor < 5:
                            self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem(str(Perevibor) + " Перевыбора"))
                        elif Perevibor == 5:
                            self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem(str(Perevibor) + " Перевыборов"))
                        if dattaFormEG < (9 + (Perevibor * 7)):
                            self.tableWidget.item(tablerow, 11).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                            self.tableWidget.item(tablerow, 12).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                        else:
                            self.tableWidget.item(tablerow, 11).setForeground(QColor("#ff0000"))  # цвет текста!!!!!!!!!!!
                            self.tableWidget.item(tablerow, 12).setForeground(QColor("#ff0000"))  # цвет текста!!!!!!!!!!!

                    if QtWidgets.QTableWidgetItem(row[80]).text() != "":
                        Data_EGS = QtWidgets.QTableWidgetItem(row[80]).text()
                        Data_EGS = datetime.strptime(Data_EGS, '%d.%m.%Y')
                        datta2 = cal.add_working_days(Data_EGS, 1)
                        datta21 = datta2.strftime("%d.%m.%Y")
                        datta10 = str(datta21)
                        self.tableWidget.setItem(tablerow, 11, QtWidgets.QTableWidgetItem(datta10))
                        self.tableWidget.setItem(tablerow, 10, QtWidgets.QTableWidgetItem("Ждем 2 дня по ЭГ"))
                        GU1Status = 3

                        datenow = QDate.currentDate().toPyDate()  # сегодня
                        delta = cal.get_working_days_delta(datta2, datenow)  # считает даты в рабочем
                        delta2str = str(delta)

                    if Data_EG != "":
                        cal = Russia()
                        Data_EG_1 = datetime.strptime(Data_EG, '%d.%m.%Y')
                        datenow = QDate.currentDate().toPyDate()  # сегодня
                        datta = cal.get_working_days_delta(Data_EG_1, datenow) + 1  # считает даты в рабочем
                        datta11 = str(datta)
                        self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem(datta11))
                        if datta < 2:
                            self.tableWidget.item(tablerow, 11).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                            self.tableWidget.item(tablerow, 12).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                        else:
                            self.tableWidget.item(tablerow, 11).setForeground(QColor("#ff0000"))  # цвет текстаыы!!!!!!!!!!!
                            self.tableWidget.item(tablerow, 12).setForeground(QColor("#ff0000"))  # цвет текста!!!!!!!!!!!

                    if GU1Status == 3 and datta >= 3:
                        self.tableWidget.setItem(tablerow, 10, QtWidgets.QTableWidgetItem("Можно запускать приказ"))
                        GU1Status = 4

                        datta10 = datetime.strptime(datta10, '%d.%m.%Y')
                        datta = cal.get_working_days_delta(datta10, datenow) + 1
                        datta10data = cal.add_working_days(datta10, 2)
                        datta10data2 = datta10data.strftime("%d.%m.%Y")
                        datta10 = str(datta10data2)
                        self.tableWidget.setItem(tablerow, 11, QtWidgets.QTableWidgetItem(datta10))

                        datta11 = str(datta)
                        self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem(datta11))
                        if datta < 3:
                            self.tableWidget.item(tablerow, 11).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                            self.tableWidget.item(tablerow, 12).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                        else:
                            self.tableWidget.item(tablerow, 11).setForeground(QColor("#ff0000"))  # цвет текста!!!!!!!!!!!
                            self.tableWidget.item(tablerow, 12).setForeground(QColor("#ff0000"))  # цвет текста!!!!!!!!!!!

                    Data_Pricaza = QtWidgets.QTableWidgetItem(row[25]).text().split()
                    if  Data_Pricaza == ['Не', 'зарегистрирован']:
                        self.tableWidget.setItem(tablerow, 10, QtWidgets.QTableWidgetItem("Приказ на регистрации"))
                        self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem(""))

                    if Data_Pricaza != [] and Data_Pricaza != ['Не', 'зарегистрирован']:
                        try:
                            Data_Pricaza = Data_Pricaza[2].replace("г.", "")
                        except:
                            Data_Pricaza = Data_Pricaza[0]
                        self.tableWidget.setItem(tablerow, 10, QtWidgets.QTableWidgetItem("Ждем договор 3 дня"))
                        self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem(Data_Pricaza))
                        Data_Pricaza = datetime.strptime(Data_Pricaza, '%d.%m.%Y')
                        datta = cal.get_working_days_delta(Data_Pricaza, datenow) + 1
                        datta10data = cal.add_working_days(Data_Pricaza, 2)
                        datta10data2 = datta10data.strftime("%d.%m.%Y")
                        datta10 = str(datta10data2)
                        datta11 = str(datta)
                        self.tableWidget.setItem(tablerow, 11, QtWidgets.QTableWidgetItem(datta10))
                        self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem(datta11))
                        GU1Status = 6
                    if Data_Pricaza != [] and int(datta11) > 3 and GU1Status == 6:
                        self.tableWidget.setItem(tablerow, 10, QtWidgets.QTableWidgetItem("Ждем договор 20 дней"))
                        #Data_Pricaza = datetime.strptime(Data_Pricaza, '%d.%m.%Y')
                        datta = cal.get_working_days_delta(Data_Pricaza, datenow) + 1
                        datta10data = cal.add_working_days(Data_Pricaza, 19)
                        datta10data2 = datta10data.strftime("%d.%m.%Y")
                        datta10 = str(datta10data2)
                        datta11 = str(datta)
                        self.tableWidget.setItem(tablerow, 11, QtWidgets.QTableWidgetItem(datta10))
                        self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem(datta11))


                    if row[81] != datta10 and row[82] != datta11:
                        curTo1g = connection.cursor()
                        curTo1g.execute('UPDATE GU1Group SET DeadLine = ? WHERE Регистрационныйномер = ?',
                                        (datta10, row[1]))
                        curTo1g.execute('UPDATE GU1Group SET DeadLineDay = ? WHERE Регистрационныйномер = ?',
                                        (datta11, row[1]))
                        connection.commit()

                    if FullFioLogin == "Group1End":
                        self.tableWidget.setItem(tablerow, 10, QtWidgets.QTableWidgetItem(pole10))
                        self.tableWidget.setItem(tablerow, 11, QtWidgets.QTableWidgetItem(""))
                        self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem(""))


                    tablerow += 1

            for i in range(sqllong2):  # sqllong2 = все ГУ
                # z = self.tableWidget.item(i, 2).text()
                pushButton = QtWidgets.QPushButton("Открыть")
                # pushButton.setStyleSheet("border-image: url('YSourceGitHub/Pictures/migrate.png')") # Картинка в кнопке

                pushButton.clicked.connect(lambda ch, nomerGU=self.tableWidget.item(i, 3): self.TakeNomerGU(nomerGU.text()))
                pushButton.clicked.connect(lambda ch, btn=pushButton: self.gotoGroup1GU())
                self.tableWidget.setCellWidget(i, 0, pushButton)
            self.tableWidget.setSortingEnabled(True)
        except BufferError:
            connection.close()
            self.DoTable()


    def DoTableEndGU(self):
        global FindNumGU
        global FullFioLogin
        global TipDop
        global sqllong3
        global curPage
        global OtPage
        global DoPage
        global sqllongPages
        global UseTableGU
        global UPR
        cal = Russia()
        UseTableGU = "Завершенные"
        self.tableWidget.clear()
        self.tableWidget.setSortingEnabled(True)
        maxPage = 30
        try:
            if curPage > 1:
                curPage = curPage
            else:
                curPage = 1
        except:
            curPage = 1

        self.FilterStatus.show()
        self.Back.show()
        self.Next.show()
        self.Long_Pages.show()
        self.tableWidget.show()
        self.Welcome_image.hide()
        self.tableWidget.setColumnCount(14)
        self.tableWidget.setColumnWidth(0, 70)
        self.tableWidget.setColumnWidth(1, 150)
        self.tableWidget.setColumnWidth(2, 50)
        self.tableWidget.setColumnWidth(3, 70)
        self.tableWidget.setColumnWidth(4, 70)
        self.tableWidget.setColumnWidth(5, 120)
        self.tableWidget.setColumnWidth(6, 120)
        self.tableWidget.setColumnWidth(7, 120)
        self.tableWidget.setColumnWidth(8, 120)
        self.tableWidget.setColumnWidth(9, 80)
        self.tableWidget.setColumnWidth(10, 80)
        self.tableWidget.setColumnWidth(11, 80)
        self.tableWidget.setColumnWidth(12, 80)
        self.tableWidget.setColumnWidth(13, 80)
        self.tableWidget.setColumnWidth(14, 80)
        self.tableWidget.setHorizontalHeaderLabels(
            ["", "Комментарий 3", "Тип ГУ", "Номер ГУ", "Дата ГУ", "РАЛ", "Исполнитель 3гр",
             "Наименование", "Эксперт", "Тер", "Приказ 1гр.", "Приказ 2гр.", "Приказ 3гр.", "Итог Вывод", "Итог Дней"])

        if self.FilterGU.currentText() == "Все ГУ из портала К":
            Filter = ""
        else:
            Filter = self.FilterGU.currentText()
        if Filter != '':
            maxPage = 1
            curPage = 1
        if self.FilterGU.itemText(0) != "Все ГУ из портала К":
            self.FilterGU.clear()
        TakeStatusGU = self.FilterStatus.currentText()

        connection = sqlite3.connect("SourceGitHub/DB/DinamicBaseNew.db")
        kud = connection.cursor()
        kudall = connection.cursor()
        FioTake = self.Ispolnitel.currentText()

        results = kud.execute("SELECT * FROM GU1Group WHERE (ДатаПриказа3гр != '' AND ДатаПриказа3гр is not NULL) AND ОтделУП = ?",
            (UPR,))
        tablerow = 0
        self.FilterGU.clear()
        self.FilterGU.addItems([""])
        for row in results:
            self.FilterGU.addItems([QtWidgets.QTableWidgetItem(row[1]).text()])  # Добавить ГУ в поиск
            tablerow += 1

        if TipDop == 10 and FioTake == "Выбрать исполнителя":
            kud.execute("SELECT COUNT(0) FROM GU1Group WHERE (ДатаПриказа3гр != '' AND ДатаПриказа3гр is not NULL) AND ОтделУП = ?", (UPR, ))
        else:
            kud.execute("SELECT COUNT(0) FROM GU1Group WHERE Исполнитель3Гр = ? AND (ДатаПриказа3гр != '' AND ДатаПриказа3гр is not NULL) AND ОтделУП = ?",
                        (FioTake, UPR,))
        if Filter != '':
            kud.execute("SELECT COUNT(0) FROM GU1Group WHERE Регистрационныйномер = ? AND (ДатаПриказа3гр != '' AND ДатаПриказа3гр is not NULL) AND ОтделУП = ?", (Filter, UPR,))

        sqllong = kud.fetchone()  # Находим длинну базы данных в кортеже
        sqllong2 = sqllong[0]  # Вытаскиваем длину баззы данных из первого картежа
        if TipDop == 10 and FioTake == "Выбрать исполнителя":
            results = kud.execute("SELECT * FROM GU1Group WHERE (ДатаПриказа3гр != '' AND ДатаПриказа3гр is not NULL) AND ОтделУП = ?", (UPR, ))
        else:
            results = kud.execute("SELECT * FROM GU1Group WHERE Исполнитель3Гр = ? AND (ДатаПриказа3гр != '' AND ДатаПриказа3гр is not NULL) AND ОтделУП = ?",
                                  (FioTake, UPR,))
        if Filter != '':
            results = kud.execute("SELECT * FROM GU1Group WHERE Регистрационныйномер = ? AND (ДатаПриказа3гр != '' AND ДатаПриказа3гр is not NULL) AND ОтделУП = ?",
                                  (Filter, UPR,))
        sqllong3 = int(sqllong2)
        self.SqlLongDetectNew()
        if sqllong3 < maxPage:
            maxPage = sqllong3
        elif curPage == int(sqllongPages):
            maxPage = sqllong3 - maxPage * (int(sqllongPages) - 1)
        if sqllong3 == 1:
            maxPage = 1

        self.tableWidget.setRowCount(maxPage)

        if TipDop == 10 and FioTake == "Выбрать исполнителя":
            results2 = kud.execute(
                "SELECT * FROM (SELECT *, ROW_NUMBER() OVER (ORDER BY Датарегистрации DESC) AS row FROM GU1Group WHERE (ДатаПриказа3гр != '' AND ДатаПриказа3гр is not NULL) AND ОтделУП = ?) a WHERE row > ? AND row <= ?",
                (UPR, OtPage, DoPage))
        else:
            results2 = kud.execute(
                "SELECT * FROM (SELECT *, ROW_NUMBER() OVER (ORDER BY Датарегистрации DESC) AS row FROM GU1Group WHERE Исполнитель3Гр = ? AND (ДатаПриказа3гр != '' AND ДатаПриказа3гр is not NULL) AND ОтделУП = ?) a WHERE row > ? AND row <= ?",
                (FioTake, UPR, OtPage, DoPage))
        if Filter != '':
            results2 = kud.execute(
                "SELECT * FROM (SELECT *, ROW_NUMBER() OVER (ORDER BY Датарегистрации DESC) AS row FROM GU1Group WHERE Регистрационныйномер = ? AND (ДатаПриказа3гр != '' AND ДатаПриказа3гр is not NULL) AND ОтделУП = ?) a WHERE row > ? AND row <= ?",
                (Filter, UPR, OtPage, DoPage))
        tablerow = 0

        for row in results2:
            self.tableWidget.setItem(tablerow, 1, QtWidgets.QTableWidgetItem(row[107]))
            self.tableWidget.setItem(tablerow, 2, QtWidgets.QTableWidgetItem(row[0]))
            self.tableWidget.setItem(tablerow, 3, QtWidgets.QTableWidgetItem(row[1]))
            self.tableWidget.setItem(tablerow, 4, QtWidgets.QTableWidgetItem(row[2]))
            self.tableWidget.setItem(tablerow, 5, QtWidgets.QTableWidgetItem(row[3]))
            ShortFIO = QtWidgets.QTableWidgetItem(row[108]).text()
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.tableWidget.setItem(tablerow, 6, QtWidgets.QTableWidgetItem(ShortFIO))
            self.tableWidget.setItem(tablerow, 7, QtWidgets.QTableWidgetItem(row[4]))
            if QtWidgets.QTableWidgetItem(row[20]).text() != "":
                ShortEXP = QtWidgets.QTableWidgetItem(row[20]).text()
                ShortEXP = ShortEXP.split()
                ShortEXP = ShortEXP[0] + " " + ShortEXP[1][0:1] + "." + ShortEXP[2][0:1] + "."
            else:
                ShortEXP = ""
            self.tableWidget.setItem(tablerow, 8, QtWidgets.QTableWidgetItem(ShortEXP))
            TU = QtWidgets.QTableWidgetItem(row[91]).text()
            if TU == "Центральный федеральный округ":
                TU = "ЦФО"
            elif TU == "Приволжский федеральный округ":
                TU = "ПФО"
            elif TU == "Сибирский федеральный округ":
                TU = "СФО"
            elif TU == "Северо-Западный федеральный округ":
                TU = "СЗФО"
            elif TU == "Дальневосточный федеральный округ":
                TU = "ДФО"
            elif TU == "Уральский федеральный округ":
                TU = "УФО"
            elif TU == "Южный и Северо-Кавказский федеральные округа":
                TU = "ЮСФО"
            elif TU == "Выбрать территориальное Управление":
                TU = "Не выбрано"
            self.tableWidget.setItem(tablerow, 9, QtWidgets.QTableWidgetItem(TU))
            Data_Pricaza = QtWidgets.QTableWidgetItem(row[25]).text().split()
            try:
                Data_Pricaza = Data_Pricaza[2].replace("г.", "")
            except:
                Data_Pricaza = Data_Pricaza[0]
            self.tableWidget.setItem(tablerow, 10, QtWidgets.QTableWidgetItem(Data_Pricaza))
            self.tableWidget.setItem(tablerow, 11, QtWidgets.QTableWidgetItem(row[100]))
            self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem(row[101]))
            self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem(""))
            DataGU = row[2]
            DataPricaz3GR = row[101]
            Data_DataGU = datetime.strptime(DataGU, '%d.%m.%Y')
            Data_DataPricaz3GR = datetime.strptime(DataPricaz3GR, '%d.%m.%Y')
            datta = cal.get_working_days_delta(Data_DataGU, Data_DataPricaz3GR) + 1  # считает даты в рабочем
            datta13 = str(datta)
            self.tableWidget.setItem(tablerow, 14, QtWidgets.QTableWidgetItem(datta13))

            tablerow += 1


        for i in range(maxPage):  # sqllong2 = все ГУ
            # z = self.tableWidget.item(i, 2).text()
            pushButton = QtWidgets.QPushButton("Открыть")
            # pushButton.setStyleSheet("border-image: url('YSourceGitHub/Pictures/migrate.png')") # Картинка в кнопке

            pushButton.clicked.connect(lambda ch, nomerGU=self.tableWidget.item(i, 3): self.TakeNomerGU(nomerGU.text()))
            pushButton.clicked.connect(lambda ch, btn=pushButton: self.gotoGroup1GU())
            self.tableWidget.setCellWidget(i, 0, pushButton)

    def DoTableEndGU2GR(self):
        global FindNumGU
        global FullFioLogin
        global TipDop
        global sqllong3
        global curPage
        global OtPage
        global DoPage
        global sqllongPages
        global UseTableGU
        global UPR
        cal = Russia()
        UseTableGU = "Завершенные 2гр"
        self.tableWidget.clear()
        self.tableWidget.setSortingEnabled(True)
        maxPage = 30
        try:
            if curPage > 1:
                curPage = curPage
            else:
                curPage = 1
        except:
            curPage = 1

        self.FilterStatus.show()
        self.Back.show()
        self.Next.show()
        self.Long_Pages.show()
        self.tableWidget.show()
        self.Welcome_image.hide()
        self.tableWidget.setColumnCount(14)
        self.tableWidget.setColumnWidth(0, 70)
        self.tableWidget.setColumnWidth(1, 150)
        self.tableWidget.setColumnWidth(2, 50)
        self.tableWidget.setColumnWidth(3, 70)
        self.tableWidget.setColumnWidth(4, 70)
        self.tableWidget.setColumnWidth(5, 120)
        self.tableWidget.setColumnWidth(6, 120)
        self.tableWidget.setColumnWidth(7, 120)
        self.tableWidget.setColumnWidth(8, 120)
        self.tableWidget.setColumnWidth(9, 80)
        self.tableWidget.setColumnWidth(10, 80)
        self.tableWidget.setColumnWidth(11, 80)
        self.tableWidget.setColumnWidth(12, 100)
        self.tableWidget.setColumnWidth(13, 80)
        self.tableWidget.setHorizontalHeaderLabels(
            ["", "Комментарий 2гр", "Тип ГУ", "Номер ГУ", "Дата ГУ", "РАЛ", "Исполнитель 2гр",
             "Наименование", "Эксперт", "Тер", "Приказ 1гр.", "Приказ 2гр.", "Итог Вывод", "Итог Дней"])

        if self.FilterGU.currentText() == "Все ГУ из портала К":
            Filter = ""
        else:
            Filter = self.FilterGU.currentText()
        if Filter != '':
            maxPage = 1
            curPage = 1
        if self.FilterGU.itemText(0) != "Все ГУ из портала К":
            self.FilterGU.clear()
        TakeStatusGU = self.FilterStatus.currentText()

        connection = sqlite3.connect("SourceGitHub/DB/DinamicBaseNew.db")
        kud = connection.cursor()
        kudall = connection.cursor()
        FioTake = self.Ispolnitel.currentText()

        results = kud.execute("SELECT * FROM GU1Group WHERE (ДатаПриказа2гр != '' AND ДатаПриказа2гр is not NULL) AND ОтделУП = ?",
            (UPR,))
        tablerow = 0
        self.FilterGU.clear()
        self.FilterGU.addItems([""])
        for row in results:
            self.FilterGU.addItems([QtWidgets.QTableWidgetItem(row[1]).text()])  # Добавить ГУ в поиск
            tablerow += 1

        if TipDop == 10 and FioTake == "Выбрать исполнителя":
            kud.execute("SELECT COUNT(0) FROM GU1Group WHERE (ДатаПриказа2гр != '' AND ДатаПриказа2гр is not NULL) AND ОтделУП = ?", (UPR, ))
        else:
            kud.execute("SELECT COUNT(0) FROM GU1Group WHERE Исполнитель3Гр = ? AND (ДатаПриказа2гр != '' AND ДатаПриказа2гр is not NULL) AND ОтделУП = ?",
                        (FioTake, UPR,))
        if Filter != '':
            kud.execute("SELECT COUNT(0) FROM GU1Group WHERE Регистрационныйномер = ? AND (ДатаПриказа2гр != '' AND ДатаПриказа2гр is not NULL) AND ОтделУП = ?", (Filter, UPR,))

        sqllong = kud.fetchone()  # Находим длинну базы данных в кортеже
        sqllong2 = sqllong[0]  # Вытаскиваем длину баззы данных из первого картежа
        if TipDop == 10 and FioTake == "Выбрать исполнителя":
            results = kud.execute("SELECT * FROM GU1Group WHERE (ДатаПриказа2гр != '' AND ДатаПриказа2гр is not NULL) AND ОтделУП = ?", (UPR, ))
        else:
            results = kud.execute("SELECT * FROM GU1Group WHERE Исполнитель3Гр = ? AND (ДатаПриказа2гр != '' AND ДатаПриказа2гр is not NULL) AND ОтделУП = ?",
                                  (FioTake, UPR,))
        if Filter != '':
            results = kud.execute("SELECT * FROM GU1Group WHERE Регистрационныйномер = ? AND (ДатаПриказа2гр != '' AND ДатаПриказа2гр is not NULL) AND ОтделУП = ?",
                                  (Filter, UPR,))
        sqllong3 = int(sqllong2)
        self.SqlLongDetectNew()
        if sqllong3 < maxPage:
            maxPage = sqllong3
        elif curPage == int(sqllongPages):
            maxPage = sqllong3 - maxPage * (int(sqllongPages) - 1)
        if sqllong3 == 1:
            maxPage = 1

        self.tableWidget.setRowCount(maxPage)

        if TipDop == 10 and FioTake == "Выбрать исполнителя":
            results2 = kud.execute(
                "SELECT * FROM (SELECT *, ROW_NUMBER() OVER (ORDER BY Датарегистрации DESC) AS row FROM GU1Group WHERE (ДатаПриказа2гр != '' AND ДатаПриказа2гр is not NULL) AND ОтделУП = ?) a WHERE row > ? AND row <= ?",
                (UPR, OtPage, DoPage))
        else:
            results2 = kud.execute(
                "SELECT * FROM (SELECT *, ROW_NUMBER() OVER (ORDER BY Датарегистрации DESC) AS row FROM GU1Group WHERE Исполнитель3Гр = ? AND (ДатаПриказа2гр != '' AND ДатаПриказа2гр is not NULL) AND ОтделУП = ?) a WHERE row > ? AND row <= ?",
                (FioTake, UPR, OtPage, DoPage))
        if Filter != '':
            results2 = kud.execute(
                "SELECT * FROM (SELECT *, ROW_NUMBER() OVER (ORDER BY Датарегистрации DESC) AS row FROM GU1Group WHERE Регистрационныйномер = ? AND (ДатаПриказа2гр != '' AND ДатаПриказа2гр is not NULL) AND ОтделУП = ?) a WHERE row > ? AND row <= ?",
                (Filter, UPR, OtPage, DoPage))
        tablerow = 0

        for row in results2:
            self.tableWidget.setItem(tablerow, 1, QtWidgets.QTableWidgetItem(row[107]))
            self.tableWidget.setItem(tablerow, 2, QtWidgets.QTableWidgetItem(row[0]))
            self.tableWidget.setItem(tablerow, 3, QtWidgets.QTableWidgetItem(row[1]))
            self.tableWidget.setItem(tablerow, 4, QtWidgets.QTableWidgetItem(row[2]))
            self.tableWidget.setItem(tablerow, 5, QtWidgets.QTableWidgetItem(row[3]))
            if row[112] == "" or row[112] is None:
                ShortFIO = QtWidgets.QTableWidgetItem(row[7]).text()
            else:
                ShortFIO = QtWidgets.QTableWidgetItem(row[112]).text()
            ShortFIO = ShortFIO.split()
            if ShortFIO != []:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            else:
                ShortFIO = ""
            self.tableWidget.setItem(tablerow, 6, QtWidgets.QTableWidgetItem(ShortFIO))
            self.tableWidget.setItem(tablerow, 7, QtWidgets.QTableWidgetItem(row[4]))
            if QtWidgets.QTableWidgetItem(row[20]).text() != "":
                ShortEXP = QtWidgets.QTableWidgetItem(row[20]).text()
                ShortEXP = ShortEXP.split()
                ShortEXP = ShortEXP[0] + " " + ShortEXP[1][0:1] + "." + ShortEXP[2][0:1] + "."
            else:
                ShortEXP = ""
            self.tableWidget.setItem(tablerow, 8, QtWidgets.QTableWidgetItem(ShortEXP))
            TU = QtWidgets.QTableWidgetItem(row[91]).text()
            if TU == "Центральный федеральный округ":
                TU = "ЦФО"
            elif TU == "Приволжский федеральный округ":
                TU = "ПФО"
            elif TU == "Сибирский федеральный округ":
                TU = "СФО"
            elif TU == "Северо-Западный федеральный округ":
                TU = "СЗФО"
            elif TU == "Дальневосточный федеральный округ":
                TU = "ДФО"
            elif TU == "Уральский федеральный округ":
                TU = "УФО"
            elif TU == "Южный и Северо-Кавказский федеральные округа":
                TU = "ЮСФО"
            elif TU == "Выбрать территориальное Управление":
                TU = "Не выбрано"
            self.tableWidget.setItem(tablerow, 9, QtWidgets.QTableWidgetItem(TU))
            Data_Pricaza = QtWidgets.QTableWidgetItem(row[25]).text().split()
            try:
                Data_Pricaza = Data_Pricaza[2].replace("г.", "")
            except:
                try:
                    Data_Pricaza = Data_Pricaza[0]
                except:
                    Data_Pricaza = ""
            self.tableWidget.setItem(tablerow, 10, QtWidgets.QTableWidgetItem(Data_Pricaza))
            self.tableWidget.setItem(tablerow, 11, QtWidgets.QTableWidgetItem(row[100]))
            self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem(row[99]))
            DataGU = row[2]
            DataPricaz2GR = row[100]
            Data_DataGU = datetime.strptime(DataGU, '%d.%m.%Y')
            Data_DataPricaz2GR = datetime.strptime(DataPricaz2GR, '%d.%m.%Y')
            datta = cal.get_working_days_delta(Data_DataGU, Data_DataPricaz2GR) + 1  # считает даты в рабочем
            datta13 = str(datta)
            self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem(datta13))

            tablerow += 1


        for i in range(maxPage):  # sqllong2 = все ГУ
            # z = self.tableWidget.item(i, 2).text()
            pushButton = QtWidgets.QPushButton("Открыть")
            # pushButton.setStyleSheet("border-image: url('YSourceGitHub/Pictures/migrate.png')") # Картинка в кнопке

            pushButton.clicked.connect(lambda ch, nomerGU=self.tableWidget.item(i, 3): self.TakeNomerGU(nomerGU.text()))
            pushButton.clicked.connect(lambda ch, btn=pushButton: self.gotoGroup1GU())
            self.tableWidget.setCellWidget(i, 0, pushButton)



    def DoTable2G(self):
        global FindNumGU
        global FullFioLogin
        global TipDop
        global sqllong3
        global curPage
        global OtPage
        global DoPage
        global sqllongPages
        global UseTableGU
        global UPR
        UseTableGU = "2 Группа"
        self.tableWidget.clear()
        self.tableWidget.setSortingEnabled(True)
        maxPage = 30
        try:
            if curPage > 1:
                curPage = curPage
            else:
                curPage = 1
        except:
            curPage = 1

        self.FilterStatus.show()
        self.Back.show()
        self.Next.show()
        self.Long_Pages.show()
        self.tableWidget.show()
        self.Welcome_image.hide()
        self.tableWidget.setColumnCount(16)
        self.tableWidget.setColumnWidth(0, 70)
        self.tableWidget.setColumnWidth(1, 150)
        self.tableWidget.setColumnWidth(2, 50)
        self.tableWidget.setColumnWidth(3, 70)
        self.tableWidget.setColumnWidth(4, 70)
        self.tableWidget.setColumnWidth(5, 120)
        self.tableWidget.setColumnWidth(6, 150)
        self.tableWidget.setColumnWidth(7, 120)
        self.tableWidget.setColumnWidth(8, 70)
        self.tableWidget.setColumnWidth(9, 70)
        self.tableWidget.setColumnWidth(10, 70)
        self.tableWidget.setColumnWidth(11, 80)
        self.tableWidget.setColumnWidth(12, 150)
        self.tableWidget.setColumnWidth(13, 80)
        self.tableWidget.setColumnWidth(14, 40)
        self.tableWidget.setColumnWidth(15, 70)
        self.tableWidget.setHorizontalHeaderLabels(
            ["", "Комментарий", "Тип ГУ", "Номер ГУ", "Дата ГУ", "Исполнитель",
             "Наименование", "Эксперт", "Приказ 1гр.", "Договор", "Дата ЭЗ", "Вывод ЭЗ", "Статус ГУ", "Дедлайн", "Дней",  "Итог"])

        if self.FilterGU.currentText() == "Все ГУ из портала К":
            Filter = ""
        else:
            Filter = self.FilterGU.currentText()
        if Filter != '':
            maxPage = 1
            curPage = 1
        if self.FilterGU.itemText(0) != "Все ГУ из портала К":
            self.FilterGU.clear()
        TakeStatusGU = self.FilterStatus.currentText()

        connection = sqlite3.connect("SourceGitHub/DB/DinamicBaseNew.db")
        kud = connection.cursor()
        kudall = connection.cursor()
        FioTake = self.Ispolnitel.currentText()

        results = kud.execute("SELECT * FROM GU1Group WHERE (ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND (ТипГУ = 'РОА' OR ТипГУ = 'АК') AND (ДатаПриказа2гр = '' or ДатаПриказа2гр is NULL) AND ОтделУП = ?",
            (UPR,))
        tablerow = 0
        self.FilterGU.clear()
        self.FilterGU.addItems([""])
        for row in results:
            self.FilterGU.addItems([QtWidgets.QTableWidgetItem(row[1]).text()])  # Добавить ГУ в поиск
            tablerow += 1

        if TipDop == 10 and FioTake == "Выбрать исполнителя":
            kud.execute("SELECT COUNT(0) FROM GU1Group WHERE (СтатусГУ = 'В работе' or СтатусГУ = 'Приостановка') AND (ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND (ТипГУ = 'РОА' OR ТипГУ = 'АК') AND (ДатаПриказа2гр = '' or ДатаПриказа2гр is NULL) AND ОтделУП = ?", (UPR, ))
        else:
            kud.execute("SELECT COUNT(0) FROM GU1Group WHERE ((Исполнитель2Гр = ? AND (ДатаприказаосоставеЭГ != '' AND "
                        "ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND "
                        "(ТипГУ = 'РОА' OR ТипГУ = 'АК') AND (ДатаПриказа2гр = '' or ДатаПриказа2гр is NULL) AND "
                        "ОтделУП =  ?) AND СтатусГУ != 'Отзыв ГУ' AND СтатусГУ != 'Приказ об отказе') or (( "
                        "(ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ is not NULL AND "
                        "ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND (ТипГУ = 'РОА' OR ТипГУ = 'АК') AND "
                        "(ДатаПриказа2гр = '' or ДатаПриказа2гр is NULL) AND ОтделУП =  ?) AND "
                        "(СтатусГУ = 'Отзыв ГУ' AND ВозвратПроверка = ? AND ВозвратРешение = 'На проверке'))",
                        (FioTake, UPR, UPR, FioTake))
        if Filter != '':
            kud.execute("SELECT COUNT(0) FROM GU1Group WHERE Регистрационныйномер = ? AND (ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND (ТипГУ = 'РОА' OR ТипГУ = 'АК') AND (ДатаПриказа2гр = '' or ДатаПриказа2гр is NULL) AND ОтделУП = ?", (Filter, UPR,))

        sqllong = kud.fetchone()  # Находим длинну базы данных в кортеже
        sqllong2 = sqllong[0]  # Вытаскиваем длину баззы данных из первого картежа
        if TipDop == 10 and FioTake == "Выбрать исполнителя":
            results = kud.execute("SELECT * FROM GU1Group WHERE (СтатусГУ = 'В работе' or СтатусГУ = 'Приостановка') AND (ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND (ТипГУ = 'РОА' OR ТипГУ = 'АК') AND (ДатаПриказа2гр = '' or ДатаПриказа2гр is NULL) AND ОтделУП = ?", (UPR, ))
        else:
            results = kud.execute("SELECT * FROM GU1Group WHERE ((Исполнитель2Гр = ? AND (ДатаприказаосоставеЭГ != '' AND "
                                  "ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND "
                                  "(ТипГУ = 'РОА' OR ТипГУ = 'АК') AND (ДатаПриказа2гр = '' or ДатаПриказа2гр is NULL) AND "
                                  "ОтделУП =  ?) AND СтатусГУ != 'Отзыв ГУ' AND СтатусГУ != 'Приказ об отказе') or (( "
                                  "(ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ is not NULL AND "
                                  "ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND (ТипГУ = 'РОА' OR ТипГУ = 'АК') AND "
                                  "(ДатаПриказа2гр = '' or ДатаПриказа2гр is NULL) AND ОтделУП =  ?) AND "
                                  "(СтатусГУ = 'Отзыв ГУ' AND ВозвратПроверка = ? AND ВозвратРешение = 'На проверке'))",
                                  (FioTake, UPR, UPR, FioTake))
        if Filter != '':
            results = kud.execute("SELECT * FROM GU1Group WHERE Регистрационныйномер = ? AND (ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND (ТипГУ = 'РОА' OR ТипГУ = 'АК') AND (ДатаПриказа2гр = '' or ДатаПриказа2гр is NULL) AND ОтделУП = ?",
                                  (Filter, UPR,))
        sqllong3 = int(sqllong2)
        self.SqlLongDetectNew()
        if sqllong3 < maxPage:
            maxPage = sqllong3
        elif curPage == int(sqllongPages):
            maxPage = sqllong3 - maxPage * (int(sqllongPages) - 1)
        if sqllong3 == 1:
            maxPage = 1

        self.tableWidget.setRowCount(maxPage)

        if TipDop == 10 and FioTake == "Выбрать исполнителя":
            results2 = kud.execute(
                "SELECT * FROM (SELECT *, ROW_NUMBER() OVER (ORDER BY Датарегистрации DESC) AS row FROM GU1Group WHERE (СтатусГУ = 'В работе' or СтатусГУ = 'Приостановка') AND (ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND (ТипГУ = 'РОА' OR ТипГУ = 'АК') AND (ДатаПриказа2гр = '' OR ДатаПриказа2гр is NULL) AND ОтделУП = ?) a WHERE row > ? AND row <= ?",
                (UPR, OtPage, DoPage))
        else:
            results2 = kud.execute(
                "SELECT * FROM (SELECT *, ROW_NUMBER() OVER (ORDER BY Датарегистрации DESC) AS row FROM GU1Group WHERE ((Исполнитель2Гр = ? AND (ДатаприказаосоставеЭГ != '' AND "
                                  "ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND "
                                  "(ТипГУ = 'РОА' OR ТипГУ = 'АК') AND (ДатаПриказа2гр = '' or ДатаПриказа2гр is NULL) AND "
                                  "ОтделУП =  ?) AND СтатусГУ != 'Отзыв ГУ' AND СтатусГУ != 'Приказ об отказе') or (( "
                                  "(ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ is not NULL AND "
                                  "ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND (ТипГУ = 'РОА' OR ТипГУ = 'АК') AND "
                                  "(ДатаПриказа2гр = '' or ДатаПриказа2гр is NULL) AND ОтделУП =  ?) AND "
                                  "(СтатусГУ = 'Отзыв ГУ' AND ВозвратПроверка = ? AND ВозвратРешение = 'На проверке'))) a WHERE row > ? AND row <= ?",
                (FioTake, UPR, UPR, FioTake, OtPage, DoPage))
        if Filter != '':
            results2 = kud.execute(
                "SELECT * FROM (SELECT *, ROW_NUMBER() OVER (ORDER BY Датарегистрации DESC) AS row FROM GU1Group WHERE Регистрационныйномер = ? AND (ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND (ТипГУ = 'РОА' OR ТипГУ = 'АК') AND (ДатаПриказа2гр = '' OR ДатаПриказа2гр is NULL) AND ОтделУП = ?) a WHERE row > ? AND row <= ?",
                (Filter, UPR, OtPage, DoPage))
        tablerow = 0

        for row in results2:
            self.tableWidget.setItem(tablerow, 2, QtWidgets.QTableWidgetItem(row[0]))
            self.tableWidget.setItem(tablerow, 3, QtWidgets.QTableWidgetItem(row[1]))
            self.tableWidget.setItem(tablerow, 4, QtWidgets.QTableWidgetItem(row[2]))
            if FullFioLogin != "NULL" or FullFioLogin != [] or FullFioLogin != "":
                if row[112] == "" or row[112] is None:
                    ShortFIO = QtWidgets.QTableWidgetItem(row[7]).text()
                else:
                    ShortFIO = QtWidgets.QTableWidgetItem(row[112]).text()
                ShortFIO = ShortFIO.split()
                if ShortFIO != []:
                    ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
                else:
                    ShortFIO = ""
            else:
                ShortFIO = ""
            self.tableWidget.setItem(tablerow, 1, QtWidgets.QTableWidgetItem(row[113]))
            self.tableWidget.setItem(tablerow, 5, QtWidgets.QTableWidgetItem(ShortFIO))
            self.tableWidget.setItem(tablerow, 6, QtWidgets.QTableWidgetItem(row[4]))
            self.tableWidget.setItem(tablerow, 9, QtWidgets.QTableWidgetItem(row[84]))
            self.tableWidget.setItem(tablerow, 11, QtWidgets.QTableWidgetItem(row[98]))
            if QtWidgets.QTableWidgetItem(row[20]).text() != "":
                ShortEXP = QtWidgets.QTableWidgetItem(row[20]).text()
                ShortEXP = ShortEXP.split()
                ShortEXP = ShortEXP[0] + " " + ShortEXP[1][0:1] + "." + ShortEXP[2][0:1] + "."
            else:
                ShortEXP = ""
            self.tableWidget.setItem(tablerow, 7, QtWidgets.QTableWidgetItem(ShortEXP))

            TakeVivodEZNew = row[97]
            TakeVivodEZ1C = row[26]
            if TakeVivodEZNew == "" or TakeVivodEZNew == None:
                if TakeVivodEZ1C != "" and TakeVivodEZ1C != None:
                    TakeVivodEZ1C = TakeVivodEZ1C.split()
                    try:
                        TakeVivodEZ1C = TakeVivodEZ1C[2].replace("г.", "")
                        self.tableWidget.setItem(tablerow, 10, QtWidgets.QTableWidgetItem(TakeVivodEZ1C))
                        TrueDataEZ = TakeVivodEZ1C
                    except:
                        self.tableWidget.setItem(tablerow, 10, QtWidgets.QTableWidgetItem(""))
                        TrueDataEZ = ""
                else:
                    TrueDataEZ = ""
            else:
                self.tableWidget.setItem(tablerow, 10, QtWidgets.QTableWidgetItem(TakeVivodEZNew))
                TrueDataEZ = TakeVivodEZNew

            Data_Pricaza = QtWidgets.QTableWidgetItem(row[25]).text().split()
            try:
                Data_Pricaza = Data_Pricaza[2].replace("г.", "")
            except:
                Data_Pricaza = Data_Pricaza[0]
            self.tableWidget.setItem(tablerow, 8, QtWidgets.QTableWidgetItem(Data_Pricaza))

            DataPricaz2GRNew = row[100]
            DataPricaz2GR1Cbad = row[31]
            DataPricaz2GR1Cbad2 = row[44]
            DataPricaz2GR1Cgood = row[30]
            if DataPricaz2GRNew == "" or DataPricaz2GRNew == None:
                if DataPricaz2GR1Cgood != "" and DataPricaz2GR1Cgood != None:
                    DataPricaz2GR1Cgood = DataPricaz2GR1Cgood.split()
                    try:
                        DataPricaz2GR1Cgood = DataPricaz2GR1Cgood[2].replace("г.", "")
                        self.tableWidget.setItem(tablerow, 15, QtWidgets.QTableWidgetItem(DataPricaz2GR1Cgood))
                    except:
                        self.tableWidget.setItem(tablerow, 15, QtWidgets.QTableWidgetItem(""))
                elif DataPricaz2GR1Cbad != "" and DataPricaz2GR1Cbad != None:
                    DataPricaz2GR1Cbad = DataPricaz2GR1Cbad.split()
                    try:
                        DataPricaz2GR1Cbad = DataPricaz2GR1Cbad[2].replace("г.", "")
                        self.tableWidget.setItem(tablerow, 15, QtWidgets.QTableWidgetItem("Отказ\n" + DataPricaz2GR1Cbad))
                    except:
                        self.tableWidget.setItem(tablerow, 15, QtWidgets.QTableWidgetItem(""))
                elif DataPricaz2GR1Cbad2 != "" and DataPricaz2GR1Cbad2 != None:
                    DataPricaz2GR1Cbad2 = DataPricaz2GR1Cbad2.split()
                    try:
                        DataPricaz2GR1Cbad2 = DataPricaz2GR1Cbad2[2].replace("г.", "")
                        self.tableWidget.setItem(tablerow, 15, QtWidgets.QTableWidgetItem("Отказ\n" + DataPricaz2GR1Cbad2))
                    except:
                        self.tableWidget.setItem(tablerow, 15, QtWidgets.QTableWidgetItem(""))
            else:
                self.tableWidget.setItem(tablerow, 15, QtWidgets.QTableWidgetItem(DataPricaz2GRNew))

            cal = Russia()
            datenow = QDate.currentDate().toPyDate()  # сегодня
            Vivod_EZ = QtWidgets.QTableWidgetItem(row[98]).text()
            GU1Status = 0
            DataDogovor = QtWidgets.QTableWidgetItem(row[84]).text()


            if TrueDataEZ == "" or TrueDataEZ is None:
                if DataDogovor == "" or DataDogovor is None:
                    GU1Status = 0
                    self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem("Нет Даты Договора"))
                    self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem("Прошло"))
                    Data_Pricaz1GR_date = datetime.strptime(Data_Pricaza, '%d.%m.%Y')
                    datta = cal.get_working_days_delta(Data_Pricaz1GR_date, datenow) + 1  # считает даты в рабочем
                    datta13 = str(datta)
                    self.tableWidget.setItem(tablerow, 14, QtWidgets.QTableWidgetItem(datta13))
                else:
                    GU1Status = 1
                    self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem("Ждем ЭЗ"))
                    self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem("Прошло"))
                    Data_DataDogovor_date = datetime.strptime(DataDogovor, '%d.%m.%Y')
                    datta = cal.get_working_days_delta(Data_DataDogovor_date, datenow) + 1  # считает даты в рабочем
                    datta13 = str(datta)
                    self.tableWidget.setItem(tablerow, 14, QtWidgets.QTableWidgetItem(datta13))

            if TrueDataEZ != "" and TrueDataEZ is not None:
                if (Vivod_EZ == "" or Vivod_EZ is None):
                    GU1Status = 2
                    self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem("ЭЗ еще не рассмотрено"))
                elif Vivod_EZ == "Положительный":
                    GU1Status = 3
                    self.tableWidget.setItem(tablerow, 12,QtWidgets.QTableWidgetItem("ЭЗ рассмотрено, но ещё отработано"))
                elif Vivod_EZ == "Отрицательный":
                    GU1Status = 4
                    self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem("ЭЗ рассмотрено, но ещё отработано"))
                Data_EZ_date = datetime.strptime(TrueDataEZ, '%d.%m.%Y')
                datta = cal.add_working_days(Data_EZ_date, 5)
                datta2 = datta.strftime("%d.%m.%Y")
                datta12 = str(datta2)
                self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem(datta12))
                datta3 = cal.get_working_days_delta(Data_EZ_date, datenow) + 1  # считает даты в рабочем
                datta13 = str(datta3)
                self.tableWidget.setItem(tablerow, 14, QtWidgets.QTableWidgetItem(datta13))
                if datta3 < 5:
                    self.tableWidget.item(tablerow, 13).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                    self.tableWidget.item(tablerow, 14).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                else:
                    self.tableWidget.item(tablerow, 13).setForeground(QColor("#ff0000"))  # цвет текста!!!!!!!!!!!
                    self.tableWidget.item(tablerow, 14).setForeground(QColor("#ff0000"))  # цвет текста!!!!!!!!!!!

            tablerow += 1


        for i in range(maxPage):  # sqllong2 = все ГУ
            # z = self.tableWidget.item(i, 2).text()
            pushButton = QtWidgets.QPushButton("Открыть")
            # pushButton.setStyleSheet("border-image: url('YSourceGitHub/Pictures/migrate.png')") # Картинка в кнопке

            pushButton.clicked.connect(lambda ch, nomerGU=self.tableWidget.item(i, 3): self.TakeNomerGU(nomerGU.text()))
            pushButton.clicked.connect(lambda ch, btn=pushButton: self.gotoGroup1GU())
            self.tableWidget.setCellWidget(i, 0, pushButton)

    def DoTable3G(self):
        global FindNumGU
        global FullFioLogin
        global TipDop
        global sqllong3
        global curPage
        global OtPage
        global DoPage
        global sqllongPages
        global UseTableGU
        global UPR
        global Allchek
        global ALLUser
        global globalNomerGU
        global PortalKuStatus777
        UseTableGU = "3 Группа"
        self.tableWidget.clear()
        self.tableWidget.setSortingEnabled(True)
        maxPage = 30
        try:
            if curPage > 1:
                curPage = curPage
            else:
                curPage = 1
        except:
            curPage = 1

        self.FilterStatus.show()
        self.Back.show()
        self.Next.show()
        self.Long_Pages.show()
        self.tableWidget.show()
        self.Welcome_image.hide()
        self.tableWidget.setColumnCount(16)
        self.tableWidget.setColumnWidth(0, 70)
        self.tableWidget.setColumnWidth(1, 160)
        self.tableWidget.setColumnWidth(2, 50)
        self.tableWidget.setColumnWidth(3, 70)
        self.tableWidget.setColumnWidth(4, 70)
        self.tableWidget.setColumnWidth(5, 130)
        self.tableWidget.setColumnWidth(6, 120)
        self.tableWidget.setColumnWidth(7, 150)
        self.tableWidget.setColumnWidth(8, 120)
        self.tableWidget.setColumnWidth(9, 70)
        self.tableWidget.setColumnWidth(10, 70)
        self.tableWidget.setColumnWidth(11, 70)
        self.tableWidget.setColumnWidth(12, 130)
        self.tableWidget.setColumnWidth(13, 70)
        self.tableWidget.setColumnWidth(14, 45)
        self.tableWidget.setColumnWidth(15, 70)
        self.tableWidget.setHorizontalHeaderLabels(
            ["", "Комментарий", "Тип ГУ", "Номер ГУ", "Дата ГУ", "Номер в РАЛ", "Исполнитель 3гр",
             "Наименование", "Эксперт", "Дата Акта", "Вывод", "Риск", "Статус ГУ", "Дедлайн", "Дней", "Итог"])

        if self.FilterGU.currentText() == "Все ГУ из портала К":
            Filter = ""
        else:
            Filter = self.FilterGU.currentText()
        if Filter != '':
            maxPage = 1
            curPage = 1
        if self.FilterGU.itemText(0) != "Все ГУ из портала К":
            self.FilterGU.clear()
        TakeStatusGU = self.FilterStatus.currentText()

        connection = sqlite3.connect("SourceGitHub/DB/DinamicBaseNew.db")
        kud = connection.cursor()
        kudall = connection.cursor()
        FioTake = self.Ispolnitel.currentText()

        tablerow = 0
        self.FilterGU.clear()
        self.FilterGU.addItems([""])
        results = kud.execute(
            "SELECT * FROM GU1Group WHERE ((ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND ОтделУП = ? AND (ДатаПриказа3гр = '' OR ДатаПриказа3гр is NULL)) OR (ДатаПриказа3гр != '' AND ДатаПриказа3гр is not NULL AND ТУ = 'Центральный федеральный округ' AND СопроводИтог = '') OR (ПриостановкаАКилиРОА = 'Да')",
            (UPR,))
        for row in results:
            self.FilterGU.addItems([QtWidgets.QTableWidgetItem(row[1]).text()])  # Добавить ГУ в поиск
            tablerow += 1

        if TipDop == 10 and Allchek == "Нерасписанныая 3":
            kud.execute(
                "SELECT COUNT(0) FROM GU1Group WHERE ((((ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND (ТипГУ = 'РОА' OR ТипГУ = 'АК') AND (ДатаПриказа2гр != '' and ДатаПриказа2гр is NOT NULL) AND Итог2Гр != 'Отриц. отказ') or ((ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND (ТипГУ != 'РОА' AND ТипГУ != 'АК'))) AND (Исполнитель3Гр = '' OR Исполнитель3Гр is NULL) AND ОтделУП = ?) AND (ДатаПриказа3гр = '' OR ДатаПриказа3гр is NULL) AND (Договор1Г != '' AND Договор1Г is NOT NULL)",
                (UPR,))
        elif TipDop == 10 and FioTake == "Выбрать исполнителя" and Allchek == "Расписанная 3":
            kud.execute(
                "SELECT COUNT(0) FROM GU1Group WHERE ((Исполнитель3Гр != '' AND Исполнитель3Гр is not NULL) AND ОтделУП = ?) AND ((ДатаПриказа3гр = '' OR ДатаПриказа3гр is NULL) OR (ДатаПриказа3гр != '' AND ДатаПриказа3гр is not NULL AND ТУ = 'Центральный федеральный округ' AND СопроводИтог = '') OR (ПриостановкаАКилиРОА = 'Да'))",
                (UPR,))
        else:
            kud.execute(
                "SELECT COUNT(0) FROM GU1Group WHERE (Исполнитель3Гр = ? AND ОтделУП = ?) AND ((ДатаПриказа3гр = '' OR ДатаПриказа3гр is NULL) OR (ДатаПриказа3гр != '' AND ДатаПриказа3гр is not NULL AND ТУ = 'Центральный федеральный округ' AND СопроводИтог = '') OR (ПриостановкаАКилиРОА = 'Да'))",
                (FioTake, UPR,))
        if Filter != '':
            kud.execute(
                "SELECT COUNT(0) FROM GU1Group WHERE (Регистрационныйномер = ? AND ОтделУП = ?) AND ((ДатаПриказа3гр = '' OR ДатаПриказа3гр is NULL) OR (ДатаПриказа3гр != '' AND ДатаПриказа3гр is not NULL AND ТУ = 'Центральный федеральный округ' AND СопроводИтог = '') OR (ПриостановкаАКилиРОА = 'Да'))",
                (Filter, UPR,))

        sqllong = kud.fetchone()  # Находим длинну базы данных в кортеже
        sqllong2 = sqllong[0]  # Вытаскиваем длину баззы данных из первого картежа
        if TipDop == 10 and Allchek == "Нерасписанныая 3":
            results = kud.execute(
                "SELECT * FROM GU1Group WHERE ((((ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND (ТипГУ = 'РОА' OR ТипГУ = 'АК') AND (ДатаПриказа2гр != '' and ДатаПриказа2гр is NOT NULL) AND Итог2Гр != 'Отриц. отказ') or ((ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ != '' AND ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != 'Не зарегистрирован') AND (ТипГУ != 'РОА' AND ТипГУ != 'АК'))) AND (Исполнитель3Гр = '' OR Исполнитель3Гр is NULL) AND ОтделУП = ?) AND (ДатаПриказа3гр = '' OR ДатаПриказа3гр is NULL) AND (Договор1Г != '' AND Договор1Г is NOT NULL)",
                (UPR,))
        elif TipDop == 10 and FioTake == "Выбрать исполнителя" and Allchek == "Расписанная 3":
            results = kud.execute(
                "SELECT * FROM GU1Group WHERE ((Исполнитель3Гр != '' AND Исполнитель3Гр is not NULL) AND ОтделУП = ?) AND ((ДатаПриказа3гр = '' OR ДатаПриказа3гр is NULL) OR (ДатаПриказа3гр != '' AND ДатаПриказа3гр is not NULL AND ТУ = 'Центральный федеральный округ' AND СопроводИтог = '') OR (ПриостановкаАКилиРОА = 'Да'))",
                (UPR,))
        else:
            results = kud.execute(
                "SELECT * FROM GU1Group WHERE (Исполнитель3Гр = ? AND ОтделУП = ?) AND ((ДатаПриказа3гр = '' OR ДатаПриказа3гр is NULL) OR (ДатаПриказа3гр != '' AND ДатаПриказа3гр is not NULL AND ТУ = 'Центральный федеральный округ' AND СопроводИтог = '') OR (ПриостановкаАКилиРОА = 'Да'))",
                (FioTake, UPR, ))
        if Filter != '':
            results = kud.execute(
                "SELECT * FROM GU1Group WHERE (Регистрационныйномер = ? AND ОтделУП = ?) AND ((ДатаПриказа3гр = '' OR ДатаПриказа3гр is NULL) OR (ДатаПриказа3гр != '' AND ДатаПриказа3гр is not NULL AND ТУ = 'Центральный федеральный округ' AND СопроводИтог = '') OR (ПриостановкаАКилиРОА = 'Да'))",
                (Filter, UPR,))
        sqllong3 = int(sqllong2)
        self.SqlLongDetectNew()
        if sqllong3 < maxPage:
            maxPage = sqllong3
        elif curPage == int(sqllongPages):
            maxPage = sqllong3 - maxPage * (int(sqllongPages) - 1)
        if sqllong3 == 1:
            maxPage = 1

        self.tableWidget.setRowCount(maxPage)

        if TipDop == 10 and Allchek == "Нерасписанныая 3":
            results2 = kud.execute(
                'SELECT * FROM ( SELECT *, ROW_NUMBER() OVER (ORDER BY Датарегистрации DESC) AS row FROM GU1Group WHERE ((((ДатаприказаосоставеЭГ != "" AND ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != "Не зарегистрирован") AND (ТипГУ = "РОА" OR ТипГУ = "АК") AND (ДатаПриказа2гр != "" and ДатаПриказа2гр is NOT NULL) AND Итог2Гр != "Отриц. отказ") or ((ДатаприказаосоставеЭГ != "" AND ДатаприказаосоставеЭГ != "" AND ДатаприказаосоставеЭГ is not NULL AND ДатаприказаосоставеЭГ != "Не зарегистрирован") AND (ТипГУ != "РОА" AND ТипГУ != "АК"))) AND (Исполнитель3Гр = "" OR Исполнитель3Гр is NULL) AND ОтделУП = ?) AND (ДатаПриказа3гр = "" OR ДатаПриказа3гр is NULL) AND (Договор1Г != "" AND Договор1Г is NOT NULL)) a WHERE row > ? AND row <= ?',
                (UPR, OtPage, DoPage))
        elif TipDop == 10 and FioTake == "Выбрать исполнителя" and Allchek == "Расписанная 3":
            results2 = kud.execute(
                'SELECT * FROM ( SELECT *, ROW_NUMBER() OVER (ORDER BY Датарегистрации DESC) AS row FROM GU1Group WHERE ((Исполнитель3Гр != "" AND Исполнитель3Гр is not NULL) AND ОтделУП = ?) AND ((ДатаПриказа3гр = "" OR ДатаПриказа3гр is NULL) OR (ДатаПриказа3гр != "" AND ДатаПриказа3гр is not NULL AND ТУ = "Центральный федеральный округ" AND СопроводИтог = "") OR (ПриостановкаАКилиРОА = "Да"))) a WHERE row > ? AND row <= ?',
                (UPR, OtPage, DoPage))
        else:
            results2 = kud.execute(
                'SELECT * FROM ( SELECT *, ROW_NUMBER() OVER (ORDER BY Датарегистрации DESC) AS row FROM GU1Group WHERE (Исполнитель3Гр = ? AND ОтделУП = ?) AND ((ДатаПриказа3гр = "" OR ДатаПриказа3гр is NULL) OR (ДатаПриказа3гр != "" AND ДатаПриказа3гр is not NULL AND ТУ = "Центральный федеральный округ" AND СопроводИтог = "") OR (ПриостановкаАКилиРОА = "Да"))) a WHERE row > ? AND row <= ?',
                (FioTake, UPR, OtPage, DoPage))
        if Filter != '':
            results2 = kud.execute(
                'SELECT * FROM ( SELECT *, ROW_NUMBER() OVER (ORDER BY Датарегистрации DESC) AS row FROM GU1Group WHERE (Регистрационныйномер = ? AND ОтделУП = ?) AND ((ДатаПриказа3гр = "" OR ДатаПриказа3гр is NULL) OR (ДатаПриказа3гр != "" AND ДатаПриказа3гр is not NULL AND ТУ = "Центральный федеральный округ" AND СопроводИтог = "") OR (ПриостановкаАКилиРОА = "Да"))) a WHERE row > ? AND row <= ?',
                (Filter, UPR, OtPage, DoPage))
        tablerow = 0

        for row in results2:

            Etap3gr = QtWidgets.QTableWidgetItem(row[114]).text()
            TU = QtWidgets.QTableWidgetItem(row[91]).text()
            if TU == "Центральный федеральный округ":
                TU = "ЦФО"
            elif TU == "Приволжский федеральный округ":
                TU = "ПФО"
            elif TU == "Сибирский федеральный округ":
                TU = "СФО"
            elif TU == "Северо-Западный федеральный округ":
                TU = "СЗФО"
            elif TU == "Дальневосточный федеральный округ":
                TU = "ДФО"
            elif TU == "Уральский федеральный округ":
                TU = "УФО"
            elif TU == "Южный и Северо-Кавказский федеральные округа":
                TU = "ЮСФО"
            elif TU == "Выбрать территориальное Управление":
                TU = ""
            DopZapr = QtWidgets.QTableWidgetItem(row[121]).text()
            NaSK = QtWidgets.QTableWidgetItem(row[119]).text()
            if DopZapr == "Да":
                DopZapr = " Дозапрос"
            else:
                DopZapr = ""
            if NaSK == "Да":
                NaSK = " На СК"
            else:
                NaSK = ""

            ItogPricaz3GR = QtWidgets.QTableWidgetItem(row[101]).text()
            if ItogPricaz3GR != "" and ItogPricaz3GR is not None:
                ItogPricaz3GR_True = datetime.strptime(ItogPricaz3GR, '%d.%m.%Y')
            SoprovodItog = QtWidgets.QTableWidgetItem(row[123]).text()

            if (Etap3gr == "" or Etap3gr is None):
                Etap3gr = "(Не выбрано)\n"
            elif Etap3gr == "Приостановка 353":
                Etap3gr = "Приостановка по 353\n"
            else:
                Etap3gr = "(" + Etap3gr + ") " + "\n" + TU + DopZapr + NaSK
            if (ItogPricaz3GR != "" and ItogPricaz3GR is not None) and SoprovodItog == "":
                Etap3gr = "Нужно направить Сопровод " + TU + "\n"

            self.tableWidget.setItem(tablerow, 2, QtWidgets.QTableWidgetItem(row[0]))
            self.tableWidget.setItem(tablerow, 3, QtWidgets.QTableWidgetItem(row[1]))
            self.tableWidget.setItem(tablerow, 4, QtWidgets.QTableWidgetItem(row[2]))

            try:
                ShortFIO = QtWidgets.QTableWidgetItem(row[108]).text()
                ShortFIO = ShortFIO.split()
                if ShortFIO != []:
                    ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
                else:
                    ShortFIO = ""
            except:
                ShortFIO = QtWidgets.QTableWidgetItem(row[108]).text()

            Komment = QtWidgets.QTableWidgetItem(row[107]).text()
            if Komment == "" or Komment is None:
                Komment = ""
            self.tableWidget.setItem(tablerow, 1, QtWidgets.QTableWidgetItem(Etap3gr + "| " + Komment))
            self.tableWidget.setItem(tablerow, 6, QtWidgets.QTableWidgetItem(ShortFIO))
            self.tableWidget.setItem(tablerow, 5, QtWidgets.QTableWidgetItem(row[3]))
            self.tableWidget.setItem(tablerow, 7, QtWidgets.QTableWidgetItem(row[4]))
            if QtWidgets.QTableWidgetItem(row[20]).text() != "":
                ShortEXP = QtWidgets.QTableWidgetItem(row[20]).text()
                ShortEXP = ShortEXP.split()
                ShortEXP = ShortEXP[0] + " " + ShortEXP[1][0:1] + "." + ShortEXP[2][0:1] + "."
            else:
                ShortEXP = ""
            self.tableWidget.setItem(tablerow, 8, QtWidgets.QTableWidgetItem(ShortEXP))
            self.tableWidget.setItem(tablerow, 10, QtWidgets.QTableWidgetItem(row[109]))
            self.tableWidget.setItem(tablerow, 11, QtWidgets.QTableWidgetItem(row[110]))
            self.tableWidget.setItem(tablerow, 9, QtWidgets.QTableWidgetItem(row[111]))

            cal = Russia()
            DataDogovor = QtWidgets.QTableWidgetItem(row[84]).text()
            #DataZapPrik = QtWidgets.QTableWidgetItem(row[115]).text()
            DataActa = QtWidgets.QTableWidgetItem(row[111]).text()
            PriostAKorROA = QtWidgets.QTableWidgetItem(row[120]).text()
            try:
                DataActa_date_True = datetime.strptime(DataActa, '%d.%m.%Y')
                if QtWidgets.QTableWidgetItem(row[121]).text() == "Да" or QtWidgets.QTableWidgetItem(row[119]).text() == "Да":
                    datta = cal.add_working_days(DataActa_date_True, 20)
                    DataActa_date = datta
                else:
                    DataActa_date = datetime.strptime(DataActa, '%d.%m.%Y')
            except:
                pass

            datenow = QDate.currentDate().toPyDate()  # сегодня
            Data_Pricaza = QtWidgets.QTableWidgetItem(row[25]).text().split()
            try:
                Data_Pricaza = Data_Pricaza[2].replace("г.", "")
            except:
                Data_Pricaza = Data_Pricaza[0]


            if DataActa == "" or DataActa is None:
                if DataDogovor == "" or DataDogovor is None:
                    GU1Status = 0
                    self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem("Нет Даты Договора"))
                    self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem("Прошло"))
                    Data_Pricaz1GR_date = datetime.strptime(Data_Pricaza, '%d.%m.%Y')
                    datta = cal.get_working_days_delta(Data_Pricaz1GR_date, datenow) + 1  # считает даты в рабочем
                    datta13 = str(datta)
                    self.tableWidget.setItem(tablerow, 14, QtWidgets.QTableWidgetItem(datta13))
                else:
                    GU1Status = 1
                    self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem("Ждем Акт"))
                    self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem("Прошло"))
                    DataDogovor_date = datetime.strptime(DataDogovor, '%d.%m.%Y')
                    datta = cal.get_working_days_delta(DataDogovor_date, datenow) + 1  # считает даты в рабочем
                    datta13 = str(datta)
                    self.tableWidget.setItem(tablerow, 14, QtWidgets.QTableWidgetItem(datta13))

            Etap3gr = QtWidgets.QTableWidgetItem(row[114]).text()
            if (DataActa != "" and DataActa is not None) and (Etap3gr != "Приказ на согласовании" and Etap3gr != "Приказ зарегистрирован"):
                GU1Status = 2
                datta = cal.add_working_days(DataActa_date, 10)
                datta2 = datta.strftime("%d.%m.%Y")
                datta13 = str(datta2)
                if (Etap3gr == "" or Etap3gr is None):
                    Etap3gr = "Акт на рассмотрении"
                self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem("Акт на рассмотрении"))
                self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem(datta13))
                datta4 = cal.get_working_days_delta(DataActa_date_True, datenow) + 1  # считает даты в рабочем
                datta14 = str(datta4)
                self.tableWidget.setItem(tablerow, 14, QtWidgets.QTableWidgetItem(datta14))
                if datta4 < 10:
                    self.tableWidget.item(tablerow, 13).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                    self.tableWidget.item(tablerow, 14).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                else:
                    self.tableWidget.item(tablerow, 13).setForeground(QColor("#ff0000"))  # цвет текста!!!!!!!!!!!
                    self.tableWidget.item(tablerow, 14).setForeground(QColor("#ff0000"))  # цвет текста!!!!!!!!!!!

            if (DataActa != "" and DataActa is not None) and (Etap3gr == "Приказ на согласовании" or Etap3gr == "Приказ зарегистрирован"):
                GU1Status = 3
                try:
                    #DataZapPrik_date = datetime.strptime(DataZapPrik, '%d.%m.%Y')
                    datta = cal.add_working_days(DataActa_date, 13)
                    #datta = cal.add_working_days(DataZapPrik_date, 3)
                    datta2 = datta.strftime("%d.%m.%Y")
                    datta13 = str(datta2)
                except:
                    #DataZapPrik_date = datetime.strptime(str(datenow), '%d-%m-%Y') # если ошибка в дате приказа 3 гркппы. Похоже возникает когда выбирают сразу приказ зарегистрирован
                    datta = cal.add_working_days(DataActa_date, 13)
                    datta2 = datta.strftime("%d.%m.%Y")
                    datta13 = str(datta2)
                self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem(Etap3gr))
                self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem(datta13))
                datta4 = cal.get_working_days_delta(DataActa_date_True, datenow) + 1  # считает даты в рабочем
                datta14 = str(datta4)
                self.tableWidget.setItem(tablerow, 14, QtWidgets.QTableWidgetItem(datta14))
                if datta4 < 13:
                    self.tableWidget.item(tablerow, 13).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                    self.tableWidget.item(tablerow, 14).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                else:
                    self.tableWidget.item(tablerow, 13).setForeground(QColor("#ff0000"))  # цвет текста!!!!!!!!!!!
                    self.tableWidget.item(tablerow, 14).setForeground(QColor("#ff0000"))  # цвет текста!!!!!!!!!!!

            if (DataActa != "" and DataActa is not None) and (Etap3gr == "Приостановка 353" or Etap3gr == "Приостановка 353"):
                GU1Status = 3
                self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem(Etap3gr))
                self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem("Прошло"))
                datta4 = cal.get_working_days_delta(DataActa_date_True, datenow) + 1  # считает даты в рабочем
                datta14 = str(datta4)
                self.tableWidget.setItem(tablerow, 14, QtWidgets.QTableWidgetItem(datta14))
                self.tableWidget.item(tablerow, 13).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                self.tableWidget.item(tablerow, 14).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!

            if (ItogPricaz3GR != "" and ItogPricaz3GR is not None) and SoprovodItog == "":
                GU1Status = 3
                self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem("Направьте сопровод"))
                self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem("Прошло"))
                datta4 = cal.get_working_days_delta(ItogPricaz3GR_True, datenow) + 1  # считает даты в рабочем
                datta14 = str(datta4)
                self.tableWidget.setItem(tablerow, 14, QtWidgets.QTableWidgetItem(datta14))
                self.tableWidget.item(tablerow, 13).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                self.tableWidget.item(tablerow, 14).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!

            if (ItogPricaz3GR != "" and ItogPricaz3GR is not None) and PriostAKorROA == "Да" and SoprovodItog != "":
                GU1Status = 3
                TipGU = QtWidgets.QTableWidgetItem(row[0]).text()
                if TipGU != "АК":
                    TipGU = "РОА"
                self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem("Приостановка по " + TipGU))
                self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem("Прошло"))
                datta4 = cal.get_working_days_delta(ItogPricaz3GR_True, datenow) + 1  # считает даты в рабочем
                datta14 = str(datta4)
                self.tableWidget.setItem(tablerow, 14, QtWidgets.QTableWidgetItem(datta14))
                self.tableWidget.item(tablerow, 13).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                self.tableWidget.item(tablerow, 14).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!

            tablerow += 1
            Allchek = ""

        for i in range(maxPage):  # sqllong2 = все ГУ
            # z = self.tableWidget.item(i, 2).text()
            pushButton = QtWidgets.QPushButton("Открыть")
            pushButton2 = QtWidgets.QPushButton("Сохранить")
            # pushButton.setStyleSheet("border-image: url('YSourceGitHub/Pictures/migrate.png')") # Картинка в кнопке

            pushButton.clicked.connect(lambda ch, nomerGU=self.tableWidget.item(i, 3): self.TakeNomerGU(nomerGU.text()))
            pushButton.clicked.connect(lambda ch, btn=pushButton: self.gotoGroup1GU())

            pushButton2.clicked.connect(lambda ch, nomerGU=self.tableWidget.item(i, 3): self.TakeNomerGU(nomerGU.text()))
            pushButton2.clicked.connect(lambda ch, pkud=i: self.TakePortalKuStatus(pkud))
            pushButton2.clicked.connect(lambda ch, btn=pushButton: self.ChangePortal3GR())

            Ispolnitel3Gr = QtWidgets.QTableWidgetItem(row[108]).text()
            if Ispolnitel3Gr != "" and Ispolnitel3Gr is not None and Ispolnitel3Gr != [] and Ispolnitel3Gr != "Исполнитель не выбран" and DataActa != "" and DataActa is not None:
                self.tableWidget.setCellWidget(i, 0, pushButton)
            else:

                self.tableWidget.setCellWidget(i, 0, pushButton2)
                self.tableWidget.setCellWidget(i, 1, pushButton)

                NewIspolnitel3gr = QtWidgets.QComboBox()
                ALLUserLong = len(ALLUser)
                NewIspolnitel3gr.addItems([""])
                for ii in range(ALLUserLong):
                    BigALLUser = ALLUser[ii][0] + " " + ALLUser[ii][1] + " " + ALLUser[ii][2]
                    ShortFIO = BigALLUser.split()
                    ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
                    NewIspolnitel3gr.addItems([BigALLUser])
                self.tableWidget.setCellWidget(i, 6, NewIspolnitel3gr)

                DataActa = QtWidgets.QTextEdit()
                self.tableWidget.setCellWidget(i, 9, DataActa)

                ActVivod = QtWidgets.QComboBox()
                ActVivod.addItems(["", "Положительный", "Отрицательный"])
                self.tableWidget.setCellWidget(i, 10, ActVivod)

                RiskActa = QtWidgets.QComboBox()
                RiskActa.addItems(["", "Низкий", "Средний", "Высокий"])
                self.tableWidget.setCellWidget(i, 11, RiskActa)






    def ChangePortal3GR(self):
        global globalNomerGU
        global PortalKuStatus777


        GR3DataActa = self.tableWidget.cellWidget(PortalKuStatus777, 9).toPlainText()
        GR3VivodActa = self.tableWidget.cellWidget(PortalKuStatus777, 10).currentText()
        GR3RiskActa = self.tableWidget.cellWidget(PortalKuStatus777, 11).currentText()
        GR3Ispolnitel = self.tableWidget.cellWidget(PortalKuStatus777, 6).currentText()

        connection = sqlite3.connect("SourceGitHub/DB/DinamicBaseNew.db")
        kud = connection.cursor()
        kud.execute('UPDATE GU1Group SET Исполнитель3Гр = ?, ДатаАкта = ?, ВыводАкта = ?, РискАкта = ? WHERE Регистрационныйномер = ?', (GR3Ispolnitel, GR3DataActa, GR3VivodActa, GR3RiskActa, globalNomerGU, ))
        connection.commit()
        connection.close()

    def NextPage(self):
        global sqllong3
        global curPage
        global OtPage
        global DoPage
        global sqllongPages
        global TextsqllongPages
        global UseTableGU
        global NextPage
        global Allchek
        if UseTableGU == "Старый портал":
            try:
                if int(sqllongPages) != 0 and int(sqllongPages) != 2 and int(sqllongPages) != curPage and int(TextsqllongPages) != curPage:
                    curPage = curPage + 1
                    self.DoPortalKu()
            except:
                return
                print("Вылетело на NextPage")
        elif UseTableGU == "2 Группа":
            try:
                if int(sqllongPages) != curPage and int(TextsqllongPages) != curPage:
                    curPage = curPage + 1
                    self.DoTable2G()
            except:
                return
                print("Вылетело на NextPage")
        elif UseTableGU == "3 Группа":
            try:
                if int(sqllongPages) != curPage and int(TextsqllongPages) != curPage:
                    curPage = curPage + 1
                    if NextPage == "Нерасписанныая 3":
                        Allchek = "Нерасписанныая 3"
                    self.DoTable3G()
            except:
                return
                print("Вылетело на NextPage")


    def BackPage(self):
        global sqllong3
        global curPage
        global OtPage
        global DoPage
        global sqllongPages
        global UseTableGU
        global NextPage
        global Allchek

        if UseTableGU == "Старый портал":
            try:
                if curPage == 1:
                    curPage = 1
                else:
                    curPage = curPage - 1
                self.DoPortalKu()
            except:
                print("Вылетело на BackPage")
        elif UseTableGU == "2 Группа":
            try:
                if curPage == 1:
                    curPage = 1
                else:
                    curPage = curPage - 1
                self.DoTable2G()
            except:
                print("Вылетело на BackPage")
        elif UseTableGU == "3 Группа":
            try:
                if curPage == 1:
                    curPage = 1
                else:
                    curPage = curPage - 1
                if NextPage == "Нерасписанныая 3":
                    Allchek = "Нерасписанныая 3"
                self.DoTable3G()
            except:
                print("Вылетело на BackPage")

    def SqlLongDetectNew(self):
        global sqllong3
        global curPage
        global OtPage
        global DoPage
        global sqllongPages
        global TextsqllongPages

        if curPage == 1:
            OtPage = 0
            DoPage = 30
        elif curPage == 2:
            OtPage = 30
            DoPage = 60
        elif curPage > 2:
            OtPage = 30 * (curPage - 1)
            DoPage = OtPage + 30

        sqllongPages = sqllong3 / 30
        if (isinstance(sqllongPages, float)):
            sqllongPages = int(sqllongPages + 1)
        TextsqllongPages = str(sqllongPages)
        self.Long_Pages.setText(str(curPage) + " из " + TextsqllongPages)




    def SqlLongDetect(self):
        global sqllong3
        global curPage
        global OtPage
        global DoPage
        global sqllongPages
        global TextsqllongPages
        if curPage == 1:
            OtPage = 0
            DoPage = 30
        elif curPage == 2:
            OtPage = 30
            DoPage = 60
        elif curPage > 2:
            OtPage = 30 * (curPage - 1)
            DoPage = OtPage + 30


        sqllongPages = sqllong3 / 30

        if int(sqllongPages) == 0:
            sqllongPages = 1

        if (isinstance(sqllongPages, float)):
            sqllongPages = int(sqllongPages)
            if sqllongPages == 1:
                sqllongPages = sqllongPages + 1
                TextsqllongPages = str(int(sqllongPages))
            else:
                TextsqllongPages = str(int(sqllongPages) + 1)
        elif sqllongPages == 1:
            TextsqllongPages = "1"
        sqllongPages = int(sqllongPages)
        sqllong2 = 30 * curPage

        self.Long_Pages.setText(str(curPage) + " из " + TextsqllongPages)
        sqllongPages = int(sqllongPages) + 1
        return

    def DoPortalKu(self):
        global FindNumGU
        global FullFioLogin
        global TipDop
        global sqllong3
        global curPage
        global OtPage
        global DoPage
        global sqllongPages
        global UseTableGU
        global UPR
        UPRK = UPR
        if UPR == 'Управление аккредитации' or UPR == 'Все Управления':
            UPRK = "О%"

        UseTableGU = "Старый портал"
        self.tableWidget.clear()
        self.tableWidget.setSortingEnabled(False)
        maxPage = 30
        try:
            if curPage > 1:
                curPage = curPage
            else:
                curPage = 1
        except:
            curPage = 1

        self.FilterStatus.show()
        self.Back.show()
        self.Next.show()
        self.Long_Pages.show()
        self.tableWidget.show()
        self.Welcome_image.hide()

        try:
            self.tableWidget.setColumnCount(17)
            self.tableWidget.setColumnWidth(0, 70)
            self.tableWidget.setColumnWidth(1, 250)
            self.tableWidget.setColumnWidth(2, 70)
            self.tableWidget.setColumnWidth(3, 70)
            self.tableWidget.setColumnWidth(4, 70)
            self.tableWidget.setColumnWidth(5, 150)
            self.tableWidget.setColumnWidth(6, 150)
            self.tableWidget.setColumnWidth(7, 150)
            self.tableWidget.setColumnWidth(8, 85)
            self.tableWidget.setColumnWidth(9, 85)
            self.tableWidget.setColumnWidth(10, 85)
            self.tableWidget.setColumnWidth(11, 85)
            self.tableWidget.setColumnWidth(12, 85)
            self.tableWidget.setColumnWidth(13, 85)
            self.tableWidget.setColumnWidth(14, 85)
            self.tableWidget.setColumnWidth(15, 85)
            self.tableWidget.setColumnWidth(16, 150)
            self.tableWidget.setHorizontalHeaderLabels(["Сохранить", "Комментарий", "Тип ГУ", "Номер ГУ", "Дата ГУ", "Заявитель", "Исполнитель по ГУ", "Эксперт ГУ", "Решение 1 гр", "Приост. 2 гр.", "Решение 2 гр", "Дата Акта", "Вывод", "Приост. 3 гр.", "Решение 3 гр", "Дедлайн", "Группа"])
            Filter = self.FilterGU.currentText()
            if Filter != "Все ГУ из портала К" and Filter != '':
                maxPage = 1
                curPage = 1
            if  self.FilterGU.itemText(0) != "Все ГУ из портала К":
                self.FilterGU.clear()
            TakeStatusGU = self.FilterStatus.currentText()

            connection = sqlite3.connect("SourceGitHub/DB/PortalKu.db")
            kud = connection.cursor()
            kudall = connection.cursor()
            FioTake = self.Ispolnitel.currentText()
            ShortFIO = FioTake.split()
            #ShortFIO = FullFioLogin.split()
            try:
                ShortFIO = ShortFIO[0] + " " + ShortFIO[1][0:1] + "." + ShortFIO[2][0:1] + "."
            except:
                ShortFIO = ShortFIO
            if Filter == "Все ГУ из портала К" or Filter == '':
                self.FilterGU.clear()
                self.FilterGU.addItems([QtWidgets.QTableWidgetItem("Все ГУ из портала К").text()])  # Добавить ГУ в поиск

            if TakeStatusGU == "Все статусы": ### Фильтр по статусам Все
                if TipDop == 10 and ShortFIO == ['Выбрать', 'исполнителя'] and (UPR != 'Управление аккредитации' and UPR != 'Все Управления'):
                    kud.execute("SELECT COUNT(0) FROM ALLIN WHERE Отдел = ? AND Группа != 'Завершенная 3 гр'", (UPRK, ))
                elif TipDop == 10 and ShortFIO == ['Выбрать', 'исполнителя'] and (UPR == 'Управление аккредитации' or UPR == 'Все Управления'):
                    kud.execute("SELECT COUNT(0) FROM ALLIN WHERE Отдел LIKE ? AND Группа != 'Завершенная 3 гр'", (UPRK,))
                elif TipDop == 10 and ShortFIO != ['Выбрать', 'исполнителя'] and (UPR == 'Управление аккредитации' or UPR == 'Все Управления'):
                    kud.execute("SELECT COUNT(0) FROM ALLIN WHERE ИсполнительпоГУ = ? AND Группа != 'Завершенная 3 гр'", (ShortFIO, ))
                else:
                    kud.execute("SELECT COUNT(0) FROM ALLIN WHERE Отдел = ? AND ИсполнительпоГУ = ? AND Группа != 'Завершенная 3 гр'", (UPRK, ShortFIO, ))
                if Filter != "Все ГУ из портала К" and Filter != '':
                    kud.execute("SELECT COUNT(0) FROM ALLIN WHERE Отдел = ? AND НомерГУ = ? AND Группа != 'Завершенная 3 гр'", (UPRK, Filter, ))

                sqllong = kud.fetchone()  # Находим длинну базы данных в кортеже
                sqllong2 = sqllong[0]  # Вытаскиваем длину баззы данных из первого картежа
                if TipDop == 10 and ShortFIO == ['Выбрать', 'исполнителя'] and (UPR != 'Управление аккредитации' and UPR != 'Все Управления'):
                    results = kud.execute("SELECT * FROM ALLIN WHERE Отдел = ? AND Группа != 'Завершенная 3 гр'", (UPRK, ))
                elif TipDop == 10 and ShortFIO == ['Выбрать', 'исполнителя'] and (UPR == 'Управление аккредитации' or UPR == 'Все Управления'):
                    results = kud.execute("SELECT * FROM ALLIN WHERE Отдел LIKE ? AND Группа != 'Завершенная 3 гр'", (UPRK, ))
                elif TipDop == 10 and ShortFIO != ['Выбрать', 'исполнителя'] and (UPR == 'Управление аккредитации' or UPR == 'Все Управления'):
                    results = kud.execute("SELECT * FROM ALLIN WHERE ИсполнительпоГУ = ? AND Группа != 'Завершенная 3 гр'",(ShortFIO,))
                else:
                    results = kud.execute("SELECT * FROM ALLIN WHERE Отдел = ? AND ИсполнительпоГУ = ? AND Группа != 'Завершенная 3 гр'", (UPRK, ShortFIO, ))
                if Filter != "Все ГУ из портала К" and Filter != '':
                    results = kud.execute("SELECT * FROM ALLIN WHERE Отдел = ? AND НомерГУ = ? AND Группа != 'Завершенная 3 гр'", (UPRK, Filter, ))


                sqllong3 = int(sqllong2)
                self.SqlLongDetect()
                if sqllong3 < maxPage:
                    maxPage = sqllong3
                elif curPage == int(sqllongPages):
                    maxPage = sqllong3 - maxPage * (int(sqllongPages) - 1)
                if sqllong3 == 1:
                    maxPage = 1

                self.tableWidget.setRowCount(maxPage)
                tablerow = 0
                for row in results:
                    if Filter == "Все ГУ из портала К" or Filter == '':
                        self.FilterGU.addItems([QtWidgets.QTableWidgetItem(row[3]).text()])  # Добавить ГУ в поиск
                    tablerow += 1

                if TipDop == 10 and ShortFIO == ['Выбрать', 'исполнителя'] and (UPR != 'Управление аккредитации' and UPR != 'Все Управления'):
                    results2 = kud.execute('SELECT * FROM (SELECT *, ROW_NUMBER() OVER (ORDER BY Сохранить DESC) AS row FROM ALLIN WHERE Отдел = ? AND Группа != "Завершенная 3 гр") a WHERE row > ? AND row <= ?', (UPRK, OtPage, DoPage))
                elif TipDop == 10 and ShortFIO == ['Выбрать', 'исполнителя'] and (UPR == 'Управление аккредитации' or UPR == 'Все Управления'):
                    results2 = kud.execute('SELECT * FROM (SELECT *, ROW_NUMBER() OVER (ORDER BY Сохранить DESC) AS row FROM ALLIN WHERE Отдел LIKE ? AND Группа != "Завершенная 3 гр") a WHERE row > ? AND row <= ?',(UPRK, OtPage, DoPage))
                elif TipDop == 10 and ShortFIO != ['Выбрать', 'исполнителя'] and (UPR == 'Управление аккредитации' or UPR == 'Все Управления'):
                    results2 = kud.execute('SELECT * FROM (SELECT *, ROW_NUMBER() OVER (ORDER BY Сохранить DESC) AS row FROM ALLIN WHERE ИсполнительпоГУ = ? AND Группа != "Завершенная 3 гр") a WHERE row > ? AND row <= ?',(ShortFIO, OtPage, DoPage))
                else:
                    results2 = kud.execute('SELECT * FROM (SELECT *, ROW_NUMBER() OVER (ORDER BY Сохранить DESC) AS row FROM ALLIN WHERE Отдел = ? AND ИсполнительпоГУ = ? AND Группа != "Завершенная 3 гр") a WHERE row > ? AND row <= ?', (UPRK, ShortFIO, OtPage, DoPage))
                if Filter != "Все ГУ из портала К" and Filter != '':
                    results2 = kud.execute('SELECT * FROM (SELECT *, ROW_NUMBER() OVER (ORDER BY Сохранить DESC) AS row FROM ALLIN WHERE Отдел = ? AND НомерГУ = ? AND Группа != "Завершенная 3 гр") a WHERE row > ? AND row <= ?', (UPRK, Filter, OtPage, DoPage))
                tablerow = 0
            elif TakeStatusGU != "Все статусы":     ### Фильтр по статусам один
                if TipDop == 10 and ShortFIO == ['Выбрать', 'исполнителя'] and (UPR != 'Управление аккредитации' and UPR != 'Все Управления'):
                    kud.execute("SELECT COUNT(0) FROM ALLIN WHERE Отдел = ? AND Группа != 'Завершенная 3 гр' and Группа = ?", (UPRK, TakeStatusGU,))
                elif TipDop == 10 and ShortFIO == ['Выбрать', 'исполнителя'] and (UPR == 'Управление аккредитации' or UPR == 'Все Управления'):
                    kud.execute("SELECT COUNT(0) FROM ALLIN WHERE Отдел LIKE ? AND Группа != 'Завершенная 3 гр' and Группа = ?", (UPRK, TakeStatusGU,))
                elif TipDop == 10 and ShortFIO != ['Выбрать', 'исполнителя'] and (UPR == 'Управление аккредитации' or UPR == 'Все Управления'):
                    kud.execute("SELECT COUNT(0) FROM ALLIN WHERE ИсполнительпоГУ = ? AND Группа != 'Завершенная 3 гр' and Группа = ?", (ShortFIO, TakeStatusGU,))
                else:
                    kud.execute("SELECT COUNT(0) FROM ALLIN WHERE Отдел = ? AND ИсполнительпоГУ = ? AND Группа != 'Завершенная 3 гр' and Группа = ?", (UPRK, ShortFIO, TakeStatusGU, ))
                if Filter != "Все ГУ из портала К" and Filter != '':
                    kud.execute("SELECT COUNT(0) FROM ALLIN WHERE Отдел = ? AND НомерГУ = ? AND Группа != 'Завершенная 3 гр'",
                                (UPRK, Filter, ))

                sqllong = kud.fetchone()  # Находим длинну базы данных в кортеже
                sqllong2 = sqllong[0]  # Вытаскиваем длину баззы данных из первого картежа
                if TipDop == 10 and ShortFIO == ['Выбрать', 'исполнителя'] and (UPR != 'Управление аккредитации' and UPR != 'Все Управления'):
                    results = kud.execute("SELECT * FROM ALLIN WHERE Отдел = ? AND Группа != 'Завершенная 3 гр' and Группа = ?", (UPRK, TakeStatusGU,))
                elif TipDop == 10 and ShortFIO == ['Выбрать', 'исполнителя'] and (UPR == 'Управление аккредитации' or UPR == 'Все Управления'):
                    results = kud.execute("SELECT * FROM ALLIN WHERE Отдел LIKE ? AND Группа != 'Завершенная 3 гр' and Группа = ?", (UPRK, TakeStatusGU,))
                elif TipDop == 10 and ShortFIO != ['Выбрать', 'исполнителя'] and (UPR == 'Управление аккредитации' or UPR == 'Все Управления'):
                    results = kud.execute("SELECT * FROM ALLIN WHERE ИсполнительпоГУ = ? AND Группа != 'Завершенная 3 гр' and Группа = ?",(ShortFIO, TakeStatusGU,))
                else:
                    results = kud.execute("SELECT * FROM ALLIN WHERE Отдел = ? AND ИсполнительпоГУ = ? AND Группа != 'Завершенная 3 гр' and Группа = ?", (UPRK, ShortFIO, TakeStatusGU, ))
                if Filter != "Все ГУ из портала К" and Filter != '':
                    results = kud.execute("SELECT * FROM ALLIN WHERE Отдел = ? AND НомерГУ = ? AND Группа != 'Завершенная 3 гр'",
                                          (UPRK, Filter, ))

               #self.FilterGU.addItems(                    [QtWidgets.QTableWidgetItem("Все ГУ из портала К").text()])  # Добавить ГУ в поиск
                sqllong3 = int(sqllong2)
                self.SqlLongDetect()
                if sqllong3 < maxPage:
                    maxPage = sqllong3
                elif curPage == sqllongPages:
                    maxPage = maxPage * int(sqllongPages) - sqllong3
                if sqllong3 == 1:
                    maxPage = 1

                self.tableWidget.setRowCount(maxPage)
                tablerow = 0
                for row in results:
                    self.FilterGU.addItems([QtWidgets.QTableWidgetItem(row[3]).text()])  # Добавить ГУ в поиск
                    tablerow += 1

                if TipDop == 10 and ShortFIO == ['Выбрать', 'исполнителя']and (UPR != 'Управление аккредитации' and UPR != 'Все Управления'):
                    results2 = kud.execute('SELECT * FROM ( SELECT *, ROW_NUMBER() OVER (ORDER BY Сохранить DESC) AS row FROM ALLIN WHERE Отдел = ? AND Группа != "Завершенная 3 гр" and Группа = ?) a WHERE row > ? AND row <= ?',
                        (UPRK, TakeStatusGU, OtPage, DoPage,))
                elif TipDop == 10 and ShortFIO == ['Выбрать', 'исполнителя'] and (UPR == 'Управление аккредитации' or UPR == 'Все Управления'):
                    results2 = kud.execute('SELECT * FROM ( SELECT *, ROW_NUMBER() OVER (ORDER BY Сохранить DESC) AS row FROM ALLIN WHERE Отдел LIKE ? AND Группа != "Завершенная 3 гр" and Группа = ?) a WHERE row > ? AND row <= ?',
                        (UPRK, TakeStatusGU, OtPage, DoPage,))
                elif TipDop == 10 and ShortFIO != ['Выбрать', 'исполнителя'] and (UPR == 'Управление аккредитации' or UPR == 'Все Управления'):
                    results2 = kud.execute('SELECT * FROM ( SELECT *, ROW_NUMBER() OVER (ORDER BY Сохранить DESC) AS row FROM ALLIN WHERE ИсполнительпоГУ = ? AND Группа != "Завершенная 3 гр" and Группа = ?) a WHERE row > ? AND row <= ?',
                        (ShortFIO, TakeStatusGU, OtPage, DoPage))
                else:
                    results2 = kud.execute('SELECT * FROM ( SELECT *, ROW_NUMBER() OVER (ORDER BY Сохранить DESC) AS row FROM ALLIN WHERE Отдел = ? AND ИсполнительпоГУ = ? AND Группа != "Завершенная 3 гр" and Группа = ?) a WHERE row > ? AND row <= ?',
                        (UPRK, ShortFIO, TakeStatusGU, OtPage, DoPage))
                if Filter != "Все ГУ из портала К" and Filter != '':
                    results2 = kud.execute(
                        'SELECT * FROM ( SELECT *, ROW_NUMBER() OVER (ORDER BY Сохранить DESC) AS row FROM ALLIN WHERE Отдел = ? AND НомерГУ = ? AND Группа != "Завершенная 3 гр") a WHERE row > ? AND row <= ?',
                        (UPRK, Filter, OtPage, DoPage))
                tablerow = 0


            for row in results2:
                self.tableWidget.setItem(tablerow, 1, QtWidgets.QTableWidgetItem(row[1]))
                self.tableWidget.setItem(tablerow, 2, QtWidgets.QTableWidgetItem(row[2]))
                self.tableWidget.setItem(tablerow, 3, QtWidgets.QTableWidgetItem(row[3]))
                self.tableWidget.setItem(tablerow, 4, QtWidgets.QTableWidgetItem(row[4].replace("/", ".")))
                self.tableWidget.setItem(tablerow, 5, QtWidgets.QTableWidgetItem(row[5]))
                self.tableWidget.setItem(tablerow, 6, QtWidgets.QTableWidgetItem(row[6]))
                self.tableWidget.setItem(tablerow, 7, QtWidgets.QTableWidgetItem(row[7]))
                self.tableWidget.setItem(tablerow, 8, QtWidgets.QTableWidgetItem(row[8]))
                self.tableWidget.setItem(tablerow, 9, QtWidgets.QTableWidgetItem(row[9]))
                self.tableWidget.setItem(tablerow, 10, QtWidgets.QTableWidgetItem(row[10]))
                self.tableWidget.setItem(tablerow, 11, QtWidgets.QTableWidgetItem(row[11]))
                self.tableWidget.setItem(tablerow, 12, QtWidgets.QTableWidgetItem(row[12]))
                self.tableWidget.setItem(tablerow, 13, QtWidgets.QTableWidgetItem(row[13]))
                self.tableWidget.setItem(tablerow, 14, QtWidgets.QTableWidgetItem(row[14]))
                self.tableWidget.setItem(tablerow, 15, QtWidgets.QTableWidgetItem(row[15]))
                self.tableWidget.setItem(tablerow, 16, QtWidgets.QTableWidgetItem(row[16]))

                tablerow += 1


            for i in range(maxPage):  # sqllong2 = все ГУ
                pushButton = QtWidgets.QPushButton("Сохранить")
                self.tableWidget.setCellWidget(i, 0, pushButton)


                NewIspolnitel = QtWidgets.QComboBox()
                NewIspolnitel.addItems(["Иванов А.А.", "Дорджиев Д.Д.",
                                        "Шкабура В.В.", "Валендер С.Л.",
                                        "Золотаревский С.Ю.", "Алексеева Я.Д.",
                                        "Бухарова А.В.", "Мухтасимова Р.Р.",
                                        "Байкал Е.А.", "Бобров М.А.",
                                        "Белогуров С.И.", "Бражникова К.Д.", "", "Касаткина А.В.", "Альмашова А.В.", "Голицын А.А."])
                NewIspolnitel.isEditable()
                WhotIsp = self.tableWidget.item(i, 6).text()
                if WhotIsp == "Иванов А.А.":
                    NewIspolnitel.setCurrentIndex(0)
                elif WhotIsp == "Дорджиев Д.Д." or WhotIsp == "(ОЕИ) Дорджиев Д.Д." :
                    NewIspolnitel.setCurrentIndex(1)
                elif WhotIsp == "Шкабура В.В.":
                    NewIspolnitel.setCurrentIndex(2)
                elif WhotIsp == "Валендер С.Л.":
                    NewIspolnitel.setCurrentIndex(3)
                elif WhotIsp == "Золотаревский С.Ю.":
                    NewIspolnitel.setCurrentIndex(4)
                elif WhotIsp == "Алексеева Я.Д." or WhotIsp == "(ОЕИ) Алексеева Я.Д." or WhotIsp == "(ИЛ) Алексеева Я.Д.":
                    NewIspolnitel.setCurrentIndex(5)
                elif WhotIsp == "Бухарова А.В.":
                    NewIspolnitel.setCurrentIndex(6)
                elif WhotIsp == "Мухтасимова Р.Р." or WhotIsp == "Мухтасимова  Р.Р.":
                    NewIspolnitel.setCurrentIndex(7)
                elif WhotIsp == "Байкал Е.А.":
                    NewIspolnitel.setCurrentIndex(8)
                elif WhotIsp == "Бобров М.А.":
                    NewIspolnitel.setCurrentIndex(9)
                elif WhotIsp == "Белогуров С.И.":
                    NewIspolnitel.setCurrentIndex(10)
                elif WhotIsp == "Бражникова К.Д.":
                    NewIspolnitel.setCurrentIndex(11)
                elif WhotIsp == "Касаткина А.В.":
                    NewIspolnitel.setCurrentIndex(13)
                elif WhotIsp == "Альмашова А.В.":
                    NewIspolnitel.setCurrentIndex(14)
                elif WhotIsp == "Голицын А.А.":
                    NewIspolnitel.setCurrentIndex(15)
                else:
                    NewIspolnitel.setCurrentIndex(12)
                if TipDop != 10:
                    NewIspolnitel.setEnabled(False)
                self.tableWidget.setCellWidget(i, 6, NewIspolnitel)

                pushCB = QtWidgets.QComboBox()
                pushCB.addItems(["Расписанные 1 гр", "Расписанные 2 гр", "Расписанные 3 гр", "Нерасписанные 3 гр", "Завершенная 3 гр", "Приостановка"])

                What = self.tableWidget.item(i, 16).text()
                if What == "Расписанные 1 гр":
                    pushCB.setCurrentIndex(0)
                elif What == "Расписанные 2 гр":
                    pushCB.setCurrentIndex(1)
                elif What == "Расписанные 3 гр":
                    pushCB.setCurrentIndex(2)
                elif What == "Нерасписанные 3 гр":
                    pushCB.setCurrentIndex(3)
                elif What == "Завершенная 3 гр":
                    pushCB.setCurrentIndex(4)
                elif What == "Приостановка":
                    pushCB.setCurrentIndex(5)
                if TipDop != 10:
                    pushCB.setEnabled(False)
                self.tableWidget.setCellWidget(i, 16, pushCB)

                '''d = QDate(2020, 1, 1)
                DateACT = QtWidgets.QDateEdit()
                DateACT.setCalendarPopup(True)
                DateACT.setMinimumDate(d)
                self.tableWidget.setCellWidget(i, 7, DateACT)'''

                ActVivod = QtWidgets.QComboBox()
                ActVivod.addItems(["Полож.", "Отриц", ""])
                WhatActVivod = self.tableWidget.item(i, 12).text()
                if WhatActVivod == "Полож.":
                    ActVivod.setCurrentIndex(0)
                elif WhatActVivod == "Отриц":
                    ActVivod.setCurrentIndex(1)
                else:
                    ActVivod.setCurrentIndex(2)
                if TipDop != 10:
                    ActVivod.setEnabled(False)
                self.tableWidget.setCellWidget(i, 12, ActVivod)

                '''ActRisk = QtWidgets.QComboBox()
                ActRisk.addItems(["Низкий", "Средний", "Высокий", ""])
                WhatActRisk = self.tableWidget.item(i, 9).text()
                if WhatActRisk == "Низкий":
                    ActRisk.setCurrentIndex(0)
                elif WhatActRisk == "Средний":
                    ActRisk.setCurrentIndex(1)
                elif WhatActRisk == "Высокий":
                    ActRisk.setCurrentIndex(2)
                else:
                    ActRisk.setCurrentIndex(3)
                if TipDop != 10:
                    ActRisk.setEnabled(False)
                self.tableWidget.setCellWidget(i, 9, ActRisk)'''


                Komment = QtWidgets.QTextEdit()
                Komment.setText(self.tableWidget.item(i, 1).text())

                self.tableWidget.setCellWidget(i, 1, Komment)

                Date1GR = QtWidgets.QTextEdit()
                Date1GR.setText(self.tableWidget.item(i, 8).text())
                self.tableWidget.setCellWidget(i, 8, Date1GR)

                PrioastDate2GR = QtWidgets.QTextEdit()
                PrioastDate2GR.setText(self.tableWidget.item(i, 9).text())
                self.tableWidget.setCellWidget(i, 9, PrioastDate2GR)

                Date2GR = QtWidgets.QTextEdit()
                Date2GR.setText(self.tableWidget.item(i, 10).text())
                self.tableWidget.setCellWidget(i, 10, Date2GR)

                DateACT = QtWidgets.QTextEdit()
                DateACT.setText(self.tableWidget.item(i, 11).text())
                if TipDop != 10:
                    DateACT.setEnabled(False)
                self.tableWidget.setCellWidget(i, 11, DateACT)
                Date1GR = QtWidgets.QTextEdit()

                PrioastDate3GR = QtWidgets.QTextEdit()
                PrioastDate3GR.setText(self.tableWidget.item(i, 13).text())
                self.tableWidget.setCellWidget(i, 13, PrioastDate3GR)

                Date3GR = QtWidgets.QTextEdit()
                Date3GR.setText(self.tableWidget.item(i, 14).text())
                self.tableWidget.setCellWidget(i, 14, Date3GR)

                #Deadline = QtWidgets.QTextEdit()
                #Deadline.setText(self.tableWidget.item(i, 15).text())
                #self.tableWidget.setCellWidget(i, 15, Deadline)

                cal = Russia()
                DateACTzz = self.tableWidget.item(i, 11).text()
                datenow = QDate.currentDate().toPyDate()
                if (DateACTzz != "" and DateACTzz is not None):
                    DataActa_date = datetime.strptime(DateACTzz, '%d.%m.%Y')
                    datta = cal.add_working_days(DataActa_date, 15)
                    datta2 = datta.strftime("%d.%m.%Y")
                    datta13 = str(datta2)
                    datta = cal.get_working_days_delta(DataActa_date, datenow) + 1  # считает даты в рабочем
                    datta4 = str(datta)
                    self.tableWidget.setItem(i, 15, QtWidgets.QTableWidgetItem(datta13))
                    if int(datta4) < 15:
                        self.tableWidget.item(i, 15).setForeground(QColor("#00aa00"))  # цвет текста!!!!!!!!!!!
                    else:
                        self.tableWidget.item(i, 15).setForeground(QColor("#ff0000"))  # цвет текста!!!!!!!!!!!




                #pushCB.currentIndexChanged[str].connect(lambda ch, nomerGU=self.tableWidget.item(i, 3): self.TakeNomerGU(nomerGU.text()))
                #pushCB.currentIndexChanged[int].connect(lambda ch, PortalKuStat=self.tableWidget.item(i, 11): self.TakePortalKuStatus(PortalKuStat.text()))
                #pushCB.currentIndexChanged[int].connect(lambda ch, PortalKuStat = self.tableWidget.cellWidget(i, 11).currentIndex(): self.TakePortalKuStatus(PortalKuStat))
                #pushCB.currentIndexChanged[str].connect(lambda ch, btn=pushCB: self.ChangePortalKu())

                pushButton.clicked.connect(lambda ch, nomerGU=self.tableWidget.item(i, 3): self.TakeNomerGU(nomerGU.text()))
                pushButton.clicked.connect(lambda ch, pkud = i: self.TakePortalKuStatus(pkud))
                pushButton.clicked.connect(lambda ch, btn=pushButton: self.ChangePortalKu())

            connection.commit()
            connection.close()
        except AssertionError:
            connection.close()
            self.DoPortalKu()

    def ChangePortalKu(self):
        global globalNomerGU
        global PortalKuStatus777
        PortalKuKomment = self.tableWidget.cellWidget(PortalKuStatus777, 1).toPlainText()
        PortalKuIspolnitel = self.tableWidget.cellWidget(PortalKuStatus777, 6).currentText()
        PortalDate1GR = self.tableWidget.cellWidget(PortalKuStatus777, 8).toPlainText()
        PortalPrioastDate2GR = self.tableWidget.cellWidget(PortalKuStatus777, 9).toPlainText()
        PortalDate2GR = self.tableWidget.cellWidget(PortalKuStatus777, 10).toPlainText()
        PortalAct = self.tableWidget.cellWidget(PortalKuStatus777, 11).toPlainText()
        PortalKuActVivod = self.tableWidget.cellWidget(PortalKuStatus777, 12).currentText()
        PortalPrioastDate3GR = self.tableWidget.cellWidget(PortalKuStatus777, 13).toPlainText()
        PortalDate3GR = self.tableWidget.cellWidget(PortalKuStatus777, 14).toPlainText()
        #PortalDeadline = self.tableWidget.cellWidget(PortalKuStatus777, 15).toPlainText()
        PortalKuStat = self.tableWidget.cellWidget(PortalKuStatus777, 16).currentText()

        connection = sqlite3.connect("SourceGitHub/DB/PortalKu.db")
        kud = connection.cursor()
        kud.execute('UPDATE ALLIN SET Комментарий = ?, ИсполнительпоГУ = ?, Решение1гр = ?, Приост2гр = ?, Решение2гр = ?, Датаакта = ?, Вывод = ?, Приост3гр = ?, Решение3гр = ?, Группа = ? WHERE НомерГУ = ?', (PortalKuKomment, PortalKuIspolnitel, PortalDate1GR, PortalPrioastDate2GR, PortalDate2GR, PortalAct, PortalKuActVivod, PortalPrioastDate3GR, PortalDate3GR, PortalKuStat, globalNomerGU, ))
        connection.commit()
        connection.close()











# main
app = QApplication(sys.argv)

welcome = LoadScreen()
widget = QtWidgets.QStackedWidget()
widget.addWidget(welcome)
# widget.setFixedHeight(1080)
# widget.setFixedWidth(1920)
widget.setGeometry(-1920, 0, 5760, 1080)

widget.setAttribute(Qt.WA_TranslucentBackground, True)
# widget.setStyleSheet("""background: rgb(65, 148, 216);""")

widget.setWindowFlags(widget.windowFlags() | Qt.FramelessWindowHint)

widget.show()

try:
    sys.exit(app.exec_())
except:
    What = "Exiting"
